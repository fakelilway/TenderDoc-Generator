"""公司组件照搬测试:首格锚匹配 + 字体实化 + 整表替换,一模一样(2026-07-07 用户拍板)。"""
import io
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn


def _make_component(first_cell: str, body_text: str) -> Document:
    doc = Document()
    doc.add_paragraph("组件标题")
    t = doc.add_table(rows=2, cols=1)
    t.cell(0, 0).text = first_cell
    t.cell(1, 0).text = body_text
    return doc


def _make_host() -> Document:
    doc = Document()
    doc.add_paragraph("六、项目管理机构")
    t1 = doc.add_table(rows=2, cols=1)
    t1.cell(0, 0).text = "拟为承包本标段工程设立的组织机构以框图方式表示"
    doc.add_paragraph("投标人企业组织结构框图")
    t2 = doc.add_table(rows=2, cols=1)
    t2.cell(0, 0).text = "以框图方式表示。"
    doc.add_paragraph("（四）某正文大表,绝不能被组件替换")
    t3 = doc.add_table(rows=6, cols=2)
    t3.cell(0, 0).text = "以框图方式表示。"  # 首格撞车但行数>3,不许动
    return doc


def test_host_matching_by_first_cell_no_cross_replace() -> None:
    """首格锚精确配对:两个组件各找各的空框,不张冠李戴;>3行大表绝不匹配。"""
    from services.company_component_service import _host_tables_matching

    host = _make_host()
    pm_hits = _host_tables_matching(host, "拟为承包本标段工程设立的组织机构以框图方式表示", [])
    org_hits = _host_tables_matching(host, "以框图方式表示。", [])
    assert len(pm_hits) == 1
    assert len(org_hits) == 1
    assert pm_hits[0] is not org_hits[0]
    # 大表(6行)虽首格相同,不在命中列表
    assert host.tables[2]._tbl not in org_hits


def test_solidify_fonts_writes_defaults_without_overriding() -> None:
    """无显式字体的 run 实化成组件默认;已有显式字体的不覆盖。"""
    from docx.oxml import OxmlElement
    from services.company_component_service import _solidify_fonts

    doc = Document()
    p = doc.add_paragraph("架构框文字")
    r_plain = p.runs[0]._r  # 无 rFonts
    p2 = doc.add_paragraph()
    r_styled = p2.add_run("已定字体")._r
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(fonts)
    r_styled.insert(0, rpr)

    _solidify_fonts(doc.element.body, "Times New Roman", "宋体")
    f1 = r_plain.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert f1.get(qn("w:eastAsia")) == "宋体"
    assert f1.get(qn("w:ascii")) == "Times New Roman"
    f2 = r_styled.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert f2.get(qn("w:eastAsia")) == "黑体"  # 显式字体保留


def test_fill_components_replaces_matching_blank(monkeypatch) -> None:
    """整链:组件按首格锚替换宿主空框,内容原样进宿主,替换表进 handled。"""
    import services.company_component_service as m

    comp = _make_component("拟为承包本标段工程设立的组织机构以框图方式表示", "员工画好的架构图")
    buf = io.BytesIO()
    comp.save(buf)
    comp_bytes = buf.getvalue()

    monkeypatch.setattr(
        m, "list_components",
        lambda: [{
            "document_id": 1, "file_name": "项目管理机构.docx",
            "file_path": "x", "component_type": "项目管理机构", "anchors": [],
        }],
    )
    monkeypatch.setattr(m, "_component_docx_bytes", lambda path: comp_bytes)

    host = _make_host()
    result = m.fill_company_components(host)
    assert result["replaced"] == 1
    texts = ["".join(c.text for r in t.rows for c in r.cells) for t in host.tables]
    assert any("员工画好的架构图" in t for t in texts)  # 成品进来了
    assert len(result["handled_tables"]) == 1


def test_title_anchor_fallback_when_first_cell_differs() -> None:
    """招标空框措辞和组件首格对不上时,标题锚兜底:标题命中→其后≤3行空框照样替换。

    用户规则(2026-07-08):招标文件有要求(标题出现),固定格式直接放进去,别挑空框措辞。"""
    from services.company_component_service import _host_tables_matching

    doc = Document()
    doc.add_paragraph("六、项目管理机构")
    t = doc.add_table(rows=2, cols=1)
    t.cell(0, 0).text = "此处以图表形式说明拟派组织机构"  # 措辞与组件首格完全不同
    hits = _host_tables_matching(
        doc, "拟为承包本标段工程设立的组织机构以框图方式表示", ["项目管理机构"]
    )
    assert len(hits) == 1
    assert hits[0] is t._tbl
