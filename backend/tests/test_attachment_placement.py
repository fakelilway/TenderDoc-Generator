"""附件锚点落位引擎测试:标记 anchor 解析 + 章节定位 + 就地落位(营业执照→基本情况表后等)。"""
from docx import Document
from docx.oxml.ns import qn


def test_parse_marker_reads_anchor() -> None:
    from utils.docx_exporter import _parse_knowledge_image_marker

    m = _parse_knowledge_image_marker(
        '{{knowledge_image:document_id=5 anchor="基本情况表" caption="营业执照" width_cm=12}}'
    )
    assert m["document_id"] == 5
    assert m["anchor"] == "基本情况表"
    assert m["width_cm"] == 12
    m2 = _parse_knowledge_image_marker('{{knowledge_image:document_id=5 caption="x"}}')
    assert m2["anchor"] == ""  # 无 anchor 默认空 → 走卷尾兜底


def test_anchor_finds_basic_info_table_skipping_toc() -> None:
    """基本情况表锚点:表格优先(表头含'投标人名称')→ 绕开目录里的同名段落。"""
    from services.generation_service import _anchor_section_end_element

    doc = Document()
    doc.add_paragraph("（一）投标人基本情况表")  # 目录式同名段,应被表格优先绕过
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "投标人名称"
    t.cell(0, 1).text = "安徽正奇建设有限公司"
    doc.add_paragraph("（二）组织机构框图")
    el = _anchor_section_end_element(doc, "基本情况表")
    assert el is not None and el.tag == qn("w:tbl")  # 命中表格本身


def test_anchor_id_proof_excludes_authorization_letter() -> None:
    """法人亲签:身份证明锚点命中'附:法定代表人身份证复印件(正反面)',排除带'委托代理人'的授权委托书那份。"""
    from services.generation_service import _anchor_section_end_element

    doc = Document()
    doc.add_paragraph("附：法定代表人身份证复印件及委托代理人身份证复印件（正反面）。")
    p2 = doc.add_paragraph("附：法定代表人身份证复印件（正反面）。")
    doc.add_paragraph("三、联合体协议书")
    el = _anchor_section_end_element(doc, "法定代表人身份证明")
    assert el is p2._p  # 命中身份证明那份(节末即 p2,下一块是节标题)


def test_anchor_missing_returns_none() -> None:
    from services.generation_service import _anchor_section_end_element

    doc = Document()
    doc.add_paragraph("一些无关内容")
    assert _anchor_section_end_element(doc, "基本情况表") is None


def test_place_anchored_images_consumes_anchored_keeps_unanchored(monkeypatch) -> None:
    """带 anchor 的图就地落位(取图失败也落占位、消费标记);无 anchor 的留 prose 走卷尾。"""
    import services.generation_service as gs

    monkeypatch.setattr(gs, "_resolve_knowledge_image", lambda did: None)  # 不连 MinIO
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "投标人名称"
    t.cell(0, 1).text = "安徽正奇"
    doc.add_paragraph("（二）下一节")
    prose = (
        '{{knowledge_image:document_id=999999 anchor="基本情况表" caption="营业执照" width_cm=14}}\n'
        '{{knowledge_image:document_id=888888 caption="无锚点图" width_cm=14}}'
    )
    remaining = gs._place_anchored_images(doc, prose)
    assert "anchor=" not in remaining  # 带锚点的被消费
    assert "888888" in remaining  # 无锚点的留 prose(卷尾兜底)
    # 占位/图注落在基本情况表(表格)之后
    texts = [p.text for p in doc.paragraphs]
    assert any("营业执照" in t for t in texts)


def test_drop_empty_headings_removes_emptied_appendix() -> None:
    from services.generation_service import _drop_empty_headings

    # 锚点图全被移走后:附录及其子标题后面只剩别的标题/空行 → 整组被丢
    out = _drop_empty_headings(
        ["## 附录：资格证明材料", "### 营业执照", "### 企业资质证书", "", ""]
    )
    assert "营业执照" not in out and "资格证明材料" not in out and "资质证书" not in out
    # 真有内容的标题保留
    out2 = _drop_empty_headings(["## 有内容的节", "正文一行", "### 子节", "子节正文"])
    assert "有内容的节" in out2 and "正文一行" in out2 and "子节" in out2
