import base64
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from utils.docx_exporter import (
    VOLUME_MARKERS,
    build_export_filename,
    combine_delivery_volumes,
    markdown_to_docx,
    split_delivery_markdown,
    strip_meta_notes,
)


def test_markdown_to_docx_exports_headings_paragraphs_bullets_and_table(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "bid.docx"
    markdown = """# 技术标书

## 施工组织设计

本章节说明施工部署。

- 质量控制
- 安全文明施工

| 项目 | 措施 |
| --- | --- |
| 质量 | 三检制 |
"""

    markdown_to_docx(markdown, output_path)

    assert output_path.exists()
    document = Document(output_path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "技术标书" in texts
    assert "施工组织设计" in texts
    assert "本章节说明施工部署。" in texts
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 1).text == "三检制"


def test_markdown_to_docx_renders_fill_in_blanks_as_underlined_runs(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "fill_in.docx"
    markdown = """# 商务文件

投标人：________

| 项目 | 内容 |
| --- | --- |
| 注册地址 | ________ |
"""

    markdown_to_docx(markdown, output_path)

    document = Document(output_path)
    body = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("投标人："))
    assert any(run.underline for run in body.runs)
    table_cell = document.tables[0].cell(1, 1).paragraphs[0]
    assert any(run.underline for run in table_cell.runs)


def test_markdown_to_docx_renders_pagebreak_marker(tmp_path: Path) -> None:
    output_path = tmp_path / "pagebreak.docx"
    markdown = """# 商务文件

## 一、投标函

正文一。

<!-- tdg:pagebreak -->

## 二、授权委托书

正文二。
"""

    markdown_to_docx(markdown, output_path)

    document = Document(output_path)
    assert 'w:type="page"' in document.element.xml


def test_markdown_to_docx_renders_bidder_basic_info_locked_table(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "bidder_basic.docx"
    markdown = """# 商务文件

## （一）投标人基本情况表

{{tdg_table:bidder_basic_info company_name="安徽正奇建设有限公司" address="________" postal_code="________"}}
"""

    markdown_to_docx(markdown, output_path)

    document = Document(output_path)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 15
    assert len(table.columns) == 6
    assert table.cell(0, 0).text == "投标人名称"
    assert "安徽正奇建设有限公司" in table.cell(0, 1).text
    assert table.cell(2, 0).text == "联系方式"
    assert "投标人应提供关联企业情况" in table.cell(13, 1).text


def test_markdown_to_docx_right_aligns_signature_lines(tmp_path: Path) -> None:
    output_path = tmp_path / "signature.docx"
    markdown = """# 商务文件

投 标 人： （盖单位章）

法定代表人： （签字或盖章）

日 期： 年 月 日
"""

    markdown_to_docx(markdown, output_path)

    document = Document(output_path)
    aligned = [
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith(("投 标 人", "法定代表人", "日 期"))
    ]
    assert len(aligned) == 3
    assert all(paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT for paragraph in aligned)


def test_markdown_to_docx_inserts_knowledge_image_marker(tmp_path: Path) -> None:
    output_path = tmp_path / "bid_with_image.docx"
    tiny_png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )
    markdown = """# 资格文件

## 资格审查资料

{{knowledge_image:document_id=36 caption="一级建造师证书" width_cm=6}}
"""

    markdown_to_docx(
        markdown,
        output_path,
        image_resolver=lambda document_id: tiny_png if document_id == 36 else None,
    )

    document = Document(output_path)
    assert len(document.inline_shapes) == 1
    assert "一级建造师证书" in [paragraph.text for paragraph in document.paragraphs]


def test_markdown_to_docx_adds_cover_toc_header_and_page_numbers(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "bid.docx"
    markdown = """# 高层住宅投标文件

## 施工组织设计

施工部署与质量控制措施。
"""

    markdown_to_docx(
        markdown,
        output_path,
        title="高层住宅投标文件",
        subtitle="投标文件",
        cover=True,
        toc=True,
        header_text="高层住宅投标文件",
        page_numbers=True,
    )

    document = Document(output_path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    # Cover title and table-of-contents heading are present.
    assert "高层住宅投标文件" in texts
    assert "目录" in texts
    # TOC field is embedded in the body and updates inside Word.
    assert "TOC" in document.element.xml
    # Running header carries the document title.
    assert document.sections[0].header.paragraphs[0].text == "高层住宅投标文件"
    # Footer carries a PAGE field for page numbering.
    footer_xml = document.sections[0].footer.paragraphs[0]._p.xml
    assert "PAGE" in footer_xml


def test_split_delivery_markdown_keyword_fallback_routes_sections() -> None:
    markdown = """# 某项目投标文件

## 施工组织设计

施工部署。

## 项目管理机构

人员配置。

## 投标报价说明

报价构成。

## 资格审查资料

营业执照与资质。
"""

    volumes = split_delivery_markdown(markdown)

    assert set(volumes) == {"commercial", "technical", "pricing"}
    assert "施工组织设计" in volumes["technical"]
    assert "项目管理机构" in volumes["technical"]
    assert "投标报价说明" in volumes["pricing"]
    assert "资格审查资料" in volumes["commercial"]
    # Each volume keeps a document-level title prefix.
    assert volumes["technical"].startswith("# 某项目投标文件（技术文件）")
    assert volumes["commercial"].startswith("# 某项目投标文件（商务文件）")


def test_combine_delivery_volumes_round_trip_is_lossless() -> None:
    volumes = {
        "commercial": "# 商务文件\n\n## 资格审查资料\n\n营业执照与资质。",
        "technical": "# 技术文件\n\n## 施工组织设计\n\n施工部署。",
        "pricing": "# 报价文件\n\n## 投标报价说明\n\n报价构成。",
    }

    combined = combine_delivery_volumes("某项目投标文件", volumes, notes="## 审查修正说明\n\n- 修正项")
    recovered = split_delivery_markdown(combined)

    for label, body in volumes.items():
        assert recovered[label] == body
    # The notes section is excluded from every delivery volume.
    assert all("审查修正说明" not in body for body in recovered.values())


def test_strip_meta_notes_removes_markers_and_meta_sections() -> None:
    combined = combine_delivery_volumes(
        "某项目投标文件",
        {"technical": "# 技术文件\n\n## 施工组织设计\n\n施工部署。"},
        notes="## 人工修正意见\n\n补充保证金响应。",
    )

    stripped = strip_meta_notes(combined)

    assert "tdg:volume" not in stripped
    assert "人工修正意见" not in stripped
    assert "补充保证金响应" not in stripped
    assert "施工部署。" in stripped
    assert VOLUME_MARKERS["technical"] not in stripped


def test_build_export_filename_includes_name_and_version() -> None:
    assert (
        build_export_filename("高层 住宅/项目", version=2, kind="技术标")
        == "高层_住宅_项目_技术标_v2.docx"
    )
    assert build_export_filename("", version=0) == "投标文件_v1.docx"


def test_markdown_to_docx_uses_chinese_production_typography(tmp_path: Path) -> None:
    output_path = tmp_path / "bid.docx"
    markdown = """# 投标文件

## 第一章、施工组织设计

本工程采用分段流水施工，明确施工部署与资源配置。
"""

    markdown_to_docx(markdown, output_path)

    document = Document(output_path)
    normal = document.styles["Normal"]
    normal_rfonts = normal.element.rPr.rFonts
    # Body is 宋体 / Times New Roman at 小四 (12pt).
    assert normal_rfonts.get(qn("w:eastAsia")) == "SimSun"
    assert normal.font.size.pt == 12

    heading1 = document.styles["Heading 1"]
    assert heading1.element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert heading1.font.bold is True

    # Body paragraphs carry a 2-character first-line indent.
    body = next(p for p in document.paragraphs if p.text.startswith("本工程采用分段流水施工"))
    assert body.paragraph_format.first_line_indent is not None
    assert body.paragraph_format.first_line_indent.pt == 24


def test_markdown_to_docx_can_apply_zhengqi_bid_style_profile(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "zhengqi.docx"
    markdown = """# 正奇格式投标文件

## 第一章、总体施工组织布置及规划

本工程采用分段流水施工，明确施工部署、资源配置、质量安全环保控制措施。

| 序号 | 项目 | 内容 |
| --- | --- | --- |
| 1 | 质量 | 执行三检制和样板引路制度 |
"""

    markdown_to_docx(
        markdown,
        output_path,
        title="正奇格式投标文件",
        subtitle="投标文件",
        cover=True,
        toc=True,
        header_text="正奇格式投标文件 施工组织设计",
        page_numbers=True,
        style_profile="zhengqi",
    )

    document = Document(output_path)
    normal = document.styles["Normal"]
    normal_rfonts = normal.element.rPr.rFonts
    assert normal_rfonts.get(qn("w:eastAsia")) == "SimSun"
    assert normal.font.size.pt == 14
    assert normal.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert normal.paragraph_format.line_spacing.pt == 32

    heading1 = document.styles["Heading 1"]
    assert heading1.element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
    assert heading1.font.size.pt == 18
    assert heading1.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    cover_title = next(
        paragraph for paragraph in document.paragraphs if paragraph.text == "正奇格式投标文件"
    )
    assert cover_title.runs[0].font.size.pt == 36

    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("本工程采用")
    )
    assert body.paragraph_format.first_line_indent.pt == 28

    section = document.sections[0]
    assert round(section.top_margin.cm, 1) == 2.0
    assert round(section.left_margin.cm, 1) == 2.2
    assert section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    footer_xml = section.footer.paragraphs[0]._p.xml
    assert "PAGE" in footer_xml
    assert "NUMPAGES" in footer_xml
    assert "页/共" in section.footer.paragraphs[0].text

    table_run = document.tables[0].cell(1, 1).paragraphs[0].runs[0]
    assert table_run.font.size.pt == 14
