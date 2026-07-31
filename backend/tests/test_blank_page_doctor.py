"""空白页硬闸的单元测试(不依赖 LibreOffice,测纯函数)。

背景:用户 2026-07-30 死命令"不允许任何空白页"。巢湖 Word 真卷实测 11 页空白,
逐一治理的手段(拆段前分页/硬分页符/分节符改接续/压扁/并节)各有对应用例。
"""
from docx import Document
from docx.oxml.ns import qn

from services import blank_page_doctor as bp


def _sect_para(doc, right_margin="1118"):
    """造一个挂分节符的空段。"""
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sect = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz"); pgSz.set(qn("w:w"), "11905"); pgSz.set(qn("w:h"), "16839")
    pgMar = OxmlElement("w:pgMar")
    for k, v in (("top", "1167"), ("right", right_margin), ("bottom", "1295"), ("left", "1672")):
        pgMar.set(qn("w:" + k), v)
    sect.append(pgSz); sect.append(pgMar)
    pPr.append(sect)
    return p, sect


def test_same_page_geometry_tolerates_half_cm_margin_diff() -> None:
    """右边距差0.5cm(1118 vs 1389)要判同版式——巢湖实测就卡在这半厘米上并不了节。"""
    doc = Document()
    _, a = _sect_para(doc, "1118")
    _, b = _sect_para(doc, "1389")
    assert bp._same_page_geometry(a, b)


def test_same_page_geometry_rejects_big_diff_and_size_change() -> None:
    doc = Document()
    _, a = _sect_para(doc, "1118")
    _, b = _sect_para(doc, "1720")  # 差>300twips
    assert not bp._same_page_geometry(a, b)
    _, c = _sect_para(doc, "1118")
    c.find(qn("w:pgSz")).set(qn("w:w"), "16839")  # 横版
    assert not bp._same_page_geometry(a, c)


def test_minimize_para_height_keeps_ct_ppr_order() -> None:
    """压扁属性必须插在 sectPr 之前——追加到末尾会被 Word/LibreOffice 静默忽略(实测)。"""
    doc = Document()
    p, _ = _sect_para(doc)
    assert bp._minimize_para_height(p._p)
    pPr = p._p.find(qn("w:pPr"))
    kids = [c.tag.split("}")[1] for c in pPr]
    assert kids.index("spacing") < kids.index("sectPr")
    assert kids.index("rPr") < kids.index("sectPr")
    # 幂等:再压返回 False(防死循环)
    assert not bp._minimize_para_height(p._p)


def test_remove_break_prefers_page_break_before() -> None:
    doc = Document()
    doc.add_paragraph("上一节内容")
    target = doc.add_paragraph("六、拟分包项目情况表")
    target.paragraph_format.page_break_before = True
    children = list(doc.element.body)
    idx = children.index(target._p)
    assert bp._remove_one_break_before(children, idx) == "段前分页"
    assert target.paragraph_format.page_break_before is not True


def test_redundant_section_para_gets_merged_away() -> None:
    """分节符段压扁后仍独占一页的场景:与下节几何相同 → 整段删除并节。"""
    doc = Document()
    doc.add_paragraph("附表五表格占满一页")
    sect_p, _ = _sect_para(doc, "1118")
    doc.add_paragraph("二、其他材料")
    _sect_para(doc, "1389")  # 下一节(右边距差0.5cm,允许并)

    def _attempt():
        children = list(doc.element.body)
        idx = children.index(doc.paragraphs[2]._p)  # 二、其他材料
        return bp._remove_one_break_before(children, idx)

    # 三级阶梯:改接续 → 压扁 → 并节删除
    assert "改接续" in _attempt()
    assert "压扁" in _attempt()
    assert _attempt() == "冗余分节符段删除(并节)"
    assert sect_p._p.getparent() is None  # 真删了
