import pytest

from services import generation_service
from utils.docx_exporter import combine_delivery_volumes
from docx import Document


@pytest.fixture(autouse=True)
def _isolate_delivery_quality_eval_dir(monkeypatch, tmp_path):
    """出标后的非阻断质量打分钩子会往仓库根 eval_results/ 落 JSON;测试里把它
    重定向到 tmp,避免 export 测试污染工作区(钩子本身仍真实运行)。"""
    from services import delivery_quality

    monkeypatch.setattr(
        delivery_quality, "_EVAL_RESULTS_DIR", tmp_path / "eval_results"
    )


def test_evaluate_generation_quality_counts_placeholders() -> None:
    markdown = """# 标书

## 施工组织设计

这是一个完整段落，说明施工部署、进度安排、质量控制和安全文明施工措施。

待补充
"""

    report = generation_service.evaluate_generation_quality(markdown)

    assert report["total_paragraphs"] == 2
    assert report["needs_revision_paragraphs"] == 1
    assert report["usable_rate"] == 0.5


class _FakeCursor:
    def __init__(self, statements):
        self.statements = statements

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False

    def execute(self, statement, params=None):
        self.statements.append((statement, params))


class _FakeConnection:
    def __init__(self, statements):
        self.statements = statements

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False

    def cursor(self, *args, **kwargs):
        return _FakeCursor(self.statements)


class _FakeMinio:
    def __init__(self):
        self.uploads = []

    def upload_file(self, bucket, file_path, object_name):
        content = None
        if str(file_path).endswith(".md"):
            content = file_path.read_text(encoding="utf-8")
        self.uploads.append((bucket, object_name, content))
        return object_name


def test_export_markdown_for_project_stores_markdown_docx_and_quality(
    monkeypatch,
) -> None:
    statements = []
    fake_minio = _FakeMinio()
    monkeypatch.setattr(
        generation_service, "_connect", lambda: _FakeConnection(statements)
    )
    monkeypatch.setattr(generation_service, "minio_client", fake_minio)

    markdown = "# 项目\n\n## 施工组织设计\n\n这是完整生成段落，描述施工部署、质量、安全和进度。\n"
    quality_report = generation_service.evaluate_generation_quality(markdown)

    markdown_object, docx_object = generation_service.export_markdown_for_project(
        7,
        markdown,
        quality_report,
    )

    assert markdown_object == "projects/7/generated/bid.md"
    assert docx_object == "projects/7/generated/bid.docx"
    assert [upload[1] for upload in fake_minio.uploads] == [
        "projects/7/generated/bid.md",
        "projects/7/generated/bid.docx",
    ]
    assert any("generated_docx_path" in statement for statement, _params in statements)
    assert any(params and "generated" in params for _statement, params in statements)


def test_export_markdown_for_project_prefers_original_docx_format(
    monkeypatch,
) -> None:
    statements = []
    fake_minio = _FakeMinio()
    fake_minio.download_bytes = lambda bucket, object_name: b"docx bytes"
    captured = {}

    def fake_build_original_format_docx(tender_bytes, output_path, *, profile=None):
        captured["tender_bytes"] = tender_bytes
        captured["profile"] = profile
        output_path.write_bytes(b"original format docx")

    def fail_markdown_to_docx(*args, **kwargs):
        raise AssertionError(
            "markdown_to_docx should not be used for original DOCX tender"
        )

    monkeypatch.setattr(
        generation_service, "_connect", lambda: _FakeConnection(statements)
    )
    monkeypatch.setattr(generation_service, "minio_client", fake_minio)
    monkeypatch.setattr(
        generation_service,
        "_fetch_tender_document",
        lambda project_id: {
            "file_name": "招标文件.docx",
            "file_path": "projects/7/tender/original.docx",
            "name": "测试项目",
            "confirmed_parsed_json": {"project_name": "测试项目", "tenderer_name": "招标人"},
            "parsed_json": None,
        },
    )
    monkeypatch.setattr(
        generation_service,
        "build_original_format_docx",
        fake_build_original_format_docx,
    )
    monkeypatch.setattr(generation_service, "markdown_to_docx", fail_markdown_to_docx)
    monkeypatch.setattr(
        generation_service,
        "get_company_profile",
        lambda: {"profile": {"company_name": "安徽正奇建设有限公司"}},
    )

    generation_service.export_markdown_for_project(7, "# 项目\n", {"usable_rate": 1.0})

    assert captured["tender_bytes"] == b"docx bytes"
    assert captured["profile"]["company_name"] == "安徽正奇建设有限公司"
    assert fake_minio.uploads[-1][1] == "projects/7/generated/bid.docx"


def test_export_markdown_for_project_strips_meta_notes(monkeypatch) -> None:
    statements = []
    fake_minio = _FakeMinio()
    captured = {}

    def fake_markdown_to_docx(markdown, docx_path, **kwargs):
        captured["docx_markdown"] = markdown
        captured["title"] = kwargs.get("title")

    monkeypatch.setattr(
        generation_service, "_connect", lambda: _FakeConnection(statements)
    )
    monkeypatch.setattr(generation_service, "minio_client", fake_minio)
    monkeypatch.setattr(generation_service, "markdown_to_docx", fake_markdown_to_docx)

    markdown = combine_delivery_volumes(
        "测试项目投标文件",
        {
            "commercial": "# 商务文件\n\n法定代表人授权书等商务内容，满足资格审查要求。",
            "technical": "# 技术文件\n\n## 施工组织设计\n\n这是完整生成段落，描述施工部署、质量、安全和进度。",
            "pricing": "# 报价文件\n\n投标报价汇总表内容。",
        },
        notes="第 1 轮审查发现 2 处问题，已自动修正。",
    )
    markdown += "\n## 审查修正说明\n\n遗留的旧版审查说明段落。\n"

    generation_service.export_markdown_for_project(7, markdown, {"usable_rate": 1.0})

    uploaded_markdown = next(
        content
        for _bucket, object_name, content in fake_minio.uploads
        if object_name.endswith(".md")
    )
    for exported in (uploaded_markdown, captured["docx_markdown"]):
        assert "tdg:volume" not in exported
        assert "审查发现" not in exported
        assert "审查修正说明" not in exported
        assert "施工组织设计" in exported
        assert "投标报价汇总表内容" in exported
    assert captured["title"] == "测试项目投标文件"


def test_assemble_two_volumes_commercial_copies_format_technical_is_prose(tmp_path) -> None:
    """Editable 照抄: 商务卷=格式章原样, 技术卷=独立生成正文, 不产出报价卷。"""
    format_path = tmp_path / "format.docx"
    fmt = Document()
    fmt.add_paragraph("一、投标函")
    fmt.add_paragraph("致：（招标人），我方愿按招标文件投标。")
    table = fmt.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "投标人名称"
    fmt.save(format_path)

    markdown = combine_delivery_volumes(
        "测试项目",
        {
            "commercial": "# 商务文件\n\n资格响应合规说明。",
            "technical": "# 技术文件\n\n## 施工组织设计\n\n施工部署与技术措施详尽正文。",
        },
    )
    main_docx = tmp_path / "bid.docx"
    generation_service._assemble_two_volumes(
        str(format_path), tmp_path, 8, markdown, main_docx, "测试项目"
    )

    commercial = Document(tmp_path / "project_8_commercial.docx")
    commercial_text = "\n".join(p.text for p in commercial.paragraphs)
    technical_text = "\n".join(p.text for p in Document(tmp_path / "project_8_technical.docx").paragraphs)

    # 商务卷：照抄了格式章（表格保留）+ 合规正文
    assert len(commercial.tables) == 1
    assert "投标函" in commercial_text
    assert "资格响应合规说明" in commercial_text
    # 技术卷：是生成的施工组织设计正文，不含商务格式页
    assert "施工部署与技术措施详尽正文" in technical_text
    assert "投标函" not in technical_text
    # 不产出报价卷
    assert not (tmp_path / "project_8_pricing.docx").exists()
    # 主 bid.docx = 技术卷
    assert "施工部署与技术措施详尽正文" in "\n".join(
        p.text for p in Document(main_docx).paragraphs
    )


def test_two_volume_technical_excludes_commercial_sections(tmp_path) -> None:
    """Regression: 资格响应/投标保证金/项目管理机构 (commercial) must NOT leak into
    the technical卷. Bug was strip_meta_notes dropping tdg:volume markers before
    the split → heading heuristic misrouted commercial sections into technical."""
    format_path = tmp_path / "format.docx"
    fmt = Document()
    fmt.add_paragraph("一、投标函")
    fmt.save(format_path)

    combined = combine_delivery_volumes(
        "测试项目",
        {
            "commercial": (
                "# 测试项目 商务文件\n\n## 资格响应\n\n### 企业资质等级\n资质说明。\n\n"
                "## 投标保证金\n保证金说明。\n\n## 项目管理机构\n机构说明。"
            ),
            "technical": "# 测试项目 技术文件\n\n## 施工组织设计\n\n施工部署正文。\n\n## 其他内容\n其他正文。",
        },
    )
    # Combined still carries tdg:volume markers (as it does coming from the workflow).
    main_docx = tmp_path / "bid.docx"
    generation_service._assemble_two_volumes(
        str(format_path), tmp_path, 9, combined, main_docx, "测试项目"
    )

    technical_text = "\n".join(
        p.text for p in Document(tmp_path / "project_9_technical.docx").paragraphs
    )
    assert "施工部署正文" in technical_text
    assert "资格响应" not in technical_text
    assert "投标保证金" not in technical_text
    assert "项目管理机构" not in technical_text


def test_append_prose_falls_back_to_plaintext_when_styled_render_fails(
    tmp_path, monkeypatch
) -> None:
    """转换件(如福昕云)的内置标题样式 python-docx 按名取不到,会让 add_heading/
    _configure_styles 抛 ``KeyError: no style with name 'Heading 1'``。此时商务卷
    合规正文必须退到纯文本追加 —— 不丢内容、绝不让整卷导出失败(=出标有下载件)。
    回归:用户实测"显示生成好了→突然报 no style with name 'Heading 1'→无下载按钮"。"""
    from utils import docx_exporter

    base = tmp_path / "commercial.docx"
    doc = Document()
    doc.add_paragraph("投标人基本情况表")
    doc.save(base)

    def boom(*_a, **_k):
        raise KeyError("no style with name 'Heading 1'")

    # 函数内 `from utils.docx_exporter import _render_markdown_body` 在调用时读模块属性
    monkeypatch.setattr(docx_exporter, "_render_markdown_body", boom)

    # 不得抛异常;合规正文(带 ## 标题)仍须写入
    generation_service._append_prose_to_docx(
        base, "## 商务响应\n\n本公司承诺响应全部商务条款。"
    )

    text = "\n".join(p.text for p in Document(base).paragraphs)
    assert "投标人基本情况表" in text  # 原格式章保留
    assert "本公司承诺响应全部商务条款" in text  # 合规正文未丢
