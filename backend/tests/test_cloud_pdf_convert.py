import hashlib
import io
import urllib.parse

import httpx
from docx import Document

from services import cloud_pdf_convert


def test_sn_sorts_params_and_appends_secret() -> None:
    # 签名必须:参数排序 → urlencode → 追加 &sk=secret → md5(与福昕官方 demo 一致)
    sn = cloud_pdf_convert._sn("CID", "SEC", {"format": "word", "clientId": "CID"})
    expected = hashlib.md5(
        ("clientId=CID&format=word&sk=" + urllib.parse.quote("SEC")).encode("utf-8")
    ).hexdigest()
    assert sn == expected


def _minimal_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("投标人：测试公司")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _Resp:
    def __init__(self, json_data=None, content=b"", status=200):
        self._json = json_data
        self.content = content
        self.status_code = status

    def raise_for_status(self):
        if self.status_code >= 400:
            raise httpx.HTTPStatusError(
                "err",
                request=httpx.Request("GET", "http://x"),
                response=httpx.Response(self.status_code),
            )

    def json(self):
        return self._json


class _FakeClient:
    def __init__(self, create, task, download):
        self._create, self._task, self._download = create, task, download

    def __enter__(self):
        return self

    def __exit__(self, *a):
        return False

    def post(self, url, **kw):
        return self._create

    def get(self, url, **kw):
        return self._task if url.endswith("/task") else self._download


def test_convert_pdf_to_docx_via_foxit_happy_path(tmp_path, monkeypatch) -> None:
    fake = _FakeClient(
        create=_Resp({"code": 0, "data": {"taskInfo": {"taskId": "t1"}}}),
        task=_Resp({"code": 0, "data": {"taskInfo": {"percentage": 100, "docId": "d1"}}}),
        download=_Resp(content=_minimal_docx_bytes()),
    )
    monkeypatch.setattr(cloud_pdf_convert.httpx, "Client", lambda *a, **k: fake)

    out = str(tmp_path / "out.docx")
    result = cloud_pdf_convert.convert_pdf_to_docx_via_foxit(
        b"%PDF-1.4 fake", out, client_id="CID", secret="SEC", poll_interval_seconds=0
    )
    assert result == out
    assert any(p.text.strip() for p in Document(out).paragraphs)


def test_convert_raises_on_create_error_code(tmp_path, monkeypatch) -> None:
    fake = _FakeClient(
        create=_Resp({"code": 600000, "msg": "Missing parameter"}),
        task=_Resp({}),
        download=_Resp(content=b""),
    )
    monkeypatch.setattr(cloud_pdf_convert.httpx, "Client", lambda *a, **k: fake)
    out = str(tmp_path / "out.docx")
    try:
        cloud_pdf_convert.convert_pdf_to_docx_via_foxit(
            b"x", out, client_id="CID", secret="SEC", poll_interval_seconds=0
        )
        assert False, "应在 code!=0 时抛异常"
    except RuntimeError as exc:
        assert "600000" in str(exc)


def test_convert_raises_on_task_failure_not_spin(tmp_path, monkeypatch) -> None:
    # 任务失败要立刻抛(而非空转到 max_wait)
    fake = _FakeClient(
        create=_Resp({"code": 0, "data": {"taskInfo": {"taskId": "t1"}}}),
        task=_Resp({"code": 0, "data": {"taskInfo": {"status": "failed", "percentage": 30}}}),
        download=_Resp(content=b""),
    )
    monkeypatch.setattr(cloud_pdf_convert.httpx, "Client", lambda *a, **k: fake)
    out = str(tmp_path / "out.docx")
    try:
        cloud_pdf_convert.convert_pdf_to_docx_via_foxit(
            b"x", out, client_id="CID", secret="SEC",
            poll_interval_seconds=0, max_wait_seconds=5,
        )
        assert False, "任务失败应抛异常"
    except RuntimeError as exc:
        assert "失败" in str(exc)
