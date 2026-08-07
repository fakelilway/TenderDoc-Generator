from pathlib import Path

from docx import Document

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.original_docx_format_service import (
    _drop_spurious_stream_tables,
    _fill_known_table_cells,
    _is_blank_or_placeholder,
    _looks_like_next_chapter_page,
    _looks_like_other_volume_start,
    _table_label_value,
    build_original_format_docx,
    unfilled_known_fields,
)


def test_next_chapter_ignores_midsentence_crossreference() -> None:
    # 实测招标#122 p105：跨引用"…招标文件第二章…"不得当成新章(会切掉资格审查后半+八)
    assert not _looks_like_next_chapter_page(
        "106备注注：1.投标人应根据招标文件第二章“投标人须知”第3.5.1项的要求在本表后附材料。"
    )
    # 页首(带页码)的真章标题才算边界
    assert _looks_like_next_chapter_page("106第八章评标办法投标人应仔细阅读")
    assert _looks_like_next_chapter_page("第三章投标人须知")
    # 格式章自身不作为结束边界
    assert not _looks_like_next_chapter_page("第八章投标文件格式")


def test_other_volume_start_marks_commercial_end() -> None:
    # 技术/报价卷起始 → 商务格式章到此为止
    assert _looks_like_other_volume_start("（标段名称）施工招标投标文件（技术文件）投标人：（盖单位章）")
    assert _looks_like_other_volume_start("（标段名称）施工招标投标文件（报价文件）投标人")
    # 商务卷自身、普通表单不触发
    assert not _looks_like_other_volume_start("（标段名称）施工招标投标文件（商务文件）")
    assert not _looks_like_other_volume_start("三、联合体协议书（所有成员单位名称）自愿组成")


def _add_cell_borders(cell) -> None:
    """Give a cell real single borders (mimics pdf2docx's true-table cells)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "start", "end"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        borders.append(el)
    tc_pr.append(borders)


def test_drop_spurious_stream_tables_flattens_borderless_small_tables() -> None:
    doc = Document()
    spurious = doc.add_table(rows=1, cols=2)  # 填空行被误判:2格、无边框
    spurious.cell(0, 0).text = "1.我方已仔细研究"
    spurious.cell(0, 1).text = "标段招标文件的全部内容。"
    big = doc.add_table(rows=15, cols=10)  # 真表(基本情况表式):>2格 → 保留
    big.cell(0, 0).text = "投标人名称"

    dropped = _drop_spurious_stream_tables(doc)

    assert dropped == 1
    assert len(doc.tables) == 1  # 大表保留
    assert any(
        "我方已仔细研究" in p.text and "招标文件的全部内容" in p.text
        for p in doc.paragraphs
    )  # 假表内容已还原成连续段落


def test_drop_spurious_stream_tables_keeps_small_bordered_table() -> None:
    doc = Document()
    bordered = doc.add_table(rows=2, cols=1)  # 2格但有真边框(如项目管理机构图说明)
    _add_cell_borders(bordered.cell(0, 0))
    bordered.cell(0, 0).text = "拟为承包本标段以框图方式表示。"
    bordered.cell(1, 0).text = "说明"

    dropped = _drop_spurious_stream_tables(doc)

    assert dropped == 0
    assert len(doc.tables) == 1  # 有边框的真表不动


def test_table_label_value_maps_known_and_skips_others() -> None:
    profile = {
        "company_name": "安徽正奇建设有限公司",
        "credit_code": "91340100578516708N",
        "legal_representative": "许明英",
        "registered_capital": "10060万元人民币",
    }
    profile["project_manager_name"] = "江舟"
    assert _table_label_value("投标人名称", profile) == "安徽正奇建设有限公司"
    assert _table_label_value("统一社会信用代码", profile) == "91340100578516708N"
    assert _table_label_value("注册资本", profile) == "10060万元人民币"
    assert _table_label_value("项目经理", profile) == "江舟"
    assert _table_label_value("项目经理姓名", profile) == "江舟"
    # 另一个人 / 日期 / 无对应字段 → 不填
    assert _table_label_value("技术负责人", profile) == ""
    assert _table_label_value("成立时间", profile) == ""
    assert _table_label_value("员工总人数：", profile) == ""
    assert _table_label_value("随便什么标题", profile) == ""


def _solid_png(rgb: tuple[int, int, int], size: tuple[int, int] = (80, 80)) -> bytes:
    from io import BytesIO

    from PIL import Image

    image = Image.new("RGBA", size, (*rgb, 255))
    buffer = BytesIO()
    image.save(buffer, "PNG")
    return buffer.getvalue()


def test_image_is_seal_detects_red_stamp_only() -> None:
    from services.original_docx_format_service import _image_is_seal

    assert _image_is_seal(_solid_png((220, 20, 20))) is True  # 红章
    assert _image_is_seal(_solid_png((250, 250, 250))) is False  # 白底
    assert _image_is_seal(_solid_png((15, 15, 15))) is False  # 黑字
    assert _image_is_seal(_solid_png((20, 40, 200))) is False  # 蓝线附表图


def test_strip_seal_images_removes_seals_keeps_rest() -> None:
    """回归(实测真招标商务卷):福昕把招标人/代理红章照搬进来,乱盖一片;只删红章,
    保留文字、表格、非章图(附表线框/页面图)。"""
    from io import BytesIO

    from docx import Document
    from docx.shared import Cm

    from services.original_docx_format_service import _strip_seal_images

    doc = Document()
    doc.add_paragraph("投标人：（盖单位章）")
    doc.add_paragraph().add_run().add_picture(
        BytesIO(_solid_png((220, 20, 20))), width=Cm(3)
    )  # 红章
    doc.add_paragraph("附表二 分项工程进度率计划")
    doc.add_paragraph().add_run().add_picture(
        BytesIO(_solid_png((20, 40, 200))), width=Cm(3)
    )  # 蓝线附表图(非章)

    before = len(doc.element.findall(".//" + qn("w:drawing")))
    removed = _strip_seal_images(doc)
    after = len(doc.element.findall(".//" + qn("w:drawing")))

    assert removed == 1
    assert before - after == 1  # 只删了红章,蓝图还在
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "盖单位章" in text and "附表二" in text  # 文字完整


def test_table_label_value_expanded_aliases() -> None:
    """扩标签别名:招标各家措辞(企业名称/公司名称/法人代表/住所…)都能对上字段。"""
    profile = {
        "company_name": "安徽正奇建设有限公司",
        "legal_representative": "许明英",
        "registered_address": "安徽省合肥市…",
    }
    for label in ("企业名称", "公司名称", "单位名称", "投标人全称"):
        assert _table_label_value(label, profile) == "安徽正奇建设有限公司", label
    assert _table_label_value("法人代表", profile) == "许明英"
    assert _table_label_value("住所", profile).startswith("安徽省")


def test_unfilled_known_fields_flags_recognized_but_empty() -> None:
    """缺字段显式告警:认得标签但档案无值 → 列出来,别静默留空。"""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "投标人名称"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "开户银行"  # 档案缺 → 应被标记
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "联系电话"  # 档案缺 → 应被标记
    table.cell(2, 1).text = ""
    profile = {"company_name": "安徽正奇建设有限公司"}  # 有公司名,缺银行/电话

    missing = dict((label, key) for label, key in unfilled_known_fields(doc, profile))
    keys = set(missing.values())
    assert "bank_name" in keys and "contact_phone" in keys
    assert "company_name" not in keys  # 已填的不算缺


def test_table_label_value_broad_key_does_not_overfill() -> None:
    """回归(实测 122 商务卷):宽泛主体词("投标人"/"项目经理")含在标签里 ≠ 就该填它的
    名字。子字段标签必须留空待人工,绝不能拿公司名/项目经理名瞎填(废标级错误)。"""
    profile = {
        "company_name": "安徽正奇建设有限公司",
        "credit_code": "91340100578516708N",
        "project_manager_name": "江舟",
    }
    # ❌ 原来全被"投标人"/"项目经理"子串命中、瞎填成公司名或"江舟"
    for wrong in (
        "投标人响应资质", "投标人资格业绩", "投标人加分业绩", "投标人荣誉",
        "项目经理身份证号码", "项目经理证书名称", "项目经理证书编号", "项目经理荣誉",
    ):
        assert _table_label_value(wrong, profile) == "", wrong
    # ✓ "要名字"的复合标签仍正确填(联合体牵头人=投标人本身)
    assert (
        _table_label_value("独立投标人或联合体牵头人名称", profile)
        == "安徽正奇建设有限公司"
    )
    # ✓ 最长匹配:具体字段胜过宽泛"投标人",信用代码不再被填成公司名
    assert (
        _table_label_value("独立投标人或联合体牵头人统一社会信用代码", profile)
        == "91340100578516708N"
    )


def test_fill_known_table_cells_fills_adjacent_empty_only() -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "投标人名称"  # (0,1) 空 → 应填
    table.cell(1, 0).text = "统一社会信用代码"
    table.cell(1, 1).text = "已有值"  # 已填 → 不覆盖
    table.cell(2, 0).text = "技术负责人"  # 另一个人 → 跳过,(2,1) 保持空

    profile = {"company_name": "安徽正奇建设有限公司", "credit_code": "91X"}
    filled = _fill_known_table_cells(doc, profile)

    assert table.cell(0, 1).text == "安徽正奇建设有限公司"
    assert table.cell(1, 1).text == "已有值"  # 未被覆盖
    assert table.cell(2, 1).text.strip() == ""  # 技术负责人行未填
    assert filled == 1


def test_fill_known_table_cells_fills_through_sublabel() -> None:
    # 法定代表人 | 姓名 | [空] —— 跨过"姓名"子标签,把值填进真正的值格
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "法定代表人"
    table.cell(0, 1).text = "姓名"

    _fill_known_table_cells(doc, {"legal_representative": "许明英"})

    assert table.cell(0, 2).text.strip() == "许明英"


def test_fill_known_table_cells_skips_second_person_sublabel_row() -> None:
    # 技术负责人 是另一个人(在 skip 列),其"姓名"值格不应被填法定代表人
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "技术负责人"
    table.cell(0, 1).text = "姓名"

    _fill_known_table_cells(doc, {"legal_representative": "许明英"})

    assert table.cell(0, 2).text.strip() == ""


def test_fill_known_table_cells_noop_when_profile_empty() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "投标人名称"
    assert _fill_known_table_cells(doc, {}) == 0
    assert table.cell(0, 1).text.strip() == ""


def test_fill_known_table_cells_reconstructs_fragmented_label() -> None:
    # 中文公文表常把"投标人："逐字对齐拆成 投|标|人： 三格 + 空值格;
    # 单格匹配不到,碎标签重组后应把 company_name 填进值格。
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "投"
    table.cell(0, 1).text = "标"
    table.cell(0, 2).text = "人："
    filled = _fill_known_table_cells(doc, {"company_name": "安徽正奇建设有限公司"})
    assert table.cell(0, 3).text.strip() == "安徽正奇建设有限公司"
    assert filled == 1


def test_fill_known_table_cells_fragmented_does_not_false_match() -> None:
    # 短格拼起来不构成已知标签 → 绝不误填(防碎标签重组造出假标签)。
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "序"
    table.cell(0, 1).text = "号"
    table.cell(0, 2).text = "A1"
    filled = _fill_known_table_cells(doc, {"company_name": "安徽正奇建设有限公司"})
    assert table.cell(0, 3).text.strip() == ""
    assert filled == 0


def test_build_original_format_docx_copies_format_tables_verbatim(
    tmp_path: Path,
) -> None:
    source_path = tmp_path / "tender.docx"
    source = Document()
    source.add_paragraph("第一章 招标公告")
    source.add_paragraph("第八章 投标文件格式")
    source.add_paragraph("投标文件（商务文件）")
    source.add_paragraph("（一）投标人基本情况表")
    table = source.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "投标人名称"
    table.cell(0, 1).merge(table.cell(0, 2))
    table.cell(0, 1).text = "（投标人名称）"
    table.cell(1, 0).text = "注册地址"
    table.cell(1, 1).text = "邮政编码"
    table.cell(1, 2).text = "________"
    table.cell(2, 0).text = "备注"
    table.cell(2, 1).merge(table.cell(2, 2))
    table.cell(2, 1).text = "________"
    source.add_paragraph("第九章 评标办法")
    source.save(source_path)

    output_path = tmp_path / "copied.docx"
    build_original_format_docx(
        source_path.read_bytes(),
        output_path,
        profile={"company_name": "安徽正奇建设有限公司"},
    )

    copied = Document(output_path)
    texts = [paragraph.text for paragraph in copied.paragraphs]
    # 2026-07-31 用户拍板:卷首招标章标题剥掉,封面打头
    assert "第八章 投标文件格式" not in texts
    assert "第九章 评标办法" not in texts
    assert len(copied.tables) == 1
    copied_table = copied.tables[0]
    assert len(copied_table.rows) == 3
    assert len(copied_table.columns) == 3
    assert copied_table.cell(0, 0).text == "投标人名称"
    assert copied_table.cell(0, 1).text == "安徽正奇建设有限公司"
    assert copied_table.cell(0, 2).text == "安徽正奇建设有限公司"


def test_fill_personnel_table_fills_project_manager_row() -> None:
    from services.original_docx_format_service import _fill_personnel_table

    doc = Document()
    table = doc.add_table(rows=4, cols=6)
    # 表头第1行
    table.cell(0, 0).text = "职务"
    table.cell(0, 1).text = "姓名"
    table.cell(0, 2).text = "职称"
    table.cell(0, 3).text = "执业或职业资格证明"
    # 表头第2行(子列)
    table.cell(1, 3).text = "证书名称"
    table.cell(1, 4).text = "级别"
    table.cell(1, 5).text = "证号"
    # 第3、4行为空数据行

    profile = {"project_manager_name": "江舟", "project_manager_cert": "皖1342006200803161"}
    assert _fill_personnel_table(doc, profile) is True
    assert table.cell(2, 0).text == "项目经理"
    assert table.cell(2, 1).text == "江舟"
    assert table.cell(2, 5).text == "皖1342006200803161"


def test_fill_personnel_table_noop_without_pm() -> None:
    from services.original_docx_format_service import _fill_personnel_table

    doc = Document()
    table = doc.add_table(rows=3, cols=6)
    table.cell(0, 0).text = "职务"
    table.cell(0, 1).text = "姓名"
    table.cell(1, 5).text = "证号"
    assert _fill_personnel_table(doc, {}) is False
    assert table.cell(2, 1).text.strip() == ""


# ── #8 占位填空:省略号/点线签字线该当"空"来填,但绝不覆盖真值/孤立短横/句号 ──────
def test_is_blank_or_placeholder_treats_fill_lines_as_blank() -> None:
    # 真·空 / 下划线 / 全角空格 / 制表符:占位,单个即算
    for blank in ["", "   ", "_____", "＿＿＿", "　", "\t", "…", "‥"]:
        assert _is_blank_or_placeholder(blank) is True, blank
    # 点线/破折号签字线:连续≥2 才算占位(实测产物里的真实漏填形状)
    for leader in ["……", "....", "．．．", "- - -", "————", "··········"]:
        assert _is_blank_or_placeholder(leader) is True, leader


def test_is_blank_or_placeholder_never_eats_real_values() -> None:
    # 真值绝不能被判空(否则会被公司档案值覆盖)
    for real in [
        "安徽正奇建设集团有限公司", "江舟", "0551-65650939",
        "2022.03", "5.4m", "0-50分", "1.1.4.5", "91340000X",
    ]:
        assert _is_blank_or_placeholder(real) is False, real
    # 孤立单个 短横/点/间隔号 = 人工填的"无/不适用",不判空
    for nil in ["—", "-", "－", "·", "．", "/"]:
        assert _is_blank_or_placeholder(nil) is False, nil
    # 句号(及句号串)是真标点,永不判空
    for dot in ["。", "。。。", "以上。"]:
        assert _is_blank_or_placeholder(dot) is False, dot


def test_fill_known_table_cells_fills_ellipsis_placeholder() -> None:
    # 核心 bug:值格是省略号/点线占位(福昕产物实测存在 U+2026),旧 _TABLE_BLANK_RE
    # 只认下划线 → 判为非空 → 跳过留空。现应识别为占位并填入。
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "投标人名称"
    table.cell(0, 1).text = "…"  # 省略号占位
    table.cell(1, 0).text = "统一社会信用代码"
    table.cell(1, 1).text = "________"  # 下划线占位(旧逻辑也能填)

    filled = _fill_known_table_cells(
        doc, {"company_name": "安徽正奇建设有限公司", "credit_code": "91X"}
    )
    assert table.cell(0, 1).text == "安徽正奇建设有限公司"
    assert table.cell(1, 1).text == "91X"
    assert filled == 2


def test_fill_known_table_cells_keeps_nil_dash_answer() -> None:
    # 值格是人工填的"—"(无/不适用)→ 不是占位,绝不覆盖
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "投标人名称"
    table.cell(0, 1).text = "—"

    filled = _fill_known_table_cells(doc, {"company_name": "安徽正奇建设有限公司"})
    assert table.cell(0, 1).text == "—"  # 真答案保留
    assert filled == 0


# ── #9 接入可推导字段:项目名称/工期 从招标解析(combined_profile 中文键)填进商务卷 ──
def test_table_label_value_fills_project_derived_fields() -> None:
    # combined_profile 里项目级字段是中文键(v2_generation_service.project_fields)
    combined = {
        "company_name": "安徽正奇建设有限公司",
        "项目名称": "XX县2025年农村公路提质改造联网路工程",
        "工期": "90日历天",
        "project_manager_name": "江舟",
    }
    assert _table_label_value("项目名称", combined) == combined["项目名称"]
    assert _table_label_value("工程名称", combined) == combined["项目名称"]
    assert _table_label_value("计划工期", combined) == "90日历天"
    assert _table_label_value("工期", combined) == "90日历天"
    # 关键防撞:"项目名称"不得把"项目经理"格抢成项目名(仍填项目经理名)
    assert _table_label_value("项目经理", combined) == "江舟"
    assert _table_label_value("项目经理姓名", combined) == "江舟"


def test_fill_known_table_cells_fills_project_name_and_duration() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目名称"
    table.cell(0, 1).text = "…"  # 占位
    table.cell(1, 0).text = "计划工期"  # (1,1) 空

    filled = _fill_known_table_cells(
        doc, {"项目名称": "XX路网工程", "工期": "90日历天"}
    )
    assert table.cell(0, 1).text == "XX路网工程"
    assert table.cell(1, 1).text == "90日历天"
    assert filled == 2


# ── 投标函内联空(用户实测):工程质量/安全目标/工期 写在正文里 "标签：<tab>" ──────
def _bid_letter_para(doc):
    p = doc.add_paragraph()
    for t in ["3", ".", "工程质量：", " ", "\t", "，安全目标：", " ", "\t",
              "，工期", "：", " ", "\t", "日历天", "。"]:
        p.add_run(t)
    return p


def test_fill_inline_labeled_blanks_fills_bid_letter_prose() -> None:
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = Document()
    _bid_letter_para(doc)
    n = _fill_inline_labeled_blanks(doc, {
        "质量": "符合国家现行工程质量验收标准规范合格标准",
        "安全": "无安全责任事故发生",
        "工期": "90日历天",
    })
    norm = doc.paragraphs[0].text.replace(" ", "")
    assert n == 3
    assert "工程质量：符合国家现行工程质量验收标准规范合格标准" in norm
    assert "安全目标：无安全责任事故发生" in norm
    assert "工期：90日历天。" in norm
    assert "日历天日历天" not in norm  # 单位不重复


def test_fill_inline_blanks_ignores_unlabeled_tabs() -> None:
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = Document()
    p = doc.add_paragraph()
    for t in ["目录", "\t", "第1页"]:
        p.add_run(t)
    assert _fill_inline_labeled_blanks(doc, {"质量": "X"}) == 0
    assert doc.paragraphs[0].text == "目录\t第1页"  # 无标签的 tab 不动


def test_fill_inline_blanks_does_not_overwrite_existing_value() -> None:
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = Document()
    p = doc.add_paragraph()
    for t in ["工期：", "90日历天", "。"]:  # 已有值、无 tab 槽
        p.add_run(t)
    assert _fill_inline_labeled_blanks(doc, {"工期": "45日历天"}) == 0
    assert "90日历天" in doc.paragraphs[0].text  # 原值保留,不被覆盖


def test_fill_inline_blanks_signature_block() -> None:
    """用户定:投标人(盖单位章)填公司名;法定代表人(签字或盖章)处留白给人工签。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = Document()
    p = doc.add_paragraph()
    for t in ["投标人：", " ", "\t", "（盖单位章）", "法定代表人：", " ", "\t", "（签字）"]:
        p.add_run(t)
    _fill_inline_labeled_blanks(
        doc, {"company_name": "安徽正奇建设有限公司", "legal_representative": "许明英"}
    )
    txt = doc.paragraphs[0].text
    assert "安徽正奇建设有限公司" in txt  # 投标人(盖单位章)填公司名
    assert "许明英" not in txt  # 法定代表人签字处留白,不打名字
    assert "（盖单位章）" in txt and "（签字）" in txt  # 标记原样保留


def test_fill_inline_blanks_skips_委托代理人_ambiguous() -> None:
    """法定代表人或委托代理人:谁签不定 → 不自动填(endswith 不命中 法定代表人)。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = Document()
    p = doc.add_paragraph()
    for t in ["法定代表人或委托代理人：", "\t", "（签字）"]:
        p.add_run(t)
    assert _fill_inline_labeled_blanks(doc, {"legal_representative": "许明英"}) == 0
    assert "\t" in doc.paragraphs[0].text  # 留空待人工


def test_fill_inline_blanks_run_split_independent() -> None:
    """不依赖 run 边界:福昕/pdf2docx 把"标签：\t"切成一个 run 还是几个 run 都能填。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    prof = {"工期": "90日历天", "company_name": "正奇建设"}
    for label, name, runs in [
        ("同run", "工期", ["工期：\t日历天"]),
        ("冒号tab同run", "工期", ["工期", "：\t", "日历天"]),
        ("投标人名称同run", "投标人名称", ["投标人名称：\t，"]),  # 正文表单名称位:可填
    ]:
        doc = Document(); p = doc.add_paragraph()
        for t in runs: p.add_run(t)
        n = _fill_inline_labeled_blanks(doc, prof)
        assert n == 1, f"{label} 应填1处, 实际{n}"
        assert "\t" not in doc.paragraphs[0].text, f"{label} tab未被填"
    # 用户拍板:签署块"投标人：\t（盖单位章）"填公司名,标记保留,槽位填掉
    doc = Document(); p = doc.add_paragraph(); p.add_run("投标人：\t（盖单位章）")
    assert _fill_inline_labeled_blanks(doc, prof) == 1
    assert "正奇建设" in doc.paragraphs[0].text
    assert "（盖单位章）" in doc.paragraphs[0].text
    assert "\t" not in doc.paragraphs[0].text
    # 多字段跨混合 run 切分,一段填 3 处
    doc = Document(); p = doc.add_paragraph()
    for t in ["工程质量：\t，安全目标：", "\t", "，工期：\t日历天。"]:
        p.add_run(t)
    n = _fill_inline_labeled_blanks(
        doc, {"质量": "合格", "安全": "无事故", "工期": "90日历天"}
    )
    assert n == 3
    assert doc.paragraphs[0].text == "工程质量：合格，安全目标：无事故，工期：90日历天。"


def test_fill_inline_blanks_fills_id_proof_form() -> None:
    """法定代表人身份证明:投标人名称空槽→公司名;'姓名：'空槽→法人(因含'的法定代表人'
    语境);'性别：'等 PII 留空。复现用户截图(screenshot 2)的真实排版。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = Document()
    doc.add_paragraph().add_run("投标人名称：")  # 空槽,接段末
    p2 = doc.add_paragraph()
    p2.add_run("姓名：性别：")  # 姓名空槽,紧跟同级小标签 性别：
    p2.add_run("\t年龄：\t职务：\t系（投标人名称）的法定代表人。")
    prof = {"company_name": "安徽正奇建设有限公司", "legal_representative": "许明英"}
    _fill_inline_labeled_blanks(doc, prof)
    assert doc.paragraphs[0].text == "投标人名称：安徽正奇建设有限公司"
    t2 = doc.paragraphs[1].text
    assert "姓名：许明英" in t2  # 法人名填进姓名位
    assert "性别：\t" in t2  # PII 留空待人工
    assert t2.count("许明英") == 1  # 法人名只出现一次(不漏填到职务/年龄位)


def test_fill_inline_blanks_姓名_without_id_proof_context_is_skipped() -> None:
    """普通'姓名：'(无'的法定代表人'语境,如人员表)绝不填法人名——防 122 卷歧义误填。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = Document()
    doc.add_paragraph().add_run("姓名：性别：")
    assert _fill_inline_labeled_blanks(
        doc, {"legal_representative": "许明英", "company_name": "正奇"}
    ) == 0
    assert "许明英" not in doc.paragraphs[0].text


def test_known_replacements_fills_本人姓名_in_poa() -> None:
    """授权委托书'本人（姓名）系…的法定代表人'→ 填法人名;但只锚定'本人'前缀。"""
    from services.original_docx_format_service import _known_replacements
    repl = _known_replacements({"legal_representative": "许明英"})
    assert repl.get("本人（姓名）") == "本人许明英"
    assert "（姓名）" not in repl  # 不全局替换 → 人员表每行（姓名）安全


def test_inline_value_性别年龄_only_in_id_proof_context() -> None:
    """法定代表人身份证明的 性别/年龄(来自法人身份证OCR)仅在身份证明语境填,防误填人员表。"""
    from services.original_docx_format_service import _inline_value_for
    prof = {"法人性别": "女", "法人年龄": "50"}
    assert _inline_value_for("性别", prof, id_proof_context=True) == "女"
    assert _inline_value_for("年龄", prof, id_proof_context=True) == "50"
    # 非身份证明语境(如人员表)绝不填
    assert _inline_value_for("性别", prof, id_proof_context=False) == ""
    assert _inline_value_for("年龄", prof, id_proof_context=False) == ""
    # 职务无据 → 任何语境都不填(留人工)
    assert _inline_value_for("职务", prof, id_proof_context=True) == ""


def test_format_range_keeps_cover_page(tmp_path) -> None:
    """商务卷页范围必须从'投标人：（盖单位章）'封面起,不能把封面当过场页跳掉。

    真实病例(埇桥商务卷):封面页字少、无正文标记,历史 _skip_toc_pages 跳过它、
    从目录起 → 整卷丢封面。封面含'盖单位章'(目录绝不含),据此保留。
    """
    import fitz
    from services.original_docx_format_service import _find_format_page_range_in_pdf

    F = "china-s"  # fitz 内置简体中文字体,否则中文渲染成方块无法提取
    doc = fitz.open()
    # p0: 封面(章标题 + 投标文件 + 盖单位章)
    p0 = doc.new_page()
    p0.insert_text((72, 100), "第九章 投标文件格式", fontname=F)
    p0.insert_text((72, 300), "投 标 文 件（商务及技术文件）", fontname=F)
    p0.insert_text((72, 500), "投标人： ____（盖单位章） 年 月 日", fontname=F)
    # p1: 目录(点线,无盖单位章)
    p1 = doc.new_page()
    p1.insert_text((72, 100), "目 录", fontname=F)
    p1.insert_text((72, 140), "一、投标函及投标函附录 ........................ 1", fontname=F)
    p1.insert_text((72, 170), "二、授权委托书 ........................ 2", fontname=F)
    p1.insert_text((72, 200), "三、联合体协议书 ........................ 3", fontname=F)
    # p2: 投标函正文
    p2 = doc.new_page()
    p2.insert_text((72, 100), "一、投标函", fontname=F)
    p2.insert_text((72, 140), "（招标人名称）：我方已仔细研究本次施工招标文件的全部内容", fontname=F)
    path = tmp_path / "tender.pdf"
    doc.save(str(path))
    doc.close()

    rng = _find_format_page_range_in_pdf(str(path))
    assert rng is not None
    assert rng[0] == 0, f"应从封面(第0页)起,实际从第{rng[0]}页起(封面被跳掉)"


def test_native_docx_path_skips_foxit_artifact_healers(tmp_path: Path) -> None:
    """原生 Word 招标不许跑福昕伪影 healer。

    2026-07-29 实测:split_paragraphs(治福昕"一句话劈成几段")在原生 Word 上会把
    "第八章 投标文件格式"和下一行"投标文件（商务文件）"错并成一段。原样复制路径
    必须绕开这类 healer,否则招标原文被改写。
    """
    source_path = tmp_path / "tender.docx"
    source = Document()
    source.add_paragraph("第一章 招标公告")
    source.add_paragraph("第八章 投标文件格式")
    source.add_paragraph("投标文件（商务文件）")
    source.add_paragraph("（一）投标人基本情况表")
    source.save(str(source_path))

    out = tmp_path / "format.docx"
    build_original_format_docx(
        source_path.read_bytes(), out, profile={"company_name": "安徽正奇建设有限公司"}
    )
    texts = [p.text.strip() for p in Document(str(out)).paragraphs if p.text.strip()]
    # 2026-07-31 用户拍板:卷首"第X章 投标文件格式"是招标章标题,必须剥掉,封面打头
    assert "第八章 投标文件格式" not in texts
    assert "投标文件（商务文件）" in texts


def test_fill_format_docx_native_vs_foxit_healer_sets() -> None:
    """两条路各跑各的 healer 集合:原生版只含通用三项,不含任何福昕伪影修复。"""
    from services.docx_format_doctor import _HEALERS, _NATIVE_DOCX_HEALERS

    native = {n for n, _ in _NATIVE_DOCX_HEALERS}
    assert native == {
        "template_header_lines",
        "section_title_page_breaks",
        "underline_slots",
        "filler_blank_runs",
    }
    # 福昕专治项绝不能混进原生集合
    assert "split_paragraphs" not in native
    assert "split_paragraphs" in {n for n, _ in _HEALERS}


def test_declared_id_scans_are_attached_after_form_tail(monkeypatch) -> None:
    """表单写着"附：法定代表人身份证正反面扫描件"就必须真附图,且插在落款之后。

    2026-07-30 用户实测:身份证明/承诺书页脚的"附：…扫描件"声明被原样照抄,图从没插过。
    """
    from services import original_docx_format_service as svc

    doc = Document()
    doc.add_paragraph("特此证明。")
    doc.add_paragraph("附：法定代表人身份证正反面扫描件")
    doc.add_paragraph("投  标  人：安徽正奇建设有限公司（盖单位章）")  # 落款标签带排版空格
    doc.add_paragraph("日      期：2026年7月30日")
    doc.add_paragraph("三、联合体协议书")

    import services.asset_resolver as ar
    monkeypatch.setattr(
        ar, "pick_id_card_documents",
        lambda name: [{"document_id": 101, "side": "front"}, {"document_id": 102, "side": "back"}]
        if name == "许明英" else [],
    )
    calls = []

    import services.generation_service as gs
    def fake_insert(anchor, document, doc_id, caption, width):
        calls.append((doc_id, caption, "".join(anchor.itertext())))
        return anchor
    monkeypatch.setattr(gs, "_insert_image_after", fake_insert)

    n = svc._attach_declared_id_scans(doc, {"legal_representative": "许明英"})
    assert n == 2
    assert [c[0] for c in calls] == [101, 102]
    assert "法定代表人身份证（正面）" == calls[0][1]
    # 锚点必须越过带排版空格的落款行(投 标 人/日 期),图跟在整个表单之后
    assert calls[0][2].startswith("日")  # 首图锚在"日 期"行之后,不是楔在声明行下


def test_declared_id_scans_missing_person_leaves_line_untouched(monkeypatch) -> None:
    """取不到图/没选派 → 不插不删,声明行留给人工。"""
    from services import original_docx_format_service as svc

    doc = Document()
    doc.add_paragraph("本页后附项目经理身份证正反面扫描件")
    # 没选派项目经理(profile 无 pm_resume) → 跳过
    assert svc._attach_declared_id_scans(doc, {"legal_representative": "许明英"}) == 0
    assert "本页后附项目经理身份证正反面扫描件" in doc.paragraphs[0].text


def test_declared_scans_skip_agent_lines_and_set_letdown_flag(monkeypatch) -> None:
    """授权委托书的声明行(含"代理人")一律跳过;法人证插好后要立让位标志。

    2026-07-30 用户拍板:授权委托书后不附法人身份证。老锚点链会误命中委托书里
    "附：法定代表人身份证明"行,靠 _legal_id_inline 标志让它整段让位。
    """
    from services import original_docx_format_service as svc
    import services.asset_resolver as ar
    import services.generation_service as gs

    monkeypatch.setattr(
        ar, "pick_id_card_documents",
        lambda name: [{"document_id": 7, "side": "both"}],
    )
    monkeypatch.setattr(gs, "_insert_image_after", lambda a, d, i, c, w: a)

    # 授权委托书场景:声明里带"代理人" → 跳过、不立标志
    doc = Document()
    doc.add_paragraph("附：法定代表人身份证明  代理人身份证正反面扫描件")
    prof = {"legal_representative": "许明英"}
    assert svc._attach_declared_id_scans(doc, prof) == 0
    assert "_legal_id_inline" not in prof

    # 身份证明表场景:正常插,并立让位标志
    doc2 = Document()
    doc2.add_paragraph("附：法定代表人身份证正反面扫描件")
    prof2 = {"legal_representative": "许明英"}
    assert svc._attach_declared_id_scans(doc2, prof2) == 1
    assert prof2.get("_legal_id_inline") is True


def test_leading_chapter_title_is_stripped_but_body_mentions_kept() -> None:
    """卷首"第九章 投标文件格式"必须删(招标自己的章标题);正文里引用它的句子不动。

    2026-07-31 用户实测:成品第一页顶着招标章标题("这不是吓我的头吗")。
    """
    from services.original_docx_format_service import _strip_leading_chapter_title

    doc = Document()
    doc.add_paragraph("第九章  投标文件格式")
    doc.add_paragraph("")
    doc.add_paragraph("巢湖市 2026 年农村公路养护工程 标段招标")
    doc.add_paragraph("投标文件")
    doc.add_paragraph("投标人应按照第九章 投标文件格式的要求编写。")  # 正文引用,不许动

    assert _strip_leading_chapter_title(doc) == 1
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts[0].startswith("巢湖市")
    assert any("投标人应按照第九章" in t for t in texts)

    # 分两段写的也逮
    doc2 = Document()
    doc2.add_paragraph("第八章")
    doc2.add_paragraph("投标文件格式")
    doc2.add_paragraph("封面内容")
    assert _strip_leading_chapter_title(doc2) == 2
    assert [p.text for p in doc2.paragraphs if p.text.strip()] == ["封面内容"]


def test_template_resume_wholesale_replace_and_experience_fill(monkeypatch) -> None:
    """整表照搬(2026-08-01 用户拍板):有成品模版 → 整表替换招标空白表,
    拟任职务按选派角色改,经历空行填勾选业绩;无模版 → 退回字段填空。"""
    from copy import deepcopy
    from services import original_docx_format_service as svc
    import services.curated_resume_service as crs

    # 成品模版(秦蒙蒙,模版里预填"项目经理")
    tpl = Document()
    t = tpl.add_table(rows=6, cols=4)
    t.cell(0, 0).text = "姓  名"; t.cell(0, 1).text = "秦蒙蒙"
    t.cell(1, 0).text = "拟在本标段工程任职"; t.cell(1, 1).text = "项目经理"
    for i, h in enumerate(("时间", "参加过的类似工程项目名称", "担任职务", "发包人及联系电话")):
        t.cell(2, i).text = h
    # 第3、4行为经历空行;第5行表尾
    t.cell(5, 0).text = "获奖情况"
    monkeypatch.setattr(crs, "get_template_table_el", lambda name: deepcopy(t._tbl) if name == "秦蒙蒙" else None)

    host = Document()
    host.add_paragraph("（六）拟委任的项目经理和项目总工资历表")
    ht = host.add_table(rows=2, cols=2)
    ht.cell(0, 0).text = "拟在本标段工程任职"

    prof = {"similar_projects_td": [{
        "project_name": "澳门路道路工程", "project_manager": "别人", "tech_leader": "秦蒙蒙",
        "start_date": "2023.07.11", "end_date": "2025.02.27",
        "owner_name": "包河区重点工程建设管理中心", "owner_phone": "0551-63357659",
    }]}
    new_tb = svc._replace_with_template_resume(
        host, ht, {"姓名": "秦蒙蒙", "拟任职务": "项目技术负责人"}, prof
    )
    assert new_tb is not None
    texts = " ".join(c.text for row in new_tb.rows for c in row.cells)
    assert "秦蒙蒙" in texts
    # 拟任职务按选派角色改成了总工(模版预填的是项目经理)
    assert "项目总工" in new_tb.rows[1].cells[1].text
    # 经历行填了勾选业绩:担任职务=他在那个工程里的角色(总工),发包人带电话
    assert "澳门路道路工程" in texts and "包河区重点工程建设管理中心/0551-63357659" in texts
    assert "2023.07.11-2025.02.27" in texts

    # 无模版 → None(调用方退回字段填空)
    assert svc._replace_with_template_resume(host, ht, {"姓名": "无名氏"}, {}) is None


def test_new_format_adaptations_waixian_trial() -> None:
    """外地格式试跑(2026-08-05)的四处适配:质量评定/委托同志/法人姓名token/附件正则。"""
    from services import original_docx_format_service as svc

    # ① 质量评定内联空
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("3.标段工程交工验收的质量评定：")
    s1 = p.add_run("      "); s1.font.underline = True
    p.add_run("；竣工验收的质量评定：")
    s2 = p.add_run("      "); s2.font.underline = True
    p.add_run("。")
    svc._fill_inline_labeled_blanks(doc, {"质量": "合格"})
    assert doc.paragraphs[0].text.count("合格") == 2

    # ② 我方将委托__同志
    d2 = Document()
    d2.add_paragraph("8.我方将委托        同志作为本工程的项目经理。")
    d2.add_paragraph("9.我方将委托        同志作为本工程的现场总工程师。")
    n = svc._fill_entrust_lines(d2, {"pm_resume": {"姓名": "李刚"}, "tech_resume": {"姓名": "许明英"}})
    assert n == 2
    assert "委托李刚同志" in d2.paragraphs[0].text.replace(" ", "")
    assert "委托许明英同志" in d2.paragraphs[1].text.replace(" ", "")

    # ③ （法定代表人姓名）token
    d3 = Document()
    d3.add_paragraph("姓名：（法定代表人姓名）性别：     年龄：")
    svc._replace_known_fields(d3, {"legal_representative": "许明英", "company_name": "安徽正奇建设有限公司"})
    assert "姓名：许明英" in d3.paragraphs[0].text

    # ④ 照单附件正则容纳长前缀
    assert svc._ATTACH_SCAN_RE.search("附：投标人的法定代表人有效期内居民身份证扫描件")
