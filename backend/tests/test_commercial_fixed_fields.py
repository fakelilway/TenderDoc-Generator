"""商务标固定字段规则测试。

重点守护 2026-07-11 的决策:「质量」是招标优先、固定值兜底——
招标解析出的质量标准不能再被固定值"合格"无条件覆盖(员工反馈的真实事故)。
"""

from services.commercial_fixed_fields import (
    COMMERCIAL_FIXED_FIELDS,
    TENDER_FIRST_FALLBACK_FIELDS,
    apply_fixed_fields,
)


def test_quality_from_tender_is_preserved():
    """招标解析出了质量标准 → 保留,不被固定值覆盖。"""
    profile = {"质量": "符合设计要求及验收规范,确保省优质工程"}
    apply_fixed_fields(profile)
    assert profile["质量"] == "符合设计要求及验收规范,确保省优质工程"


def test_quality_falls_back_to_fixed_when_missing():
    """招标里没解析到质量标准 → 兜底填固定值"合格"。"""
    for empty in ("", "  ", None):
        profile = {"质量": empty}
        apply_fixed_fields(profile)
        assert profile["质量"] == "合格"
    # 键完全不存在也要兜底
    profile = {}
    apply_fixed_fields(profile)
    assert profile["质量"] == "合格"


def test_quality_not_in_unconditional_fixed_fields():
    """「质量」必须待在招标优先清单里,不能回到无条件固定清单(防止改回老毛病)。"""
    assert "质量" not in COMMERCIAL_FIXED_FIELDS
    assert TENDER_FIRST_FALLBACK_FIELDS.get("质量") == "合格"


def test_truly_fixed_fields_still_override():
    """全固定字段(公司名/安全等)仍然无条件覆盖,行为不变。"""
    profile = {
        "company_name": "别的公司名",
        "安全": "招标里写的安全目标",
        "质量": "创优",
    }
    apply_fixed_fields(profile)
    assert profile["company_name"] == "安徽正奇建设有限公司"
    assert profile["安全"] == "无安全事故"
    # 同时确认质量不受牵连
    assert profile["质量"] == "创优"


def test_qualification_grade_is_full_list():
    """企业资质等级=全量10项清单(2026-07-12用户定稿),不许退回只写总承包一项。"""
    v = COMMERCIAL_FIXED_FIELDS["qualification_grade"]
    for must in (
        "公路工程施工总承包贰级", "市政公用工程施工总承包贰级",
        "公路交通工程（公路安全设施）专业承包贰级", "公路路面工程专业承包贰级",
        "公路路基工程专业承包贰级", "环保工程专业承包贰级",
        "城市及道路照明工程专业承包贰级", "施工劳务序列不分等级",
        "路基路面养护甲级资质", "交通安全设施养护资质",
    ):
        assert must in v, f"资质清单缺:{must}"


def test_affiliated_companies_fills_basic_info_table():
    """投标人基本情况表:企业资质等级填全量清单、关联企业情况填股东股权三条(多行)。"""
    from docx import Document

    from services.original_docx_format_service import _fill_known_table_cells

    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "企业资质等级"
    t.cell(1, 0).text = "投标人关联企业情况"
    profile = apply_fixed_fields({})
    _fill_known_table_cells(doc, profile)

    grade = t.cell(0, 1).text
    assert "路基路面养护甲级资质" in grade and "施工劳务序列不分等级" in grade
    rel = t.cell(1, 1).text
    assert "江舟:94.83%" in rel and "许明英:5.16%" in rel
    assert "（2）" in rel and "（3）" in rel
    assert "\n" in rel  # (1)(2)(3) 分行,不挤成一坨


def test_affiliated_cell_rewritten_even_with_preprinted_prompt():
    """关联企业格:招标预印提示文字占着(泗沙路实测叠行乱码+没填) → 整格重写成定稿全文。"""
    from docx import Document

    from services.original_docx_format_service import _fill_basic_info_subfields

    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "投标人名称"
    t.cell(1, 0).text = "员工总人数：26"
    t.cell(2, 0).text = "投标人关联企业情况"
    vc = t.cell(2, 1)
    vc.text = "投标人应提供关联企业情况，包括："
    vc.add_paragraph("（1） 投标人的所有股东名称及相应股权（出资额）比例；如投标人为上市%以上的")
    vc.add_paragraph("（2） 投标人投资（控股）或管理的下属企业名称、持有股权（出资额）比例；")

    profile = apply_fixed_fields({})
    filled = _fill_basic_info_subfields(doc, profile)
    assert filled >= 1
    txt = t.cell(2, 1).text
    assert "江舟:94.83%" in txt and "许明英:5.16%" in txt
    assert "上市%以上" not in txt  # 福昕转坏的残句被定稿替掉
    assert len(t.cell(2, 1).paragraphs) == 1  # 预印多段清干净
    # 幂等:再跑一遍不重复追加
    assert _fill_basic_info_subfields(doc, profile) >= 0
    assert t.cell(2, 1).text == txt


def test_bank_permit_group_keeps_both_pages():
    """开户许可证=正页+注意事项两页一本,组上限必须≥2且不按one去重(泗沙路实测丢页)。"""
    from services.v2_generation_service import _EVIDENCE_GROUPS

    entry = next(g for g in _EVIDENCE_GROUPS if g[0] == "基本账户开户许可证")
    assert entry[2] >= 2 and entry[4] != "one"
