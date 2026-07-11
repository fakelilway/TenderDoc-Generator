"""类似项目按人自动填表测试:节归属识别 + 信息表克隆/裁剪 + 汇总表扩行 + 垃圾填充屏蔽。"""
from docx import Document
from docx.oxml.ns import qn


# 招标"信息表"竖表的标准行标签(截图口径)
_DETAIL_LABELS = [
    "项目名称\n（合同名称）", "项目所在地", "发包人名称", "发包人地址", "发包人电话",
    "合同价格", "开工日期", "交工日期", "承担的工作", "工程质量",
    "项目经理", "项目总工", "监理单位及联系电话", "总监理工程师及电话", "项目描述", "备注",
]


def _record(i: int, manager: str = "李刚", tech: str = "") -> dict:
    return {
        "project_name": f"测试业绩{i}号工程",
        "location": f"某地{i}",
        "owner_name": f"发包人{i}",
        "owner_address": f"地址{i}",
        "owner_phone": f"0551-000{i}",
        "contract_price": f"{i}000000元",
        "start_date": f"202{i}.1.1",
        "end_date": f"202{i}.12.1",
        "work_scope": "施工总承包",
        "quality": "合格",
        "project_manager": manager,
        "tech_leader": tech,
        "supervisor_org": f"监理公司{i}/0551-111{i}",
        "chief_supervisor": f"总监{i}/0551-222{i}",
        "description": f"工程内容{i}",
    }


def _add_detail_table(doc: Document) -> None:
    t = doc.add_table(rows=len(_DETAIL_LABELS), cols=2)
    for ri, label in enumerate(_DETAIL_LABELS):
        t.cell(ri, 0).text = label
    t.cell(len(_DETAIL_LABELS) - 1, 1).text = "资格审查业绩"  # 备注招标预印


def _add_summary_table(doc: Document, data_rows: int = 3) -> None:
    t = doc.add_table(rows=1 + data_rows, cols=3)
    t.cell(0, 0).text = "业绩序号"
    t.cell(0, 1).text = "项目名称（合同名称）"
    t.cell(0, 2).text = "备注"


def _build_six_section_doc() -> Document:
    """迷你商务卷:投标人/项目经理/项目总工 三节(各 汇总表+1张信息表)。"""
    doc = Document()
    doc.add_paragraph("（三）投标人近年完成的类似项目情况表（资格审查）")
    _add_summary_table(doc)
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：投标人应根据招标文件要求在本表后附相关证明材料。")
    doc.add_paragraph("1.项目经理近年完成的类似项目情况表（资格审查）")
    _add_summary_table(doc)
    doc.add_paragraph("项目经理近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：应附证明材料。")
    doc.add_paragraph("2.项目总工近年完成的类似项目情况表（资格审查）")
    _add_summary_table(doc)
    doc.add_paragraph("项目总工近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：应附证明材料。")
    return doc


def _detail_tables(doc: Document) -> list:
    from services.original_docx_format_service import _is_similar_project_detail_table

    return [t for t in doc.tables if _is_similar_project_detail_table(t)]


def _cell_of(table, label_key: str) -> str:
    for row in table.rows:
        if label_key in row.cells[0].text.replace("\n", "").replace(" ", ""):
            return row.cells[-1].text.strip()
    raise AssertionError(f"表里没有标签 {label_key}")


def test_fill_clones_detail_tables_per_record_and_extends_summary() -> None:
    """经理3条业绩:信息表 1张→克隆成3张 逐条原样填;汇总表填3行。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = _build_six_section_doc()
    records = [_record(1), _record(2), _record(3)]
    profile = {
        "similar_projects_pm": records,
        "selected_performance": [{"name": r["project_name"]} for r in records],
        # 全部人工手选:经理表也要显式勾选才填(2026-07-11拍板,不默认全填名下)
        "selected_pm_performance": [{"name": r["project_name"]} for r in records],
        "similar_projects_td": [],
    }
    result = fill_similar_project_sections(doc, profile)

    details = _detail_tables(doc)
    # 投标人节 3 张 + 项目经理节 3 张 + 总工节 1 张(留白不裁)
    assert len(details) == 7
    bidder_tables = details[:3]
    for i, table in enumerate(bidder_tables, start=1):
        assert _cell_of(table, "项目名称") == f"测试业绩{i}号工程"
        assert _cell_of(table, "发包人名称") == f"发包人{i}"
        assert _cell_of(table, "合同价格") == f"{i}000000元"
        assert _cell_of(table, "项目经理") == "李刚"
        assert _cell_of(table, "备注") == "资格审查业绩"  # 招标预印,不动
    # 总工节信息表整表留白(没有他名下的记录,绝不拿经理的凑)
    td_table = details[-1]
    assert _cell_of(td_table, "项目名称") == ""
    assert _cell_of(td_table, "项目经理") == ""
    assert result["detail_filled"] > 0
    # 汇总表:投标人节首个汇总表填了3行项目名
    summary = doc.tables[0]
    names = [summary.cell(ri, 1).text.strip() for ri in range(1, len(summary.rows))]
    assert names[:3] == ["测试业绩1号工程", "测试业绩2号工程", "测试业绩3号工程"]
    # 总工节汇总表留白且在 handled 名单里(通用填表器不得再碰)
    td_summary = [
        t for t in doc.tables
        if t.rows and "业绩序号" in t.cell(0, 0).text
    ][-1]
    assert all(
        not td_summary.cell(ri, 1).text.strip() for ri in range(1, len(td_summary.rows))
    )
    assert td_summary._tbl in result["handled_tables"]


def test_summary_table_grows_rows_when_needed() -> None:
    """业绩多于汇总表空行 → 克隆末行扩行,9条全列上。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("（三）投标人近年完成的类似项目情况表（资格审查）")
    _add_summary_table(doc, data_rows=3)
    records = [_record(i) for i in range(1, 10)]
    profile = {
        "similar_projects_pm": records,
        "selected_performance": [{"name": r["project_name"]} for r in records],
    }
    fill_similar_project_sections(doc, profile)
    summary = doc.tables[0]
    names = [summary.cell(ri, 1).text.strip() for ri in range(1, len(summary.rows))]
    assert len([n for n in names if n]) == 9
    assert summary.cell(9, 0).text.strip() == "9"  # 序号连续


def test_prunes_extra_blank_detail_tables() -> None:
    """模板给了3张空信息表、只选1条业绩 → 裁掉2张空表。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：附证明。")
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：附证明。")
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：附证明。")
    records = [_record(1)]
    profile = {
        "similar_projects_pm": records,
        "selected_performance": [{"name": records[0]["project_name"]}],
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == "测试业绩1号工程"


def test_bidder_table_follows_selected_performance_not_manager() -> None:
    """投标人表跟"用户选的业绩"走:从全部信息表记录按项目名补全全字段,与经理是谁无关。

    给一个跟选中业绩无关的 similar_projects_pm(赵六名下),证明投标人表压根不看它,
    只认 selected_performance + similar_projects_all。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    all_records = [
        _record(1, manager="张三"),
        _record(2, manager="李四"),
        _record(3, manager="王五"),
    ]
    profile = {
        "similar_projects_all": all_records,
        "selected_performance": [{"name": "测试业绩2号工程"}],  # 只选第2条
        "similar_projects_pm": [_record(9, manager="赵六")],  # 无关经理,不该被用
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == "测试业绩2号工程"
    assert _cell_of(details[0], "发包人名称") == "发包人2"
    assert _cell_of(details[0], "合同价格") == "2000000元"
    assert _cell_of(details[0], "项目经理") == "李四"  # 取自被选项目自己的记录


def test_manager_selection_does_not_touch_bidder_table() -> None:
    """选了项目经理但没选业绩 → 投标人信息表留白(经理不再顶替业绩选择)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    profile = {
        "similar_projects_all": [_record(1), _record(2)],
        "similar_projects_pm": [_record(1), _record(2)],  # 经理名下有业绩
        # 注意:没有 selected_performance 键 → 用户没选业绩
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == ""  # 留白,不被经理业绩顶上


def test_manual_added_name_without_record_fills_name_only() -> None:
    """手动加的业绩(信息表里没有)→ 只填项目名称,其余留白待人工。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    profile = {
        "similar_projects_pm": [_record(1)],
        "selected_performance": [
            {"name": _record(1)["project_name"]},
            {"name": "手动补选的台账工程"},
        ],
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 2
    assert _cell_of(details[1], "项目名称") == "手动补选的台账工程"
    assert _cell_of(details[1], "发包人名称") == ""


def test_generic_filler_skips_similar_detail_table() -> None:
    """通用填表器绕开信息表竖表:项目名称/电话/项目经理不再被本项目值污染。"""
    from services.original_docx_format_service import _fill_known_table_cells

    doc = Document()
    _add_detail_table(doc)
    profile = {
        "项目名称": "本次招标工程名(不该出现)",
        "contact_phone": "15000000000",
        "project_manager_name": "王某某",
        "company_name": "安徽正奇建设有限公司",
    }
    _fill_known_table_cells(doc, profile)
    table = doc.tables[0]
    assert _cell_of(table, "项目名称") == ""
    assert _cell_of(table, "发包人电话") == ""
    assert _cell_of(table, "项目经理") == ""


def test_no_records_is_noop_keeps_legacy_behavior() -> None:
    """四个数据源全空(没导库/没选派/没选业绩) → 整个填表器不动文档(老路径照旧)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = _build_six_section_doc()
    result = fill_similar_project_sections(doc, {})
    assert result["handled_tables"] == set()
    assert len(_detail_tables(doc)) == 3  # 没克隆没裁剪


def test_selected_performance_alone_fills_even_without_pm(monkeypatch) -> None:
    """只选了业绩、没选派经理/总工 → 投标人表照填(埇桥实测:旧闸只看pm/td直接早退,
    导致选了业绩也一格不填)。"""
    from services import similar_project_fill_service as m

    doc = Document()
    doc.add_paragraph("（四）近年完成的类似项目情况表")  # 埇桥式裸标题,无人称前缀
    _add_detail_table(doc)
    monkeypatch.setattr(
        m.similar_project_info_service,
        "list_similar_project_records",
        lambda: [_record(1)],
        raising=False,
    ) if hasattr(m, "similar_project_info_service") else None
    profile = {
        "selected_performance": [{"name": _record(1)["project_name"]}],
        "similar_projects_all": [_record(1)],
        # 注意:没有 similar_projects_pm / similar_projects_td
    }
    result = m.fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == "测试业绩1号工程"
    assert _cell_of(details[0], "发包人名称") == "发包人1"
    assert result["detail_filled"] > 0


def test_cloned_units_carry_title_and_note_and_pagebreak() -> None:
    """克隆的信息表带标题段+注段,且前面有分页符(一表一页)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：附证明。")
    records = [_record(1), _record(2)]
    profile = {
        "similar_projects_pm": records,
        "selected_performance": [{"name": r["project_name"]} for r in records],
    }
    fill_similar_project_sections(doc, profile)
    texts = [p.text.strip() for p in doc.paragraphs]
    assert texts.count("投标人近年完成的类似项目信息表（资格审查）") == 2
    assert texts.count("注：附证明。") == 2
    # 克隆的标题段带"段前分页"(一表一页;不加独立分页空段,防浮框叠加挤出空白页)
    cloned_titles = [
        p for p in doc.paragraphs
        if p.text.strip() == "投标人近年完成的类似项目信息表（资格审查）"
        and p._p.find(qn("w:pPr")) is not None
        and p._p.find(qn("w:pPr")).find(qn("w:pageBreakBefore")) is not None
    ]
    assert len(cloned_titles) == 1


def test_adjacent_same_role_sections_each_get_full_records() -> None:
    """相邻两个同归属节(资格审查+详细评审都是投标人表)必须各填全套记录,
    不能并成一组把记录拆开(对抗验证确认的分组bug)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("注：附证明。")
    doc.add_paragraph("第三部分 商务评分材料")  # 普通标题,不切归属
    doc.add_paragraph("投标人近年完成的类似项目信息表（详细评审）")
    _add_detail_table(doc)
    doc.add_paragraph("注：附证明。")
    records = [_record(1), _record(2), _record(3)]
    profile = {
        "similar_projects_pm": records,
        "selected_performance": [{"name": r["project_name"]} for r in records],
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 6  # 两节各3张
    names = [_cell_of(t, "项目名称") for t in details]
    assert names == [f"测试业绩{i}号工程" for i in (1, 2, 3)] * 2


def test_manual_cleared_selection_leaves_tables_blank() -> None:
    """用户手动清空业绩(selected_performance 存在但为空)→ 整节留白,不回填经理全量。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    profile = {
        "similar_projects_pm": [_record(1)],
        "selected_performance": [],  # 手动清空
    }
    result = fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == ""
    assert details[0]._tbl in result["handled_tables"]  # 仍受保护不被乱填


def test_prune_keeps_preprinted_table() -> None:
    """预印了内容的多余信息表不删(那是招标给的示例,不是我们的空模板)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    preprinted = doc.tables[-1]
    for row in preprinted.rows:
        if "发包人名称" in row.cells[0].text:
            row.cells[-1].text = "招标预印示例发包人"
    records = [_record(1)]
    profile = {
        "similar_projects_pm": records,
        "selected_performance": [{"name": records[0]["project_name"]}],
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 2  # 预印表保留
    assert _cell_of(details[1], "发包人名称") == "招标预印示例发包人"


def test_multi_pair_row_fills_correct_cells() -> None:
    """一行两对字段(开工日期|_|交工日期|_)各填各的,不把开工日期写进交工日期的格。"""
    from services.similar_project_fill_service import _fill_detail_table

    doc = Document()
    t = doc.add_table(rows=3, cols=4)
    t.cell(0, 0).text = "项目名称"
    t.cell(1, 0).text = "开工日期"
    t.cell(1, 2).text = "交工日期"
    t.cell(2, 0).text = "项目经理"
    filled = _fill_detail_table(t, _record(1))
    assert t.cell(1, 1).text.strip() == "2021.1.1"
    assert t.cell(1, 3).text.strip() == "2021.12.1"
    assert filled >= 3


def test_pm_selection_subset_fills_only_selected() -> None:
    """经理名下3条、用户只勾1条 → 经理节只填勾中的那条(多余空表裁掉)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("项目经理近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    records = [_record(1), _record(2), _record(3)]
    profile = {
        "similar_projects_pm": records,
        "selected_pm_performance": [{"name": "测试业绩2号工程"}],
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == "测试业绩2号工程"
    assert _cell_of(details[0], "发包人名称") == "发包人2"


def test_pm_selection_order_follows_user() -> None:
    """经理业绩勾选顺序=出表顺序(勾3再勾1,表就先3后1)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("项目经理近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    records = [_record(1), _record(2), _record(3)]
    profile = {
        "similar_projects_pm": records,
        "selected_pm_performance": [
            {"name": "测试业绩3号工程"},
            {"name": "测试业绩1号工程"},
        ],
    }
    fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert [_cell_of(t, "项目名称") for t in details] == [
        "测试业绩3号工程",
        "测试业绩1号工程",
    ]


def test_td_selection_cleared_leaves_blank() -> None:
    """总工业绩人工清空([]) → 总工节留白且受保护,不回填名下全量。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("项目总工近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    profile = {
        "similar_projects_td": [_record(1, tech="许明英")],
        "selected_td_performance": [],  # 人工清空
    }
    result = fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == ""
    assert details[0]._tbl in result["handled_tables"]


def test_pm_no_selection_leaves_blank_all_manual() -> None:
    """没勾经理业绩(键缺失/None) → 留白(全部人工手选,不默认全填;用户2026-07-11拍板)。"""
    from services.similar_project_fill_service import fill_similar_project_sections

    doc = Document()
    doc.add_paragraph("项目经理近年完成的类似项目信息表（资格审查）")
    _add_detail_table(doc)
    profile = {
        "similar_projects_pm": [_record(1), _record(2)],
        "selected_pm_performance": None,  # 显式 None 等价键缺失=没勾
    }
    result = fill_similar_project_sections(doc, profile)
    details = _detail_tables(doc)
    assert len(details) == 1
    assert _cell_of(details[0], "项目名称") == ""  # 留白,不自动灌名下全量
    assert details[0]._tbl in result["handled_tables"]  # 受保护不被通用填表器乱填


def test_pick_by_selection_exact_name_beats_norm_collision() -> None:
    """两条只差括号内容的业绩(归一化后同名)各勾各的,不许都填成同一条的数据。"""
    from services.similar_project_fill_service import _pick_by_selection

    a = dict(_record(1), project_name="养护工程（一标段）")
    b = dict(_record(2), project_name="养护工程（二标段）")
    picked = _pick_by_selection(
        [a, b],
        [{"name": "养护工程（二标段）"}, {"name": "养护工程（一标段）"}],
    )
    assert [r["owner_name"] for r in picked] == ["发包人2", "发包人1"]


def test_parse_similar_projects_docx_roundtrip(tmp_path) -> None:
    """导入脚本解析:17行竖表 → 全字段记录(含换行切开的'项目经理(或注册建造师)')。"""
    import sys
    from pathlib import Path

    scripts_dir = str(Path(__file__).resolve().parents[1] / "scripts")
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    from import_similar_project_info import parse_similar_projects

    doc = Document()
    labels_values = [
        ("序号", "179"),
        ("项目名称", "合肥市某养护工程"),
        ("项目所在地", "巢湖市"),
        ("发包人名称", "合肥市公路管理局"),
        ("发包人地址", "合肥市梅山路77号"),
        ("发包人电话", "0551-65117872"),
        ("合同价格", "12127175.92元"),
        ("开工日期", "2020.9.5"),
        ("交工日期", "2022.1.12"),
        ("承担的工作", "施工总承包"),
        ("工程质量", "合格"),
        ("项目经理(或注册建\n造师)", "李刚"),
        ("技术负责人", ""),
        ("监理单位及联系电话", "某监理/0551-1"),
        ("总监理工程师及电话", "某总监/0551-2"),
        ("项目描述", "路面修复"),
        ("备注", "□资格审查用业绩"),
    ]
    t = doc.add_table(rows=len(labels_values), cols=2)
    for ri, (label, value) in enumerate(labels_values):
        t.cell(ri, 0).text = label
        t.cell(ri, 1).text = value
    path = tmp_path / "similar.docx"
    doc.save(str(path))

    records = parse_similar_projects(path)
    assert len(records) == 1
    r = records[0]
    assert r.project_name == "合肥市某养护工程"
    assert r.project_manager == "李刚"
    assert r.ledger_no == "179"
    assert r.project_year == 2022  # 以交工日期为准
    assert r.amount_wan == 1212.72
    assert r.owner_phone == "0551-65117872"
