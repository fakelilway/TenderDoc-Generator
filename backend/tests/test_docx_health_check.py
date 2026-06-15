"""docx_health_check.score_docx 的确定性体检测试。

现场用 python-docx 构造若干小 .docx 到 tmp_path，验证各打分项与硬否决。
"""

from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402

from services.docx_health_check import score_docx  # noqa: E402


def _add_kv_table(document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value


def _build_doc(path: Path, *, headings: int, body_chars: int, table_rows, extra_paras=()):
    document = Document()
    for i in range(headings):
        document.add_heading(f"第{i + 1}章 章节标题{i + 1}", level=1)
        # 给每章塞点正文以撑字数
        document.add_paragraph("正" * max(1, body_chars // max(1, headings)))
    for para in extra_paras:
        document.add_paragraph(para)
    if table_rows:
        _add_kv_table(document, table_rows)
    document.save(str(path))
    return path


# --------------------------------------------------------------------------- #
# (a) 含 {{foo}} 残留 -> hard_block True 且分 <= 40
# --------------------------------------------------------------------------- #


def test_curly_residue_triggers_hard_block(tmp_path: Path) -> None:
    path = tmp_path / "bad_curly.docx"
    _build_doc(
        path,
        headings=10,  # 章节满分
        body_chars=20000,  # 篇幅满分
        table_rows=[
            ("投标人", "某某建设工程有限公司"),
            ("法定代表人", "张三"),
            ("统一社会信用代码", "91320000MA1ABCDE7K"),
            ("安全生产许可证", "(苏)JZ安许证字[2020]001"),
            ("项目经理", "李四"),
        ],
        extra_paras=["本段落含残留模板标记 {{foo}} 未被替换。"],
    )
    result = score_docx(path)
    assert result["hard_block"] is True
    assert result["quality_score"] <= 40
    assert result["items"]["residue"]["score"] == 0.0


def test_knowledge_image_placeholder_hard_blocks(tmp_path: Path) -> None:
    path = tmp_path / "bad_image.docx"
    _build_doc(
        path,
        headings=8,
        body_chars=16000,
        table_rows=[("投标人", "某公司")],
        extra_paras=["{{knowledge_image:document_id=12 caption=营业执照}}"],
    )
    result = score_docx(path)
    assert result["hard_block"] is True
    assert result["quality_score"] <= 40


def test_meta_phrase_hard_blocks(tmp_path: Path) -> None:
    path = tmp_path / "bad_meta.docx"
    _build_doc(
        path,
        headings=8,
        body_chars=16000,
        table_rows=[("投标人", "某公司")],
        extra_paras=["作为您的投标助手，根据您的要求，以下是施工方案。"],
    )
    result = score_docx(path)
    assert result["hard_block"] is True
    assert result["quality_score"] <= 40
    assert result["items"]["residue"]["score"] == 0.0


# --------------------------------------------------------------------------- #
# (b) 空表格 cell 多 -> 表格填充率低
# --------------------------------------------------------------------------- #


def test_empty_cells_lower_table_fill(tmp_path: Path) -> None:
    path = tmp_path / "empty_cells.docx"
    _build_doc(
        path,
        headings=6,
        body_chars=3000,
        table_rows=[
            ("投标人", "________"),
            ("法定代表人", "________"),
            ("统一社会信用代码", "________"),
            ("安全生产许可证", "________"),
            ("项目经理", "________"),
        ],
    )
    result = score_docx(path)
    fill = result["items"]["table_fill"]
    # 5 行 2 列共 10 个 cell，值列 5 个都是下划线占位 -> 填充率约 0.5
    assert fill["score"] < 0.6
    # 无 profile 时，必填字段全是下划线空位 -> 完整度低
    assert result["items"]["required_fields"]["score"] < 0.5
    assert result["hard_block"] is False


# --------------------------------------------------------------------------- #
# (c) 正常填充的小文档 -> 各项分合理
# --------------------------------------------------------------------------- #


def test_well_filled_doc_scores_reasonably(tmp_path: Path) -> None:
    path = tmp_path / "good.docx"
    _build_doc(
        path,
        headings=10,
        body_chars=20000,
        table_rows=[
            ("投标人", "某某建设工程有限公司"),
            ("法定代表人", "张三"),
            ("统一社会信用代码", "91320000MA1ABCDE7K"),
            ("安全生产许可证", "(苏)JZ安许证字[2020]001"),
            ("项目经理", "李四"),
        ],
    )
    result = score_docx(path)
    assert result["hard_block"] is False
    assert result["quality_score"] >= 80
    items = result["items"]
    assert items["required_fields"]["score"] == 1.0
    assert items["table_fill"]["score"] == 1.0
    assert items["section_coverage"]["score"] == 1.0
    assert items["residue"]["score"] == 1.0
    assert items["length"]["score"] == 1.0
    # 权重之和 = 100
    assert sum(it["weight"] for it in items.values()) == 100


def test_score_is_deterministic(tmp_path: Path) -> None:
    path = tmp_path / "det.docx"
    _build_doc(
        path,
        headings=8,
        body_chars=16000,
        table_rows=[("投标人", "某公司"), ("法定代表人", "王五")],
    )
    first = score_docx(path)
    second = score_docx(path)
    assert first["quality_score"] == second["quality_score"]


def test_profile_checks_field_values(tmp_path: Path) -> None:
    path = tmp_path / "profile.docx"
    _build_doc(
        path,
        headings=8,
        body_chars=16000,
        table_rows=[
            ("投标人", "实达建设有限公司"),
            ("法定代表人", "赵六"),
        ],
    )
    # profile 给了 5 个字段，但文档只体现了 2 个 -> required_fields < 1.0
    profile = {
        "company_name": "实达建设有限公司",
        "legal_representative": "赵六",
        "credit_code": "91320000MA1ZZZZ9X",
        "safety_license_no": "(苏)JZ安许证字[2021]099",
        "project_manager_name": "孙七",
    }
    result = score_docx(path, profile=profile)
    rf = result["items"]["required_fields"]["score"]
    assert 0.3 < rf < 0.6  # 2/5 命中


def test_length_baseline_saturates(tmp_path: Path) -> None:
    path = tmp_path / "len.docx"
    _build_doc(path, headings=4, body_chars=7000, table_rows=[("投标人", "某公司")])
    # baseline 给 10000，目标 70%=7000，正好满分（饱和封顶 1.0）
    result = score_docx(path, baseline_chars=10000)
    assert result["items"]["length"]["score"] == 1.0
    # baseline 调高到 100000 -> 同样字数远不够 -> 篇幅项明显 < 1.0
    result2 = score_docx(path, baseline_chars=100000)
    assert result2["items"]["length"]["score"] < 0.2
