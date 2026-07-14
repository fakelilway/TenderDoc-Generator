"""员工反馈第2/3/6条(面子工程)的格式修复测试。

2: 填空时冒号后的长空格一并吃掉,值紧贴冒号,不再被空格推开。
3: 表格行禁止跨页拆分(cantSplit) + 标题与下文同页(keepNext)。
6: 章(二级标题)另起一页,不再与上一章续写。
"""

from docx import Document
from docx.oxml.ns import qn


def _page_break_count(document) -> int:
    return len([
        br for br in document.element.body.iter(qn("w:br"))
        if br.get(qn("w:type")) == "page"
    ])


# ── 第2条:填空前后空格 ────────────────────────────────────────────────


def test_inline_fill_eats_long_leading_spaces() -> None:
    """冒号后连排长空格视为书写槽,填值后值紧贴冒号,长空格消失。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("工程质量：　　　　　；安全目标：无事故")  # 5个全角空格槽
    _fill_inline_labeled_blanks(doc, {"质量": "优良"})
    assert doc.paragraphs[0].text == "工程质量：优良；安全目标：无事故"


def test_inline_fill_eats_spaces_before_underscore_slot() -> None:
    """空格+下划线混合槽:整段(空格含)被值替换。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("工程质量：  ＿＿＿＿；")
    _fill_inline_labeled_blanks(doc, {"质量": "合格"})
    assert doc.paragraphs[0].text == "工程质量：合格；"


def test_inline_fill_keeps_single_space_gap() -> None:
    """单个空格是正常间隔,保留(不破坏原排版习惯)。"""
    from services.original_docx_format_service import _fill_inline_labeled_blanks

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("工程质量： （按招标要求）")
    _fill_inline_labeled_blanks(doc, {"质量": "合格"})
    assert doc.paragraphs[0].text == "工程质量： 合格（按招标要求）"


# ── 第6条:章另起一页 ──────────────────────────────────────────────────


def test_chapter_headings_start_new_page() -> None:
    """二级标题(章)前自动分页;开头第一块内容前不分(防空白页)。"""
    from utils.docx_exporter import _render_markdown_body

    doc = Document()
    md = "# 某项目 技术文件\n\n## 第一章 总体布置\n正文A\n\n## 第二章 施工方案\n正文B\n"
    _render_markdown_body(doc, md, "zhengqi")
    # 卷标题后第一章前 1 次 + 第二章前 1 次 = 2
    assert _page_break_count(doc) == 2


def test_first_chapter_at_doc_start_no_break() -> None:
    """文档一开头就是章标题 → 不插分页(当前页还没内容)。"""
    from utils.docx_exporter import _render_markdown_body

    doc = Document()
    _render_markdown_body(doc, "## 第一章\n正文\n", "zhengqi")
    assert _page_break_count(doc) == 0


def test_explicit_pagebreak_not_doubled() -> None:
    """显式分页标记后紧跟章标题 → 不再叠加第二个分页。"""
    from utils.docx_exporter import _render_markdown_body

    doc = Document()
    md = "正文…\n\n<!-- tdg:pagebreak -->\n\n## 新章\n内容\n"
    _render_markdown_body(doc, md, "zhengqi")
    assert _page_break_count(doc) == 1  # 只有显式那一个


def test_third_level_heading_no_break() -> None:
    """三级标题(节)不分页,只有章分页。"""
    from utils.docx_exporter import _render_markdown_body

    doc = Document()
    md = "## 第一章\n正文\n\n### 1.1 小节\n内容\n\n### 1.2 小节\n内容\n"
    _render_markdown_body(doc, md, "zhengqi")
    assert _page_break_count(doc) == 0


# ── 第3条:防内容被拆页 ────────────────────────────────────────────────


def test_heading_styles_keep_with_next() -> None:
    """标题样式带"与下文同页",标题不孤悬页底。"""
    from utils.docx_exporter import _configure_styles

    doc = Document()
    _configure_styles(doc, "zhengqi")
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        assert doc.styles[name].paragraph_format.keep_with_next is True


def test_table_rows_cant_split() -> None:
    """heal_table_row_integrity 给所有行加 cantSplit;幂等(第二次跑=0)。"""
    from services.docx_format_doctor import heal_table_row_integrity

    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "姓名"
    fixed = heal_table_row_integrity(doc)
    assert fixed == 3
    for row in t.rows:
        trPr = row._tr.find(qn("w:trPr"))
        assert trPr is not None and trPr.find(qn("w:cantSplit")) is not None
    assert heal_table_row_integrity(doc) == 0  # 幂等


def test_assembled_doctor_includes_row_integrity() -> None:
    """成品级体检注册了表格行防拆页。"""
    from services.docx_format_doctor import run_format_doctor_assembled

    doc = Document()
    doc.add_table(rows=2, cols=2)
    report = run_format_doctor_assembled(doc)
    assert report.get("table_row_integrity") == 2


def test_extract_credit_requirement_items_from_144():
    """招标1.4.4信誉要求逐条抽取(供信誉情况表左列)。"""
    from services.v2_generation_service import _extract_credit_requirement_items

    tender = (
        "1.4.3 财务要求:无。\n"
        "1.4.4 信誉要求:(1)没有处于被责令停业,投标资格被取消或者财产被接管、冻结和破产状态;\n"
        "（2）近三年没有骗取中标或者严重违约以及重大工程质量问题；\n"
        "(3) 不属于失信被执行人。\n"
        "1.4.5 项目经理资格要求:二级建造师。\n"
    )
    items = _extract_credit_requirement_items(tender)
    assert len(items) == 3
    assert items[0].startswith("没有处于被责令停业")
    assert "失信被执行人" in items[2]
    assert _extract_credit_requirement_items("没有相关章节") == []


def test_credit_table_filled_from_items():
    """(五)投标人的信誉情况表:空表按1.4.4逐条填,行不够自动扩;有预印内容不碰。"""
    from docx import Document

    from services.original_docx_format_service import _fill_credit_status_table

    doc = Document()
    t = doc.add_table(rows=3, cols=2)  # 表头+2空行,3条要求→自动扩1行
    t.cell(0, 0).text = "项 目"
    t.cell(0, 1).text = "投标人情况说明"
    profile = {"credit_requirement_items": ["没有处于被责令停业状态", "近三年无重大质量问题", "不属于失信被执行人"]}
    filled = _fill_credit_status_table(doc, profile)
    assert filled == 3
    assert len(t.rows) == 4
    assert t.cell(1, 0).text.strip() == "没有处于被责令停业状态"
    assert t.cell(1, 1).text.strip() == "无此类情形"
    assert t.cell(3, 0).text.strip() == "不属于失信被执行人"

    # 有预印内容的表不碰
    doc2 = Document()
    t2 = doc2.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "项 目"
    t2.cell(0, 1).text = "投标人情况说明"
    t2.cell(1, 0).text = "招标预印的一行"
    assert _fill_credit_status_table(doc2, profile) == 0
    assert t2.cell(1, 1).text.strip() == ""


def test_honor_certs_not_anchored_to_credit_table():
    """荣誉/信用证书组不锚信誉情况表(泗沙路用户点名"信誉表后不跟奖项")。"""
    from services.v2_generation_service import _EVIDENCE_GROUPS

    entry = next(g for g in _EVIDENCE_GROUPS if g[0] == "企业荣誉与信誉证明")
    assert entry[3] == ""  # 空锚=卷尾


def test_single_resume_table_cloned_for_two_roles(monkeypatch):
    """单表双人(泗沙路"项目经理和项目总工资历表"):克隆一张,表1经理表2总工,
    各自证件紧跟各自表,并设让位标志。"""
    from docx import Document
    from docx.text.paragraph import Paragraph

    from services import generation_service as g
    from services import v2_generation_service as v2
    from services.original_docx_format_service import _fill_resume_tables

    def fake_insert(anchor, doc, doc_id, caption, width):
        p = doc.add_paragraph()._p
        p.getparent().remove(p)
        anchor.addnext(p)
        Paragraph(p, doc).add_run(f"[证{doc_id}:{caption}]")
        return p

    monkeypatch.setattr(g, "_insert_image_after", fake_insert)
    monkeypatch.setattr(
        v2, "person_cert_documents",
        lambda role, name: [(101 if name == "江甜甜" else 201, f"{role}（{name}）身份证")],
    )

    doc = Document()
    doc.add_paragraph("（六）拟委任的项目经理和项目总工资历表")
    t = doc.add_table(rows=3, cols=4)
    t.cell(0, 0).text = "姓名"
    t.cell(1, 0).text = "技术职称"
    t.cell(2, 0).text = "拟在本标段工程任职"
    doc.add_paragraph("注：本表应填写项目经理和项目总工相关情况。")

    profile = {}
    _fill_resume_tables(
        doc,
        {"姓名": "江甜甜", "拟任职务": "项目经理"},
        {"姓名": "许明英", "拟任职务": "项目技术负责人"},
        profile=profile,
    )
    tables = [tb for tb in doc.tables]
    assert len(tables) == 2  # 克隆出第二张
    assert tables[0].cell(0, 1).text.strip() == "江甜甜"
    assert tables[1].cell(0, 1).text.strip() == "许明英"
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    i_pm_cert = texts.index("[证101:项目经理（江甜甜）身份证]")
    i_td_cert = texts.index("[证201:项目技术负责人（许明英）身份证]")
    i_note2 = len(texts) - 1 - texts[::-1].index("注：本表应填写项目经理和项目总工相关情况。")
    assert i_pm_cert < i_note2 <= i_td_cert  # 经理证件在总工那套(注2)之前,总工证件在其后
    assert set(profile.get("_personnel_certs_inline_roles") or []) == {"项目经理", "项目技术负责人"}


def test_tech_only_single_resume_table_not_cloned():
    """总工**专属**单表(标题只提总工)绝不克隆、绝不灌项目经理(防员工第11条回归)。"""
    from docx import Document

    from services.original_docx_format_service import _fill_resume_tables

    doc = Document()
    doc.add_paragraph("（六）拟委任的项目总工资历表")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "姓名"
    t.cell(1, 0).text = "拟在本标段工程任职"

    profile = {}
    _fill_resume_tables(
        doc,
        {"姓名": "江甜甜", "拟任职务": "项目经理"},
        {"姓名": "许明英", "拟任职务": "项目技术负责人"},
        profile=profile,
    )
    assert len(doc.tables) == 1  # 不克隆
    assert doc.tables[0].cell(0, 1).text.strip() == "许明英"  # 总工的表填总工


def test_credit_items_extracted_even_when_toc_lists_144_first():
    """目次里也有"1.4.4 信誉要求……页码"行:抽取要跳过它,取正文里能出条目的那处。"""
    from services.v2_generation_service import _extract_credit_requirement_items

    tender = (
        "目 次\n1.4.4 信誉要求 12\n1.4.5 项目经理资格 13\n"
        "……正文……\n"
        "1.4.4 信誉要求:(1)没有处于被责令停业状态;（2）不属于失信被执行人。\n"
        "1.4.5 项目经理资格要求:二级建造师。\n"
    )
    items = _extract_credit_requirement_items(tender)
    assert len(items) == 2 and "失信被执行人" in items[1]
