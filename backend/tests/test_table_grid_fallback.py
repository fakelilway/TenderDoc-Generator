"""回归:底板缺"Table Grid"样式(福昕/pdf2docx转换件)时,渲染表格不崩、图片照插。

历史 bug:#151 商务卷 0 图——福昕底板无"Table Grid"样式,table.style 赋值抛 KeyError
把整段渲染带崩→兜底吞掉证件插图。本测试守住 _apply_table_grid 的兜底。
"""

from __future__ import annotations

from docx import Document

from utils.docx_exporter import _apply_table_grid, _render_markdown_body


def _strip_table_grid(doc) -> None:
    """删掉文档里的 Table Grid 样式,模拟福昕转换件。"""
    try:
        el = doc.styles["Table Grid"].element
        el.getparent().remove(el)
    except KeyError:
        pass


def test_apply_table_grid_no_crash_without_style() -> None:
    doc = Document()
    _strip_table_grid(doc)
    table = doc.add_table(rows=2, cols=2)
    _apply_table_grid(table)  # 不该抛错
    # 退化为直接写边框 XML
    assert table._tbl.tblPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders"
    ) is not None


def test_render_table_and_image_survives_missing_table_grid(monkeypatch) -> None:
    # 渲染含表格 + 图片标记的 markdown,在缺样式的底板上不崩、且插入图片
    monkeypatch.setattr(
        "utils.docx_exporter._resolve_image_bytes_noop", lambda *a, **k: None, raising=False
    )
    doc = Document()
    _strip_table_grid(doc)
    md = (
        "## 资格审查\n\n"
        "| 项目 | 值 |\n| --- | --- |\n| 公司 | 正奇 |\n\n"
        "{{knowledge_image:document_id=999 caption=\"证件\"}}\n"
    )
    inserted = {"n": 0}

    def resolver(doc_id):
        # 返回一个最小合法 PNG 字节
        import base64
        inserted["n"] += 1
        return base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
        )

    # 不该抛错(表格样式缺失也能渲染)
    _render_markdown_body(doc, md, "zhengqi", resolver)
    assert inserted["n"] == 1  # 图片解析器被调用,图已插
