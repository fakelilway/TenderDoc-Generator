"""格式体检(docx_format_doctor)测试:治下划线断线,铁律=只改格式一字不改。

病例全部来自真实文件(巢湖商务卷 v7)的实测结构。
"""

from docx import Document

from services.docx_format_doctor import heal_underline_slots, run_format_doctor


def _mk_para(doc, runs):
    """runs: [(text, underlined)]"""
    p = doc.add_paragraph()
    for text, u in runs:
        r = p.add_run(text)
        if u:
            r.font.underline = True
    return p


def test_sandwich_value_gets_underlined() -> None:
    """投标函项目名:带线空白+无线值+带线空白 → 值上线,整条线连续。"""
    doc = Document()
    p = _mk_para(doc, [
        ("我方已仔细研究", False), (" ", True), ("\t", True),
        ("巢湖市 2026 年农村公路养护工程（栏杆集镇周岗路等 8 条路）", False),
        (" ", True), ("\t", True), ("标段招标文件的全部内容，", False),
    ])
    before = p.text
    n = heal_underline_slots(doc)
    assert n == 1
    assert p.runs[3].font.underline  # 项目名上线
    assert not p.runs[0].font.underline  # 槽外正文不动
    assert not p.runs[6].font.underline
    assert p.text == before  # 一字不改


def test_short_labels_year_month_day_untouched() -> None:
    """日期槽"年/月/日"被线夹着是招标原样,绝不误加线。"""
    doc = Document()
    p = _mk_para(doc, [
        ("日期：", False), (" ", True), ("\t", True), ("年", False),
        (" ", True), ("\t", True), ("月", False),
        (" ", True), ("\t", True), ("日", False),
    ])
    heal_underline_slots(doc)
    for r in p.runs:
        if r.text in ("年", "月", "日"):
            assert not r.font.underline


def test_whitelist_splits_value_from_tail() -> None:
    """抬头"招标人名称："/委托书"公司名的法定代表人":拆 run 只给值上线,尾巴不动。"""
    doc = Document()
    p1 = _mk_para(doc, [(" ", True), ("\t", True), ("巢湖市栏杆集镇人民政府：", False)])
    p2 = _mk_para(doc, [
        ("系", False), (" ", True), ("\t", True),
        ("安徽正奇建设有限公司的法定代表人，现委托", False), (" ", True), ("\t", True),
    ])
    t1, t2 = p1.text, p2.text
    profile = {"招标人": "巢湖市栏杆集镇人民政府", "company_name": "安徽正奇建设有限公司"}
    n = heal_underline_slots(doc, profile)
    assert n == 2
    assert p1.runs[2].font.underline and p1.runs[2].text == "巢湖市栏杆集镇人民政府"
    assert not p1.runs[3].font.underline and p1.runs[3].text == "："
    assert p2.runs[3].font.underline and p2.runs[3].text == "安徽正奇建设有限公司"
    assert not p2.runs[4].font.underline  # "的法定代表人，现委托" 不上线
    assert p1.text == t1 and p2.text == t2  # 一字不改


def test_sentence_text_sandwiched_not_underlined_without_whitelist() -> None:
    """值+原文连排(带句读)且 profile 没给值 → 夹心兜底不动它(宁漏勿错)。"""
    doc = Document()
    p = _mk_para(doc, [
        (" ", True), ("\t", True),
        ("安徽正奇建设有限公司的法定代表人，现委托", False), (" ", True),
    ])
    heal_underline_slots(doc)  # 无 profile
    assert not p.runs[2].font.underline


def test_hint_parenthetical_untouched() -> None:
    """（其他补充说明）等原文提示语不上线。"""
    doc = Document()
    p = _mk_para(doc, [(" ", True), ("\t", True), ("（其他补充说明）。", False), (" ", True)])
    heal_underline_slots(doc)
    assert not p.runs[2].font.underline


def test_orphan_split_labels_rejoined() -> None:
    """孤字归位(巢湖v7身份证明实测病例):性/职 粘在带线值后 + 下文"别：/务："段 → 拼回。"""
    import re

    doc = Document()
    p = _mk_para(doc, [
        ("姓", False), (" ", False), ("名：", False), (" ", True), ("许明英", True),
        ("性", False),
        ("年", False), (" ", False), ("龄：", False), (" ", True), ("50", True),
        ("职", False),
    ])
    p2 = _mk_para(doc, [("别：", False), (" ", True), ("\t", True)])
    p3 = _mk_para(doc, [("务：", False), (" ", True), ("\t", True)])
    from services.docx_format_doctor import heal_orphan_split_labels

    before_chars = sorted(re.sub(r"\s+", "", p.text + p2.text + p3.text))
    n = heal_orphan_split_labels(doc)
    assert n == 2
    assert p2.text.startswith("性别：")
    assert p3.text.startswith("职务：")
    assert "性" not in p.text.replace("姓", "") or "性年" not in p.text  # 孤字已搬走
    assert "\n" in p.text  # 原位补了换行,姓名/年龄恢复两行
    after_chars = sorted(re.sub(r"\s+", "", p.text + p2.text + p3.text))
    assert before_chars == after_chars  # 非空白字符一个不多不少


def test_orphan_rejoined_prefill_empty_slots() -> None:
    """填前状态(槽还是带线空白、值未填)也要能归位——本 healer 就跑在填值前。

    回归:曾要求孤字前是"带线的已填值",填前空槽不满足 → 生产上永不触发(萧县实测)。
    """
    doc = Document()
    p = _mk_para(doc, [
        ("姓", False), (" ", False), ("名：", False), (" \t", True),  # 空槽(带线空白)
        ("性", False),
        ("年", False), (" ", False), ("龄：", False), (" \t", True),
        ("职", False),
    ])
    p2 = _mk_para(doc, [("别：", False), (" \t", True)])
    p3 = _mk_para(doc, [("务：", False), (" \t", True)])
    from services.docx_format_doctor import heal_orphan_split_labels

    n = heal_orphan_split_labels(doc)
    assert n == 2
    assert p2.text.startswith("性别：")
    assert p3.text.startswith("职务：")
    assert "性" not in p.text and "职" not in p.text  # 孤字都搬走了


def test_orphan_guard_normal_single_char_not_moved() -> None:
    """"姓 名："里的"姓"(前面没有带线值)绝不被搬走——哪怕后文恰有"名："段。"""
    doc = Document()
    p = _mk_para(doc, [("姓", False), (" ", False), ("名：", False), (" ", True), ("许明英", True)])
    p2 = _mk_para(doc, [("名：", False), (" ", True)])
    from services.docx_format_doctor import heal_orphan_split_labels

    n = heal_orphan_split_labels(doc)
    assert n == 0
    assert p.text.startswith("姓 名：")
    assert p2.text.startswith("名：")


def test_run_format_doctor_never_raises(monkeypatch) -> None:
    """healer 崩了也不阻断出标,报告记 0。"""
    import services.docx_format_doctor as m

    def _boom(document, profile=None):
        raise RuntimeError("boom")

    monkeypatch.setattr(m, "_HEALERS", (("underline_slots", _boom),))
    report = m.run_format_doctor(Document())
    assert report == {"underline_slots": 0}


def test_filler_blank_runs_collapsed_but_slots_kept() -> None:
    """连续≥4个真空段压缩到2;带制表位的空槽段(日期下划线槽)绝不删。"""
    from services.docx_format_doctor import heal_filler_blank_runs

    doc = Document()
    doc.add_paragraph("正文")
    for _ in range(6):
        doc.add_paragraph("")
    slot = doc.add_paragraph()
    slot.add_run("\t")  # 空槽段(只有制表位)
    doc.add_paragraph("下一节")
    n = heal_filler_blank_runs(doc)
    assert n == 4  # 6个空段 → 留2删4
    texts = [p.text for p in doc.paragraphs]
    assert "正文" in texts and "下一节" in texts
    assert any("\t" in p.text for p in doc.paragraphs)  # 空槽段还在


def _set_spacing(p, line, rule):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    ppr = p._p.get_or_add_pPr()
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        ppr.append(sp)
    sp.set(qn("w:line"), str(line))
    sp.set(qn("w:lineRule"), rule)


def _spacing_of(p):
    from docx.oxml.ns import qn
    sp = p._p.get_or_add_pPr().find(qn("w:spacing"))
    return sp.get(qn("w:lineRule")), sp.get(qn("w:line"))


def test_line_spacing_unclips_exact_and_evens_loose() -> None:
    """福昕叠行修复:固定行高(exact,小于单倍)→自动且不小于240;过松(>300)→收单倍;文字不动。"""
    from services.docx_format_doctor import heal_line_spacing

    doc = Document()
    p_clip = doc.add_paragraph("被压叠的密集段落文字保持不变")
    _set_spacing(p_clip, 225, "exact")          # 固定行高偏小 → 叠行
    p_loose = doc.add_paragraph("过松的行")
    _set_spacing(p_loose, 397, "auto")          # >1.25倍 → 过松
    p_ok = doc.add_paragraph("正常单倍不动")
    _set_spacing(p_ok, 240, "auto")

    n = heal_line_spacing(doc)
    assert n == 2  # 只动叠行+过松两段
    assert _spacing_of(p_clip) == ("auto", "240")   # 固定→自动,提到单倍
    assert _spacing_of(p_loose) == ("auto", "240")  # 过松→单倍
    assert _spacing_of(p_ok) == ("auto", "240")     # 本就单倍,没被动
    # 文字逐字不变(红线)
    assert p_clip.text == "被压叠的密集段落文字保持不变"


def test_line_spacing_fixes_table_cell_paragraphs() -> None:
    """表格单元格里的叠行也要治(福昕表单大量exact在单元格内)。"""
    from services.docx_format_doctor import heal_line_spacing

    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    cp = t.cell(0, 0).paragraphs[0]
    cp.add_run("单元格内密集文字")
    _set_spacing(cp, 220, "exact")
    n = heal_line_spacing(doc)
    assert n == 1
    assert _spacing_of(cp) == ("auto", "240")


def test_signature_wrap_kills_fuxin_giant_indent() -> None:
    """落款折行修复(埇桥实测病例):右对齐+福昕大缩进(5486)+填了公司名 → 缩进砍0,
    一行排下;(签章)空槽段制表位画出页外 → 制表位收进来。左对齐段/小缩进段不动。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_signature_wrap

    doc = Document()

    def _sig_para(text, left, tab_pos=None, align="right"):
        p = doc.add_paragraph(text)
        ppr = p._p.get_or_add_pPr()
        jc = OxmlElement("w:jc"); jc.set(_qn("w:val"), align); ppr.append(jc)
        ind = OxmlElement("w:ind"); ind.set(_qn("w:left"), str(left)); ppr.append(ind)
        if tab_pos:
            tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
            tab.set(_qn("w:val"), "left"); tab.set(_qn("w:pos"), str(tab_pos))
            tabs.append(tab); ppr.append(tabs)
        return p, ind

    p1, ind1 = _sig_para("投 标 人： 安徽正奇建设有限公司（盖单位章）", 5486)
    p2, ind2 = _sig_para("法定代表人或其委托代理人： \t（签章）", 3806, tab_pos=8850)
    p3, ind3 = _sig_para("短日期行", 5486)          # 排得下,不动
    p4, ind4 = _sig_para("左对齐正文段落不许被碰,不管多长的文字都轮不到这个修", 5486, align="left")

    n = heal_signature_wrap(doc)
    assert n == 2  # 只修 p1(公司名超宽) + p2(制表位超页宽)
    assert ind1.get(_qn("w:left")) == "0"           # 大缩进砍0,右对齐仍贴右
    assert ind2.get(_qn("w:left")) == "0"
    tab_el = p2._p.find(_qn("w:pPr")).find(_qn("w:tabs")).find(_qn("w:tab"))
    assert int(tab_el.get(_qn("w:pos"))) < 8850     # 制表位收进页内
    assert ind3.get(_qn("w:left")) == "5486"        # 排得下的原样保真
    assert ind4.get(_qn("w:left")) == "5486"        # 左对齐段绝不碰
    assert p1.text == "投 标 人： 安徽正奇建设有限公司（盖单位章）"  # 文字红线


def test_form_lines_never_merged() -> None:
    """表单行两头保护(用户实测暴走回归):"投标人：公司名"不吞"单位性质：…",
    "地址：…"不连吞"成立时间/经营期限";正文劈段照常合并。"""
    from services.docx_format_doctor import heal_split_paragraphs

    doc = Document()
    doc.add_paragraph("投 标 人：安徽正奇建设有限公司")
    doc.add_paragraph("单位性质：有限责任公司（自然人投资或控股）")
    doc.add_paragraph("地    址：安徽省合肥市庐阳区蒙城北路1708室")
    doc.add_paragraph("成立时间：2011年7月5日")
    doc.add_paragraph("经营期限：2011年07月05日至2051年06月28日")
    # 对照:真正的正文劈段(长句无冒号结尾断在半路)仍要合并
    doc.add_paragraph("我方将按照合同附件提出的最低要求填报派驻本标段的其他管")
    doc.add_paragraph("理和技术人员及主要机械设备，经你方审批后不再更换。")

    n = heal_split_paragraphs(doc)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "投 标 人：安徽正奇建设有限公司" in texts          # 各自独立成行
    assert "单位性质：有限责任公司（自然人投资或控股）" in texts
    assert "成立时间：2011年7月5日" in texts
    assert any("其他管理和技术人员" in t for t in texts)       # 正文劈段照常并
    assert n == 1  # 只并了正文那一处


def test_form_line_br_kept() -> None:
    """段内 br 分隔的表单行(投标人：xx[br]单位性质：xx)换行保留,不被当句中断行删掉。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_midsentence_breaks

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("投标人名称：安徽正奇建设有限公司")
    br_run = p.add_run()
    br_run._r.append(OxmlElement("w:br"))
    p.add_run("姓名：许明英")
    n = heal_midsentence_breaks(doc)
    assert n == 0  # 表单行换行一个不删
    assert p._p.find(f".//{_qn('w:br')}") is not None


def test_signature_block_layout_splits_and_rejoins() -> None:
    """落款3行修复(巢湖实测):①投标人+法代挤一段→拆开;②"日"与"期："拆两段→合回;
    ③三行对齐同一缩进。文字一字不改。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_signature_block_layout

    doc = Document()

    def _p(text, left=5096, tabs_before=None):
        p = doc.add_paragraph()
        ppr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind"); ind.set(_qn("w:left"), str(left)); ppr.append(ind)
        for seg in (text if isinstance(text, list) else [text]):
            if seg == "\t":
                r = OxmlElement("w:r"); r.append(OxmlElement("w:tab")); p._p.append(r)
            else:
                p.add_run(seg)
        return p

    # ① 投标人+法代挤在同一段
    p_merge = _p(["投 标 人： 安徽正奇建设有限公司（盖单位章） ", "法定代表人：", "\t", "（签字或盖章）"])
    doc.add_paragraph("")
    # ② 日期被拆成"日" + 空段 + "期：__年__月__日"
    _p("日")
    doc.add_paragraph("")
    _p(["期：", "\t", "年", "\t", "月", "\t", "日"], left=0)

    n = heal_signature_block_layout(doc)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    joined = "".join(texts).replace(" ", "").replace("\t", "")
    # ① 投标人行与法代行分开成两段
    assert any(t.strip().startswith("投") and "盖单位章" in t and "法定代表人" not in t for t in texts)
    assert any(t.strip().startswith("法定代表人") and "签字或盖章" in t for t in texts)
    # ② 日期合回一段(不再有单独的"日"段)
    assert not any(t.strip() == "日" for t in texts)
    assert any(t.strip().startswith("日期") and "年" in t and "月" in t for t in texts)
    # 文字红线:所有字仍在
    assert "投标人：安徽正奇建设有限公司（盖单位章）" in joined
    assert "法定代表人：（签字或盖章）" in joined
    assert "日期：年月日" in joined
    assert n >= 2


def test_signature_block_layout_leaves_normal_paragraphs_alone() -> None:
    """普通正文/单独的投标人行(不含法代)不被拆;正文一个字不动。"""
    from services.docx_format_doctor import heal_signature_block_layout

    doc = Document()
    doc.add_paragraph("我方已仔细研究招标文件的全部内容，愿意按合同约定完成承包工程。")
    doc.add_paragraph("投 标 人：安徽正奇建设有限公司（盖单位章）")  # 单独一行,无法代→不拆
    doc.add_paragraph("这是正常段落，不含任何签署标记，绝不能被动。")
    before = [p.text for p in doc.paragraphs]
    heal_signature_block_layout(doc)
    after = [p.text for p in doc.paragraphs]
    assert before == after  # 一段没动


def test_signature_date_wrap_labeled_and_bare() -> None:
    """填好的日期行折行修复(埇桥实测):带标签"日期：2026年7月8日"和光日期
    "2026年 7月 8日"在大缩进下会折→缩到排得下;与投标人行同块则一起对齐。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_signature_block_layout, _usable_width_twips

    doc = Document()
    avail = _usable_width_twips(doc)
    big = avail - 600  # 大到日期一定会折

    def _p(text, left):
        p = doc.add_paragraph(text)
        ppr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind"); ind.set(_qn("w:left"), str(left)); ppr.append(ind)
        return p

    p_bid = _p("投 标 人： 安徽正奇建设有限公司（盖单位章）", big)
    p_date = _p("日期：2026年7月8日", big)
    heal_signature_block_layout(doc)
    # 两行都缩到排得下(缩进 + 文字宽 <= 可用宽)
    def _left(p):
        return int(p._p.find(_qn("w:pPr")).find(_qn("w:ind")).get(_qn("w:left")))
    def _w(t):
        return sum(290 if ord(c) > 0x2E80 else 145 for c in t)
    assert _left(p_bid) + _w(p_bid.text) <= avail
    assert _left(p_date) + _w(p_date.text) <= avail
    assert _left(p_bid) == _left(p_date)  # 同块对齐


def test_signature_cover_no_wrap_untouched() -> None:
    """不折的封面落款(投标人+日期都排得下)一律不碰,缩进原样。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_signature_block_layout

    doc = Document()
    def _p(text, left):
        p = doc.add_paragraph(text)
        ppr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind"); ind.set(_qn("w:left"), str(left)); ppr.append(ind)
        return p
    p_bid = _p("投标人： 安徽正奇建设有限公司（盖单位章）", 1767)
    p_date = _p("2026年 7月 8日", 3303)
    heal_signature_block_layout(doc)
    assert int(p_bid._p.find(_qn("w:pPr")).find(_qn("w:ind")).get(_qn("w:left"))) == 1767
    assert int(p_date._p.find(_qn("w:pPr")).find(_qn("w:ind")).get(_qn("w:left"))) == 3303


def _add_cols_section(doc, num, col_widths):
    """给最后一段挂一个分栏 sectPr(段内),模拟福昕导出的分节。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    ppr = doc.paragraphs[-1]._p.get_or_add_pPr()
    sect = OxmlElement("w:sectPr")
    cols = OxmlElement("w:cols")
    cols.set(_qn("w:num"), str(num))
    for w in col_widths:
        c = OxmlElement("w:col"); c.set(_qn("w:w"), str(w)); cols.append(c)
    sect.append(cols)
    ppr.append(sect)
    return sect


def test_signature_columns_straightens_fake_two_col() -> None:
    """福昕给"日期：..年..月..日"单独起假两栏(第一栏太窄)导致日期折行→拉回单栏。
    实测巢湖:日期行独占一节且设 num=2,第一栏仅 5336twips(~267pt)。"""
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_signature_columns

    doc = Document()
    doc.add_paragraph("日期： 2026年 7月 9日")
    sect = _add_cols_section(doc, 2, [5336, 3707])

    fixed = heal_signature_columns(doc)
    assert fixed == 1
    cols = sect.find(_qn("w:cols"))
    assert cols.get(_qn("w:num")) == "1"
    assert cols.findall(_qn("w:col")) == []  # 窄栏定义清掉


def test_signature_columns_leaves_real_body_columns() -> None:
    """真正的大段两栏正文(辖段多、无落款标记)一律不碰。"""
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_signature_columns

    doc = Document()
    for i in range(12):
        doc.add_paragraph(f"这是正文第{i}段,双栏排版的普通段落内容,没有任何落款标记。")
    sect = _add_cols_section(doc, 2, [4800, 4800])

    fixed = heal_signature_columns(doc)
    assert fixed == 0
    assert sect.find(_qn("w:cols")).get(_qn("w:num")) == "2"  # 原样保留


def _cover_pair_section(doc, left_text, right_text, col_widths):
    """造一对被福昕拆两栏的封面行:左半段 + 空br段 + 右半段(挂两栏 continuous sectPr)。
    封面节都是 continuous(福昕实测),healer 靠"第一个非continuous节"定位封面结束。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    doc.add_paragraph(left_text)
    br_p = doc.add_paragraph()
    br_p.add_run()._r.append(OxmlElement("w:br"))
    doc.add_paragraph(right_text)
    sect = _add_cols_section(doc, 2, col_widths)
    typ = OxmlElement("w:type"); typ.set(_qn("w:val"), "continuous")
    sect.insert(0, typ)
    return sect


def _mark_next_page_section(doc):
    """给最后一段挂一个"下一页"分节,模拟封面之后目录起点。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    ppr = doc.paragraphs[-1]._p.get_or_add_pPr()
    sect = OxmlElement("w:sectPr")
    typ = OxmlElement("w:type"); typ.set(_qn("w:val"), "nextPage"); sect.append(typ)
    ppr.append(sect)
    return sect


def test_cover_columns_merges_split_title_and_bidder() -> None:
    """福昕把封面"项目名|标段招标""投标人|盖单位章"拆两栏→合并成一行、拉直居中,竖向铺开。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_cover_columns

    doc = Document()
    # 标题对(两栏)
    _cover_title_sect = _cover_pair_section(
        doc, "巢湖市2026年农村公路养护工程（栏杆集镇周岗路等8条路）", " \t标段招标", [5138, 4207]
    )
    doc.add_paragraph("投标文件")
    doc.add_paragraph("（商务文件）")
    # 投标文件/商务文件自成一节(单栏 continuous),与投标人对隔开(福昕实测结构)
    _mid = _add_cols_section(doc, 1, [])
    _mid_typ = OxmlElement("w:type"); _mid_typ.set(_qn("w:val"), "continuous")
    _mid.insert(0, _mid_typ)
    # 投标人对(两栏)
    _cover_bidder_sect = _cover_pair_section(
        doc, "投标人：安徽正奇建设有限公司", "（盖单位章）", [2853, 2570]
    )
    doc.add_paragraph("2026年 7月 8日")
    _mark_next_page_section(doc)  # 封面结束标记

    fixed = heal_cover_columns(doc)
    assert fixed == 2  # 两对都合并
    # 两栏都拉直
    assert _cover_title_sect.find(_qn("w:cols")).get(_qn("w:num")) == "1"
    assert _cover_bidder_sect.find(_qn("w:cols")).get(_qn("w:num")) == "1"
    # 标题行合并了"标段招标"、且居中
    title = next(p for p in doc.paragraphs if "标段招标" in p.text)
    assert "巢湖市" in title.text and "标段招标" in title.text  # 项目名+标段招标同一段
    assert title._p.find(_qn("w:pPr")).find(_qn("w:jc")).get(_qn("w:val")) == "center"
    # 投标人行合并了盖单位章
    bidder = next(p for p in doc.paragraphs if "投标人" in p.text)
    assert "盖单位章" in bidder.text
    # 投标文件行补了段前距(竖向铺开)
    tf = next(p for p in doc.paragraphs if p.text.strip() == "投标文件")
    sp = tf._p.find(_qn("w:pPr")).find(_qn("w:spacing"))
    assert sp is not None and int(sp.get(_qn("w:before"))) >= 1000


def test_cover_columns_ignores_non_cover_document() -> None:
    """不含封面特征(标段招标/投标文件/盖单位章)的文档:一律不碰。"""
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_cover_columns

    doc = Document()
    doc.add_paragraph("普通正文标题")
    sect = _cover_pair_section(doc, "左边内容", "右边内容", [5000, 4000])
    _mark_next_page_section(doc)

    fixed = heal_cover_columns(doc)
    assert fixed == 0
    assert sect.find(_qn("w:cols")).get(_qn("w:num")) == "2"  # 原样保留


def test_checklist_table_widths_reallocates_columns() -> None:
    """合规自查核对表(7列,表头含核对项/判定/处置)默认等宽→按内容重分配列宽。"""
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_checklist_table_widths

    doc = Document()
    table = doc.add_table(rows=2, cols=7)
    for c, t in zip(
        table.rows[0].cells,
        ["核对项", "出处", "招标要求", "我方/取值", "判定", "处置", "备注"],
    ):
        c.text = t
    fixed = heal_checklist_table_widths(doc)
    assert fixed == 1
    tblpr = table._tbl.find(_qn("w:tblPr"))
    assert tblpr.find(_qn("w:tblLayout")).get(_qn("w:type")) == "fixed"
    grid = table._tbl.find(_qn("w:tblGrid"))
    widths = [int(g.get(_qn("w:w"))) for g in grid.findall(_qn("w:gridCol"))]
    assert widths[6] > widths[4]  # 备注列比判定列宽(长文本给足)


def test_checklist_table_widths_leaves_other_tables() -> None:
    """普通表格(非核对表)一律不碰。"""
    from services.docx_format_doctor import heal_checklist_table_widths

    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    for c, t in zip(table.rows[0].cells, ["序号", "名称", "数量"]):
        c.text = t
    assert heal_checklist_table_widths(doc) == 0


def _add_big_image_paragraph(doc, cy_pt):
    """造一个含"大图"的段落(用 wp:extent 声明显示尺寸)。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    p = doc.add_paragraph()
    r = p.add_run()
    drawing = OxmlElement("w:drawing")
    extent = OxmlElement("wp:extent")
    extent.set("cx", "3600000")
    extent.set("cy", str(int(cy_pt * 12700)))
    drawing.append(extent)
    r._r.append(drawing)
    return p


def test_evidence_caption_binding_sets_keepnext() -> None:
    """业绩扫描大图 + 紧跟的证据图注 → 给图段加 keepNext(图与图注同页)。"""
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_evidence_caption_binding

    doc = Document()
    p_img = _add_big_image_paragraph(doc, 500)
    doc.add_paragraph("2022年农村公路建设项目（萧县村道安全防护工程）三标段-合同（1）")

    fixed = heal_evidence_caption_binding(doc)
    assert fixed == 1
    ppr = p_img._p.find(_qn("w:pPr"))
    assert ppr is not None and ppr.find(_qn("w:keepNext")) is not None


def test_evidence_caption_binding_ignores_non_caption() -> None:
    """大图后面不是证据图注(普通正文)→ 不加 keepNext。"""
    from docx.oxml.ns import qn as _qn
    from services.docx_format_doctor import heal_evidence_caption_binding

    doc = Document()
    p_img = _add_big_image_paragraph(doc, 500)
    doc.add_paragraph("这是一段普通的正文说明文字，与证据材料无关。")

    fixed = heal_evidence_caption_binding(doc)
    assert fixed == 0
    ppr = p_img._p.find(_qn("w:pPr"))
    assert ppr is None or ppr.find(_qn("w:keepNext")) is None


def test_template_header_lines_removed_conservatively() -> None:
    """招标模板页眉行(整卷反复)删掉;正文引用/低频行/别的示范文本一律不动。"""
    from docx import Document

    from services.docx_format_doctor import heal_template_header_lines

    doc = Document()
    for _ in range(4):
        doc.add_paragraph("公路养护施工招标示范文本（2023 年版）")  # 页眉行(福昕带空格)
        doc.add_paragraph("这一页的正文内容保留。")
    doc.add_paragraph("本招标文件依据公路养护施工招标示范文本（2023年版）编制。")  # 正文引用
    doc.add_paragraph("公路工程施工招标示范文本（2018年版）")  # 只出现1次,不够3次
    doc.add_paragraph("投标保函示范文本（独立保函）")  # 无年版,不匹配

    healed = heal_template_header_lines(doc)
    texts = [p.text for p in doc.paragraphs]
    assert healed == 4
    assert not any(t.replace(" ", "") == "公路养护施工招标示范文本（2023年版）" for t in texts)
    assert sum("正文内容保留" in t for t in texts) == 4
    assert any("依据公路养护施工招标示范文本" in t for t in texts)  # 引用句保留
    assert any("2018年版" in t for t in texts)  # 低频不删
    assert any("独立保函" in t for t in texts)


def test_template_header_line_with_sectpr_only_cleared() -> None:
    """带分节符的页眉行:只清文字,段和分节符保留(防乱版)。"""
    from docx import Document
    from docx.oxml.ns import qn

    from services.docx_format_doctor import heal_template_header_lines

    doc = Document()
    paras = [doc.add_paragraph("公路养护施工招标示范文本（2023年版）") for _ in range(3)]
    # 给第2条塞进分节符
    pPr = paras[1]._p.get_or_add_pPr()
    pPr.append(paras[1]._p.makeelement(qn("w:sectPr"), {}))
    before = len(doc.paragraphs)

    healed = heal_template_header_lines(doc)
    assert healed == 3
    assert len(doc.paragraphs) == before - 2  # 删2条,分节符那条保留
    kept = [p for p in doc.paragraphs if p._p.find(qn("w:pPr")) is not None
            and p._p.find(qn("w:pPr")).find(qn("w:sectPr")) is not None]
    assert kept and kept[0].text.strip() == ""  # 文字清了,分节符还在


def test_section_titles_get_page_breaks_but_toc_untouched() -> None:
    """正文章节标题(一、/（一）)补段前分页;目录区同级连排标题绝不加(泗沙路实测结构)。"""
    from docx import Document
    from docx.oxml.ns import qn

    from services.docx_format_doctor import heal_section_title_page_breaks

    doc = Document()
    # 目录区:同级标题连排
    doc.add_paragraph("一、投标函及投标函附录")
    doc.add_paragraph("二、授权委托书或法定代表人身份证明")
    doc.add_paragraph("三、联合体协议书（不适用）")
    # 正文:大节+子节+表格分隔的子节
    doc.add_paragraph("正文过渡内容,说明目录结束了。")
    doc.add_paragraph("一、投标函及投标函附录")
    doc.add_paragraph("（一）投 标 函")
    doc.add_paragraph("致:某某招标人,我方愿意参加投标。")
    doc.add_paragraph("（一）投标人基本情况表")
    doc.add_table(rows=1, cols=2).cell(0, 0).text = "投标人名称"
    doc.add_paragraph("（二）投标人企业组织机构框图")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "以框图方式表示。"

    healed = heal_section_title_page_breaks(doc)

    def has_pb(p):
        pPr = p._p.find(qn("w:pPr"))
        return pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None

    paras = {p.text.strip(): p for p in doc.paragraphs if p.text.strip()}
    assert not has_pb(paras["二、授权委托书或法定代表人身份证明"])  # 目录行不动
    assert not has_pb(paras["三、联合体协议书（不适用）"])
    assert has_pb(paras["一、投标函及投标函附录"]) or True  # 同文重名取后者,单独查下面
    # 正文大节(前面是正文内容,后面是不同级"（一）") → 加
    body_titles = [p for p in doc.paragraphs if p.text.strip() == "一、投标函及投标函附录"]
    assert not has_pb(body_titles[0]) and has_pb(body_titles[1])
    # （一）投 标 函 紧跟在大节"一、投标函及投标函附录"后面(中间无正文)→ **不加分页**,
    # 否则大节标题独占一张几乎全空的页(2026-07-29 用户实测"五、项目管理机构"单独一页)。
    assert not has_pb(paras["（一）投 标 函"])
    # （一）投标人基本情况表 与（二）框图 之间隔着表格 → 不算目录连排,都加
    assert has_pb(paras["（一）投标人基本情况表"])
    assert has_pb(paras["（二）投标人企业组织机构框图"])
    assert healed >= 3


def test_two_level_toc_not_split() -> None:
    """两级目录(一、下面挂（一）（二）)是连排≥3的标题串,一整串都不许加分页(对抗审查修正)。"""
    from docx import Document
    from docx.oxml.ns import qn

    from services.docx_format_doctor import heal_section_title_page_breaks

    doc = Document()
    doc.add_paragraph("一、投标函及投标函附录")
    doc.add_paragraph("（一）投标函")
    doc.add_paragraph("（二）投标函附录")
    doc.add_paragraph("二、授权委托书")
    doc.add_paragraph("正文开始了,上面是目录。")
    doc.add_paragraph("一、投标函及投标函附录")
    doc.add_paragraph("致:招标人,以下是正文。")

    heal_section_title_page_breaks(doc)

    def has_pb(p):
        pPr = p._p.find(qn("w:pPr"))
        return pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None

    paras = list(doc.paragraphs)
    assert not any(has_pb(p) for p in paras[:4])  # 两级目录整串不动
    body_title = [p for p in paras if p.text.strip() == "一、投标函及投标函附录"][-1]
    assert has_pb(body_title)  # 正文标题照加


def test_template_header_cleared_from_real_header_part() -> None:
    """招标本身是 Word 时,模板页眉是真页眉部件(w:hdr),也必须清掉。

    2026-07-29 巢湖实测:Word 招标 48 个节的页眉都印着"公路养护施工招标示范文本
    （2023年版）",正文扫描够不着 → 投标文件每页顶上都带着招标书的页眉。
    """
    from docx import Document

    from services.docx_format_doctor import heal_template_header_lines

    doc = Document()
    hdr = doc.sections[0].header
    hdr.paragraphs[0].text = "公路养护施工招标示范文本（2023年版）"

    assert heal_template_header_lines(doc) >= 1
    assert doc.sections[0].header.paragraphs[0].text.strip() == ""


def test_template_header_healer_keeps_other_headers() -> None:
    """别的页眉(页码、公司自己的)一律不碰。"""
    from docx import Document

    from services.docx_format_doctor import heal_template_header_lines

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "安徽正奇建设有限公司  第 1 页"

    heal_template_header_lines(doc)
    assert "安徽正奇建设有限公司" in doc.sections[0].header.paragraphs[0].text


def test_idproof_split_columns_merged_back() -> None:
    """身份证明两栏拆段并回:性别归姓名行尾、职务归年龄行尾,孤段删除,一字不丢。"""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from services.docx_format_doctor import heal_idproof_column_pairs

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("姓 名： 许明英")
    br_r = p.add_run(); br_r._r.append(OxmlElement("w:br"))
    p.add_run("年 龄： 50")
    doc.add_paragraph("性别： 女")
    doc.add_paragraph("职务： 总经理")
    doc.add_paragraph("系　安徽正奇建设有限公司的法定代表人。")  # 正文,不许动

    assert heal_idproof_column_pairs(doc) == 2
    texts = [x.text for x in doc.paragraphs if x.text.strip()]
    assert len(texts) == 2  # 合并段 + 正文
    merged = texts[0]
    assert "许明英" in merged and "性　别：女" in merged
    assert "50" in merged and "职　务：总经理" in merged
    # 性别在换行前(姓名行),职务在换行后(年龄行)
    line1, line2 = merged.split("\n")
    assert "性　别" in line1 and "职　务" in line2
    assert "法定代表人" in texts[1]
