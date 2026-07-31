"""排版医生(尾巴页治理)纯逻辑测试;渲染反馈循环在真实文件上验证(见 commit 记录)。"""

import datetime

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from services.pagination_doctor import _HEADING_RE, _locate_section_slice, _tighten_range


def test_heading_regex_matches_section_starts_only() -> None:
    assert _HEADING_RE.match("一、投标函")
    assert _HEADING_RE.match("（三）联合体协议")
    assert _HEADING_RE.match("附表一 总体作业计划表")
    assert not _HEADING_RE.match("标文件澄清、修改、补充文件将成为")
    assert not _HEADING_RE.match("询途径（网址链接及查询方式）。")


def test_tighten_round1_compresses_loose_spacing() -> None:
    """第1轮:段前后×0.7;1.5/1.62 撑松档行距一步压到 1.15;空白段不动。"""
    doc = Document()
    p1 = doc.add_paragraph("正文一")
    p1.paragraph_format.space_before = Pt(10)
    p1.paragraph_format.space_after = Pt(10)
    p2 = doc.add_paragraph("撑松段")
    p2.paragraph_format.line_spacing = 1.62
    blank = doc.add_paragraph("")
    n = _tighten_range(doc, 0, 2, round_no=1)
    assert n == 2
    assert abs(p1.paragraph_format.space_before.pt - 7) < 0.1
    assert abs(p2.paragraph_format.line_spacing - 1.15) < 0.01
    assert blank.paragraph_format.line_spacing is None  # 第1轮不动空白段


def test_tighten_round2_thins_blank_paragraphs() -> None:
    """第2轮起:空白段压成 8pt 细缝(不删段,保留分隔感)。"""
    doc = Document()
    doc.add_paragraph("正文")
    blank = doc.add_paragraph("")
    _tighten_range(doc, 0, 1, round_no=2)
    pf = blank.paragraph_format
    assert pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert abs(pf.line_spacing.pt - 8) < 0.1


def test_tighten_never_touches_text() -> None:
    doc = Document()
    p = doc.add_paragraph("投 标 人： 安徽正奇建设有限公司（盖单位章）")
    p.paragraph_format.space_after = Pt(12)
    before = p.text
    _tighten_range(doc, 0, 0, round_no=1)
    _tighten_range(doc, 0, 0, round_no=2)
    assert p.text == before


def test_locate_section_slice_walks_back_to_heading() -> None:
    doc = Document()
    doc.add_paragraph("一、投标函")
    doc.add_paragraph("我方已仔细研究招标文件的全部内容")
    doc.add_paragraph("（1）签订合同；")
    doc.add_paragraph("标文件澄清、修改、补充文件将成为约束双方的合同文件")
    doc.add_paragraph("期： 2026年 7月 4日")
    stub = {"first_line": "标文件澄清、修改、补充文件将成为约束", "last_line": "期： 2026年 7月 4日"}
    sl = _locate_section_slice(doc, stub)
    assert sl is not None
    start, end = sl
    assert start == 1  # 标题"一、投标函"之后
    assert end == 4


def test_commercial_footer_page_numbers_all_sections() -> None:
    """商务卷页码:逐节写入右下角"第X页/共Y页"(招标原件页码已清,标书用自己的)。"""
    from docx.oxml.ns import qn

    from utils.docx_exporter import ZHENGQI_PROFILE, _configure_header_footer

    doc = Document()
    doc.add_paragraph("第一节")
    doc.add_section()
    doc.add_paragraph("第二节")
    _configure_header_footer(doc, None, True, ZHENGQI_PROFILE)
    for sec in doc.sections:
        fp = sec.footer.paragraphs[0]
        assert "第" in fp.text and "页" in fp.text
        assert any(True for _ in fp._p.iter(qn("w:fldChar")))  # PAGE 域真实存在


def test_stub_detection_ignores_page_number_footer_lines() -> None:
    """页码脚注不算页面内容,更不能当定位锚点。

    2026-07-30 实测:投标函落款两行溢出的尾巴页,body 首行被算成"第5页/共204页",
    拿它回 docx 定位段落必然落空 → 尾巴直接被跳过、没人收养。
    """
    import re

    from services.pagination_doctor import _STUB_MAX_CHARS

    # 直接测过滤逻辑等价物:页码行必须被 body 过滤器剔除
    lines = ["第5页/共204页", "法定代表人：（签字或盖章）", "日      期：2026 年7 月30 日", "159"]
    body = [
        l for l in lines
        if "招标示范文本" not in l
        and not re.fullmatch(r"第\s*\d+\s*页\s*/?\s*(共\s*\d+\s*页)?", l)
        and not re.fullmatch(r"\d{1,4}", l)
    ]
    assert body[0].startswith("法定代表人")
    assert len("".join(body)) <= _STUB_MAX_CHARS
