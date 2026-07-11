"""员工意见第10条:经理/总工个人"近年完成的类似项目信息表"后附业绩证明扫描。

覆盖:新锚点(停止词收边+跳目录取末匹配)定位、_place_anchored_images 真插图落位、
角色证明链 markdown 按勾选口径出图(None=名下全部/[]=无/列表=只勾中的)。
"""
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn


def _tiny_png() -> bytes:
    from PIL import Image

    buf = BytesIO()
    Image.new("RGB", (2, 2), (200, 30, 30)).save(buf, format="PNG")
    return buf.getvalue()


def _summary_table(doc):
    t = doc.add_table(rows=2, cols=3)
    t.cell(0, 0).text = "业绩序号"
    t.cell(0, 1).text = "项目名称（合同名称）"
    t.cell(0, 2).text = "备注"


def _detail_table(doc):
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "项目名称\n（合同名称）"
    t.cell(1, 0).text = "发包人名称"
    t.cell(2, 0).text = "项目经理"


def _build_three_role_doc() -> Document:
    """目录条目 + 投标人/项目经理(2张克隆信息表)/项目总工 三节 + 信誉节标题。"""
    doc = Document()
    doc.add_paragraph("项目经理近年完成的类似项目情况表\t12")  # 目录条目,不许命中
    doc.add_paragraph("（三）投标人近年完成的类似项目情况表（资格审查）")
    _summary_table(doc)
    doc.add_paragraph("投标人近年完成的类似项目信息表（资格审查）")
    _detail_table(doc)
    doc.add_paragraph("1.项目经理近年完成的类似项目情况表（资格审查）")
    _summary_table(doc)
    doc.add_paragraph("项目经理近年完成的类似项目信息表（资格审查）")
    _detail_table(doc)
    doc.add_paragraph("注：应附证明材料。")
    doc.add_paragraph("项目经理近年完成的类似项目信息表（资格审查）")
    _detail_table(doc)
    doc.add_paragraph("注：项目经理节最后一段。")
    doc.add_paragraph("2.项目总工近年完成的类似项目情况表（资格审查）")
    _summary_table(doc)
    doc.add_paragraph("项目总工近年完成的类似项目信息表（资格审查）")
    _detail_table(doc)
    doc.add_paragraph("注：总工节最后一段。")
    doc.add_paragraph("（四）投标人信誉情况表")
    return doc


def _body_index(doc, el) -> int:
    return list(doc.element.body.iterchildren()).index(el)


def _para_index(doc, text: str) -> int:
    from docx.text.paragraph import Paragraph

    for i, child in enumerate(doc.element.body.iterchildren()):
        if child.tag == qn("w:p") and Paragraph(child, doc).text.strip() == text:
            return i
    raise AssertionError(f"没找到段落: {text}")


def test_pm_anchor_lands_at_pm_section_end() -> None:
    """经理锚点=经理节最后一个元素(跳过目录/克隆标题,停在总工标题前)。"""
    from services.generation_service import _anchor_section_end_element

    doc = _build_three_role_doc()
    el = _anchor_section_end_element(doc, "项目经理类似项目表")
    assert el is not None
    assert _body_index(doc, el) == _para_index(doc, "注：项目经理节最后一段。")


def test_td_anchor_lands_at_td_section_end() -> None:
    """总工锚点=总工节最后一个元素(停在信誉标题前)。"""
    from services.generation_service import _anchor_section_end_element

    doc = _build_three_role_doc()
    el = _anchor_section_end_element(doc, "项目总工类似项目表")
    assert el is not None
    assert _body_index(doc, el) == _para_index(doc, "注：总工节最后一段。")


def test_company_anchor_lands_after_company_detail_tables() -> None:
    """公司业绩锚点=公司整节末尾(信息表之后、经理标题之前),不楔在情况表和信息表中间。"""
    from services.generation_service import _anchor_section_end_element

    doc = _build_three_role_doc()
    el = _anchor_section_end_element(doc, "类似项目情况表")
    assert el is not None
    idx = _body_index(doc, el)
    info_title = _para_index(doc, "投标人近年完成的类似项目信息表（资格审查）")
    pm_title = _para_index(doc, "1.项目经理近年完成的类似项目情况表（资格审查）")
    assert info_title < idx < pm_title  # 在公司信息表之后、经理节之前


def test_personnel_cert_groups_order_and_webcheck_guard(monkeypatch) -> None:
    """资历表附件序=身份证→毕业证→建造师→建造师网查→安全B→安全网查→社保→职称;
    网查截图不许被建造师/交安组当证书本体抢先插。"""
    from services import asset_resolver
    from services import knowledge_service
    from services import v2_generation_service as v2

    monkeypatch.setattr(
        asset_resolver, "pick_id_card_documents",
        lambda name: [{"document_id": 1, "side": "both"}],
    )
    refs = [
        {"document_id": 2, "owner_name": "李刚", "certificate_type": "毕业证", "file_type": "jpg", "image_insertable": True},
        {"document_id": 3, "owner_name": "李刚", "certificate_type": "一级建造师证", "file_type": "jpg", "image_insertable": True},
        {"document_id": 4, "owner_name": "李刚", "certificate_type": "建造师证网查截图", "file_type": "jpg", "image_insertable": True},
        {"document_id": 5, "owner_name": "李刚", "certificate_type": "交安证", "file_type": "jpg", "image_insertable": True},
        {"document_id": 6, "owner_name": "李刚", "certificate_type": "安全B证网查截图", "file_type": "jpg", "image_insertable": True},
        {"document_id": 7, "owner_name": "李刚", "certificate_type": "社保", "file_type": "jpg", "image_insertable": True},
        {"document_id": 8, "owner_name": "李刚", "certificate_type": "职称证书", "file_type": "jpg", "image_insertable": True},
    ]
    monkeypatch.setattr(
        knowledge_service, "list_knowledge_image_references", lambda *a, **k: refs
    )
    md = v2._one_person_cert_markdown("项目经理", "李刚", "项目经理资历表")
    import re as _re

    ids = [int(x) for x in _re.findall(r"document_id=(\d+)", md)]
    assert ids == [1, 2, 3, 4, 5, 6, 7, 8]  # 身份证最前,其后严格按用户定的附件序


def test_place_anchored_images_inserts_into_pm_section(monkeypatch) -> None:
    """带经理锚点的业绩图真插在经理节尾、总工标题之前;标记从 prose 里拿掉。"""
    from services import generation_service as g

    monkeypatch.setattr(g, "_resolve_knowledge_image", lambda _id: _tiny_png())
    doc = _build_three_role_doc()
    md = (
        "## 附录：项目经理（李刚）类似业绩证明材料（中标通知书·合同·交工验收，插于其类似项目信息表后）\n"
        "### 项目经理类似业绩 1：测试业绩1号工程\n"
        '{{knowledge_image:document_id=11 anchor="项目经理类似项目表" caption="测试业绩1号工程-中标通知书" width_cm=14}}\n'
        '{{knowledge_image:document_id=12 anchor="项目经理类似项目表" caption="测试业绩1号工程-合同" width_cm=14}}\n'
        "这行普通正文要保留。\n"
    )
    remain = g._place_anchored_images(doc, md)
    assert "knowledge_image" not in remain
    assert "这行普通正文要保留。" in remain
    assert "附录：项目经理" not in remain  # 图全被移走,空标题一并丢弃

    body = list(doc.element.body.iterchildren())
    drawings = [
        i for i, ch in enumerate(body)
        if ch.tag == qn("w:p") and ch.findall(".//" + qn("w:drawing"))
    ]
    assert len(drawings) == 2
    pm_end = _para_index(doc, "注：项目经理节最后一段。")
    td_title = _para_index(doc, "2.项目总工近年完成的类似项目情况表（资格审查）")
    for i in drawings:
        assert pm_end < i < td_title


def _wire_role_md(monkeypatch, chosen, records):
    """给 _role_performance_evidence_markdown 打桩:选派李刚 + 勾选/名下记录 + 证明行。"""
    from services import project_service
    from services import similar_project_info_service as spi
    from services import v2_generation_service as v2

    monkeypatch.setattr(
        project_service,
        "get_selected_personnel",
        lambda pid: {"selected": {"project_manager": {"name": "李刚"}}},
    )
    monkeypatch.setattr(
        project_service,
        "get_selected_role_performance",
        lambda pid, role: {"selected": chosen},
    )
    monkeypatch.setattr(spi, "records_for_manager", lambda name: records)
    monkeypatch.setattr(
        v2,
        "_query_performance_evidence_rows",
        lambda: [
            (11, "测试业绩1号工程", "中标通知书", "2023", 1),
            (12, "测试业绩1号工程", "交工验收", "2023", 1),
            (21, "测试业绩2号工程", "中标通知书", "2022", 1),
        ],
    )
    return v2


def test_role_evidence_follows_checked_subset(monkeypatch) -> None:
    """勾了1条 → 只出这条的证明图,锚点/人名都对。"""
    v2 = _wire_role_md(monkeypatch, [{"name": "测试业绩1号工程"}], [])
    md = v2._role_performance_evidence_markdown(7, "pm")
    assert 'anchor="项目经理类似项目表"' in md
    assert "项目经理（李刚）类似业绩证明材料" in md
    assert "document_id=11" in md and "document_id=12" in md
    assert "document_id=21" not in md  # 没勾的不出


def test_role_evidence_none_means_no_images(monkeypatch) -> None:
    """没勾(None) → 跟表一致:留白不附图(全部人工手选,用户2026-07-11拍板)。"""
    records = [
        {"project_name": "测试业绩1号工程"},
        {"project_name": "测试业绩2号工程"},
    ]
    v2 = _wire_role_md(monkeypatch, None, records)
    assert v2._role_performance_evidence_markdown(7, "pm") == ""


def test_role_evidence_cleared_or_no_person_is_empty(monkeypatch) -> None:
    """人工清空([]) 或未选派人 → 不出图。"""
    v2 = _wire_role_md(monkeypatch, [], [])
    assert v2._role_performance_evidence_markdown(7, "pm") == ""

    from services import project_service

    monkeypatch.setattr(
        project_service, "get_selected_personnel", lambda pid: {"selected": {}}
    )
    assert v2._role_performance_evidence_markdown(7, "pm") == ""
