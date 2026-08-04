"""联合体/保证金留白 + 投标日期=今天 + 资料库图表注入 测试。"""

from __future__ import annotations

import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt

from services import original_docx_format_service as o

import pytest


@pytest.fixture(autouse=True)
def _no_curated_template(monkeypatch):
    """本文件的简历表测试全部针对"字段填空白表"路径:屏蔽成品模版整表照搬
    (开发机连着真库,李刚等16人有真模版,不屏蔽会把宿主表整个替换掉)。"""
    import services.curated_resume_service as crs

    monkeypatch.setattr(crs, "get_template_table_el", lambda name: None)



def test_bond_left_blank_but_consortium_lead_name_fills() -> None:
    # 保证金 留白
    assert o._label_to_profile_key("投标保证金金额") == ""
    # 但"独立投标人或联合体牵头人名称"对单独投标人就该填公司名(不能误伤)
    assert o._label_to_profile_key("独立投标人或联合体牵头人名称") == "company_name"
    assert o._label_to_profile_key("投标人名称") == "company_name"


def test_inline_value_skips_bond() -> None:
    assert o._inline_value_for("投标保证金金额", {"company_name": "正奇"}) == ""


def test_consortium_agreement_section_blank() -> None:
    """联合体协议书"整章"留白:章节内的投标人名/日期都不自动填。"""
    doc = Document()
    doc.add_paragraph("联合体协议书")
    doc.add_paragraph("投标人：____")  # 在留白章节内 → 不填
    doc.add_paragraph("日期：    年   月   日")
    doc.add_paragraph("（三）下一章节")  # 编号章节 → 留白区结束
    doc.add_paragraph("投标人：____")  # 章节外 → 该填
    profile = {"company_name": "安徽正奇建设有限公司"}
    o._fill_inline_labeled_blanks(doc, profile)
    o._fill_bid_date_today(doc)
    assert "安徽正奇" not in doc.paragraphs[1].text  # 协议书内投标人 留白
    assert "2026" not in doc.paragraphs[2].text  # 协议书内日期 留白
    assert "安徽正奇" in doc.paragraphs[4].text  # 章节外投标人 正常填


def test_bid_date_today_split_runs() -> None:
    doc = Document()
    p = doc.add_paragraph()
    for t in ["日期：", "  ", "年", "  ", "月", "  ", "日"]:
        p.add_run(t)
    n = o._fill_bid_date_today(doc)
    d = datetime.date.today()
    assert n >= 1
    assert f"{d.year}年{d.month}月{d.day}日" in doc.paragraphs[0].text


def test_bid_date_today_single_run() -> None:
    doc = Document()
    doc.add_paragraph("投标日期：____年__月__日")
    o._fill_bid_date_today(doc)
    d = datetime.date.today()
    assert f"{d.year}年{d.month}月{d.day}日" in doc.paragraphs[0].text


def test_establish_date_not_touched_by_bid_date() -> None:
    doc = Document()
    doc.add_paragraph("成立日期：____年__月__日")
    assert o._fill_bid_date_today(doc) == 0


def test_inject_org_tables_does_not_duplicate_org_chart(monkeypatch) -> None:
    """回归:组织机构图不再按锚点注入——公司组件库已把成品整表替换进招标那张空框里。

    2026-07-29 用户实测:两处各插一次,商务卷里出现两张一模一样的项目管理机构图
    ("同一个表你每次生成都要放两遍进去")。与拟分包表同一个病因。
    """
    src = Document()
    tbl = src.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "项目经理"
    buf = BytesIO()
    src.save(buf)
    src_bytes = buf.getvalue()

    from services import generation_service as g

    monkeypatch.setattr(
        g, "_fetch_org_source_docx",
        lambda kw: src_bytes if kw == "项目管理机构" else None,
    )

    target = Document()
    target.add_paragraph("（一）项目管理机构组织机构图")
    before = len(target.tables)
    assert g._inject_org_tables(target) == 0
    assert len(target.tables) == before


def test_inject_org_tables_does_not_inject_fenbao(monkeypatch) -> None:
    """回归:拟分包表不再注入(福昕格式章已带招标那张,避免重复变丑)。

    即便资料库有拟分包 docx、目标里也有"拟分包"章节,也不应再注入第二张拟分包表。
    """
    src = Document()
    src.add_table(rows=21, cols=4).cell(0, 0).text = "拟分包的工程项目"
    buf = BytesIO()
    src.save(buf)
    fenbao_bytes = buf.getvalue()

    from services import generation_service as g

    # 只有"拟分包"源可取(组织机构图源返回 None);若拟分包仍是注入源,这里会注入1张
    monkeypatch.setattr(
        g, "_fetch_org_source_docx",
        lambda kw: fenbao_bytes if "拟分包" in kw else None,
    )
    target = Document()
    target.add_paragraph("六、拟分包项目情况表")
    before = len(target.tables)
    inserted = g._inject_org_tables(target)
    assert inserted == 0  # 拟分包不再是注入源
    assert len(target.tables) == before


def test_strip_tender_page_numbers() -> None:
    from services.original_docx_format_service import _strip_tender_page_numbers
    doc = Document()
    doc.add_paragraph("正文内容")
    doc.add_paragraph("171")  # 招标原件页码
    doc.add_paragraph("第 5 页")
    n = _strip_tender_page_numbers(doc)
    texts = [p.text for p in doc.paragraphs]
    assert "171" not in texts and "第 5 页" not in texts
    assert "正文内容" in texts


def test_pm_resume_table_fill_scoped() -> None:
    from services.original_docx_format_service import _fill_pm_resume_table
    doc = Document()
    t = doc.add_table(rows=2, cols=4)
    t.cell(0, 0).text = "姓名"
    t.cell(0, 2).text = "职称"
    t.cell(1, 0).text = "拟在本标段工程担任职务"
    # 另一张表(无"拟在本标段")不应被填
    other = doc.add_table(rows=1, cols=2)
    other.cell(0, 0).text = "姓名"
    resume = {"姓名": "李刚", "职称": "工程师", "拟任职务": "项目经理"}
    _fill_pm_resume_table(doc, resume)
    assert t.cell(0, 1).text == "李刚"
    assert t.cell(1, 1).text == "项目经理"
    assert other.cell(0, 1).text.strip() == ""  # 非简历表不动


def test_org_table_anchor_skips_toc() -> None:
    from services import generation_service as g
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("五、项目管理机构")  # 目录条目
    doc.add_paragraph("投标函正文……（实质内容，退出目录）" + "x" * 30)
    real = doc.add_paragraph("五、项目管理机构")  # 正文真章节
    el = g._find_section_paragraph(doc, ("项目管理机构",))
    assert el is real._p  # 取正文章节,不取目录条目


def test_bond_mention_in_paragraph_does_not_blank_quality_duration() -> None:
    """投标函正文顺带提到'投标保证金'时,同段的质量/工期空仍要照填(修留白过宽回归)。"""
    from docx import Document as _D
    from services.original_docx_format_service import _fill_inline_labeled_blanks
    doc = _D()
    p = doc.add_paragraph()
    for t in ["3.质量标准：", " ", "\t", "；工期：", " ", "\t", "日历天。我方按时提交投标保证金。"]:
        p.add_run(t)
    _fill_inline_labeled_blanks(doc, {"质量": "合格", "工期": "90日历天"})
    txt = doc.paragraphs[0].text
    assert "合格" in txt and "90日历天" in txt


def test_authorization_letter_fills_legal_rep_only() -> None:
    """授权委托书:本人___（姓名）系→法人名;代理人（姓名）留白;不误填随处的（姓名）。"""
    from docx import Document as _D
    from services.original_docx_format_service import _fill_authorization_letter
    doc = _D()
    p = doc.add_paragraph()
    for t in ["本人 ", "\t", "（姓名）系 ", "\t", "（投标人名称）的法定代表人，现委托 ", "\t", "（姓名）为我方代理人。"]:
        p.add_run(t)
    # 另一段普通"（姓名）"不该被填(无授权委托书语境)
    other = doc.add_paragraph("项目经理（姓名）：")
    _fill_authorization_letter(doc, {"legal_representative": "许明英"})
    t0 = doc.paragraphs[0].text
    assert "本人许明英（姓名）系" in t0  # 法人名填在第一个（姓名）前
    assert t0.count("许明英") == 1  # 代理人那个（姓名）没被填
    assert "许明英" not in other.text  # 无语境的（姓名）不动


def test_personnel_table_fills_pm_and_tech_rows() -> None:
    from docx import Document as _D
    from services.original_docx_format_service import _fill_personnel_table
    doc = _D()
    t = doc.add_table(rows=4, cols=5)
    for c, h in enumerate(["职务", "姓名", "职称", "证书名称", "证号"]):
        t.cell(0, c).text = h
        t.cell(1, c).text = h
    prof = {
        "project_manager_name": "江甜甜", "project_manager_cert": "皖234", "project_manager_title": "工程师",
        "tech_director_name": "李刚", "tech_director_cert": "皖134", "tech_director_title": "工程师",
    }
    _fill_personnel_table(doc, prof)
    assert [t.cell(2, c).text for c in (0, 1, 4)] == ["项目经理", "江甜甜", "皖234"]
    assert [t.cell(3, c).text for c in (0, 1, 4)] == ["项目技术负责人", "李刚", "皖134"]


def test_evidence_groups_drop_system_patent() -> None:
    from services.v2_generation_service import _EVIDENCE_GROUPS
    titles = [g[0] for g in _EVIDENCE_GROUPS]
    assert "管理体系认证证书" not in titles  # 体系删了
    assert "专利与工法证书" not in titles  # 专利工法删了
    # 基本情况表后只剩这4类
    basic = [g[0] for g in _EVIDENCE_GROUPS if g[3] == "基本情况表"]
    assert set(basic) == {"营业执照", "企业资质证书", "安全生产许可证", "基本账户开户许可证"}


def test_performance_table_fills_selected(monkeypatch) -> None:
    """投标人业绩情况表(业绩序号|项目名称|备注)按选中业绩填项目名,去重,保真。"""
    from docx import Document as _D
    d = _D(); t = d.add_table(rows=4, cols=3)
    for c, v in enumerate(["业绩序号", "项目名称（合同名称）", "备注"]):
        t.cell(0, c).text = v
    t.cell(1, 0).text = "1"; t.cell(2, 0).text = "2"; t.cell(3, 0).text = "……"
    prof = {"selected_performance": [
        {"name": "项目A"}, {"name": "项目A"}, {"name": "项目B"},
    ]}
    assert o._fill_performance_table(d, prof) == 2  # 去重后2个
    assert t.cell(1, 1).text == "项目A"
    assert t.cell(2, 1).text == "项目B"
    # 保真:首数据行已填则整表跳过
    d2 = _D(); t2 = d2.add_table(rows=2, cols=3)
    for c, v in enumerate(["业绩序号", "项目名称（合同名称）", "备注"]):
        t2.cell(0, c).text = v
    t2.cell(1, 1).text = "已填"
    assert o._fill_performance_table(d2, prof) == 0
    assert t2.cell(1, 1).text == "已填"


def test_performance_table_empty_when_no_selection() -> None:
    from docx import Document as _D
    d = _D(); t = d.add_table(rows=2, cols=3)
    for c, v in enumerate(["业绩序号", "项目名称", "备注"]):
        t.cell(0, c).text = v
    assert o._fill_performance_table(d, {}) == 0


def test_bid_date_fills_split_ri_qi_variant() -> None:
    """福昕把"日期"劈成'日'段+'期：…'段的落款(实测v7有10处)→ 也要填今天。"""
    import datetime
    doc = Document()
    doc.add_paragraph("日")
    p = doc.add_paragraph()
    for t in ["期：", " ", "\t", "年", " ", "\t", "月", " ", "\t", "日"]:
        p.add_run(t)
    n = o._fill_bid_date_today(doc, today=datetime.date(2026, 7, 4))
    assert n >= 3
    assert "2026" in p.text and "7月" in p.text.replace(" ", "") and "4日" in p.text.replace(" ", "")


def test_bid_date_label_ri_not_mistaken_for_unit() -> None:
    """签名和日期同段时,"日 期"标签的"日"不是日期单位——别把数字塞到标签前(实测p#182回归)。"""
    import datetime
    doc = Document()
    p = doc.add_paragraph()
    for t in ["投 标 人： 公司（盖单位章）", " ", "日", " 期：", " ", "年", " ", "月", " ", "日"]:
        p.add_run(t)
    o._fill_bid_date_today(doc, today=datetime.date(2026, 7, 4))
    assert "）2026" not in p.text and "）4" not in p.text  # 标签前没被塞数字
    compact = p.text.replace(" ", "")
    assert "2026年" in compact and "7月" in compact and "4日" in compact


def test_bid_date_leaves_unknown_data_dates() -> None:
    """开立时间等数据型日期(非落款)不是"今天",保持留白。"""
    import datetime
    doc = Document()
    doc.add_paragraph("开立时间： \t年 \t月 \t日")
    n = o._fill_bid_date_today(doc, today=datetime.date(2026, 7, 4))
    assert n == 0
    assert "2026" not in doc.paragraphs[0].text


def test_textbox_placeholder_replaced() -> None:
    """浮动文本框里的（招标人名称）占位符也要被替换(正文替换够不着,实测p62)。"""
    from docx.oxml import parse_xml
    doc = Document()
    xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml"><w:r><w:pict><v:shape><v:textbox>'
        "<w:txbxContent><w:p><w:r><w:t>（招标人名称）</w:t></w:r></w:p></w:txbxContent>"
        "</v:textbox></v:shape></w:pict></w:r></w:p>"
    )
    doc.element.body.append(parse_xml(xml))
    n = o._fill_textbox_placeholders(doc, {"招标人": "巢湖市栏杆集镇人民政府"})
    assert n == 1
    from docx.oxml.ns import qn
    texts = [t.text for t in doc.element.body.iter(qn("w:t"))]
    assert any("巢湖市栏杆集镇人民政府" in (t or "") for t in texts)
    assert not any("（招标人名称）" in (t or "") for t in texts)

def _mk_resume_table(doc, heading: str):
    """建一张带标题的简历表(锚='拟在本标段'),返回表对象。"""
    doc.add_paragraph(heading)
    t = doc.add_table(rows=2, cols=4)
    t.cell(0, 0).text = "姓名"
    t.cell(0, 2).text = "职称"
    t.cell(1, 0).text = "拟在本标段工程担任职务"
    return t


def test_resume_tables_third_person_table_left_blank() -> None:
    """员工反馈第11条:认不出角色的第三张简历表必须留空,绝不默认灌项目经理信息。"""
    from services.original_docx_format_service import _fill_resume_tables
    doc = Document()
    t_pm = _mk_resume_table(doc, "项目经理简历表")
    t_tech = _mk_resume_table(doc, "项目技术负责人简历表")
    t_other = _mk_resume_table(doc, "安全员简历表")  # 系统不认识的角色
    pm = {"姓名": "江舟", "职称": "高级工程师", "拟任职务": "项目经理"}
    tech = {"姓名": "王俊明", "职称": "高级工程师", "拟任职务": "项目技术负责人"}
    _fill_resume_tables(doc, pm, tech)
    assert t_pm.cell(0, 1).text == "江舟"
    assert t_tech.cell(0, 1).text == "王俊明"
    assert t_other.cell(0, 1).text.strip() == ""   # 第三张表留空给人工
    assert t_other.cell(1, 1).text.strip() == ""


def test_resume_tables_unlabeled_pair_still_inferred() -> None:
    """原有推断保留:项目经理表带字样+末张无名表归总工;全无字样时首张归PM末张归总工。"""
    from services.original_docx_format_service import _fill_resume_tables
    pm = {"姓名": "江舟", "拟任职务": "项目经理"}
    tech = {"姓名": "王俊明", "拟任职务": "项目技术负责人"}

    # 场景A:PM表带字样,第二张无字样 → 归总工(老兜底)
    doc = Document()
    t1 = _mk_resume_table(doc, "项目经理简历表")
    t2 = _mk_resume_table(doc, "主要人员简历表")
    _fill_resume_tables(doc, pm, tech)
    assert t1.cell(0, 1).text == "江舟" and t2.cell(0, 1).text == "王俊明"

    # 场景B:三张全无字样 → 首张PM,末张总工,中间留空
    doc = Document()
    ta = _mk_resume_table(doc, "简历表一")
    tb = _mk_resume_table(doc, "简历表二")
    tc = _mk_resume_table(doc, "简历表三")
    _fill_resume_tables(doc, pm, tech)
    assert ta.cell(0, 1).text == "江舟"
    assert tb.cell(0, 1).text.strip() == ""       # 中间那张不再被灌PM
    assert tc.cell(0, 1).text == "王俊明"

    # 场景C:只选了项目经理,单张无字样表 → 照旧填PM(不回归)
    doc = Document()
    td = _mk_resume_table(doc, "简历表")
    _fill_resume_tables(doc, pm, None)
    assert td.cell(0, 1).text == "江舟"


def test_subcontract_table_filled_with_none_keeps_all_rows() -> None:
    """无分包计划 → 首行填"无"、合计填"无",**行数一行不动**(保持招标原表形态)。"""
    doc = Document()
    t = doc.add_table(rows=21, cols=4)
    for i, h in enumerate(("拟分包的工程项目", "主要工程内容", "预计造价（万元）", "备 注")):
        t.cell(0, i).text = h
    for r in range(1, 21):
        t.cell(r, 3).text = "注：若无分包计划，则投标人应在本表填写“无”"
    t.cell(20, 0).text = "拟分包工程造价合计（万元）"

    assert o._fill_subcontract_table(doc) == 1
    tb = doc.tables[0]
    # 行数必须与招标原表一致——绝不删行(2026-07-29 用户明确:"他招标文件多少行你就
    # 搞多少行")。原先占两页是前面的空段把表推下去了,与行数无关。
    assert len(tb.rows) == 21
    assert tb.cell(1, 0).text.strip() == "无"
    assert tb.cell(1, 2).text.strip() == "无"
    assert "合计" in tb.rows[-1].cells[0].text
    assert tb.rows[-1].cells[2].text.strip() == "无"
    # 备注列的招标原注不许被改
    assert "若无分包计划" in tb.cell(1, 3).text


def test_subcontract_table_with_real_plan_is_untouched() -> None:
    """真有分包内容时一个字都不动(别把人家填好的计划抹了)。"""
    doc = Document()
    t = doc.add_table(rows=4, cols=4)
    for i, h in enumerate(("拟分包的工程项目", "主要工程内容", "预计造价（万元）", "备 注")):
        t.cell(0, i).text = h
    t.cell(1, 0).text = "交通安全设施"
    t.cell(3, 0).text = "拟分包工程造价合计（万元）"

    assert o._fill_subcontract_table(doc) == 0
    assert len(doc.tables[0].rows) == 4
    assert doc.tables[0].cell(1, 0).text == "交通安全设施"


def test_inline_blank_before_unit_word_gets_filled() -> None:
    """"工期：＿＿日历天"这种"槽后跟单位词"的空必须填上。

    2026-07-29 用户实测:同一句里质量、安全(槽后是"；")填上了,唯独工期因为后面跟
    "日历天"被误判成"已有真实值"而跳过,交出去空着。
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("3. 工程质量：")
    p.add_run("      ").font.underline = True
    p.add_run("；安全目标：")
    p.add_run("      ").font.underline = True
    p.add_run("；工期：")
    p.add_run("      ").font.underline = True
    p.add_run("日历天。")

    o._fill_inline_labeled_blanks(doc, {"质量": "合格", "安全": "无安全事故", "工期": "90日历天"})
    assert doc.paragraphs[0].text == "3. 工程质量：合格；安全目标：无安全事故；工期：90日历天。"


def test_bid_date_written_into_underlined_slot_not_the_spacer() -> None:
    """日期值要写进**带下划线**的空槽,不能写进槽与"年"之间那个普通空格 run。

    2026-07-29 用户截图:值写进夹缝空格 → 带线空槽原样留着、整行变长,右对齐落款行
    的字挤成一坨。
    """
    import datetime

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("日      期：")
    slots = []
    for unit in ("年", "月", "日"):
        slot = p.add_run("      ")
        slot.font.underline = True
        slots.append(slot)
        p.add_run(" ")  # 槽与单位之间的普通空格(招标原件就长这样)
        p.add_run(unit)

    o._fill_bid_date_today(doc)
    today = datetime.date.today()
    assert [s.text for s in slots] == [str(today.year), str(today.month), str(today.day)]
    # 夹缝空格必须原样留着(值没被写进它)
    assert [r.text for r in p.runs if not r.font.underline and r.text == " "]


def test_appending_prose_does_not_inflate_tender_table_rows() -> None:
    """追加商务正文不许把招标格式章的行距撑高。

    2026-07-29 实测:追加正文时会把 Normal 样式改成正奇正文(固定行距32pt),
    招标表格里没写死行距的空格子跟着从 23pt 涨到 35pt,招标本来一页放得下的
    拟分包表(21行)被挤成两页。改样式前必须先把已有段落的行距固化到段落级。
    """
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING

    from services.generation_service import _freeze_existing_paragraph_spacing

    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    cell_para = t.cell(1, 0).paragraphs[0]
    body_para = doc.add_paragraph("招标原文一句话")
    # 有人显式写死了行距的段落:必须原样保留
    fixed = doc.add_paragraph("已写死行距的段落")
    fixed.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    assert _freeze_existing_paragraph_spacing(doc) >= 2

    # 冻结后再把 Normal 改成固定 32pt —— 已有段落不该跟着变
    normal = doc.styles["Normal"].paragraph_format
    normal.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.line_spacing = Pt(32)

    for para in (cell_para, body_para):
        assert para.paragraph_format.line_spacing_rule is not None
        assert para.paragraph_format.line_spacing != Pt(32)
    assert fixed.paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE


def test_establish_date_written_into_underlined_slot_not_the_spacer() -> None:
    """成立时间的值也要写进带下划线的空槽——与落款日期同坑同方子。

    2026-07-30 用户实测身份证明页:值被写进槽与"年"之间的普通空格,带点线的
    空槽原样留着,Word 里点线和数字叠在一起。
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("成立时间：")
    slots = []
    for unit in ("年", "月", "日"):
        s = p.add_run("            ")
        s.font.underline = True
        slots.append(s)
        p.add_run(" ")  # 夹缝空格
        p.add_run(unit)

    n = o._fill_establish_segmented(doc, {"establish_date": "2011-07-05"})
    assert n == 3
    assert [s.text for s in slots] == ["2011", "7", "5"]


def test_inline_fill_keeps_gap_before_next_label_on_same_line() -> None:
    """"邮政编码：__电话：__"同一行两个标签:值比槽短要补空格保住间距。

    2026-07-30 身份证明页实测:"许明英"把"姓名"的槽整个吃掉,和"性 别"挤成一团。
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("邮政编码：")
    slot = p.add_run("            ")
    slot.font.underline = True
    p.add_run("联系电话：")
    slot2 = p.add_run("             ")
    slot2.font.underline = True

    o._fill_inline_labeled_blanks(
        doc, {"postal_code": "230041", "contact_phone": "13905510000"}
    )
    text = doc.paragraphs[0].text
    assert "230041" in text and "13905510000" in text
    import re as _re
    assert _re.search(r"230041\s+联", text), f"值和下一标签挤在一起了: {text!r}"


def test_token_replacement_shrinks_leading_slot_to_keep_line_width() -> None:
    """（投标人名称）→公司全名变长时,前面的空白槽等量瘦身,行宽不变不折行。

    值写槽家族第三处(2026-07-30):"系＿＿（投标人名称）的法定代表人"替换后
    整行变长,Word 里"法定代表人。"被挤到下一行。
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("系")
    slot = p.add_run("                    ")  # 20格空白槽
    slot.font.underline = True
    p.add_run("（投标人名称）的法定代表人。")

    o._replace_known_fields(doc, {"company_name": "安徽正奇建设有限公司"})
    text = doc.paragraphs[0].text
    assert "安徽正奇建设有限公司的法定代表人。" in text
    # 公司名(10字)比token(7字)长3字 → 槽从20缩到17,总行宽与模板一致
    grow = len("安徽正奇建设有限公司") - len("（投标人名称）")
    assert len(slot.text) == 20 - grow


def test_insert_image_survives_zero_dpi_jpeg(monkeypatch) -> None:
    """JFIF头 横向DPI=0 的图必须能插成功,不许静默变占位符。

    2026-07-31 实测(doc 7719 逐字节解剖):病图 JFIF density=(0,2)、unit=1(dpi),
    python-docx horz_dpi=0 → 算宽度除零,旧代码静默吞异常放占位。
    """
    import io
    from PIL import Image as PILImage
    from services import generation_service as gs

    buf = io.BytesIO()
    PILImage.new("RGB", (400, 300), (120, 120, 120)).save(buf, "JPEG", dpi=(96, 96))
    raw = bytearray(buf.getvalue())
    jfif = raw.find(b"JFIF\x00")
    assert jfif > 0
    dens = jfif + 7  # units(1) Xdensity(2) Ydensity(2)
    raw[dens:dens + 5] = b"\x01\x00\x00\x00\x02"  # unit=dpi, X=0, Y=2 —— 与真病图同款
    zero_dpi = bytes(raw)

    from docx import Document
    from docx.shared import Cm
    d0 = Document()
    try:
        d0.add_paragraph().add_run().add_picture(io.BytesIO(zero_dpi), width=Cm(14))
        raise AssertionError("预期除零未发生,夹具失效")
    except ZeroDivisionError:
        pass

    monkeypatch.setattr(gs, "_resolve_knowledge_image", lambda _id: zero_dpi)
    doc = Document()
    anchor = doc.add_paragraph("锚点")._p
    gs._insert_image_after(anchor, doc, 1, "G343-交工验收（2）", 14.0)
    texts = [p.text for p in doc.paragraphs]
    assert not any("未能插入" in t for t in texts), "仍然变成了占位符"
    from docx.oxml.ns import qn
    assert doc.element.body.findall('.//' + qn('a:blip')), "没有图"
