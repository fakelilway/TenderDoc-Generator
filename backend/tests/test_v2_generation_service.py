from __future__ import annotations

import pytest

from agents.content_writer_agent import NodeFillResult
from services import v2_generation_service
from services.v2_audit_service import AuditResult, audit_format_layer
from schemas.tender import FormatOutlineNode, TenderRequirements


def test_audit_built_format_docx_flags_empty_and_passes_content(tmp_path) -> None:
    from docx import Document

    # 空文档(只有一个空段落)→ 判失败
    empty = tmp_path / "empty.docx"
    Document().save(str(empty))
    assert v2_generation_service._audit_built_format_docx(str(empty))

    # 有文字 → 通过
    with_text = tmp_path / "text.docx"
    d1 = Document()
    d1.add_paragraph("投标人：安徽正奇建设有限公司")
    d1.save(str(with_text))
    assert v2_generation_service._audit_built_format_docx(str(with_text)) == []

    # 只有表格(无正文)也算有内容 → 通过
    with_table = tmp_path / "table.docx"
    d2 = Document()
    d2.add_table(rows=2, cols=3)
    d2.save(str(with_table))
    assert v2_generation_service._audit_built_format_docx(str(with_table)) == []

    # 打不开的文件 → 判失败
    assert v2_generation_service._audit_built_format_docx(str(tmp_path / "nope.docx"))


def test_sections_from_confirmed_outline_maps_titles_and_focus() -> None:
    confirmed = [
        {"title": "总体施工组织布置及规划", "focus_points": ["施工部署", "总体目标"]},
        {"title": "主要工程施工方案与技术措施", "focus_points": []},
        {"title": "工期保证体系及措施", "focus_points": ["进度计划"], "target_chars": 2000},
        {"title": "质量管理体系及保证措施", "focus_points": ["三检制"]},
    ]
    sections = v2_generation_service._sections_from_confirmed_outline(confirmed)
    assert sections is not None
    assert [s["title"] for s in sections] == [
        "总体施工组织布置及规划",
        "主要工程施工方案与技术措施",
        "工期保证体系及措施",
        "质量管理体系及保证措施",
    ]
    # focus_points become the per-section must-cover guidance
    assert "施工部署" in sections[0]["must_cover"]
    assert "三检制" in sections[3]["must_cover"]
    # explicit target honored; default applied otherwise
    assert sections[2]["target_chars"] == 2000
    assert sections[0]["target_chars"] == 2200  # _CONFIRMED_OUTLINE_TARGET_CHARS


def test_sections_from_confirmed_outline_honors_small_and_only_empty_falls_back() -> None:
    # Empty/None → None so caller uses legacy logic.
    assert v2_generation_service._sections_from_confirmed_outline(None) is None
    assert v2_generation_service._sections_from_confirmed_outline([]) is None
    # A deliberately simple confirmed outline (even a single section) is HONORED
    # now — 招标各不相同，简单目录也合法。
    small = v2_generation_service._sections_from_confirmed_outline(
        [{"title": "施工组织设计", "focus_points": ["总体部署"]}]
    )
    assert small is not None
    assert [s["title"] for s in small] == ["施工组织设计"]
    assert "总体部署" in small[0]["must_cover"]


def test_v2_format_audit_rejects_flattened_form_tables() -> None:
    report = audit_format_layer(
        pages=[
            (
                "（一）投标人基本情况表",
                "投标人名称\n注册地址 邮政编码\n联系方式 联系人 电话\n________",
            )
        ],
        filled_pages=[
            (
                "（一）投标人基本情况表",
                "投标人名称\n注册地址 邮政编码\n联系方式 联系人 电话\n安徽正奇建设有限公司",
            )
        ],
    )

    assert not report.passed
    assert any("表格格式被拍扁" in issue.problem for issue in report.issues)


def test_v2_format_audit_accepts_markdown_table_layout() -> None:
    report = audit_format_layer(
        pages=[
            (
                "（一）投标人基本情况表",
                "投标人名称\n注册地址 邮政编码\n联系方式 联系人 电话\n________",
            )
        ],
        filled_pages=[
            (
                "（一）投标人基本情况表",
                "| 投标人名称 | 安徽正奇建设有限公司 |\n"
                "| --- | --- |\n"
                "| 注册地址 | ________ |\n"
                "| 联系方式 | ________ |",
            )
        ],
    )

    assert report.passed


def test_v2_format_audit_rejects_reconstructed_table_marker() -> None:
    report = audit_format_layer(
        pages=[
            (
                "（一）投标人基本情况表",
                "投标人名称\n注册地址 邮政编码\n联系方式 联系人 电话\n________",
            )
        ],
        filled_pages=[
            (
                "（一）投标人基本情况表",
                '{{tdg_table:bidder_basic_info company_name="安徽正奇建设有限公司"}}',
            )
        ],
    )

    assert not report.passed
    assert any("不是招标文件原样复制" in issue.problem for issue in report.issues)


def test_v2_original_format_fails_when_content_writer_fails(monkeypatch) -> None:
    requirements = TenderRequirements(project_name="测试项目")

    def fail_writer(**_kwargs):
        raise RuntimeError("llm down")

    monkeypatch.setattr(v2_generation_service, "fill_technical_volume", fail_writer)

    with pytest.raises(ValueError, match="施工方案正文生成失败"):
        v2_generation_service.generate_v2_bid_package(
            requirements,
            {},
            company_name="安徽正奇建设有限公司",
            tender_text="第八章 投标文件格式",
            original_format_docx_available=True,
        )


def test_v2_does_not_turn_bid_letter_with_bill_text_into_table() -> None:
    rendered = v2_generation_service._render_locked_format_content(
        "一、投标函",
        "致：（招标人）\n我方已仔细研究招标文件和工程量清单。",
        "致：长丰县罗塘乡人民政府\n我方已仔细研究招标文件和工程量清单。",
        {},
    )

    assert not rendered.startswith("| 项目 | 内容")
    assert "工程量清单" in rendered


def test_v2_format_audit_rejects_missing_required_figures() -> None:
    report = audit_format_layer(
        pages=[("项目管理机构组织机构图", "拟为本标段工程设立的组织机构以框图方式表示。")],
        filled_pages=[("项目管理机构组织机构图", "拟设项目经理、技术负责人及各职能部门。")],
    )

    assert not report.passed
    assert any("图表/图片要求未落实" in issue.problem for issue in report.issues)


def test_kb_qualification_tables_filters_noise_and_builds_tables(monkeypatch) -> None:
    """C1:业绩/人员汇总表——构表正确、脏人名被过滤掉。"""
    monkeypatch.setattr(
        "services.knowledge_service.list_performance_records",
        lambda limit=15: [
            {"name": "某二级公路改建工程", "amount": "1500万元", "year": "2024", "type": "公路工程"},
        ],
    )
    monkeypatch.setattr(
        "services.knowledge_service.list_key_personnel",
        lambda limit=40: [
            {"name": "江舟", "certs": "一级建造师证"},
            {"name": "微信图片", "certs": "职称证书"},  # 噪声,应被过滤
            {"name": "施工员吴", "certs": "职称证书"},  # 噪声,应被过滤
        ],
    )

    md = v2_generation_service._kb_qualification_tables_markdown()

    assert "| 项目名称 | 中标金额 | 年份 | 类型 |" in md
    assert "某二级公路改建工程" in md and "1500万元" in md
    assert "| 姓名 | 持有证书 |" in md
    assert "江舟" in md
    assert "微信图片" not in md  # 噪声人名被过滤
    assert "施工员吴" not in md


def test_kb_qualification_tables_empty_when_no_data(monkeypatch) -> None:
    monkeypatch.setattr(
        "services.knowledge_service.list_performance_records", lambda limit=15: []
    )
    monkeypatch.setattr(
        "services.knowledge_service.list_key_personnel", lambda limit=40: []
    )
    assert v2_generation_service._kb_qualification_tables_markdown() == ""


# 用 CompanyProfile 真实字段名核对 profile 匹配:旧代码取 schema 里不存在的键名
# (project_manager/pm_certificate/qualification_level/safety_license/business_license),
# 这些字段恒空。下面两条测试在旧键名下会失败,修复后通过。
_PROFILE_FIELDS = {
    "company_name": "安徽正奇建设有限公司",
    "qualification_grade": "建筑工程施工总承包一级",
    "safety_license_no": "(皖)JZ安许证字[2021]007",
    "credit_code": "91340000MA2ABCDE3K",
    "project_manager_name": "李明",
    "project_manager_cert": "皖一级建造师A123456",
}


def test_match_profile_field_uses_real_schema_keys() -> None:
    assert (
        v2_generation_service._match_profile_field("资质要求响应", _PROFILE_FIELDS)
        == "建筑工程施工总承包一级"
    )
    assert (
        v2_generation_service._match_profile_field("安全生产许可证", _PROFILE_FIELDS)
        == "(皖)JZ安许证字[2021]007"
    )
    assert (
        v2_generation_service._match_profile_field("营业执照", _PROFILE_FIELDS)
        == "91340000MA2ABCDE3K"
    )


def test_enrich_commercial_fills_project_manager_from_profile() -> None:
    requirements = TenderRequirements(project_name="测试项目")
    # header-only 触发 enrich 路径
    enriched = v2_generation_service._enrich_commercial_markdown(
        "## 商务文件\n", requirements, _PROFILE_FIELDS
    )
    assert "李明" in enriched  # project_manager_name 被正确取出
    assert "皖一级建造师A123456" in enriched  # project_manager_cert


def test_enrich_commercial_uses_llm_response_when_tender_present(monkeypatch) -> None:
    """商务通读招标:有 tender_text 且 LLM 出响应时,用 LLM 响应、跳过模板浅资格响应。"""
    from schemas.tender import RequirementItem
    from services import commercial_response_service

    monkeypatch.setattr(
        commercial_response_service,
        "generate_commercial_responses",
        lambda requirements, tender_text, profile: (
            "\n## 附录：商务响应（AI 通读招标文件生成）\n\n"
            "| 招标资格要求 | 我方情况 |\n| --- | --- |\n| 公路三级 | 公路二级，满足 |\n"
        ),
    )
    requirements = TenderRequirements(
        project_name="测试项目",
        qualification_list=[RequirementItem(title="施工资质", description="公路三级及以上")],
    )
    enriched = v2_generation_service._enrich_commercial_markdown(
        "## 商务文件\n", requirements, _PROFILE_FIELDS, tender_text="招标全文：公路三级及以上。"
    )
    assert "AI 通读招标文件生成" in enriched
    assert "公路二级，满足" in enriched
    # 用了 LLM 响应 → 不再走模板浅响应那句
    assert "依据招标文件解析自动生成的资格响应要点" not in enriched


def test_enrich_commercial_appends_hardcheck_when_project_id(monkeypatch) -> None:
    """商务硬校验:有 project_id 时,把核对引擎渲染的硬校验表 append 进商务卷;无则不调。"""
    from services import tender_spec_service

    monkeypatch.setattr(
        tender_spec_service,
        "build_conformance_hardcheck_markdown",
        lambda pid: (
            "\n## 附录：资格符合性与投标函一致性核对（系统硬校验）\n\n"
            "| 核对项 | 判定 |\n| --- | --- |\n| 工期一致性 | ✅ 符合 |\n"
        ),
    )
    requirements = TenderRequirements(project_name="测试项目")
    enriched = v2_generation_service._enrich_commercial_markdown(
        "## 商务文件\n", requirements, _PROFILE_FIELDS, project_id=131
    )
    assert "系统硬校验" in enriched
    assert "工期一致性" in enriched
    # 无 project_id → 不调硬校验
    enriched_no_pid = v2_generation_service._enrich_commercial_markdown(
        "## 商务文件\n", requirements, _PROFILE_FIELDS
    )
    assert "系统硬校验" not in enriched_no_pid


def test_inject_project_images_targets_section_and_appendix(monkeypatch) -> None:
    """② 本项目插入图:按 target_section 插到对应节;无目标/未匹配的进末尾"本项目附图"。"""
    from services import v2_generation_service as v2
    from services import knowledge_service

    monkeypatch.setattr(
        knowledge_service,
        "list_project_insert_images",
        lambda pid: [
            {"document_id": 11, "file_name": "航拍图.jpg", "target_section": "施工总平面", "caption": "本项目航拍图"},
            {"document_id": 12, "file_name": "区位图.jpg", "target_section": "", "caption": "区位图"},
        ],
    )
    md = "## 一、工程概况\n\n概况。\n\n## 二、施工总平面布置\n\n平面。\n"
    out = v2._inject_project_images(md, project_id=700)

    # 航拍图插进"施工总平面布置"节(target 是其子串)
    assert "{{knowledge_image:document_id=11" in out
    plane_block = out.split("## 二、施工总平面布置")[1].split("##")[0]
    assert "document_id=11" in plane_block
    # 区位图无目标 → 末尾"本项目附图"
    assert "## 本项目附图" in out
    assert "document_id=12" in out.split("## 本项目附图")[1]
    # 无 project_id → 原样返回
    assert v2._inject_project_images(md, None) == md


def test_qualification_evidence_groups_all_company_certs(monkeypatch) -> None:
    """A2:资格证明材料成组插**全**公司证件——分组/组上限/去重/排除人员证件/specialty进caption。"""
    import re

    from services import knowledge_service, v2_generation_service

    refs = (
        [{"document_id": 1, "document_category": "公司证件",
          "certificate_type": "营业执照", "specialty": ""}]
        + [{"document_id": 1, "document_category": "公司证件",  # 重复 doc_id
            "certificate_type": "营业执照", "specialty": ""}]
        + [{"document_id": 100 + i, "document_category": "公司证件",
            "certificate_type": "资质证书", "specialty": f"专业{i}"} for i in range(20)]
        + [{"document_id": 9, "document_category": "公司证件",
            "certificate_type": "安全生产许可证", "specialty": ""}]
        + [{"document_id": 500, "document_category": "人员证件",  # 非公司证件,不选
            "certificate_type": "二级建造师证", "specialty": ""}]
    )
    monkeypatch.setattr(
        knowledge_service, "list_knowledge_image_references", lambda *a, **k: list(refs)
    )
    md = v2_generation_service._qualification_evidence_markdown()
    ids = re.findall(r"document_id=(\d+)", md)

    assert "1" in ids and "9" in ids            # 营业执照/安许选中
    assert "500" not in ids                     # 人员证件不选
    assert ids.count("1") == 1                  # doc_id 去重
    qual = [i for i in ids if 100 <= int(i) <= 119]
    assert len(qual) == 16                       # 资质证书组上限 16(20里取16)
    assert "### 企业资质证书" in md and "### 营业执照" in md
    assert "资质证书（专业0）" in md             # specialty 进 caption


def test_build_performance_evidence_chain_priority_and_caps() -> None:
    """A3b:业绩证明按项目成组,完整链(中标+交工)优先、年份新→旧;每类有上限。"""
    import re

    from services.v2_generation_service import _build_performance_evidence_md

    rows = (
        [(1, "A路2024", "中标通知书", "2024", "0"),
         (2, "A路2024", "合同", "2024", "0"),
         (3, "A路2024", "交工验收", "2024", "0")]
        + [(4, "B路2025", "合同", "2025", "0")]          # 更新但无链
        + [(5, "C路2023", "中标通知书", "2023", "0"),
           (6, "C路2023", "交工验收", "2023", "0")]
        + [(100 + i, "A路2024", "交工验收", "2024", str(i + 1)) for i in range(6)]  # 超 cap4
    )
    md = _build_performance_evidence_md(list(rows), limit_projects=6)
    secs = re.findall(r"### 类似业绩 \d+：(.+)", md)

    assert secs[0] == "A路2024" and secs[1] == "C路2023"   # 完整链优先,年份新先
    assert "B路2025" in secs                                # 无链也收,排后
    assert "A路2024-中标通知书" in md and "A路2024-交工验收（4）" in md
    assert "A路2024-交工验收（5）" not in md                 # 交工验收上限 4
    assert md.count('document_id=') >= 7


def test_performance_evidence_empty_when_no_rows() -> None:
    from services.v2_generation_service import _build_performance_evidence_md
    assert _build_performance_evidence_md([], 6) == ""


class _FakeChunk:
    """模拟 retriever.RetrievalResult(只用到 chunk_id/content/metadata)。"""

    def __init__(self, chunk_id, content, metadata):
        self.chunk_id = chunk_id
        self.content = content
        self.metadata = metadata


def test_flatten_retrieved_chunks_recovers_title_keyed_corpus() -> None:
    """回归(喂料致命 bug):retrieved 按中文章节标题归类,旧代码取
    'technical'/'施工组织' 永远落空 → 公司施组语料被静默丢弃、技术卷沦为空写。
    flatten 必须与键名解耦地汇总语料,并按 chunk_id 去重。"""
    a = _FakeChunk(1, "深基坑支护施工工艺……", {"document_category": "施工方案"})
    b = _FakeChunk(2, "营业执照", {"document_category": "公司证件"})
    retrieved = {
        "工程概况与项目特点分析": [a, b],
        "主要工程施工方案与技术措施": [a],  # 同片段跨章节重复,须去重
    }

    flat = v2_generation_service._flatten_retrieved_chunks(retrieved)
    assert [c.chunk_id for c in flat] == [1, 2]  # 去重 + 保序

    # 坐实历史 bug:旧坏 key 在该字典里永远取不到东西
    assert retrieved.get("technical") is None
    assert retrieved.get("施工组织") is None

    # 空/None 安全
    assert v2_generation_service._flatten_retrieved_chunks({}) == []
    assert v2_generation_service._flatten_retrieved_chunks(None) == []


def test_knowledge_chunk_payload_preserves_metadata() -> None:
    """回归:写作 prompt 靠 metadata['document_category']=='施工方案' 才会把片段
    标成【公司同类施工方案】;旧代码只留 content → 深度标注永不触发。"""
    payload = v2_generation_service._knowledge_chunk_payload(
        _FakeChunk(9, "钢筋加工与安装……", {"document_category": "施工方案", "specialty": "公路工程"})
    )
    assert payload["content"].startswith("钢筋加工")
    assert payload["metadata"]["document_category"] == "施工方案"

    # dict / str 两种入参也兼容
    assert v2_generation_service._knowledge_chunk_payload(
        {"content": "x", "metadata": {"document_category": "施工方案"}}
    ) == {"content": "x", "metadata": {"document_category": "施工方案"}}
    assert v2_generation_service._knowledge_chunk_payload("纯文本片段") == {
        "content": "纯文本片段",
        "metadata": {},
    }


def test_plan_chunk_reaches_writer_prompt_with_label() -> None:
    """端到端接线:retrieved(标题归类)→ flatten → payload → 写作 prompt,
    施工方案语料应被标成【公司同类施工方案】且正文出现在'知识库参考'里,
    证明料一路没丢、metadata 没被吞——技术卷不再是'未匹配到相关知识片段'。"""
    from prompts.generator_prompt import build_node_fill_prompt

    retrieved = {
        "主要工程施工方案与技术措施": [
            _FakeChunk(1, "路面基层水泥稳定碎石分层摊铺、碾压工艺……", {"document_category": "施工方案"}),
        ],
    }
    chunks = v2_generation_service._flatten_retrieved_chunks(retrieved)
    payloads = [v2_generation_service._knowledge_chunk_payload(c) for c in chunks]

    messages = build_node_fill_prompt(
        node_title="主要工程施工方案与技术措施",
        project_name="某农村公路提质改造工程",
        requirements={},
        company_name="安徽正奇建设有限公司",
        knowledge_chunks=payloads,
    )
    user_prompt = messages[-1]["content"]
    assert "【公司同类施工方案" in user_prompt
    assert "路面基层水泥稳定碎石" in user_prompt
    assert "未匹配到相关知识片段" not in user_prompt
