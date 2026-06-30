import pytest
from docx import Document

from services.generation_service import _append_docx


def test_append_docx_merges_paragraphs_and_tables(tmp_path) -> None:
    """技术卷末尾拼附表:正文 + 附表段落/表格都在,表格存活。"""
    base = tmp_path / "technical.docx"
    appendix = tmp_path / "appendix.docx"

    b = Document()
    b.add_paragraph("施工组织设计正文")
    b.save(str(base))

    a = Document()
    a.add_paragraph("附表一 施工总体计划表")
    a.add_table(rows=3, cols=4)
    a.save(str(appendix))

    _append_docx(base, str(appendix))

    merged = Document(str(base))
    texts = [p.text for p in merged.paragraphs if p.text.strip()]
    assert "施工组织设计正文" in texts
    assert "附表一 施工总体计划表" in texts
    assert len(merged.tables) == 1  # 附表的真表格被保留


def test_number_company_appendix_prepends_label(tmp_path) -> None:
    """公司定稿表标题(无编号)补成"附表X 名称",且仍是 Heading 样式(才进目录)。"""
    from services.appendix_service import _number_company_appendix

    src = tmp_path / "劳动力计划表.docx"
    d = Document()
    d.add_heading("劳动力计划表", level=2)
    d.save(str(src))

    out = _number_company_appendix(src, "三", "劳动力计划表", tmp_path / "out.docx")
    heads = [
        p.text.strip()
        for p in Document(str(out)).paragraphs
        if p.style.name.startswith("Heading") and p.text.strip()
    ]
    assert "附表三 劳动力计划表" in heads
    # 已带编号则不重复补
    out2 = _number_company_appendix(out, "三", "劳动力计划表", tmp_path / "out2.docx")
    heads2 = [p.text.strip() for p in Document(str(out2)).paragraphs if p.text.strip()]
    assert "附表三 劳动力计划表" in heads2
    assert "附表三 附表三 劳动力计划表" not in heads2


def test_build_appendix_docx_numbers_company_tables(tmp_path) -> None:
    """装配技术卷附表时,公司定稿表必须带上招标"附表X"编号(回归:函数曾定义却没调用)。"""
    pytest.importorskip("docxcompose")
    from services import appendix_service

    if appendix_service.match_company_appendix("劳动力计划表") is None:
        pytest.skip("公司定稿附表资产缺失,跳过")
    required = [("二", "施工总平面图"), ("三", "劳动力计划表")]
    out = appendix_service.build_appendix_docx(required, str(tmp_path / "ap.docx"))
    heads = [
        p.text.strip()
        for p in Document(str(out)).paragraphs
        if p.style.name.startswith("Heading") and p.text.strip()
    ]
    assert any(h.startswith("附表二") and "施工总平面图" in h for h in heads)
    assert any(h.startswith("附表三") and "劳动力计划表" in h for h in heads)
