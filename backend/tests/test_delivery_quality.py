"""delivery_quality 接入层的非阻断打分测试。

现场用 python-docx 造小商务卷(含身份字段表)+ 技术卷(长正文多标题)到 tmp_path,
mock get_company_profile 返回带 profile 的 dict,验证:
  - score_delivery_files 返回含 quality_score 的 dict 且落 eval_results;
  - 必填字段从商务卷读到(profile 值出现在 commercial → required_fields 高);
  - 非阻断:score_delivery 抛异常时,export 钩子不会让调用方报错。
"""

from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402

from services import delivery_quality  # noqa: E402

_PROFILE = {
    "company_name": "安徽正奇建设有限公司",
    "legal_representative": "王建国",
    "credit_code": "91340000MA2ABCDE3K",
    "safety_license_no": "(皖)JZ安许证字[2021]007",
    "project_manager_name": "李明",
}


def _build_commercial(path: Path) -> Path:
    """商务卷:含投标人基本情况表,字段值取自 _PROFILE(模拟已填真值)。"""
    document = Document()
    document.add_heading("投标人基本情况表", level=1)
    rows = [
        ("投标人", _PROFILE["company_name"]),
        ("法定代表人", _PROFILE["legal_representative"]),
        ("统一社会信用代码", _PROFILE["credit_code"]),
        ("安全生产许可证", _PROFILE["safety_license_no"]),
        ("项目经理", _PROFILE["project_manager_name"]),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    document.save(str(path))
    return path


def _build_technical(path: Path, *, headings: int = 10, body_chars: int = 18000) -> Path:
    """技术卷:长正文多标题,撑章节齐全度与篇幅。"""
    document = Document()
    for i in range(headings):
        document.add_heading(f"第{i + 1}章 施工组织设计章节{i + 1}", level=1)
        document.add_paragraph("正" * max(1, body_chars // max(1, headings)))
    document.save(str(path))
    return path


@pytest.fixture(autouse=True)
def _patch_profile_and_eval_dir(monkeypatch, tmp_path):
    """默认:profile 返回带 _PROFILE 的外层包装;eval_results 重定向到 tmp。"""
    monkeypatch.setattr(
        "services.company_profile_service.get_company_profile",
        lambda: {"profile": dict(_PROFILE), "updated_at": None},
    )
    # requirements 走启发式基准(无 project_id 或 DB 不可用时本就退 None)。
    monkeypatch.setattr(delivery_quality, "_load_requirements", lambda pid: None)
    monkeypatch.setattr(
        delivery_quality, "_EVAL_RESULTS_DIR", tmp_path / "eval_results"
    )


def test_score_delivery_files_returns_score_and_writes_eval(tmp_path: Path) -> None:
    commercial = _build_commercial(tmp_path / "commercial.docx")
    technical = _build_technical(tmp_path / "technical.docx")

    result = delivery_quality.score_delivery_files(
        commercial_path=commercial,
        technical_path=technical,
        project_id=42,
    )

    assert "quality_score" in result
    assert isinstance(result["quality_score"], (int, float))
    assert 0 <= result["quality_score"] <= 100
    assert "items" in result and "hard_block" in result

    eval_file = tmp_path / "eval_results" / "project_42.json"
    assert eval_file.exists()
    import json

    saved = json.loads(eval_file.read_text(encoding="utf-8"))
    assert saved["quality_score"] == result["quality_score"]


def test_required_fields_read_from_commercial_volume(tmp_path: Path) -> None:
    """profile 真值落在商务卷 → required_fields 应满分(全 5 字段命中)。"""
    commercial = _build_commercial(tmp_path / "commercial.docx")
    technical = _build_technical(tmp_path / "technical.docx")

    result = delivery_quality.score_delivery_files(
        commercial_path=commercial,
        technical_path=technical,
        project_id=1,
    )

    rf = result["items"]["required_fields"]
    assert rf["score"] == 1.0
    assert "商务卷" in rf["detail"]


def test_required_fields_low_when_commercial_missing_values(tmp_path: Path) -> None:
    """商务卷里不含 profile 真值 → required_fields 偏低,证明确实从商务卷核对。"""
    document = Document()
    document.add_heading("投标人基本情况表", level=1)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "投标人"
    table.rows[0].cells[1].text = "________"
    table.rows[1].cells[0].text = "法定代表人"
    table.rows[1].cells[1].text = "________"
    empty_commercial = tmp_path / "commercial_empty.docx"
    document.save(str(empty_commercial))
    technical = _build_technical(tmp_path / "technical.docx")

    result = delivery_quality.score_delivery_files(
        commercial_path=empty_commercial,
        technical_path=technical,
        project_id=2,
    )
    # profile 5 字段全不在文档里 → 0/5 命中。
    assert result["items"]["required_fields"]["score"] == 0.0


def test_write_eval_false_skips_file(tmp_path: Path) -> None:
    technical = _build_technical(tmp_path / "technical.docx")
    delivery_quality.score_delivery_files(
        technical_path=technical, project_id=99, write_eval=False
    )
    assert not (tmp_path / "eval_results" / "project_99.json").exists()


def test_export_hook_is_non_blocking_when_scoring_raises(monkeypatch, tmp_path) -> None:
    """score_delivery 抛异常时,export 钩子吞掉异常、出标流程继续。"""
    from services import generation_service

    # 让底层打分必抛。
    monkeypatch.setattr(
        "services.docx_health_check.score_delivery",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("boom")),
    )
    # 现场造一个 bid.docx(非两卷模式),让钩子能找到可打分的卷。
    technical = _build_technical(tmp_path / "project_7_bid.docx")
    assert technical.exists()

    # 钩子内部任何异常都不应冒泡。
    generation_service._score_delivery_quality(7, tmp_path)


def test_export_hook_skips_when_no_volumes(monkeypatch, tmp_path) -> None:
    """临时目录里没有任何卷文件时,钩子安静跳过,不报错、不调打分。"""
    from services import generation_service

    called = {"n": 0}

    def _spy(**kwargs):
        called["n"] += 1
        return {"quality_score": 1.0}

    monkeypatch.setattr(delivery_quality, "score_delivery_files", _spy)
    generation_service._score_delivery_quality(7, tmp_path)
    assert called["n"] == 0
