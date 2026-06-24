from __future__ import annotations

from io import BytesIO

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from api.deps import _raise_http_error
from schemas.auth import UserProfile
from services import auth_service, cert_archive_service

router = APIRouter()


class PersonReassignRequest(BaseModel):
    document_ids: list[int]
    target_person: str


class CompanyRetypeRequest(BaseModel):
    document_ids: list[int]
    target_type: str


class RenameRequest(BaseModel):
    title: str


@router.get("/api/cert-archive/person")
def get_person_archive(
    _current_user: UserProfile = Depends(auth_service.require_knowledge_view),
) -> dict:
    try:
        return cert_archive_service.list_person_archive()
    except Exception as error:  # noqa: BLE001
        _raise_http_error(error)


@router.get("/api/cert-archive/company")
def get_company_archive(
    _current_user: UserProfile = Depends(auth_service.require_knowledge_view),
) -> dict:
    try:
        return cert_archive_service.list_company_archive()
    except Exception as error:  # noqa: BLE001
        _raise_http_error(error)


@router.post("/api/cert-archive/person/reassign")
def reassign_person(
    request: PersonReassignRequest,
    _current_user: UserProfile = Depends(auth_service.require_knowledge_edit),
) -> dict:
    try:
        changed = cert_archive_service.reassign_person_cert(
            request.document_ids, request.target_person
        )
        return {"ok": True, "changed": changed}
    except Exception as error:  # noqa: BLE001
        _raise_http_error(error)


@router.post("/api/cert-archive/company/retype")
def retype_company(
    request: CompanyRetypeRequest,
    _current_user: UserProfile = Depends(auth_service.require_knowledge_edit),
) -> dict:
    try:
        changed = cert_archive_service.retype_company_cert(
            request.document_ids, request.target_type
        )
        return {"ok": True, "changed": changed}
    except Exception as error:  # noqa: BLE001
        _raise_http_error(error)


@router.get("/api/cert-archive/test-insert/{document_id}")
def test_insert_docx(
    document_id: int,
    _current_user: UserProfile = Depends(auth_service.require_knowledge_view),
) -> StreamingResponse:
    """自助验证:把这张证件的**真实图片**插进一个 DOCX 并下载,亲眼确认不是只写名字。"""
    from docx import Document
    from docx.shared import Inches

    from services import asset_resolver

    try:
        asset = asset_resolver.resolve_asset(document_id=document_id)
        if not asset.get("matched"):
            raise ValueError("没找到这张证件的图片")
        blob = asset_resolver.read_asset_bytes(document_id)
        if not blob:
            raise ValueError("读到空图片")

        doc = Document()
        doc.add_heading("图片插入测试", level=1)
        doc.add_paragraph(f"材料名称：{asset['asset_name']}")
        doc.add_paragraph(
            f"归属：{asset.get('owner_name', '')}  类型：{asset.get('asset_type', '')}"
        )
        from utils.image_orient import upright_image_bytes

        doc.add_picture(BytesIO(upright_image_bytes(blob)), width=Inches(5.5))
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
    except Exception as error:  # noqa: BLE001
        _raise_http_error(error)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=test_insert_{document_id}.docx"
        },
    )


@router.patch("/api/cert-archive/cert/{document_id}")
def rename_cert(
    document_id: int,
    request: RenameRequest,
    _current_user: UserProfile = Depends(auth_service.require_knowledge_edit),
) -> dict:
    try:
        cert_archive_service.rename_cert(document_id, request.title)
        return {"ok": True}
    except Exception as error:  # noqa: BLE001
        _raise_http_error(error)
