"""最小闭环验证:资料库真实图片 → 读字节 → 插入 DOCX → 校验 word/media 里有图。

用法:
  # 按人名+类型(默认演示:第一张身份证)
  PYTHONPATH=backend backend/venv/bin/python backend/scripts/test_asset_insert.py
  PYTHONPATH=backend backend/venv/bin/python backend/scripts/test_asset_insert.py --owner 江舟 --type 身份证
  PYTHONPATH=backend backend/venv/bin/python backend/scripts/test_asset_insert.py --type 营业执照
  PYTHONPATH=backend backend/venv/bin/python backend/scripts/test_asset_insert.py --document-id 290

验收:打开 test_asset_insert.docx 能看到真实图片,且本脚本输出 closed_loop_ok=True。
"""

from __future__ import annotations

import argparse
import json
import sys
from io import BytesIO
from pathlib import Path

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from docx import Document  # noqa: E402
from docx.shared import Inches  # noqa: E402

from services import asset_resolver  # noqa: E402
from utils.docx_image_check import build_insert_report  # noqa: E402


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="资料库图片插入 DOCX 最小闭环")
    p.add_argument("--owner", default=None, help="owner_name,如 江舟")
    p.add_argument("--type", dest="asset_type", default="身份证", help="asset_type,如 身份证/营业执照")
    p.add_argument("--owner-type", default=None, help="personnel / company")
    p.add_argument("--document-id", type=int, default=None)
    p.add_argument("--side", default=None, help="身份证选面: front / back / both(走OCR判面)")
    p.add_argument("--out", default="test_asset_insert.docx")
    return p.parse_args()


def main() -> None:
    args = parse_args()

    # 身份证 + 指定面 → 走 OCR 判面,避免抓错正反面
    if args.side and (args.asset_type == "身份证") and args.owner and not args.document_id:
        asset = asset_resolver.resolve_id_card(args.owner, side=args.side)
    else:
        asset = asset_resolver.resolve_asset(
            owner_name=args.owner,
            asset_type=args.asset_type,
            owner_type=args.owner_type,
            document_id=args.document_id,
        )
    print("AssetResolver 输出:")
    print(json.dumps(asset, ensure_ascii=False, indent=2))

    if not asset.get("matched"):
        print("\n❌ 没解析到资产,闭环失败。")
        sys.exit(1)

    errors: list[str] = []
    expected = [{"asset_name": asset["asset_name"], "asset_type": asset["asset_type"], "status": "pending"}]

    doc = Document()
    doc.add_heading("图片插入测试", level=1)
    doc.add_paragraph(f"材料名称：{asset['asset_name']}")
    doc.add_paragraph(f"归属：{asset.get('owner_name','')}  类型：{asset.get('asset_type','')}")

    try:
        blob = asset_resolver.read_asset_bytes(asset["document_id"])
        if not blob:
            raise ValueError("读到空字节")
        # 真插入 image part(不是写文件名)
        doc.add_picture(BytesIO(blob), width=Inches(5.5))
        expected[0]["status"] = "inserted"
        print(f"\n✅ 已插入真实图片 {len(blob)} 字节")
    except Exception as exc:  # noqa: BLE001
        errors.append(f"{asset['asset_name']}: {exc}")
        print(f"\n❌ 插图失败: {exc}")

    out_path = str(BACKEND_DIR.parent / args.out)
    doc.save(out_path)

    report = build_insert_report(out_path, expected, errors)
    print("\n插图 report:")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"\n生成文件: {out_path}")
    sys.exit(0 if report["closed_loop_ok"] else 2)


if __name__ == "__main__":
    main()
