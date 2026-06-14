"""对标真实中标标书的质量基准脚本（Phase 0）。

用法：
    # 单文件指标
    python backend/scripts/benchmark_vs_baseline.py stats <doc.pdf|doc.docx>

    # 生成稿 vs 基线对比（可选传 requirements.json 算评分点/废标项响应率）
    python backend/scripts/benchmark_vs_baseline.py compare \
        --generated out.docx --baseline data/baseline/真实投标文件（标书）/1.真实投标文件.PDF \
        [--requirements parsed.json] [--json report.json]

指标：页数、字数、表格数、章节数、残留下划线占位数、评分点/废标项响应率。
精确页数需要 LibreOffice（soffice）；自动探测 PATH 与 macOS app 包路径，
找不到时 DOCX 页数记为 None，其余指标照常输出。
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, asdict
from pathlib import Path

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from utils.keywords import extract_keywords  # noqa: E402

_UNDERSCORE_RUN = re.compile(r"_{3,}")
_HEADING_NUM = re.compile(r"^\s*(?:第[一二三四五六七八九十百\d]+[章节]|[一二三四五六七八九十]+、|\d+[\.、])")


def find_soffice() -> str | None:
    for candidate in ("soffice", "libreoffice"):
        found = shutil.which(candidate)
        if found:
            return found
    mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    return mac if os.path.exists(mac) else None


@dataclass
class DocMetrics:
    path: str
    kind: str  # pdf | docx
    page_count: int | None
    char_count: int
    table_count: int
    heading_count: int
    underscore_blanks: int


def _pdf_metrics(path: Path) -> DocMetrics:
    import fitz

    doc = fitz.open(str(path))
    try:
        texts = [doc[i].get_text() for i in range(doc.page_count)]
        full = "\n".join(texts)
        page_count = doc.page_count
    finally:
        doc.close()
    compact = re.sub(r"\s+", "", full)
    headings = sum(1 for line in full.splitlines() if _HEADING_NUM.match(line))
    # PDF tables are not structurally recoverable here; report 0 (use DOCX side for density).
    return DocMetrics(
        path=str(path),
        kind="pdf",
        page_count=page_count,
        char_count=len(compact),
        table_count=0,
        heading_count=headings,
        underscore_blanks=len(_UNDERSCORE_RUN.findall(full)),
    )


def _docx_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _docx_page_count(path: Path, soffice: str | None) -> int | None:
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, str(path)],
                check=True,
                capture_output=True,
                timeout=180,
            )
        except (subprocess.SubprocessError, OSError):
            return None
        pdfs = list(Path(tmp).glob("*.pdf"))
        if not pdfs:
            return None
        import fitz

        doc = fitz.open(str(pdfs[0]))
        try:
            return doc.page_count
        finally:
            doc.close()


def _docx_metrics(path: Path, soffice: str | None) -> DocMetrics:
    from docx import Document

    document = Document(str(path))
    full = _docx_text(document)
    compact = re.sub(r"\s+", "", full)
    headings = sum(1 for line in full.splitlines() if _HEADING_NUM.match(line))
    return DocMetrics(
        path=str(path),
        kind="docx",
        page_count=_docx_page_count(path, soffice),
        char_count=len(compact),
        table_count=len(document.tables),
        heading_count=headings,
        underscore_blanks=len(_UNDERSCORE_RUN.findall(full)),
    )


def doc_metrics(path: Path, soffice: str | None) -> DocMetrics:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf_metrics(path)
    if suffix in (".docx",):
        return _docx_metrics(path, soffice)
    raise ValueError(f"不支持的文件类型：{suffix}（仅 .pdf/.docx）")


def _doc_full_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        import fitz

        doc = fitz.open(str(path))
        try:
            return "\n".join(doc[i].get_text() for i in range(doc.page_count))
        finally:
            doc.close()
    from docx import Document

    return _docx_text(Document(str(path)))


def response_rate(text: str, items: list[dict]) -> tuple[int, int, list[str]]:
    """Fraction of requirement items whose keywords appear in ``text``.

    Mirrors scoring_agent coverage: an item is 'responded' if any of its
    extracted keywords is found in the document text.
    """
    total = len(items)
    if total == 0:
        return 0, 0, []
    missed: list[str] = []
    hit = 0
    for item in items:
        title = str(item.get("title", "") or "")
        desc = str(item.get("description", "") or "")
        keywords = extract_keywords(f"{title} {desc}", limit=8) or [title[:20]]
        covered = any(kw and re.search(re.escape(kw), text) for kw in keywords)
        if covered:
            hit += 1
        else:
            missed.append(title or desc[:20])
    return hit, total, missed


def _fmt(m: DocMetrics) -> str:
    pages = m.page_count if m.page_count is not None else "?(无soffice)"
    return (
        f"  页数={pages}  字数={m.char_count}  表格={m.table_count}  "
        f"章节={m.heading_count}  残留下划线={m.underscore_blanks}"
    )


def cmd_stats(args) -> None:
    soffice = find_soffice()
    m = doc_metrics(Path(args.path), soffice)
    print(f"{m.path} [{m.kind}]")
    print(_fmt(m))


def cmd_compare(args) -> None:
    soffice = find_soffice()
    if not soffice:
        print("⚠️  未找到 soffice，DOCX 精确页数将为空（其余指标正常）。\n")
    gen = doc_metrics(Path(args.generated), soffice)
    base = doc_metrics(Path(args.baseline), soffice)

    print("生成稿:", gen.path)
    print(_fmt(gen))
    print("基线  :", base.path)
    print(_fmt(base))

    report: dict = {"generated": asdict(gen), "baseline": asdict(base)}

    if args.requirements:
        req = json.loads(Path(args.requirements).read_text(encoding="utf-8"))
        gen_text = _doc_full_text(Path(args.generated))
        score_items = req.get("technical_score_items", [])
        invalid_items = req.get("invalid_bid_items", [])
        s_hit, s_total, s_missed = response_rate(gen_text, score_items)
        i_hit, i_total, i_missed = response_rate(gen_text, invalid_items)
        print("\n响应率（生成稿）:")
        if s_total:
            print(f"  评分点: {s_hit}/{s_total} = {s_hit / s_total:.0%}   未命中: {s_missed[:8]}")
        else:
            print("  评分点: 无（requirements 未含 technical_score_items）")
        if i_total:
            print(f"  废标项: {i_hit}/{i_total} = {i_hit / i_total:.0%}   未命中: {i_missed[:8]}")
        else:
            print("  废标项: 无")
        report["response"] = {
            "score": {"hit": s_hit, "total": s_total, "missed": s_missed},
            "invalid": {"hit": i_hit, "total": i_total, "missed": i_missed},
        }

    if args.json:
        Path(args.json).write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"\n已写入 {args.json}")


def main() -> None:
    parser = argparse.ArgumentParser(description="对标真实中标标书的质量基准脚本。")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_stats = sub.add_parser("stats", help="输出单个文件的指标")
    p_stats.add_argument("path")
    p_stats.set_defaults(func=cmd_stats)

    p_cmp = sub.add_parser("compare", help="生成稿 vs 基线对比")
    p_cmp.add_argument("--generated", required=True)
    p_cmp.add_argument("--baseline", required=True)
    p_cmp.add_argument("--requirements", help="解析出的 requirements JSON（算响应率）")
    p_cmp.add_argument("--json", help="把报告写入该 JSON 文件")
    p_cmp.set_defaults(func=cmd_compare)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
