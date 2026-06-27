from __future__ import annotations

from io import BytesIO
from pathlib import Path
import shutil
import subprocess
from tempfile import TemporaryDirectory
from typing import BinaryIO

import pdfplumber
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf", ".doc", ".docx", ".txt", ".md",
    ".jpg", ".jpeg", ".png",
    ".xlsx", ".xlsm", ".xls",
}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
LEGACY_WORD_EXTENSIONS = {".doc"}
TEXT_EXTENSIONS = {".txt", ".md"}
# 工程量清单常是 Excel:.xlsx/.xlsm 直接用 openpyxl 读,旧版 .xls 走 LibreOffice 转换。
XLSX_EXTENSIONS = {".xlsx", ".xlsm"}
LEGACY_EXCEL_EXTENSIONS = {".xls"}


def _as_bytes_io(file_data: bytes | bytearray | BinaryIO) -> BytesIO | BinaryIO:
    if isinstance(file_data, (bytes, bytearray)):
        return BytesIO(file_data)
    return file_data


def extract_text_from_pdf(file_path: str | Path | bytes | bytearray | BinaryIO) -> str:
    """Extract readable text from a PDF path or byte stream."""
    if isinstance(file_path, (str, Path)):
        path = Path(file_path)
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    else:
        data = _as_bytes_io(file_path)
        try:
            with pdfplumber.open(data) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
        except Exception:
            if hasattr(data, "seek"):
                data.seek(0)
            reader = PdfReader(data)
            pages = [page.extract_text() or "" for page in reader.pages]

    return "\n\n".join(page.strip() for page in pages if page.strip())


def extract_text_from_docx(file_path: str | Path | bytes | bytearray | BinaryIO) -> str:
    """Extract paragraph and table text from a DOCX path or byte stream."""
    document = Document(
        _as_bytes_io(file_path)
        if not isinstance(file_path, (str, Path))
        else str(file_path)
    )
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_xlsx(file_path: str | Path | bytes | bytearray | BinaryIO) -> str:
    """Extract sheet/row/cell text from an XLSX (or XLSM) path or byte stream.

    把每行单元格序列化成「单元格 | 单元格」(同 docx 表格的写法),让工程量清单这类
    Excel 表格能被当文本读、喂给 LLM 估占比。``data_only=True`` 取公式的缓存计算值
    (合价等);若工作簿从未被 Excel 算过、无缓存值,该单元格为空(只丢公式不影响整表)。
    """
    from openpyxl import load_workbook

    source = (
        str(file_path)
        if isinstance(file_path, (str, Path))
        else _as_bytes_io(file_path)
    )
    workbook = load_workbook(source, read_only=True, data_only=True)
    parts: list[str] = []
    try:
        multi_sheet = len(workbook.worksheets) > 1
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = [
                    str(value).strip()
                    for value in row
                    if value is not None and str(value).strip()
                ]
                if cells:
                    rows.append(" | ".join(cells))
            if not rows:
                continue
            if multi_sheet:
                parts.append(f"## 工作表:{sheet.title}")
            parts.extend(rows)
    finally:
        workbook.close()
    return "\n".join(parts)


def extract_text_from_legacy_xls(
    file_path: str | Path | bytes | bytearray | BinaryIO,
) -> str:
    """Convert a legacy .xls file with LibreOffice, then extract XLSX text."""
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if not converter:
        raise ValueError(
            "Legacy .xls conversion requires LibreOffice/soffice. Re-save the file "
            "as .xlsx, or install LibreOffice before parsing .xls."
        )

    with TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        source_path = tmp_path / "input.xls"
        if isinstance(file_path, (str, Path)):
            source_path.write_bytes(Path(file_path).read_bytes())
        elif isinstance(file_path, (bytes, bytearray)):
            source_path.write_bytes(bytes(file_path))
        else:
            source_path.write_bytes(file_path.read())

        subprocess.run(
            [
                converter,
                "--headless",
                "--convert-to",
                "xlsx",
                "--outdir",
                str(tmp_path),
                str(source_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        converted = source_path.with_suffix(".xlsx")
        if not converted.exists():
            matches = list(tmp_path.glob("*.xlsx"))
            if not matches:
                raise ValueError("Legacy .xls conversion did not produce an XLSX file")
            converted = matches[0]
        return extract_text_from_xlsx(converted)


def extract_text_from_txt(file_path: str | Path | bytes | bytearray | BinaryIO) -> str:
    """Extract text from a TXT path or byte stream."""
    if isinstance(file_path, (str, Path)):
        return Path(file_path).read_text(encoding="utf-8")

    data = file_path if isinstance(file_path, (bytes, bytearray)) else file_path.read()
    if isinstance(data, str):
        return data
    return bytes(data).decode("utf-8")


_RAPIDOCR_ENGINE = None
_RAPIDOCR_FAILED = False


def _get_rapidocr():
    """Cached RapidOCR engine (model load is ~1-2s; reuse across images)."""
    global _RAPIDOCR_ENGINE, _RAPIDOCR_FAILED
    if _RAPIDOCR_ENGINE is not None or _RAPIDOCR_FAILED:
        return _RAPIDOCR_ENGINE
    try:
        from rapidocr_onnxruntime import RapidOCR

        _RAPIDOCR_ENGINE = RapidOCR()
    except Exception:
        _RAPIDOCR_FAILED = True
    return _RAPIDOCR_ENGINE


def _image_bytes(
    file_path: str | Path | bytes | bytearray | BinaryIO,
) -> bytes:
    if isinstance(file_path, (bytes, bytearray)):
        return bytes(file_path)
    if isinstance(file_path, (str, Path)):
        return Path(file_path).read_bytes()
    return file_path.read()


def extract_text_from_image(
    file_path: str | Path | bytes | bytearray | BinaryIO,
) -> str:
    """OCR a scanned image (公司/人员证件等) into searchable text.

    Primary engine is RapidOCR (onnxruntime, no system binary, good Chinese on
    photos); falls back to Tesseract (pytesseract + tesseract binary) if RapidOCR
    is unavailable. Raises only when no OCR engine is usable.
    """
    data = _image_bytes(file_path)

    # Primary: RapidOCR
    engine = _get_rapidocr()
    if engine is not None:
        try:
            import io

            import numpy as np
            from PIL import Image

            arr = np.array(Image.open(io.BytesIO(data)).convert("RGB"))
            result, _ = engine(arr)
            text = "\n".join(line[1] for line in (result or [])).strip()
            if text:
                return text
        except Exception:
            pass  # fall through to Tesseract

    # Fallback: Tesseract
    try:
        import io

        import pytesseract
        from PIL import Image

        return pytesseract.image_to_string(
            Image.open(io.BytesIO(data)), lang="chi_sim+eng"
        ).strip()
    except Exception as error:
        raise ValueError(
            "Image OCR is not configured. Store this file as evidence-only or "
            "install OCR tooling (rapidocr-onnxruntime 或 tesseract) before "
            "indexing image text."
        ) from error


def extract_text_from_legacy_doc(
    file_path: str | Path | bytes | bytearray | BinaryIO,
) -> str:
    """Convert a legacy .doc file with LibreOffice, then extract DOCX text."""
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if not converter:
        raise ValueError(
            "Legacy .doc conversion requires LibreOffice/soffice. Store this file "
            "as evidence-only or install LibreOffice before indexing .doc text."
        )

    with TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        source_path = tmp_path / "input.doc"
        if isinstance(file_path, (str, Path)):
            source_path.write_bytes(Path(file_path).read_bytes())
        elif isinstance(file_path, (bytes, bytearray)):
            source_path.write_bytes(bytes(file_path))
        else:
            source_path.write_bytes(file_path.read())

        subprocess.run(
            [
                converter,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(tmp_path),
                str(source_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        converted = source_path.with_suffix(".docx")
        if not converted.exists():
            matches = list(tmp_path.glob("*.docx"))
            if not matches:
                raise ValueError("Legacy .doc conversion did not produce a DOCX file")
            converted = matches[0]
        return extract_text_from_docx(converted)


def extract_text(
    file_input: str | Path | bytes | bytearray | BinaryIO,
    filename: str | None = None,
    content_type: str | None = None,
) -> str:
    """Route an uploaded tender file to the right text extractor."""
    suffix = (
        Path(filename or str(file_input)).suffix.lower()
        if filename or isinstance(file_input, (str, Path))
        else ""
    )

    if content_type == "application/pdf" or suffix == ".pdf":
        return extract_text_from_pdf(file_input)
    if (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix == ".docx"
    ):
        return extract_text_from_docx(file_input)
    if content_type == "text/plain" or suffix in TEXT_EXTENSIONS:
        return extract_text_from_txt(file_input)
    if suffix in XLSX_EXTENSIONS or content_type in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel.sheet.macroEnabled.12",
    }:
        return extract_text_from_xlsx(file_input)
    if suffix in LEGACY_EXCEL_EXTENSIONS or content_type == "application/vnd.ms-excel":
        return extract_text_from_legacy_xls(file_input)
    if suffix in LEGACY_WORD_EXTENSIONS or content_type == "application/msword":
        return extract_text_from_legacy_doc(file_input)
    if suffix in IMAGE_EXTENSIONS or (content_type or "").startswith("image/"):
        return extract_text_from_image(file_input)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise ValueError(f"Unsupported file type. Expected one of: {supported}")
