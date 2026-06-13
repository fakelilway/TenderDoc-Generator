from __future__ import annotations

import re
import textwrap
from collections.abc import Callable
from io import BytesIO
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Production typography for Chinese bid documents:
# body uses 宋体, headings use 黑体 (bold), Latin/digits use Times New Roman.
BODY_CJK_FONT = "SimSun"  # 宋体
HEADING_CJK_FONT = "SimHei"  # 黑体
FOOTER_CJK_FONT = "NSimSun"  # 新宋体
LATIN_FONT = "Times New Roman"
BODY_SIZE_PT = 12  # 小四
HEADING_SIZES_PT = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 12}  # 三号/四号/小四
# 2-character first-line indent at the body font size.
FIRST_LINE_INDENT_PT = BODY_SIZE_PT * 2
ZHENGQI_PROFILE = "zhengqi"
KNOWLEDGE_IMAGE_MARKER_RE = re.compile(r"^\{\{knowledge_image:(?P<body>.+)\}\}$")
TDG_TABLE_MARKER_RE = re.compile(r"^\{\{tdg_table:(?P<kind>[A-Za-z_][\w-]*)(?P<body>.*?)\}\}$")
_IMAGE_MARKER_ATTR_RE = re.compile(
    r"(?P<key>[A-Za-z_][\w-]*)=(?P<value>\"[^\"]*\"|'[^']*'|[^\s]+)"
)
ImageResolver = Callable[[int], bytes | bytearray | str | Path | None]


# Headings whose title contains any of these keywords are routed to the
# commercial (商务标) volume when splitting a bid document into volumes.
COMMERCIAL_KEYWORDS = (
    "商务",
    "报价",
    "投标函",
    "投标报价",
    "开标一览",
    "资格审查",
    "资格证明",
    "营业执照",
    "财务",
    "信用",
    "承诺函",
    "声明函",
)
PRICING_KEYWORDS = (
    "报价",
    "投标报价",
    "开标一览",
    "工程量清单",
    "综合单价",
    "商务报价",
    "价格",
    "投标总价",
)


def markdown_to_docx(
    markdown_text: str,
    output_path: str | Path,
    *,
    title: str | None = None,
    subtitle: str | None = None,
    cover: bool = False,
    toc: bool = False,
    header_text: str | None = None,
    page_numbers: bool = False,
    metadata: dict[str, str] | None = None,
    style_profile: str | None = None,
    image_resolver: ImageResolver | None = None,
) -> str:
    """Render markdown into a styled DOCX file.

    The base behaviour (headings/paragraphs/lists/tables) is unchanged; the
    keyword-only arguments opt into the M58 formatting upgrades: a cover page,
    an auto-updating table of contents, running headers and page numbers.
    """
    if not markdown_text.strip():
        raise ValueError("markdown_text is empty")

    document = Document()
    _configure_styles(document, style_profile)

    if header_text or page_numbers:
        _configure_header_footer(document, header_text, page_numbers, style_profile)

    if cover:
        _add_cover_page(
            document,
            title or _extract_title(markdown_text) or "投标文件",
            subtitle,
            metadata,
            style_profile,
        )

    if toc:
        _add_table_of_contents(document)

    _render_markdown_body(document, markdown_text, style_profile, image_resolver)

    if toc or page_numbers:
        _enable_field_updates(document)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return str(path)


def markdown_to_pdf(
    markdown_text: str,
    output_path: str | Path,
    *,
    title: str | None = None,
) -> str:
    """Render markdown into a simple Chinese-readable PDF.

    DOCX remains the high-fidelity output. This PDF renderer is intentionally
    plain so tender platforms that require PDF uploads can receive a directly
    generated file without depending on LibreOffice.
    """
    if not markdown_text.strip():
        raise ValueError("markdown_text is empty")

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    font = fitz.Font("china-s")
    x = 62
    y = 64
    bottom = 790
    line_height = 21

    if title:
        y = _pdf_write_line(page, title, x, y, font, 18, bold=False)
        y += 18

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            y += 8
            continue
        if line.startswith("|"):
            line = "  ".join(cell.strip() for cell in line.strip("|").split("|"))
            if set(line.replace(" ", "")) <= {"-", ":"}:
                continue
        size = 14
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            line = line.lstrip("#").strip()
            size = {1: 18, 2: 16, 3: 15}.get(level, 14)
            y += 8
        for wrapped in _wrap_cjk_text(line, width=36 if size >= 16 else 42):
            if y > bottom:
                page = document.new_page(width=595, height=842)
                y = 64
            y = _pdf_write_line(page, wrapped, x, y, font, size, bold=False)
            y += line_height if size <= 14 else 24

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    document.close()
    return str(path)


def _pdf_write_line(page, text: str, x: float, y: float, font, size: float, bold: bool):
    page.insert_text((x, y), text, fontname="china-s", fontsize=size, fontfile=None)
    return y


def _wrap_cjk_text(text: str, width: int) -> list[str]:
    return textwrap.wrap(
        text,
        width=width,
        break_long_words=True,
        replace_whitespace=False,
        drop_whitespace=True,
    ) or [text]


def _render_markdown_body(
    document: Document,
    markdown_text: str,
    style_profile: str | None,
    image_resolver: ImageResolver | None = None,
) -> None:
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue

        if _is_pagebreak_marker(line):
            document.add_page_break()
            index += 1
            continue

        table_marker = _parse_tdg_table_marker(line)
        if table_marker:
            _add_tdg_table(document, table_marker, style_profile)
            index += 1
            continue

        image_marker = _parse_knowledge_image_marker(line)
        if image_marker:
            _add_knowledge_image(document, image_marker, image_resolver, style_profile)
            index += 1
            continue

        if _is_table_start(lines, index):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_table(document, table_lines, style_profile)
            continue

        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            title = line.lstrip("#").strip()
            if title:
                document.add_heading(title, level=level)
        elif line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_text_with_underlined_blanks(
                paragraph,
                line[2:].strip(),
                BODY_CJK_FONT,
                _body_size_pt(style_profile),
            )
        elif line[:3].isdigit() and line[3:5] in {". ", "、"}:
            paragraph = document.add_paragraph(style="List Number")
            _add_text_with_underlined_blanks(
                paragraph,
                line[5:].strip(),
                BODY_CJK_FONT,
                _body_size_pt(style_profile),
            )
        else:
            paragraph = document.add_paragraph()
            _add_text_with_underlined_blanks(
                paragraph,
                line,
                BODY_CJK_FONT,
                _body_size_pt(style_profile),
            )
            # Production body paragraphs start with a 2-character indent.
            paragraph.paragraph_format.first_line_indent = Pt(
                _body_size_pt(style_profile) * 2
            )
            _apply_form_paragraph_alignment(paragraph)
        index += 1


def _is_pagebreak_marker(line: str) -> bool:
    return line.strip() in {"<!-- tdg:pagebreak -->", "---PAGEBREAK---", "\\f"}


def _parse_tdg_table_marker(line: str) -> dict[str, object] | None:
    match = TDG_TABLE_MARKER_RE.match(line.strip())
    if not match:
        return None
    attrs = {
        attr.group("key"): _strip_attr_quotes(attr.group("value"))
        for attr in _IMAGE_MARKER_ATTR_RE.finditer(match.group("body"))
    }
    return {"kind": match.group("kind"), "attrs": attrs}


def _add_tdg_table(
    document: Document,
    marker: dict[str, object],
    style_profile: str | None,
) -> None:
    if marker.get("kind") == "bidder_basic_info":
        _add_bidder_basic_info_table(
            document,
            marker.get("attrs") if isinstance(marker.get("attrs"), dict) else {},
            style_profile,
        )
        return
    paragraph = document.add_paragraph(f"（未支持的表格类型：{marker.get('kind')}）")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _parse_knowledge_image_marker(line: str) -> dict[str, object] | None:
    match = KNOWLEDGE_IMAGE_MARKER_RE.match(line.strip())
    if not match:
        return None
    attrs = {
        attr.group("key"): _strip_attr_quotes(attr.group("value"))
        for attr in _IMAGE_MARKER_ATTR_RE.finditer(match.group("body"))
    }
    document_id = attrs.get("document_id") or attrs.get("id")
    if not document_id:
        legacy_match = re.search(r"\b(\d+)\b", match.group("body"))
        document_id = legacy_match.group(1) if legacy_match else None
    if not document_id:
        return None
    try:
        image_id = int(str(document_id))
    except ValueError:
        return None
    width_cm = _parse_marker_float(attrs.get("width_cm"), default=14.0)
    return {
        "document_id": image_id,
        "caption": str(attrs.get("caption") or "").strip(),
        "width_cm": max(4.0, min(width_cm, 16.0)),
    }


def _add_bidder_basic_info_table(
    document: Document,
    attrs: dict[str, str],
    style_profile: str | None,
) -> None:
    table = document.add_table(rows=15, cols=6)
    table.style = "Table Grid"
    table.autofit = False
    widths_cm = (2.9, 1.5, 2.0, 2.0, 2.6, 2.9)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def val(key: str) -> str:
        return str(attrs.get(key) or "________")

    def put(row: int, col: int, text: str, *, center: bool = False) -> None:
        cell = table.cell(row, col)
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        _add_text_with_underlined_blanks(
            paragraph,
            text,
            BODY_CJK_FONT,
            14 if style_profile == ZHENGQI_PROFILE else BODY_SIZE_PT,
        )

    def merge(row: int, start_col: int, end_col: int):
        return table.cell(row, start_col).merge(table.cell(row, end_col))

    def vmerge(start_row: int, end_row: int, col: int):
        return table.cell(start_row, col).merge(table.cell(end_row, col))

    # Row 0-6: mirror the tender form's visible field layout.
    put(0, 0, "投标人名称", center=True)
    merge(0, 1, 5)
    put(0, 1, val("company_name"))

    put(1, 0, "注册地址", center=True)
    merge(1, 1, 3)
    put(1, 1, val("address"))
    put(1, 4, "邮政编码", center=True)
    put(1, 5, val("postal_code"))

    vmerge(2, 3, 0)
    put(2, 0, "联系方式", center=True)
    put(2, 1, "联系人", center=True)
    merge(2, 2, 3)
    put(2, 2, val("contact_person"))
    put(2, 4, "电话", center=True)
    put(2, 5, val("phone"))
    put(3, 1, "传真", center=True)
    merge(3, 2, 3)
    put(3, 2, val("fax"))
    put(3, 4, "电子邮件", center=True)
    put(3, 5, val("email"))

    put(4, 0, "法定代表人", center=True)
    put(4, 1, "姓名", center=True)
    put(4, 2, val("legal_rep"))
    put(4, 3, "技术职称", center=True)
    put(4, 4, val("legal_rep_title"))
    put(4, 5, f"电话\n{val('legal_rep_phone')}", center=True)

    put(5, 0, "技术负责人", center=True)
    put(5, 1, "姓名", center=True)
    put(5, 2, val("technical_lead"))
    put(5, 3, "技术职称", center=True)
    put(5, 4, val("technical_lead_title"))
    put(5, 5, f"电话\n{val('technical_lead_phone')}", center=True)

    put(6, 0, "成立时间", center=True)
    merge(6, 1, 2)
    put(6, 1, val("established_date"))
    merge(6, 3, 5)
    put(6, 3, f"员工总人数：{val('employee_count')}", center=True)

    # Row 7-11: "其中" block for personnel counts.
    put(7, 0, "企业资质等级", center=True)
    merge(7, 1, 2)
    put(7, 1, val("qualification_level"))
    vmerge(7, 11, 3)
    put(7, 3, "其中", center=True)
    put(7, 4, "项目经理（或注册建造师）", center=True)
    put(7, 5, val("project_manager_name"))

    put(8, 0, "统一社会信用代码", center=True)
    merge(8, 1, 2)
    put(8, 1, val("business_license_no"))
    put(8, 4, "高级职称人员", center=True)
    put(8, 5, val("senior_title_count"))

    put(9, 0, "注册资本", center=True)
    merge(9, 1, 2)
    put(9, 1, val("registered_capital"))
    put(9, 4, "中级职称人员", center=True)
    put(9, 5, val("middle_title_count"))

    put(10, 0, "基本存款账户\n开户银行", center=True)
    merge(10, 1, 2)
    put(10, 1, val("bank_name"))
    put(10, 4, "初级职称人员", center=True)
    put(10, 5, val("junior_title_count"))

    put(11, 0, "基本存款账户\n银行账号", center=True)
    merge(11, 1, 2)
    put(11, 1, val("bank_account"))
    put(11, 4, "技工", center=True)
    put(11, 5, val("worker_count"))

    put(12, 0, "经营范围", center=True)
    merge(12, 1, 5)
    put(12, 1, val("business_scope"))

    put(13, 0, "投标人关联\n企业情况", center=True)
    merge(13, 1, 5)
    put(13, 1, "投标人应提供关联企业情况，包括：\n（1）投标人投资（控股）或管理的下属企业名称、持有股权（出资额）比例；\n（2）与投标人单位负责人（即法定代表人）为同一人的其他单位名称；\n（3）________")

    put(14, 0, "备注", center=True)
    merge(14, 1, 5)
    put(14, 1, "________")


def _strip_attr_quotes(value: str) -> str:
    if len(value) >= 2 and value[0] == value[-1] and value[0] in {"'", '"'}:
        return value[1:-1]
    return value


def _parse_marker_float(value: object, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _add_knowledge_image(
    document: Document,
    marker: dict[str, object],
    image_resolver: ImageResolver | None,
    style_profile: str | None,
) -> None:
    document_id = int(marker["document_id"])
    caption = str(marker.get("caption") or f"知识库图片资料 {document_id}")
    image_source = image_resolver(document_id) if image_resolver else None
    if not image_source:
        paragraph = document.add_paragraph(f"（图片资料未能插入：{caption}）")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    image_stream = (
        BytesIO(bytes(image_source))
        if isinstance(image_source, (bytes, bytearray))
        else str(image_source)
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 正奇风格的 Normal 是固定 32 磅行距，整页高的扫描件会向上溢出
    # 盖住前文；图片段落必须用单倍行距，并把高度限制在版心内。
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run()
    picture = run.add_picture(image_stream, width=Cm(float(marker["width_cm"])))
    max_height = Cm(20)
    if picture.height > max_height:
        scale = max_height / picture.height
        picture.width = int(picture.width * scale)
        picture.height = int(max_height)
    if caption:
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        _style_paragraph_runs(
            caption_paragraph,
            BODY_CJK_FONT,
            12 if style_profile == ZHENGQI_PROFILE else BODY_SIZE_PT,
        )


def _body_size_pt(style_profile: str | None) -> float:
    return 14 if style_profile == ZHENGQI_PROFILE else BODY_SIZE_PT


def _set_fonts(style, cjk_font: str) -> None:
    """Set Latin + East-Asian fonts on a style (python-docx needs raw XML)."""
    style.font.name = LATIN_FONT
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), cjk_font)


def _set_run_font(run, cjk_font: str, size_pt: float | None = None) -> None:
    run.font.name = LATIN_FONT
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), cjk_font)


def _style_paragraph_runs(
    paragraph,
    cjk_font: str,
    size_pt: float | None = None,
) -> None:
    for run in paragraph.runs:
        _set_run_font(run, cjk_font, size_pt)


def _add_text_with_underlined_blanks(
    paragraph,
    text: str,
    cjk_font: str,
    size_pt: float | None = None,
) -> None:
    """Render ________ placeholders as actual underlined fill-in space."""
    parts = re.split(r"([_＿]{3,})", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(" " * max(4, len(part)) if re.fullmatch(r"[_＿]{3,}", part) else part)
        _set_run_font(run, cjk_font, size_pt)
        if re.fullmatch(r"[_＿]{3,}", part):
            run.underline = True


def _apply_form_paragraph_alignment(paragraph) -> None:
    text = (paragraph.text or "").strip()
    if not text:
        return
    right_markers = (
        "投 标 人：",
        "投标人：",
        "法定代表人：",
        "法定代表人或其委托代理人：",
        "代理人：",
        "开立人：",
        "日 期：",
        "日期：",
        "年 月 日",
        "年        月        日",
    )
    if any(text.startswith(marker) for marker in right_markers) or any(
        marker in text for marker in ("（盖单位章）", "（签字或盖章）", "（签字）")
    ):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.first_line_indent = Pt(0)


def _configure_styles(document: Document, style_profile: str | None) -> None:
    is_zhengqi = style_profile == ZHENGQI_PROFILE
    normal = document.styles["Normal"]
    _set_fonts(normal, BODY_CJK_FONT)
    normal.font.size = Pt(_body_size_pt(style_profile))
    normal.font.bold = False
    if is_zhengqi:
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        normal.paragraph_format.line_spacing = Pt(32)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(3)
    else:
        normal.paragraph_format.line_spacing = 1.5

    heading_sizes = (
        {"Heading 1": 18, "Heading 2": 16, "Heading 3": 14}
        if is_zhengqi
        else HEADING_SIZES_PT
    )
    for style_name, size in heading_sizes.items():
        style = document.styles[style_name]
        _set_fonts(style, BODY_CJK_FONT if is_zhengqi else HEADING_CJK_FONT)
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(
            0, 0, 0
        )  # avoid the default blue heading colour
        style.paragraph_format.space_before = Pt(18 if is_zhengqi else 6)
        style.paragraph_format.space_after = Pt(12 if is_zhengqi else 6)
        if is_zhengqi and style_name == "Heading 1":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_zhengqi:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for style_name in ("List Bullet", "List Number"):
        try:
            _set_fonts(document.styles[style_name], BODY_CJK_FONT)
            document.styles[style_name].font.size = Pt(_body_size_pt(style_profile))
        except KeyError:
            continue

    for section in document.sections:
        if is_zhengqi:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(1.8)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(1.8)
        else:
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)


def _configure_header_footer(
    document: Document,
    header_text: str | None,
    page_numbers: bool,
    style_profile: str | None,
) -> None:
    is_zhengqi = style_profile == ZHENGQI_PROFILE
    for section in document.sections:
        if header_text:
            header_paragraph = section.header.paragraphs[0]
            header_paragraph.text = header_text
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _style_paragraph_runs(
                header_paragraph, BODY_CJK_FONT, 12 if is_zhengqi else None
            )

        if page_numbers:
            footer_paragraph = section.footer.paragraphs[0]
            footer_paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if is_zhengqi else WD_ALIGN_PARAGRAPH.CENTER
            )
            footer_paragraph.text = "第" if is_zhengqi else "第 "
            _append_field(footer_paragraph, "PAGE", placeholder="1")
            footer_paragraph.add_run("页/共" if is_zhengqi else " 页 / 共 ")
            _append_field(footer_paragraph, "NUMPAGES", placeholder="1")
            footer_paragraph.add_run("页" if is_zhengqi else " 页")
            _style_paragraph_runs(
                footer_paragraph,
                FOOTER_CJK_FONT if is_zhengqi else BODY_CJK_FONT,
                10.5 if is_zhengqi else None,
            )


def _add_cover_page(
    document: Document,
    title: str,
    subtitle: str | None,
    metadata: dict[str, str] | None,
    style_profile: str | None,
) -> None:
    is_zhengqi = style_profile == ZHENGQI_PROFILE
    for _ in range(4):
        document.add_paragraph()

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    _set_run_font(title_run, BODY_CJK_FONT, 36 if is_zhengqi else 28)

    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        _set_run_font(subtitle_run, BODY_CJK_FONT, 16)

    if metadata:
        for _ in range(2):
            document.add_paragraph()
        for key, value in metadata.items():
            row = document.add_paragraph()
            row.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = row.add_run(f"{key}：{value}")
            _set_run_font(run, BODY_CJK_FONT, 14 if is_zhengqi else 12)

    document.add_page_break()


def _add_table_of_contents(document: Document) -> None:
    heading = document.add_heading("目录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    _append_field(
        paragraph,
        instruction='TOC \\o "1-3" \\h \\z \\u',
        placeholder="（目录将在 Word 打开时自动生成；若未生成请右键“更新域”。）",
    )
    document.add_page_break()


def _enable_field_updates(document: Document) -> None:
    """Ask Word to recalculate all fields (TOC/PAGE/NUMPAGES) when opening."""
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _append_field(
    paragraph,
    instruction: str,
    placeholder: str | None = None,
):
    """Append a Word field (e.g. PAGE/TOC) to a paragraph.

    Each fldChar/instrText lives in its own run — packing them into a single
    run is invalid OOXML and breaks field evaluation in WPS/preview renderers.
    """

    def _flagged_run(tag: str, attr: str | None = None) -> None:
        run = paragraph.add_run()
        element = OxmlElement(tag)
        if attr:
            element.set(qn("w:fldCharType"), attr)
        else:
            element.set(qn("xml:space"), "preserve")
            element.text = f" {instruction} "
        run._r.append(element)

    _flagged_run("w:fldChar", "begin")
    _flagged_run("w:instrText")
    _flagged_run("w:fldChar", "separate")
    if placeholder:
        paragraph.add_run(placeholder)
    _flagged_run("w:fldChar", "end")


def _extract_title(markdown_text: str) -> str | None:
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                return title
    return None


def _is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    first = lines[index].strip()
    second = lines[index + 1].strip()
    return first.startswith("|") and second.startswith("|") and "---" in second


def _add_table(
    document: Document,
    table_lines: list[str],
    style_profile: str | None,
) -> None:
    is_zhengqi = style_profile == ZHENGQI_PROFILE
    rows = [_split_table_row(line) for line in table_lines]
    rows = [row for row in rows if row and not all(set(cell) <= {"-"} for cell in row)]
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            _add_text_with_underlined_blanks(
                paragraph,
                value,
                BODY_CJK_FONT,
                14 if is_zhengqi else BODY_SIZE_PT,
            )

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


VOLUME_MARKERS = {
    "commercial": "<!-- tdg:volume:commercial -->",
    "technical": "<!-- tdg:volume:technical -->",
    "pricing": "<!-- tdg:volume:pricing -->",
    "notes": "<!-- tdg:volume:notes -->",
}
_VOLUME_MARKER_PATTERN = re.compile(
    r"<!--\s*tdg:volume:(commercial|technical|pricing|notes)\s*-->"
)
_VOLUME_TITLE_LABELS = {
    "commercial": "商务文件",
    "technical": "技术文件",
    "pricing": "报价文件",
}
_META_NOTE_HEADINGS = ("审查修正说明", "人工修正意见")


def combine_delivery_volumes(
    doc_title: str,
    volumes: dict[str, str],
    notes: str = "",
) -> str:
    """Combine per-volume markdown into one document with explicit markers.

    Volume order follows the delivery contract: commercial, technical, pricing.
    The markers let ``split_delivery_markdown`` recover the exact volumes later
    instead of guessing by heading keywords. ``notes`` holds review/correction
    meta text that belongs to no volume and is stripped before export.
    """
    parts: list[str] = []
    if doc_title:
        parts.append(f"# {doc_title}")
        parts.append("")
    for label in ("commercial", "technical", "pricing"):
        body = (volumes.get(label) or "").strip()
        if not body:
            continue
        parts.append(VOLUME_MARKERS[label])
        parts.append("")
        parts.append(body)
        parts.append("")
    if notes.strip():
        parts.append(VOLUME_MARKERS["notes"])
        parts.append("")
        parts.append(notes.strip())
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def _split_by_markers(markdown_text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current: list[str] | None = None
    for line in markdown_text.splitlines():
        match = _VOLUME_MARKER_PATTERN.match(line.strip())
        if match:
            current = sections.setdefault(match.group(1), [])
            continue
        if current is not None:
            current.append(line)
    return {label: "\n".join(body).strip() for label, body in sections.items()}


def split_delivery_markdown(markdown_text: str) -> dict[str, str]:
    """Split bid markdown into 商务文件 / 技术文件 / 报价文件 volumes.

    Documents produced by ``combine_delivery_volumes`` carry explicit volume
    markers and split losslessly; the meta ``notes`` section is excluded from
    every volume. Legacy documents without markers fall back to the heading
    keyword heuristic.
    """
    doc_title = _extract_title(markdown_text)

    if _VOLUME_MARKER_PATTERN.search(markdown_text):
        marked = _split_by_markers(markdown_text)
        volumes: dict[str, str] = {}
        for label in ("commercial", "technical", "pricing"):
            text = marked.get(label, "")
            title_label = _VOLUME_TITLE_LABELS[label]
            if text and text.lstrip().startswith("# "):
                volumes[label] = text
            else:
                prefix = (
                    f"# {doc_title}（{title_label}）\n\n"
                    if doc_title
                    else f"# {title_label}\n\n"
                )
                volumes[label] = prefix + (
                    text or "（本卷暂无自动归类内容，请人工补充。）"
                )
        return volumes

    lines = markdown_text.splitlines()
    technical: list[str] = []
    commercial: list[str] = []
    pricing: list[str] = []
    current = technical
    seen_section = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            continue
        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip()
            seen_section = True
            if any(keyword in heading for keyword in _META_NOTE_HEADINGS):
                current = []  # meta notes belong to no volume
            elif any(keyword in heading for keyword in PRICING_KEYWORDS):
                current = pricing
            elif any(keyword in heading for keyword in COMMERCIAL_KEYWORDS):
                current = commercial
            else:
                current = technical
        elif not seen_section:
            current = technical
        current.append(line)

    volumes = {}
    for label, body in (
        ("commercial", commercial),
        ("technical", technical),
        ("pricing", pricing),
    ):
        text = "\n".join(body).strip()
        title_label = _VOLUME_TITLE_LABELS[label]
        prefix = (
            f"# {doc_title}（{title_label}）\n\n" if doc_title else f"# {title_label}\n\n"
        )
        volumes[label] = prefix + (text or "（本卷暂无自动归类内容，请人工补充。）")
    return volumes


def strip_meta_notes(markdown_text: str) -> str:
    """Remove review/correction meta sections and volume markers for export.

    Drops the marked ``notes`` section, legacy ``## 审查修正说明`` /
    ``## 人工修正意见`` heading blocks, and all ``tdg:volume`` marker comments,
    so internal workflow annotations never reach the delivered document.
    """
    kept: list[str] = []
    skipping_notes = False
    skipping_heading = False
    for line in markdown_text.splitlines():
        stripped = line.strip()
        match = _VOLUME_MARKER_PATTERN.match(stripped)
        if match:
            skipping_notes = match.group(1) == "notes"
            skipping_heading = False
            continue
        if skipping_notes:
            continue
        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip()
            skipping_heading = any(
                keyword in heading for keyword in _META_NOTE_HEADINGS
            )
            if skipping_heading:
                continue
        if skipping_heading:
            continue
        kept.append(line)
    return "\n".join(kept).strip() + "\n"


_SAFE_NAME_PATTERN = re.compile(r"[^\w.\-一-鿿]+", flags=re.UNICODE)


def build_export_filename(
    project_name: str,
    version: int = 1,
    kind: str | None = None,
    suffix: str = "docx",
) -> str:
    """Build a download filename containing the project name and version.

    Example: ``高层住宅项目_技术标_v2.docx``.
    """
    base = _SAFE_NAME_PATTERN.sub("_", (project_name or "投标文件").strip())
    base = base.strip("._") or "投标文件"
    parts = [base]
    if kind:
        parts.append(kind)
    parts.append(f"v{max(int(version), 1)}")
    return f"{'_'.join(parts)}.{suffix}"
