"""Gap evaluation between an AI-generated bid and a real bid reference (M63).

Given an AI-generated document (DOCX or Markdown) and a reference structure
(a real bid parsed into a :class:`BidTemplate`, or the template JSON), this
module reports the structural gap: missing main sections, missing construction
sub-sections, missing 施工附表, missing fixed forms, content-length difference
and manual-confirmation-point statistics. It produces a Markdown/JSON report
used as a regression metric after prompt/generator changes.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from schemas.bid_template import BidTemplate

# Markers indicating a spot the bid leaves for human confirmation / completion.
CONFIRM_MARKERS = ("人工确认", "待补充", "待确认", "待填写", "【人工", "占位")

_PREFIX_RE = re.compile(
    r"^(?:第[一二三四五六七八九十百零〇]+[章节]、?|"
    r"[一二三四五六七八九十]+、|"
    r"附表[一二三四五六七八九十]+、?|"
    r"（[一二三四五六七八九十]+）)"
)


def _normalize(text: Any) -> str:
    return re.sub(r"\s+", "", str(text or "")).lower()


def section_key(title: str) -> str:
    """Reduce a section title to its matchable core (drop numbering prefixes)."""
    core = _PREFIX_RE.sub("", str(title or "").strip())
    return _normalize(core)


def extract_docx_structure(docx_path: str | Path) -> dict[str, Any]:
    from docx import Document

    document = Document(str(docx_path))
    headings: list[str] = []
    section_lengths: dict[str, int] = {}
    manual_points: list[str] = []
    current: str | None = None
    total = 0
    content_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name and style_name.startswith("Heading"):
            headings.append(text)
            current = text
            section_lengths.setdefault(text, 0)
            continue
        content_parts.append(text)
        total += len(text)
        if current is not None:
            section_lengths[current] = section_lengths.get(current, 0) + len(text)
        if any(marker in text for marker in CONFIRM_MARKERS):
            manual_points.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                total += len(cell_text)
                content_parts.append(cell_text)

    return {
        "sections": headings,
        "section_lengths": section_lengths,
        "total_chars": total,
        "manual_confirmation_points": manual_points,
        "text": " ".join(headings + content_parts),
    }


def _split_present_missing(
    sections: list[Any],
    blob: str,
) -> tuple[list[str], list[str]]:
    present: list[str] = []
    missing: list[str] = []
    for section in sections:
        title = section.title if hasattr(section, "title") else str(section)
        key = section_key(title)
        if key and key in blob:
            present.append(title)
        else:
            missing.append(title)
    return present, missing


def _coverage(present: list[str], missing: list[str]) -> float:
    total = len(present) + len(missing)
    if total == 0:
        return 1.0
    return round(len(present) / total, 4)

