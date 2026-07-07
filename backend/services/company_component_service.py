"""公司组件库:员工做好的"整块成品"(组织结构框图/项目管理机构图…),生成时原样照搬。

业务规则(用户拍板 2026-07-07):这类表碰到了**绝不让 AI 自己画/填**,直接把员工提供的
docx 里的成品表格(文本框画的架构图)一模一样搬进商务卷对应位置。

存储:documents(project_id IS NULL, document_category='公司组件'),原件 docx 存 MinIO;
metadata_json: component_type(组件名)、anchors(锚点关键词列表,商务卷里命中即替换)。

搬运保真关键(实测):组件文本框内文字**不带显式字体**(rFonts 只有 hint),全靠继承组件
文档的默认字体(宋体/Times New Roman);深拷贝进宿主后改为继承宿主默认 → LibreOffice
渲染成豆腐块。搬运时把组件 docDefaults 的字体**实化**写进每个 run,才真正一模一样。
"""

from __future__ import annotations

import io
import logging
import re
from copy import deepcopy
from typing import Any

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

CATEGORY = "公司组件"

# 内置组件的锚点规则:组件类型 → (宿主标题/表格特征关键词)。命中标题段之后的首张表,
# 或首格文字命中的表,即为要替换的空框。
_BUILTIN_ANCHORS: dict[str, list[str]] = {
    "组织结构框图": ["组织结构框图", "组织机构框图", "以框图方式表示"],
    "项目管理机构": ["项目管理机构", "拟为承包本标段"],
}


def _connect():
    from rag.vector_store import _connect as c

    return c()


def list_components() -> list[dict[str, Any]]:
    """全部公司组件(公司级),附 document_id/file_path/anchors。"""
    with _connect() as conn, conn.cursor() as cursor:
        cursor.execute(
            """
            SELECT id, file_name, file_path, metadata_json
            FROM documents
            WHERE project_id IS NULL
              AND metadata_json->>'document_category' = %s
            ORDER BY id
            """,
            (CATEGORY,),
        )
        rows = cursor.fetchall()
    out = []
    for doc_id, file_name, file_path, meta in rows:
        meta = meta or {}
        out.append(
            {
                "document_id": int(doc_id),
                "file_name": file_name,
                "file_path": file_path,
                "component_type": meta.get("component_type") or "",
                "anchors": meta.get("anchors") or [],
            }
        )
    return out


def _component_docx_bytes(file_path: str) -> bytes:
    from core.config import settings
    from utils.minio_client import minio_client

    return minio_client.download_bytes(settings.minio_bucket, file_path)


def _solidify_fonts(el: Any, ascii_font: str, ea_font: str) -> None:
    """把组件里没写显式字体的 run 实化成组件默认字体(跨文档搬运后不再依赖宿主默认)。"""
    for r in el.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            from docx.oxml import OxmlElement

            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            from docx.oxml import OxmlElement

            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        if not fonts.get(qn("w:ascii")):
            fonts.set(qn("w:ascii"), ascii_font)
            fonts.set(qn("w:hAnsi"), ascii_font)
        if not fonts.get(qn("w:eastAsia")):
            fonts.set(qn("w:eastAsia"), ea_font)


def _default_fonts_of(doc: Document) -> tuple[str, str]:
    """组件文档 docDefaults 的 (西文, 中文) 字体;读不到用标书常规。"""
    try:
        import html

        from lxml import etree

        xml = etree.tostring(doc.styles.element).decode()
        m = re.search(r"<w:docDefaults.*?</w:docDefaults>", xml, re.S)
        if m:
            am = re.search(r'w:ascii="([^"]+)"', m.group(0))
            em = re.search(r'w:eastAsia="([^"]+)"', m.group(0))
            # tostring 会把中文字体名转成 &#23435;&#20307; 实体,必须解码回"宋体"
            return (
                html.unescape(am.group(1)) if am else "Times New Roman",
                html.unescape(em.group(1)) if em else "宋体",
            )
    except Exception:
        pass
    return ("Times New Roman", "宋体")


def _norm_cell(text: str) -> str:
    return re.sub(r"[\s　]+", "", text or "")


def _host_tables_matching(
    document: Document, comp_first_cell: str, anchors: list[str]
):
    """宿主文档里该被组件替换的空框表。

    首选锚:**组件表首格与宿主表首格文字一致**(前12字)——员工的组件就是从招标模板
    抠出来画的,首格说明文字("拟为承包本标段…"/"以框图方式表示。")原样保留,天然唯一,
    不会张冠李戴(实测:两个组件的空框首格各不相同,标题反而会被插图题注顶掉)。
    兜底锚:标题段命中 anchors 且其后首张表是短空框。只认 ≤3 行的空框,绝不动正文大表。
    """
    from docx.table import Table

    comp_key = _norm_cell(comp_first_cell)[:12]
    body = document.element.body
    hits = []
    pending_title = False
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = "".join(child.itertext()).strip()
            if text and len(text) < 40 and any(a in text for a in anchors):
                pending_title = True
            continue
        if child.tag != qn("w:tbl"):
            continue
        table = Table(child, document._body)
        if len(table.rows) > 3:
            pending_title = False
            continue
        first_cell = _norm_cell(table.rows[0].cells[0].text) if table.rows else ""
        cell_hit = bool(comp_key) and first_cell.startswith(comp_key)
        if cell_hit or pending_title:
            hits.append(child)
        pending_title = False
    return hits


def fill_company_components(document: Document) -> dict[str, Any]:
    """生成时调用:把库里每个组件按锚点搬进宿主商务卷(整表替换,字体实化)。

    返回 {"replaced": 替换数, "handled_tables": set(新表元素)}——handled 供通用
    填表器绕行(成品组件绝不允许再被填值)。库为空/无命中时不动文档。
    """
    result: dict[str, Any] = {"replaced": 0, "handled_tables": set()}
    try:
        components = list_components()
    except Exception:
        logger.warning("公司组件库读取失败,跳过组件照搬", exc_info=True)
        return result
    for comp_meta in components:
        anchors = comp_meta.get("anchors") or _BUILTIN_ANCHORS.get(
            comp_meta.get("component_type") or "", []
        )
        try:
            comp_doc = Document(io.BytesIO(_component_docx_bytes(comp_meta["file_path"])))
        except Exception:
            logger.warning(
                "组件 %s 原件读取失败,跳过", comp_meta.get("file_name"), exc_info=True
            )
            continue
        if not comp_doc.tables:
            continue
        comp_first = comp_doc.tables[0].rows[0].cells[0].text if comp_doc.tables[0].rows else ""
        targets = _host_tables_matching(document, comp_first, anchors)
        if not targets:
            continue
        ascii_f, ea_f = _default_fonts_of(comp_doc)
        comp_tbl = comp_doc.tables[0]._tbl
        for target in targets:
            new_tbl = deepcopy(comp_tbl)
            _solidify_fonts(new_tbl, ascii_f, ea_f)
            target.addnext(new_tbl)
            target.getparent().remove(target)
            result["handled_tables"].add(new_tbl)
            result["replaced"] += 1
            logger.info(
                "公司组件[%s]已照搬进商务卷(锚点:%s)",
                comp_meta.get("component_type"), anchors[0],
            )
    return result


def import_component(
    file_bytes: bytes,
    filename: str,
    component_type: str,
    anchors: list[str] | None = None,
) -> dict[str, Any]:
    """入库一个公司组件(幂等:同 component_type 旧记录先删)。"""
    from services import knowledge_service
    from services.knowledge_service import delete_knowledge_document

    # 幂等:删同类型旧组件
    for old in list_components():
        if old["component_type"] == component_type:
            try:
                delete_knowledge_document(old["document_id"])
            except Exception:
                logger.warning("旧组件删除失败(继续导入)", exc_info=True)
    resolved_anchors = anchors or _BUILTIN_ANCHORS.get(component_type) or [component_type]
    indexed = knowledge_service.index_uploaded_knowledge(
        file_bytes=file_bytes,
        filename=filename,
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        document_type="公司组件",
        document_category=CATEGORY,
        volume="商务文件",
        owner_type="公司",
        usage_scope="可用于投标",
        verified_status="已核验",  # 员工定稿成品
        image_insertable=False,
        tags=["公司组件", component_type],
        ingestion_mode="evidence_only",  # 成品不进RAG索引,只存原件
        extra_metadata={"component_type": component_type, "anchors": resolved_anchors},
    )
    return {"document_id": indexed.get("document_id"), "component_type": component_type}
