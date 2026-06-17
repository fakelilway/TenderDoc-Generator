"""落盘 .docx 标书的确定性程序化体检（自动迭代闭环的"自检目标函数"）。

不调 LLM、不联网，纯靠 python-docx 读结构 + 正则规则，对一份已生成的标书
打 0-100 的质量分。给定同一份文档，分数完全可复现，可作为自迭代的目标函数。

六项加权（权重之和 = 100）：

    required_fields   20  必填字段完整度（投标人/法定代表人/证照号/安许/项目经理）
    table_fill        15  表格填充率（非空 cell 占比）
    section_coverage  20  章节齐全度（章节数 / 基准）
    residue           15  残留物 0 容忍（{{}}/图占位/Markdown/AI 元话语）
    font_consistency  10  字体/字号一致性（显式 font.name 偏离期望占比）
    length            20  篇幅与响应（字数饱和函数，设上限防注水）

硬否决（``hard_block=True`` 且 ``quality_score`` 封顶 40）：
    - 残留 ``{{...}}`` 模板标记
    - 未消费的 ``{{knowledge_image:...}}`` 图占位 / 【图表占位】/【图片占位】
    - AI 元话语（"作为...助手""根据您的要求""待补充"等）
    - 疑似编造的身份证号（\\d{17}[\\dXx]）出现在正文

命令行：
    python -m services.docx_health_check <docx路径>
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Any

from services.bid_gap_eval import extract_docx_structure

# --------------------------------------------------------------------------- #
# 残留标记规则集（"不该出现在成品标书里"的东西）
# --------------------------------------------------------------------------- #

# 双花括号模板标记 {{...}}（含 {{knowledge_image:...}}）—— 任何残留都是硬否决。
_CURLY_RE = re.compile(r"\{\{[^{}]*\}\}")
_KNOWLEDGE_IMAGE_RE = re.compile(r"\{\{\s*knowledge_image")

# 中文图/图表占位（未被生成器消费的占位符）。
_FIGURE_PLACEHOLDER_RE = re.compile(r"【图表占位|【图片占位")

# Markdown 残留：行首 1-6 个 # 后接空白，或管道表格行（以 | 起、含至少两个 |）。
_MD_HEADING_RE = re.compile(r"^#{1,6}\s")
_MD_PIPE_TABLE_RE = re.compile(r"^\s*\|.*\|")

# AI 元话语 / 占位词（成品标书里不该出现的"自我暴露"措辞）。
_META_PHRASES = (
    "待补充",
    "待填写",
    "待确认",
    "TODO",
    "占位",
    "placeholder",
    "人工确认",
    "AI生成",
    "ai生成",
    "由AI",
    "由ai",
    "根据您的要求",
    "我们建议",
    "以下是",
    "希望以上",
)

# AI 自指措辞需精确匹配:必须是"作为…"后紧跟 AI/助手/语言模型 这类词,
# 否则会误伤正常工程文本(如"每个井段作为一个流水段""作为一名项目经理")。
# 旧版用裸子串"作为一个/作为一名"把 7 万字的真实标书误判成 AI 生成并硬否决。
_AI_SELF_REF_RE = re.compile(
    r"作为(?:一个|一名|您的)?\s*(?:AI|ai|人工智能|智能助手|语言模型|大模型)"
)

# 身份证号：18 位（前 17 位数字 + 末位数字或 X）。出现在正文即疑似编造。
_ID_CARD_RE = re.compile(r"(?<!\d)\d{17}[\dXx](?!\d)")

# 期望字体（用于一致性检查）。
_BODY_FONT = "SimSun"  # 宋体
_HEADING_FONT = "SimHei"  # 黑体
_CERT_FONT = "Times New Roman"  # 证号/英文数字
_ACCEPTABLE_FONTS = {
    _BODY_FONT,
    _HEADING_FONT,
    _CERT_FONT,
    "宋体",
    "黑体",
    "仿宋",
    "FangSong",
    "FangSong_GB2312",
    "Microsoft YaHei",
    "微软雅黑",
    "STSong",
    "DengXian",
    "等线",
    "Arial",
    "Calibri",
}

# 空 cell 判定:复用填表器的统一占位谓词(下划线/省略号/点线签字线都算"空"),
# 避免体检口径与填表口径漂移导致"填了体检仍报空"或假阴性"健康"。
from services.original_docx_format_service import _is_blank_or_placeholder

# 必填字段的字段名 -> 表格 label 关键词列表。
_REQUIRED_FIELD_LABELS: dict[str, tuple[str, ...]] = {
    "investor": ("投标人", "投标单位", "申请人", "公司名称", "企业名称"),
    "legal_rep": ("法定代表人", "法人代表"),
    "credit_code": ("统一社会信用代码", "营业执照", "信用代码", "注册号"),
    "safety_license": ("安全生产许可证", "安全许可证", "安许"),
    "project_manager": ("项目经理", "项目负责人", "拟派项目经理"),
}

# CompanyProfile 字段名 -> 必填字段 key（有 profile 时按真值核对）。
_PROFILE_FIELD_MAP: dict[str, str] = {
    "company_name": "investor",
    "legal_representative": "legal_rep",
    "credit_code": "credit_code",
    "safety_license_no": "safety_license",
    "project_manager_name": "project_manager",
}

_WEIGHTS = {
    "required_fields": 20.0,
    "table_fill": 15.0,
    "section_coverage": 20.0,
    "residue": 15.0,
    "font_consistency": 10.0,
    "length": 20.0,
}

# 章节齐全度启发式基准：>= 该数量的主章节算满分。
_SECTION_FULL_MARK = 8

# 篇幅启发式：字数 >= 该值算满分（不给 baseline 时）。
_LENGTH_FULL_MARK = 15000

# 硬否决时的质量分上限。
_HARD_BLOCK_CAP = 40.0


# --------------------------------------------------------------------------- #
# 读 docx：把段落与表格 cell 全部抽出来供规则扫描
# --------------------------------------------------------------------------- #


def _is_blank_cell(text: str) -> bool:
    """空 cell:完全空白,或只剩下划线/省略号/点线等占位(与填表器同口径)。"""
    return _is_blank_or_placeholder(text or "")


def _collect(document) -> dict[str, Any]:
    """一次遍历，收集后续所有规则需要的原始物料。"""
    para_texts: list[str] = []
    for paragraph in document.paragraphs:
        para_texts.append(paragraph.text)

    cells: list[str] = []  # 所有表格 cell 文本
    blank_cells = 0
    total_cells = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                cells.append(text)
                total_cells += 1
                if _is_blank_cell(text):
                    blank_cells += 1

    full_text = "\n".join(para_texts + cells)
    compact = re.sub(r"\s+", "", full_text)
    return {
        "para_texts": para_texts,
        "cells": cells,
        "blank_cells": blank_cells,
        "total_cells": total_cells,
        "full_text": full_text,
        "compact_chars": len(compact),
    }


# --------------------------------------------------------------------------- #
# 各项打分（每项返回 (score[0-1], detail, hard_block_bool, note|None)）
# --------------------------------------------------------------------------- #


def _score_required_fields(
    collected: dict[str, Any], profile: Any | None
) -> tuple[float, str, list[str]]:
    """必填字段完整度。

    有 profile：按 CompanyProfile 真值逐字段核对该字段值是否出现在文档里。
    无 profile：在表格里找含字段 label 的行，看其相邻 cell 是否仍是下划线空位。
    """
    notes: list[str] = []
    cells = collected["cells"]
    full_text = collected["full_text"]

    if profile is not None:
        # 把 profile 转成 dict（兼容 pydantic 模型 / dict）。
        if hasattr(profile, "model_dump"):
            pdata = profile.model_dump()
        elif isinstance(profile, dict):
            pdata = profile
        else:
            pdata = {k: getattr(profile, k, "") for k in _PROFILE_FIELD_MAP}
        filled = 0
        checked = 0
        missing: list[str] = []
        for field, key in _PROFILE_FIELD_MAP.items():
            value = str(pdata.get(field, "") or "").strip()
            if not value:
                continue  # profile 自己就没这个值，不要求文档体现
            checked += 1
            if value in full_text:
                filled += 1
            else:
                missing.append(key)
        if checked == 0:
            return 1.0, "profile 无可核对字段，默认满分", ["required_fields: profile 字段全空"]
        score = filled / checked
        detail = f"profile 核对 {filled}/{checked} 字段已落入文档"
        if missing:
            detail += f"；缺失 {missing}"
        return score, detail, notes

    # 无 profile：基于表格行的 label -> 相邻空位检测。
    # 把表格按行重建：找出含某个 label 的行，看该行是否还有非空值 cell。
    present = 0
    blank = 0
    found_keys: set[str] = set()
    for key, labels in _REQUIRED_FIELD_LABELS.items():
        # 在所有 cell 文本里找 label cell，再看同行是否有值。
        hit_label = False
        hit_value = False
        for idx, text in enumerate(cells):
            if any(lbl in text for lbl in labels):
                hit_label = True
                # 紧邻的下一个 cell 通常是值列（粗略相邻判定）。
                nxt = cells[idx + 1] if idx + 1 < len(cells) else ""
                # 同 cell 内 "label：值" 形式：只有真的含冒号才取冒号后内容。
                inline_value = ""
                if re.search(r"[:：]", text):
                    inline_value = re.sub(r"^.*?[:：]", "", text).strip()
                # 相邻 cell（值列）有非下划线内容也算填了。
                if (inline_value and not _is_blank_cell(inline_value)) or (
                    nxt and not _is_blank_cell(nxt)
                ):
                    hit_value = True
                    break
        if hit_label:
            found_keys.add(key)
            if hit_value:
                present += 1
            else:
                blank += 1

    if not found_keys:
        notes.append("required_fields: 文档表格未出现任何已知必填字段 label，无法核对，给保守满分")
        return 1.0, "未发现必填字段表格行，无法核对（保守满分）", notes
    score = present / len(found_keys)
    detail = f"无 profile：{present}/{len(found_keys)} 个必填字段表格行已填（其余仍为下划线空位）"
    return score, detail, notes


def _score_table_fill(collected: dict[str, Any]) -> tuple[float, str, list[str]]:
    total = collected["total_cells"]
    blank = collected["blank_cells"]
    if total == 0:
        return 1.0, "文档无表格，表格填充率不适用（保守满分）", ["table_fill: 无表格"]
    fill = 1.0 - blank / total
    return fill, f"表格 cell {total - blank}/{total} 非空（填充率 {fill:.0%}）", []


def _score_section_coverage(
    structure: dict[str, Any], requirements: Any | None
) -> tuple[float, str, list[str]]:
    sections = structure.get("sections", []) or []
    n = len(sections)

    # 有 requirements：用 format_outline_tree 的顶层节点数作基准。
    baseline = None
    if requirements is not None:
        tree = None
        if hasattr(requirements, "format_outline_tree"):
            tree = requirements.format_outline_tree
        elif isinstance(requirements, dict):
            tree = requirements.get("format_outline_tree")
        if tree:
            baseline = sum(len(v or []) for v in tree.values())

    if baseline and baseline > 0:
        score = min(1.0, n / baseline)
        return score, f"章节 {n}/{baseline}（基准取自 requirements 格式目录）", []

    score = min(1.0, n / _SECTION_FULL_MARK)
    return (
        score,
        f"章节 {n} 个（启发式基准 >= {_SECTION_FULL_MARK} 满分）",
        [],
    )


def _score_residue(
    collected: dict[str, Any], structure: dict[str, Any]
) -> tuple[float, str, bool, list[str]]:
    """残留物 0 容忍。命中硬类残留 -> 0 分 + hard_block。"""
    full_text = collected["full_text"]
    para_texts = collected["para_texts"]
    hits: list[str] = []
    hard = False

    curly = _CURLY_RE.findall(full_text)
    if curly:
        hard = True
        sample = curly[0][:40]
        if _KNOWLEDGE_IMAGE_RE.search(full_text):
            hits.append(f"未消费图占位 {{knowledge_image}}（如 {sample}）")
        else:
            hits.append(f"残留模板标记 {{}}（如 {sample}）")

    if _FIGURE_PLACEHOLDER_RE.search(full_text):
        hard = True
        hits.append("残留【图表/图片占位】")

    # Markdown 残留只看段落（表格 cell 里出现 | 是正常表格文本边界）。
    md_hits = 0
    for line in para_texts:
        stripped = line.strip()
        if _MD_HEADING_RE.match(stripped) or _MD_PIPE_TABLE_RE.match(stripped):
            md_hits += 1
    if md_hits:
        hits.append(f"Markdown 残留 {md_hits} 行（行首 # 或管道表格）")

    meta_found = [p for p in _META_PHRASES if p in full_text]
    if _AI_SELF_REF_RE.search(full_text):
        meta_found.append("作为…AI/助手(自指)")
    if meta_found:
        hard = True
        hits.append(f"AI 元话语/占位词 {meta_found[:5]}")

    id_cards = _ID_CARD_RE.findall(full_text)
    if id_cards:
        hard = True
        hits.append(f"疑似编造身份证号 {len(id_cards)} 处")

    if hard:
        return 0.0, "；".join(hits), True, []
    if md_hits:
        # 非硬类的 Markdown 残留：扣到 0.5。
        return 0.5, "；".join(hits), False, []
    return 1.0, "未发现残留物", False, []


def _score_font_consistency(document) -> tuple[float, str, list[str]]:
    """字体一致性：只统计显式设了 font.name 的 run，看偏离期望字体的占比。

    font.name 为 None（继承样式）的 run 不计入，避免把"继承默认"误判为缺失。
    无可判定的显式字体 run 时，给保守满分并 note 说明。
    """
    explicit = 0
    deviated = 0
    bad_samples: list[str] = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            name = run.font.name
            if not name:
                continue
            explicit += 1
            if name not in _ACCEPTABLE_FONTS:
                deviated += 1
                if name not in bad_samples:
                    bad_samples.append(name)
    # 表格里的 run 也扫一遍。
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        name = run.font.name
                        if not name:
                            continue
                        explicit += 1
                        if name not in _ACCEPTABLE_FONTS:
                            deviated += 1
                            if name not in bad_samples:
                                bad_samples.append(name)

    if explicit == 0:
        return (
            1.0,
            "无显式 font.name 的 run（全部继承样式），无法判定，给保守满分",
            ["font_consistency: 无显式字体，保守满分"],
        )
    score = 1.0 - deviated / explicit
    detail = f"显式字体 run {explicit} 个，偏离期望 {deviated} 个（{score:.0%} 一致）"
    if bad_samples:
        detail += f"；偏离字体样例 {bad_samples[:5]}"
    return score, detail, []


def _score_length(
    collected: dict[str, Any], baseline_chars: int | None
) -> tuple[float, str, list[str]]:
    chars = collected["compact_chars"]
    if baseline_chars and baseline_chars > 0:
        score = min(1.0, chars / (0.7 * baseline_chars))
        return (
            score,
            f"字数 {chars}（基线 {baseline_chars}，目标 70%={int(0.7 * baseline_chars)}，饱和封顶 1.0）",
            [],
        )
    # 分段饱和：>= _LENGTH_FULL_MARK 满分，线性折算，设上限防注水。
    score = min(1.0, chars / _LENGTH_FULL_MARK)
    return (
        score,
        f"字数 {chars}（启发式 >= {_LENGTH_FULL_MARK} 满分，饱和封顶 1.0）",
        [],
    )


# --------------------------------------------------------------------------- #
# 主入口
# --------------------------------------------------------------------------- #


def score_docx(
    path: str | Path,
    *,
    requirements: Any | None = None,
    profile: Any | None = None,
    baseline_chars: int | None = None,
) -> dict[str, Any]:
    """对一份落盘 .docx 标书做确定性体检，返回 0-100 质量分及明细。

    Parameters
    ----------
    path:
        .docx 文件路径。
    requirements:
        可选，TenderRequirements 或其 dict 形式；用其 ``format_outline_tree``
        顶层节点数作章节齐全度基准。不给则用启发式。
    profile:
        可选，CompanyProfile 或其 dict 形式；按真值核对必填字段是否落入文档。
        不给则按"表格行是否仍是下划线空位"判。
    baseline_chars:
        可选，对标真实标书的压缩字数；给则用饱和函数算篇幅项。

    Returns
    -------
    dict: ``{"quality_score", "items", "hard_block", "notes"}``
    """
    from docx import Document

    document = Document(str(path))
    collected = _collect(document)
    structure = extract_docx_structure(path)

    notes: list[str] = []
    items: dict[str, dict[str, Any]] = {}

    rf_score, rf_detail, rf_notes = _score_required_fields(collected, profile)
    notes.extend(rf_notes)
    items["required_fields"] = {
        "score": round(rf_score, 4),
        "weight": _WEIGHTS["required_fields"],
        "detail": rf_detail,
    }

    tf_score, tf_detail, tf_notes = _score_table_fill(collected)
    notes.extend(tf_notes)
    items["table_fill"] = {
        "score": round(tf_score, 4),
        "weight": _WEIGHTS["table_fill"],
        "detail": tf_detail,
    }

    sc_score, sc_detail, sc_notes = _score_section_coverage(structure, requirements)
    notes.extend(sc_notes)
    items["section_coverage"] = {
        "score": round(sc_score, 4),
        "weight": _WEIGHTS["section_coverage"],
        "detail": sc_detail,
    }

    rs_score, rs_detail, rs_hard, rs_notes = _score_residue(collected, structure)
    notes.extend(rs_notes)
    items["residue"] = {
        "score": round(rs_score, 4),
        "weight": _WEIGHTS["residue"],
        "detail": rs_detail,
    }

    fc_score, fc_detail, fc_notes = _score_font_consistency(document)
    notes.extend(fc_notes)
    items["font_consistency"] = {
        "score": round(fc_score, 4),
        "weight": _WEIGHTS["font_consistency"],
        "detail": fc_detail,
    }

    ln_score, ln_detail, ln_notes = _score_length(collected, baseline_chars)
    notes.extend(ln_notes)
    items["length"] = {
        "score": round(ln_score, 4),
        "weight": _WEIGHTS["length"],
        "detail": ln_detail,
    }

    weighted = sum(it["score"] * it["weight"] for it in items.values())
    quality_score = round(weighted, 2)  # 权重和=100,所以已是 0-100

    hard_block = bool(rs_hard)
    if hard_block:
        if quality_score > _HARD_BLOCK_CAP:
            quality_score = _HARD_BLOCK_CAP
        notes.append(f"硬否决触发：{rs_detail}（quality_score 封顶 {_HARD_BLOCK_CAP}）")

    return {
        "quality_score": quality_score,
        "items": items,
        "hard_block": hard_block,
        "notes": notes,
    }


def score_delivery(
    *,
    technical_path: str | Path | None = None,
    commercial_path: str | Path | None = None,
    requirements: Any | None = None,
    profile: Any | None = None,
    baseline_chars: int | None = None,
) -> dict[str, Any]:
    """按卷打分整套投标交付件(技术卷 + 商务卷),修正单卷打分的分母错位。

    分卷交付下(见 generation_service._assemble_two_volumes):投标人基本情况表里
    的必填字段在【商务卷】,施工组织设计正文在【技术卷】,而对外的 bid.docx 是技术卷
    副本。拿技术卷去量商务字段必然假阴性(required_fields 恒低)。本函数按卷取物料:
      - required_fields → 商务卷(缺商务卷则退回技术卷)
      - table_fill      → 两卷表格合并统计
      - section_coverage / length → 技术卷
      - residue         → 两卷取严(任一卷硬否决即硬否决)
      - font_consistency→ 仅技术卷(商务卷为原招标格式照抄,字体非本系统排版,不计)
    至少需提供 technical_path 或 commercial_path 其一。返回结构同 score_docx。
    """
    from docx import Document

    if technical_path is None and commercial_path is None:
        raise ValueError("score_delivery 需要 technical_path 或 commercial_path 至少其一")

    tech_doc = Document(str(technical_path)) if technical_path else None
    comm_doc = Document(str(commercial_path)) if commercial_path else None
    tech_collected = _collect(tech_doc) if tech_doc is not None else None
    comm_collected = _collect(comm_doc) if comm_doc is not None else None
    tech_structure = (
        extract_docx_structure(technical_path) if technical_path else {"sections": []}
    )
    comm_structure = (
        extract_docx_structure(commercial_path) if commercial_path else {"sections": []}
    )

    notes: list[str] = []
    items: dict[str, dict[str, Any]] = {}

    # required_fields → 商务卷优先(基本情况表在商务卷)
    rf_collected = comm_collected or tech_collected
    rf_vol = "商务卷" if comm_collected is not None else "技术卷(退化)"
    rf_score, rf_detail, rf_notes = _score_required_fields(rf_collected, profile)
    notes.extend(rf_notes)
    items["required_fields"] = {
        "score": round(rf_score, 4),
        "weight": _WEIGHTS["required_fields"],
        "detail": f"[{rf_vol}] {rf_detail}",
    }

    # table_fill → 两卷表格合并统计
    merged = {
        "total_cells": (tech_collected["total_cells"] if tech_collected else 0)
        + (comm_collected["total_cells"] if comm_collected else 0),
        "blank_cells": (tech_collected["blank_cells"] if tech_collected else 0)
        + (comm_collected["blank_cells"] if comm_collected else 0),
    }
    tf_score, tf_detail, tf_notes = _score_table_fill(merged)
    notes.extend(tf_notes)
    items["table_fill"] = {
        "score": round(tf_score, 4),
        "weight": _WEIGHTS["table_fill"],
        "detail": f"[两卷合并] {tf_detail}",
    }

    # section_coverage / length → 技术卷(正文所在)
    sc_score, sc_detail, sc_notes = _score_section_coverage(tech_structure, requirements)
    notes.extend(sc_notes)
    items["section_coverage"] = {
        "score": round(sc_score, 4),
        "weight": _WEIGHTS["section_coverage"],
        "detail": f"[技术卷] {sc_detail}",
    }

    ln_collected = tech_collected or comm_collected
    ln_score, ln_detail, ln_notes = _score_length(ln_collected, baseline_chars)
    notes.extend(ln_notes)
    items["length"] = {
        "score": round(ln_score, 4),
        "weight": _WEIGHTS["length"],
        "detail": f"[技术卷] {ln_detail}",
    }

    # residue → 两卷取严
    residue_parts: list[tuple[str, tuple[float, str, bool, list[str]]]] = []
    if tech_collected is not None:
        residue_parts.append(("技术卷", _score_residue(tech_collected, tech_structure)))
    if comm_collected is not None:
        residue_parts.append(("商务卷", _score_residue(comm_collected, comm_structure)))
    rs_hard = any(part[1][2] for part in residue_parts)
    rs_score = min(part[1][0] for part in residue_parts)
    rs_detail = "；".join(f"[{name}] {part[1]}" for name, part in residue_parts)
    for _, part in residue_parts:
        notes.extend(part[3])
    items["residue"] = {
        "score": round(rs_score, 4),
        "weight": _WEIGHTS["residue"],
        "detail": rs_detail,
    }

    # font_consistency → 仅技术卷(商务卷照抄原格式,不以本系统字体规范衡量)
    fc_doc = tech_doc if tech_doc is not None else comm_doc
    fc_vol = "技术卷" if tech_doc is not None else "商务卷(退化)"
    fc_score, fc_detail, fc_notes = _score_font_consistency(fc_doc)
    notes.extend(fc_notes)
    items["font_consistency"] = {
        "score": round(fc_score, 4),
        "weight": _WEIGHTS["font_consistency"],
        "detail": f"[{fc_vol}] {fc_detail}",
    }

    weighted = sum(it["score"] * it["weight"] for it in items.values())
    quality_score = round(weighted, 2)

    hard_block = bool(rs_hard)
    if hard_block:
        if quality_score > _HARD_BLOCK_CAP:
            quality_score = _HARD_BLOCK_CAP
        notes.append(f"硬否决触发:{rs_detail}(quality_score 封顶 {_HARD_BLOCK_CAP})")

    return {
        "quality_score": quality_score,
        "items": items,
        "hard_block": hard_block,
        "notes": notes,
    }


def _main(argv: list[str]) -> int:
    if not argv:
        print("用法: python -m services.docx_health_check <docx路径>", file=sys.stderr)
        return 2
    target = Path(argv[0])
    if not target.exists():
        print(f"文件不存在: {target}", file=sys.stderr)
        return 1
    result = score_docx(target)
    print(f"quality_score = {result['quality_score']}  hard_block = {result['hard_block']}")
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(_main(sys.argv[1:]))
