from __future__ import annotations

import logging
import re
from pathlib import Path
from tempfile import TemporaryDirectory

from psycopg2.extras import Json

from core.config import settings
from schemas.tender import TenderRequirements
from services.company_profile_service import get_company_profile
from services.original_docx_format_service import build_original_format_docx
from services.project_service import _connect
from utils.docx_exporter import markdown_to_docx, strip_meta_notes
from utils.minio_client import minio_client


logger = logging.getLogger(__name__)

PLACEHOLDER_WORDS = ("待补充", "TODO", "占位", "placeholder")


def export_markdown_for_project(
    project_id: int,
    markdown: str,
    quality_report: dict[str, float | int],
    *,
    original_format_path: str | None = None,
    appendix_format_path: str | None = None,
) -> tuple[str, str]:
    # Keep volume markers intact for splitting; strip meta/markers only for the
    # human-readable bid.md and the non-format whole-doc render. Stripping before
    # the split would drop the tdg:volume markers → split falls back to a heading
    # heuristic that leaks commercial sections into the technical volume.
    clean_markdown = strip_meta_notes(markdown)
    title = _extract_markdown_title(clean_markdown) or "投标文件"
    with TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        markdown_path = tmp_path / f"project_{project_id}_bid.md"
        markdown_path.write_text(clean_markdown, encoding="utf-8")
        docx_path = tmp_path / f"project_{project_id}_bid.docx"

        if original_format_path and Path(original_format_path).exists():
            # 两卷装配：商务=整本格式章照抄(整页图+填空字段，copy2 保留图片),
            # 技术=LLM 正文独立成文，报价=外部造价软件不产出。
            # 不再按页块三卷拆分——格式章即商务卷，硬拆会错分页面、丢图片。
            _assemble_two_volumes(
                original_format_path,
                tmp_path,
                project_id,
                markdown,
                docx_path,
                title,
                appendix_format_path=appendix_format_path,
            )
        elif not _try_export_original_docx_format(project_id, docx_path):
            markdown_to_docx(
                clean_markdown,
                docx_path,
                title=title,
                subtitle="投标文件",
                cover=True,
                toc=True,
                header_text=title,
                page_numbers=True,
                style_profile="zhengqi",
                image_resolver=_resolve_knowledge_image,
            )

        markdown_object = f"projects/{project_id}/generated/bid.md"
        docx_object = f"projects/{project_id}/generated/bid.docx"
        minio_client.upload_file(settings.minio_bucket, markdown_path, markdown_object)
        minio_client.upload_file(settings.minio_bucket, docx_path, docx_object)

        # Upload three independent volume DOCX files
        for vol in ("commercial", "technical", "pricing"):
            vol_path = tmp_path / f"project_{project_id}_{vol}.docx"
            if vol_path.exists():
                minio_client.upload_file(
                    settings.minio_bucket,
                    vol_path,
                    f"projects/{project_id}/generated/{vol}.docx",
                )

        # 非阻断质量打分钩子：直接用临时目录里已生成的卷(别再去 MinIO 重下)按卷
        # 体检并落 eval_results/。严格非阻断——打分失败绝不影响出标。
        _score_delivery_quality(project_id, tmp_path)

    _update_generation_paths(
        project_id,
        markdown_object,
        docx_object,
        quality_report,
    )
    return markdown_object, docx_object


def _score_delivery_quality(project_id: int, tmp_path: Path) -> None:
    """非阻断地给刚出的标书按卷打分(quality_score + hard_block)。

    两卷模式才有 commercial/technical 两个文件;非两卷模式只有 bid.docx,这时
    technical_path 传 bid.docx、commercial_path=None。整段 try/except 包裹,
    出标流程绝不能因打分崩。
    """
    try:
        from services.delivery_quality import score_delivery_files

        commercial = tmp_path / f"project_{project_id}_commercial.docx"
        technical = tmp_path / f"project_{project_id}_technical.docx"
        bid = tmp_path / f"project_{project_id}_bid.docx"

        commercial_path = commercial if commercial.exists() else None
        if technical.exists():
            technical_path = technical
        elif bid.exists():
            technical_path = bid
        else:
            technical_path = None

        if commercial_path is None and technical_path is None:
            logger.warning("质量打分跳过：项目 %s 未找到可打分的卷文件", project_id)
            return

        result = score_delivery_files(
            commercial_path=commercial_path,
            technical_path=technical_path,
            project_id=project_id,
        )
        quality_score = result.get("quality_score")
        hard_block = result.get("hard_block")
        logger.info(
            "标书质量打分 project=%s quality_score=%s hard_block=%s",
            project_id,
            quality_score,
            hard_block,
        )
        if hard_block or (isinstance(quality_score, (int, float)) and quality_score < 60):
            logger.warning(
                "标书质量偏低 project=%s quality_score=%s hard_block=%s notes=%s"
                "（仅告警，不阻断出标）",
                project_id,
                quality_score,
                hard_block,
                result.get("notes"),
            )
    except Exception:
        logger.exception("标书质量打分失败 project=%s（非阻断，出标继续）", project_id)


def _try_export_original_docx_format(project_id: int, docx_path: Path) -> bool:
    try:
        tender = _fetch_tender_document(project_id)
    except Exception:
        logger.exception("Tender document lookup unavailable; using markdown export")
        return False
    if not tender:
        return False
    filename = str(tender.get("file_name") or "")
    object_name = str(tender.get("file_path") or "")
    if not filename.lower().endswith(".docx"):
        return False
    try:
        tender_bytes = minio_client.download_bytes(settings.minio_bucket, object_name)
        profile = _export_profile_from_tender(tender)
        build_original_format_docx(tender_bytes, docx_path, profile=profile)
        return True
    except Exception:
        logger.exception("Original DOCX format export failed")
        raise ValueError("DOCX 招标文件原格式复制失败，系统不会回退生成近似格式文件。")


def _fetch_tender_document(project_id: int) -> dict[str, object] | None:
    with _connect() as conn:
        with conn.cursor() as cursor:
            cursor.execute(
                """
                SELECT d.file_name, d.file_path, p.name, p.confirmed_parsed_json, p.parsed_json
                FROM projects p
                LEFT JOIN documents d
                    ON d.project_id = p.id AND d.file_path = p.tender_file_path
                WHERE p.id = %s
                ORDER BY d.id DESC
                LIMIT 1
                """,
                (project_id,),
            )
            row = cursor.fetchone()
    if not row:
        return None
    if isinstance(row, dict):
        return row
    keys = ("file_name", "file_path", "name", "confirmed_parsed_json", "parsed_json")
    return dict(zip(keys, row))


def _export_profile_from_tender(tender: dict[str, object]) -> dict[str, object]:
    parsed = tender.get("confirmed_parsed_json") or tender.get("parsed_json") or {}
    profile: dict[str, object] = {}
    if isinstance(parsed, dict):
        profile.update(
            {
                "project_name": parsed.get("project_name") or tender.get("name") or "",
                "项目名称": parsed.get("project_name") or tender.get("name") or "",
                "tenderer_name": parsed.get("tenderer_name") or "",
                "招标人": parsed.get("tenderer_name") or "",
                "工期": parsed.get("planned_duration") or "",
                "质量": parsed.get("quality_standard") or "",
                "安全": parsed.get("safety_target") or "",
                "投标有效期": parsed.get("bid_deadline") or "",
                "投标截止时间": parsed.get("bid_deadline") or "",
            }
        )
    try:
        company_profile = get_company_profile().get("profile", {})
        if isinstance(company_profile, dict):
            profile.update(company_profile)
    except Exception:
        logger.warning("Company profile unavailable during original DOCX export")
    return profile


def evaluate_generation_quality(markdown_text: str) -> dict[str, float | int]:
    paragraphs = [
        line.strip()
        for line in markdown_text.splitlines()
        if line.strip()
        and not line.lstrip().startswith("#")
        and not _is_markdown_table_control_line(line)
        and not line.strip().startswith("{{knowledge_image:")
    ]
    total = len(paragraphs)
    needs_revision = 0
    for paragraph in paragraphs:
        lower = paragraph.lower()
        if len(paragraph) < 20 or any(
            word.lower() in lower for word in PLACEHOLDER_WORDS
        ):
            needs_revision += 1

    usable = max(total - needs_revision, 0)
    usable_rate = usable / total if total else 0.0
    return {
        "total_paragraphs": total,
        "needs_revision_paragraphs": needs_revision,
        "usable_paragraphs": usable,
        "usable_rate": round(usable_rate, 4),
    }


def _update_generation_paths(
    project_id: int,
    markdown_path: str,
    docx_path: str,
    quality_report: dict[str, float | int],
) -> None:
    with _connect() as conn:
        with conn.cursor() as cursor:
            cursor.execute(
                """
                UPDATE projects
                SET
                    generated_markdown_path = %s,
                    generated_docx_path = %s,
                    generation_quality_json = %s,
                    status = %s
                WHERE id = %s
                """,
                (
                    markdown_path,
                    docx_path,
                    Json(quality_report),
                    "generated",
                    project_id,
                ),
            )


def _extract_markdown_title(markdown_text: str) -> str | None:
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                return title
    return None


def _image_reference_query(requirements: TenderRequirements) -> str:
    descriptions = [
        item.description
        for item in [
            *requirements.qualification_list,
            *requirements.technical_score_items,
            *requirements.invalid_bid_items,
        ]
    ]
    return (
        f"{requirements.project_name} 营业执照 资质证书 安全生产许可证 "
        "建造师 身份证 建安证 交安证 职称证 社保 业绩 施工平面图 " + " ".join(descriptions)
    )


def _resolve_knowledge_image(document_id: int) -> bytes | None:
    from services import knowledge_service

    try:
        return knowledge_service.get_knowledge_document_file_bytes(document_id)
    except Exception:
        logger.exception(
            "Failed to resolve knowledge image bytes for document %s", document_id
        )
        return None


def _is_markdown_table_control_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return all(cell and set(cell) <= {"-", ":"} for cell in cells)


# 正文渲染器(markdown_to_docx)依赖的命名段落样式。
_PROSE_BASE_STYLES = ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number")


def _ensure_named_styles(doc) -> None:
    """Make the paragraph styles the markdown renderer relies on fetchable.

    Some converted DOCX (notably 福昕云 PDF→Word output) carry builtin heading/
    list styles whose raw OOXML ``w:name`` doesn't match python-docx's name
    lookup, so ``styles["Heading 1"]`` / ``add_heading`` / ``style="List Bullet"``
    all raise ``KeyError: no style with name 'Heading 1'`` even though the name
    appears in ``iter(styles)``. We can still *create* a real, fetchable style for
    each one (``add_style`` succeeds), which repairs the lookup for both
    ``_configure_styles`` and ``add_heading`` — without this, the entire two-volume
    export aborts here and no downloadable DOCX is produced.
    """
    from docx.enum.style import WD_STYLE_TYPE

    for name in _PROSE_BASE_STYLES:
        try:
            doc.styles[name]
        except KeyError:
            try:
                doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            except Exception:  # noqa: BLE001 - best effort; plaintext fallback covers it
                logger.warning("无法为 '%s' 补建命名样式;合规正文将走纯文本兜底", name)


def _append_prose_plaintext(docx_path: Path, prose_markdown: str) -> None:
    """Last-resort prose append that never touches named styles.

    Used when the styled append fails (e.g. a converted base docx whose styles
    can't be repaired). Keeps the合规正文 content (unstyled) instead of dropping
    it or失败 → 出标不被阻断、下载件照常产出。

    **仍会解析 {{knowledge_image:...}} 插图**:证件/业绩扫描是资格审查的硬内容,绝不能因为
    走了兜底就退化成一行死 token 文字(用户实测"插图没插"正是旧版样式渲染崩→兜底吞图所致)。
    """
    from docx import Document

    from utils.docx_exporter import _add_knowledge_image, _parse_knowledge_image_marker

    doc = Document(str(docx_path))
    doc.add_page_break()
    for line in prose_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        marker = _parse_knowledge_image_marker(stripped)
        if marker is not None:
            try:
                _add_knowledge_image(doc, marker, _resolve_knowledge_image, "zhengqi")
            except Exception:
                logger.warning(
                    "纯文本兜底插图失败 doc=%s", marker.get("document_id"), exc_info=True
                )
            continue
        text = stripped.lstrip("#").strip()
        if text:
            doc.add_paragraph(text)
    doc.save(str(docx_path))


# ── 附件按招标锚点就地落位 ───────────────────────────────────────────────
# 招标逐表要求"在本表后附相关证明材料"(营业执照→基本情况表后、法人身份证→身份证明后、
# 业绩→类似项目情况表后)。证据函数给 {{knowledge_image}} 标记带 anchor 关键词,这里在照抄
# 进来的格式章 DOCX 里找含该关键词的章节,把图就地插在其末尾;找不到则保留标记交原卷尾追加
# (兜底,绝不丢图)。详见记忆 attachment-placement-rules。
# anchor → (表头关键词, 段落关键词, 段落排除词)。优先匹配"含表头关键词的表格"(表格不在目录里,
# 绕开目录误命中);否则匹配"含段落关键词且不含排除词"的段落。关键词都选目录里不会出现、章节/
# 表格独有的字样(表头'投标人名称'/'投标人情况说明'、附件行'法定代表人身份证复印件')。
# 覆盖多家招标的不同措辞(萧县/长丰…):同一锚点给多个候选关键词,往通用靠。
_ANCHOR_SPECS: dict[str, tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...]]] = {
    "基本情况表": (("投标人名称",), ("投标人基本情况表",), ()),
    "类似项目情况表": (
        ("业绩序号",),
        ("近年完成的类似", "类似工程项目情况表", "类似项目情况表", "投标人业绩情况表"),
        (),
    ),
    "信誉情况表": (
        ("投标人情况说明",),
        ("投标人的信誉情况", "投标人信誉情况", "信誉情况表"),
        (),
    ),
    # 法人亲签:插在"法定代表人身份证"附件行后(复印件/正反面扫描件…),排除带"代理人"的授权委托书那份。
    "法定代表人身份证明": (
        (),
        ("法定代表人身份证复印件", "法定代表人身份证正反面", "法定代表人身份证扫描"),
        ("委托代理人", "代理人身份证"),
    ),
    # 项目经理资历表后附其证件(建造师证/身份证/职称/社保);找不到该表则走卷尾兜底。
    "项目经理资历表": (
        (),
        ("项目经理资历", "项目经理简历", "拟派项目经理", "项目负责人简历", "项目经理身份证"),
        (),
    ),
    "项目总工资历表": (
        (),
        ("项目技术负责人", "项目总工", "总工程师简历", "技术负责人简历"),
        (),
    ),
}
_SECTION_HEAD_RE = re.compile(
    r"^[（(]?[一二三四五六七八九十][)）]?[、.，]|^[（(][一二三四五六七八九十][)）]"
)


def _iter_body_blocks(doc):
    """按文档体顺序产出 (xml元素, 段落或表格对象)。"""
    from docx.oxml.ns import qn as _qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body.iterchildren():
        if child.tag == _qn("w:p"):
            yield child, Paragraph(child, doc)
        elif child.tag == _qn("w:tbl"):
            yield child, Table(child, doc)


def _block_text(block) -> str:
    from docx.table import Table

    if isinstance(block, Table):
        return " ".join(c.text for r in block.rows[:3] for c in r.cells)
    return block.text


def _anchor_section_end_element(doc, anchor: str):
    """定位 anchor 的插入点(在其后插图);找不到返回 None。

    先按表头关键词找表格(表格不在目录,绕开目录误命中)→ 紧接该表之后;否则按段落关键词
    (排除词过滤)找段落 → 走到该节末尾元素之后(跨过表单/表格,停在下一节标题前)。
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    tbl_kw, para_kw, excl = _ANCHOR_SPECS.get(anchor, ((), (anchor,), ()))
    blocks = list(_iter_body_blocks(doc))

    # 1) 表格优先:含表头关键词的表 → 紧接该表之后
    for el, blk in blocks:
        if isinstance(blk, Table) and tbl_kw and any(k in _block_text(blk) for k in tbl_kw):
            return el

    # 2) 段落:含关键词、不含排除词 → 该节末尾元素之后
    start = None
    for idx, (_el, blk) in enumerate(blocks):
        if not isinstance(blk, Paragraph):
            continue
        t = blk.text
        if not (para_kw and any(k in t for k in para_kw)):
            continue
        if excl and any(x in t for x in excl):
            continue
        start = idx
        break
    if start is None:
        return None
    end = start
    for idx in range(start + 1, min(start + 30, len(blocks))):
        _el, blk = blocks[idx]
        if isinstance(blk, Paragraph):
            t = blk.text.strip()
            if t and len(t) < 30 and _SECTION_HEAD_RE.match(t):
                break  # 撞到下一节标题,停在它之前
        end = idx
    return blocks[end][0]


def _insert_image_after(after_el, doc, document_id, caption, width_cm):
    """在 after_el 之后插入"图 + 图注"两段,返回最后插入的元素(供同锚点链式续插)。"""
    from io import BytesIO

    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.shared import Cm
    from docx.text.paragraph import Paragraph

    img = _resolve_knowledge_image(int(document_id))
    img_el = OxmlElement("w:p")
    after_el.addnext(img_el)
    img_p = Paragraph(img_el, doc)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if not img:
        img_p.add_run(f"（图片资料未能插入，请人工补充：{caption}）")
        return img_el
    try:
        from utils.image_orient import upright_image_bytes

        pic = img_p.add_run().add_picture(
            BytesIO(upright_image_bytes(bytes(img))), width=Cm(float(width_cm))
        )
        if pic.height > Cm(20):
            scale = Cm(20) / pic.height
            pic.width = int(pic.width * scale)
            pic.height = int(Cm(20))
    except Exception:
        img_p.add_run(f"（图片资料未能插入，请人工补充：{caption}）")
        return img_el
    last = img_el
    if caption:
        cap_el = OxmlElement("w:p")
        img_el.addnext(cap_el)
        cap_p = Paragraph(cap_el, doc)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap_p.add_run(caption)
        last = cap_el
    return last


def _drop_empty_headings(lines: list[str]) -> str:
    """丢掉后面没有任何内容的空标题(锚点图被移走后残留的 ## 附录 / ### 营业执照)。"""
    out: list[str] = []
    for i, line in enumerate(lines):
        if line.strip().startswith("#"):
            has_content = False
            for j in range(i + 1, len(lines)):
                t = lines[j].strip()
                if not t:
                    continue
                if t.startswith("#"):
                    break
                has_content = True
                break
            if not has_content:
                continue
        out.append(line)
    return "\n".join(out)


# 用户上传到资料库的"图/表"docx,生成时复制进标书对应章节。
# (锚点章节关键词, 资料库文件名关键词)。
# 用户定:(一)项目管理机构组织机构图 只放图3(项目级,文件名含"项目管理机构"=8_项目管理机构.docx),
# 锚点用"组织机构图"精确命中(一)那一行(别用宽词"项目管理机构"否则会落到(二)人员组成表)。
# 图4(投标人企业组织结构框图=公司级)不再自动注入(用户明确不要它出现在(一))。
_ORG_DOC_SOURCES: tuple[tuple[tuple[str, ...], str], ...] = (
    (("组织机构图",), "项目管理机构"),
    (("拟分包", "分包"), "拟分包"),
)


def _fetch_org_source_docx(filename_kw: str) -> bytes | None:
    """从资料库取公司级 .docx(文件名含关键词);取不到返回 None。"""
    try:
        from rag.vector_store import get_db_connection
        from utils.minio_client import minio_client
        from core.config import settings

        with get_db_connection() as conn, conn.cursor() as cursor:
            cursor.execute(
                "SELECT file_path FROM documents WHERE project_id IS NULL "
                "AND lower(file_type) = 'docx' AND file_name ILIKE %s "
                "ORDER BY created_at DESC, id DESC LIMIT 1",
                (f"%{filename_kw}%",),
            )
            row = cursor.fetchone()
        if not row or not row[0]:
            return None
        return minio_client.download_bytes(settings.minio_bucket, str(row[0]))
    except Exception:
        return None


def _find_section_paragraph(doc, keywords: tuple[str, ...]):
    """找含章节关键词的**正文章节标题**段落,返回其 <w:p> 元素;找不到返回 None。

    跳过目录区(目录里的"五、项目管理机构"等条目不是插图位);目录条目在前、正文章节在后,
    取最后一个匹配段即正文真章节,避免把表/图插进目录。
    """
    in_toc = False
    last_match = None
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        # 进入/退出目录区:"目录"标题起,到首个实质正文(投标函/法定代表人等)止
        if len(t) < 8 and ("目录" in t or "目 录" in t):
            in_toc = True
            continue
        if in_toc:
            # 目录条目多为"X、xxx"短行;遇到较长正文或"投标函"等正文章节起 → 出目录
            if len(t) > 24 or "投标函" in t:
                in_toc = False
            else:
                continue  # 仍在目录,条目不作锚点
        if len(t) < 40 and any(k in t for k in keywords):
            last_match = para._p
    return last_match


def _inject_org_tables(doc) -> int:
    """把资料库里的项目管理机构表/拟分包表/组织结构图(docx)复制进标书对应章节后。

    复制 docx 正文里的**表格**与**含图段落**(跳过纯文字噪声);有锚点章节插其后,无则插卷尾。
    逐源容错,绝不因此崩或丢内容。返回插入的块数。
    """
    from copy import deepcopy
    from io import BytesIO

    from docx import Document as _Doc
    from docx.oxml.ns import qn

    inserted = 0
    for anchors, fkw in _ORG_DOC_SOURCES:
        blob = _fetch_org_source_docx(fkw)
        if not blob:
            continue
        try:
            src = _Doc(BytesIO(blob))
        except Exception:
            continue
        # 取源 docx 正文里的表格 + 含图段落(<w:drawing>/<w:pict>)
        blocks = []
        for child in src.element.body:
            if child.tag == qn("w:tbl"):
                blocks.append(child)
            elif child.tag == qn("w:p") and (
                child.findall(".//" + qn("w:drawing")) or child.findall(".//" + qn("w:pict"))
            ):
                blocks.append(child)
        if not blocks:
            continue

        target_el = _find_section_paragraph(doc, anchors)
        try:
            if target_el is not None:
                for blk in blocks:
                    new_blk = deepcopy(blk)
                    target_el.addnext(new_blk)
                    target_el = new_blk
            else:
                for blk in blocks:
                    doc.element.body.append(deepcopy(blk))
            inserted += len(blocks)
        except Exception:
            logger.warning("注入资料库图表失败(%s),跳过不阻断", fkw, exc_info=True)
    if inserted:
        logger.info("资料库图表注入:%d 块(项目管理机构/拟分包/组织结构)", inserted)
    return inserted


def _place_anchored_images(doc, prose_markdown: str) -> str:
    """把 prose 里带 anchor 的 {{knowledge_image}} 就地插到格式章对应章节后;返回去掉这些标记
    (及随之变空的标题)后的 prose——剩余无锚点图仍由调用方追加到卷尾。逐图容错,绝不丢图/不崩。
    """
    from utils.docx_exporter import _parse_knowledge_image_marker

    kept: list[str] = []
    cursor: dict[str, object] = {}  # anchor -> 该锚点最后插入的元素(链式续插)
    placed = 0
    for line in prose_markdown.splitlines():
        marker = _parse_knowledge_image_marker(line.strip())
        if not (marker and marker.get("anchor")):
            kept.append(line)
            continue
        anchor = str(marker["anchor"])
        try:
            if anchor not in cursor:
                el = _anchor_section_end_element(doc, anchor)
                if el is None:
                    kept.append(line)  # 锚点没找到 → 留给卷尾兜底
                    continue
                cursor[anchor] = el
            cursor[anchor] = _insert_image_after(
                cursor[anchor], doc, marker["document_id"],
                marker.get("caption", ""), marker.get("width_cm", 14.0),
            )
            placed += 1
        except Exception:
            logger.warning("附件锚点落位单图失败,该图退回卷尾(不丢图)", exc_info=True)
            kept.append(line)
    if placed:
        logger.info("附件锚点落位:就地插入 %d 张(其余无锚点图走卷尾兜底)", placed)
    return _drop_empty_headings(kept)


def _append_prose_to_docx(docx_path: Path, prose_markdown: str) -> None:
    """Append prose content after format pages in a DOCX.

    Uses the full markdown_to_docx renderer which correctly handles
    headings (H1–H3), markdown tables, tdg:pagebreak markers,
    underlined blanks, and the zhengqi style profile (SimSun 14pt,
    SimHei headings, 32pt line spacing).

    若按样式渲染失败(如福昕转换件的内置样式 python-docx 取不到),退到纯文本
    追加,绝不让商务卷整卷导出失败 → 保证有下载件。
    """
    if not prose_markdown.strip():
        return

    from utils.docx_exporter import _render_markdown_body, _configure_styles
    from docx import Document

    try:
        doc = Document(str(docx_path))

        # 福昕等转换器产出的 docx,其内置标题/列表样式 python-docx 按名取不到
        # (styles[name]/add_heading 抛 KeyError),先补成可取的真样式,否则
        # _configure_styles 与正文渲染会在此让整卷导出失败 → 无下载件。
        _ensure_named_styles(doc)

        # _configure_styles sets zhengqi margins on ALL sections; that would clobber
        # the full-bleed (0-margin) geometry of any format-page image sections, making
        # the page images overflow/clip → blank pages in LibreOffice/Pages. Snapshot
        # existing sections and restore their geometry after styling.
        geom = [
            (
                s.page_width, s.page_height,
                s.left_margin, s.right_margin, s.top_margin, s.bottom_margin,
                s.header_distance, s.footer_distance,
            )
            for s in doc.sections
        ]
        _configure_styles(doc, "zhengqi")
        for s, g in zip(doc.sections, geom):
            (s.page_width, s.page_height,
             s.left_margin, s.right_margin, s.top_margin, s.bottom_margin,
             s.header_distance, s.footer_distance) = g

        # 带 anchor 的附件先就地落位到对应章节后(营业执照→基本情况表后…),返回去锚点后的
        # 剩余 prose;无锚点/锚点没命中的仍走下面的卷尾追加(兜底,绝不丢图)。
        prose_markdown = _place_anchored_images(doc, prose_markdown)
        # 资料库里的项目管理机构表/拟分包表/组织结构图(docx)就地复制进对应章节
        _inject_org_tables(doc)

        doc.add_page_break()
        # 传 image_resolver,让附录里的 {{knowledge_image:...}} 标记从 MinIO 取图插入(B)
        _render_markdown_body(doc, prose_markdown, "zhengqi", _resolve_knowledge_image)
        doc.save(str(docx_path))
    except Exception:
        logger.warning(
            "商务卷合规正文按样式追加失败,改用纯文本兜底(不丢内容、不阻断出标)",
            exc_info=True,
        )
        _append_prose_plaintext(docx_path, prose_markdown)


def _assemble_two_volumes(
    format_path: str,
    tmp_path: Path,
    project_id: int,
    markdown: str,
    main_docx_path: Path,
    title: str,
    *,
    appendix_format_path: str | None = None,
) -> None:
    """Two-volume delivery for editable format docs (pdf2docx / DOCX 照抄).

    - 商务卷 = the converted format chapter, copied verbatim (照抄) with known
      fields already filled, then appended compliance prose if any.
    - 技术卷 = the LLM-written prose as its OWN zhengqi-styled document (cover +
      TOC), never glued onto the commercial format pages.
    - 报价卷 = done externally in造价软件; this system does NOT produce it.

    The main bid.docx mirrors the technical volume (the part with written prose).
    """
    import shutil

    from utils.docx_exporter import split_delivery_markdown

    volumes = split_delivery_markdown(markdown)
    commercial_markdown = volumes.get("commercial", "")
    technical_markdown = volumes.get("technical", "") or markdown

    # 商务卷：照抄格式章 + 合规正文
    commercial_path = tmp_path / f"project_{project_id}_commercial.docx"
    shutil.copy2(format_path, commercial_path)
    if commercial_markdown.strip():
        _append_prose_to_docx(commercial_path, commercial_markdown)

    # 技术卷：独立成文的施工组织设计正文
    technical_path = tmp_path / f"project_{project_id}_technical.docx"
    if technical_markdown.strip():
        markdown_to_docx(
            technical_markdown,
            technical_path,
            title=title,
            subtitle="技术文件",
            cover=True,
            toc=True,
            header_text=title,
            page_numbers=True,
            style_profile="zhengqi",
            image_resolver=_resolve_knowledge_image,
        )
    else:
        shutil.copy2(format_path, technical_path)

    # 附表:把福昕转的可编辑附表(附表一~八空表)拼到技术卷末尾(另起页)。
    # best-effort——拼接失败不影响技术卷正文。
    if appendix_format_path and Path(appendix_format_path).exists():
        try:
            _append_docx(technical_path, appendix_format_path)
        except Exception:
            logger.warning("附表拼接到技术卷失败,技术卷不含附表", exc_info=True)

    shutil.copy2(technical_path, main_docx_path)


def _append_docx(base_path: Path, appendix_path: str) -> None:
    """把 appendix docx 内容拼到 base docx 末尾(另起页)。

    用 docxcompose 合并,保留附表的真实表格与图片关系(优于手动 deepcopy 元素)。
    **先写临时文件并校验(能打开 + 正文不少于合并前),通过才原子替换 base** —— docxcompose
    遇 style/numbering 冲突可能产出"打开需修复"的损坏 docx,这样坏了也不污染技术卷正文
    (校验失败抛异常,由调用方丢弃附表、保留纯技术卷)。
    """
    import os

    from docx import Document
    from docxcompose.composer import Composer

    master = Document(str(base_path))
    base_paragraphs = sum(1 for p in master.paragraphs if p.text.strip())
    master.add_page_break()
    composer = Composer(master)
    composer.append(Document(str(appendix_path)))

    tmp_out = f"{base_path}.merged.docx"
    composer.save(tmp_out)
    # 校验:合并产物必须能打开,且正文段落不少于合并前(防 docxcompose 吞内容/产损坏件)
    merged = Document(tmp_out)
    if sum(1 for p in merged.paragraphs if p.text.strip()) < base_paragraphs:
        os.remove(tmp_out)
        raise RuntimeError("附表合并产物异常(正文丢失),丢弃附表保留纯技术卷")
    os.replace(tmp_out, str(base_path))

