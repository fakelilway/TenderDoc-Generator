"""福昕 Foxit 国内 Cloud API：PDF → 可编辑 Word(docx)。

格式复制链最上层的"云转换"档。设计契约与 original_docx_format_service 里其它
builder 一致:**成功返回保存好的 docx 路径,任何失败 raise**——由 v2_generation_service
Phase 0 的 try/except 阶梯自动回退到 pdf2docx → 整页图 → 纯图。

纯 HTTP(httpx,已是依赖),不装任何 SDK,绕开"装包顶翻 httpx/numpy"的依赖坑。

鉴权(逐字照搬官方 demo foxitsoftware/ServicesAPIDemo):
  sn = md5( urlencode(sorted(本次查询参数,不含sn)) + "&sk=" + secret )
异步三段式:document/convert(taskId) → 轮询 task(percentage==100 → docId) → download。
"""

from __future__ import annotations

import hashlib
import logging
import time
import urllib.parse

import httpx

logger = logging.getLogger(__name__)

FOXIT_API_BASE = "https://servicesapi.foxitsoftware.cn/api"


def _sn(client_id: str, secret: str, params: dict[str, str]) -> str:
    """按福昕算法对一次请求的查询参数(不含 sn)签名。"""
    query = urllib.parse.urlencode(dict(sorted(params.items())))
    query += "&sk=" + urllib.parse.quote(secret)
    return hashlib.md5(query.encode("utf-8")).hexdigest()


def convert_pdf_to_docx_via_foxit(
    pdf_bytes: bytes,
    output_path: str,
    *,
    client_id: str,
    secret: str,
    config_json: str = '{"pageRange": "all"}',
    poll_interval_seconds: float = 2.0,
    max_wait_seconds: float = 300.0,
    request_timeout_seconds: float = 120.0,
) -> str:
    """把 ``pdf_bytes`` 转成可编辑 docx 存到 ``output_path``,返回该路径。失败抛异常。

    ``config_json`` 是福昕 convert 的配置串(签名要含它);PDF→Word 用 pageRange,
    无边框表可在此加 ML 识别选项。
    """
    if not client_id or not secret:
        raise RuntimeError("福昕凭证缺失(FOXIT_CLOUD_CLIENT_ID / FOXIT_CLOUD_SECRET)")

    # follow_redirects: download 会 302 跳到 pheeplatform 下载域,必须跟随。
    with httpx.Client(timeout=request_timeout_seconds, follow_redirects=True) as client:
        # ① 提交转换任务(POST /api/document/convert;format/config 在 body,clientId/sn 在 query;
        #    签名覆盖 {clientId, config, format})
        sign = {"clientId": client_id, "config": config_json, "format": "word"}
        resp = client.post(
            f"{FOXIT_API_BASE}/document/convert",
            params={"sn": _sn(client_id, secret, sign), "clientId": client_id},
            data={"format": "word", "config": config_json},
            files={"inputDocument": ("input.pdf", pdf_bytes, "application/pdf")},
        )
        resp.raise_for_status()
        body = resp.json()
        if body.get("code") != 0:
            raise RuntimeError(f"福昕 convert 失败: code={body.get('code')} {body.get('msg') or body}")
        task_id = body["data"]["taskInfo"]["taskId"]

        # ② 轮询任务直到 percentage==100。任务运行中时 /api/task 会返回 HTTP 错误
        #    且 detail 含 "The task is running"——这是"还在跑",要继续轮询,不能当失败。
        doc_id = None
        deadline = time.monotonic() + max_wait_seconds
        while time.monotonic() <= deadline:
            task_base = {"clientId": client_id, "taskId": task_id}
            task_params = {**task_base, "sn": _sn(client_id, secret, task_base)}
            try:
                tr = client.get(f"{FOXIT_API_BASE}/task", params=task_params)
                tr.raise_for_status()
                payload = tr.json()
                # 任务真失败要立刻抛(否则空转到 max_wait 才报通用超时、吞真因)
                if payload.get("code") not in (0, None):
                    raise RuntimeError(
                        f"福昕 task 失败: code={payload.get('code')} {payload.get('msg') or ''}"
                    )
                info = (payload.get("data") or {}).get("taskInfo") or {}
                if str(info.get("status") or "").lower() in ("failed", "fail", "error"):
                    raise RuntimeError(f"福昕转换任务失败: {info}")
                if info.get("percentage") == 100:
                    doc_id = info.get("docId")
                    break
            except httpx.HTTPStatusError as exc:
                detail = ""
                try:
                    detail = ((exc.response.json().get("data") or {}).get("detail") or "")
                except Exception:  # noqa: BLE001
                    pass
                if "The task is running" not in detail:
                    raise
            time.sleep(poll_interval_seconds)
        if not doc_id:
            raise RuntimeError(f"福昕转换未在 {max_wait_seconds}s 内完成")

        # ③ 下载结果 docx(签名覆盖 {clientId, docId, fileName})
        dl_base = {"clientId": client_id, "docId": doc_id, "fileName": "output.docx"}
        dl_params = {**dl_base, "sn": _sn(client_id, secret, dl_base)}
        dr = client.get(f"{FOXIT_API_BASE}/download", params=dl_params)
        dr.raise_for_status()
        with open(output_path, "wb") as fd:
            fd.write(dr.content)

    _assert_nonempty_docx(output_path)
    return output_path


def _assert_nonempty_docx(path: str) -> None:
    """空壳守卫:云返回的 docx 必须有正文或表格,否则 raise 触发回退。"""
    from docx import Document

    doc = Document(path)
    if not (any(p.text.strip() for p in doc.paragraphs) or doc.tables):
        raise RuntimeError("福昕返回的 docx 为空壳(无正文/表格)")


def _foxit_credentials() -> tuple[str, str]:
    from core.config import get_settings

    settings = get_settings()
    cid = str(getattr(settings, "foxit_cloud_client_id", "") or "")
    sec = str(getattr(settings, "foxit_cloud_secret", "") or "")
    if not cid or not sec:
        raise RuntimeError("福昕凭证未配置(FOXIT_CLOUD_CLIENT_ID / FOXIT_CLOUD_SECRET)")
    return cid, sec


def _slice_pdf_to_bytes(pdf_path: str, page_range: tuple[int, int]) -> bytes:
    """按零基、右开 ``page_range`` 切出 PDF 子集字节(只把需要的格式页传云,省额度/时延)。"""
    import fitz

    src = fitz.open(pdf_path)
    try:
        out = fitz.open()
        out.insert_pdf(src, from_page=page_range[0], to_page=page_range[1] - 1)
        data = out.tobytes()
        out.close()
        return data
    finally:
        src.close()


def convert_format_pages_via_cloud(
    tender_pdf_bytes: bytes,
    output_path: str,
    *,
    profile: dict | None = None,
) -> str:
    """商务格式章 → 福昕可编辑 docx + 字段自动填。Phase 0 最上层档,同 builder 契约。

    只把"投标文件格式"商务区页切给云;成功返回 docx 路径,任何失败 raise(由 Phase 0
    阶梯自动下沉到 pdf2docx → 整页图)。
    """
    import os
    import tempfile

    from docx import Document

    from services.original_docx_format_service import (
        _drop_spurious_stream_tables,
        _fill_authorization_letter,
        _fill_basic_info_subfields,
        _fill_bid_date_today,
        _fill_establish_segmented,
        _fill_inline_labeled_blanks,
        _fill_known_table_cells,
        _fill_performance_table,
        _fill_personnel_table,
        _fill_resume_tables,
        _find_format_page_range_in_pdf,
        _log_unfilled_fields,
        _normalize_split_labels,
        _replace_known_fields,
        _strip_seal_images,
        _strip_tender_page_numbers,
    )

    client_id, secret = _foxit_credentials()
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        tmp_pdf.write(tender_pdf_bytes)
        pdf_path = tmp_pdf.name
    try:
        page_range = _find_format_page_range_in_pdf(pdf_path)
        if not page_range:
            raise RuntimeError("未能定位“投标文件格式”商务区页范围")
        slice_bytes = _slice_pdf_to_bytes(pdf_path, page_range)
    finally:
        try:
            os.unlink(pdf_path)
        except OSError:
            pass

    convert_pdf_to_docx_via_foxit(
        slice_bytes, str(output_path), client_id=client_id, secret=secret
    )
    # 字段自动填(复用现成四件套,无 LLM,按标签精确匹配)
    doc = Document(str(output_path))
    _drop_spurious_stream_tables(doc)
    _normalize_split_labels(doc)  # 理顺福昕切开的两字标签(性 别→性别),须在填值前:标签干净才填得上
    _replace_known_fields(doc, profile or {})
    _fill_basic_info_subfields(doc, profile or {})  # 基本情况表专项:法人/技术负责人职称电话+员工总数(须在通用表格填充前)
    _fill_known_table_cells(doc, profile or {})
    _fill_inline_labeled_blanks(doc, profile or {})  # 投标函内联空:工程质量/安全目标/工期/经营期限/法人联系电话
    _fill_authorization_letter(doc, profile or {})  # 授权委托书"本人___（姓名）系"→法人名
    _fill_establish_segmented(doc, profile or {})  # 法人证明"成立时间：__年__月__日"分段填
    _fill_bid_date_today(doc)  # 投标/签署日期落款 → 标书制作当天
    _fill_personnel_table(doc, profile or {})
    _fill_performance_table(doc, profile or {})  # 投标人业绩情况表 → 填选中的类似业绩项目名
    _fill_resume_tables(  # 项目经理 + 总工简历表(各填选派那个人的台账信息)
        doc,
        (profile or {}).get("pm_resume") or {},
        (profile or {}).get("tech_resume") or {},
    )
    # 格式体检:所有值都填完后,对整份文档做"只修格式、不改文字"的合理化
    # (第一版治填空槽下划线断裂:值 run 未继承槽的下划线 → 线画一半)。
    # 传 profile 当白名单:只认我们自己填的值,防止给"年/月/日"等招标原文标签误加线。
    from services.docx_format_doctor import run_format_doctor

    run_format_doctor(doc, profile or {})
    # 福昕把招标原件每页的招标人/代理红章也照搬进来了 → 清掉。投标人章须人工手盖。
    _strip_seal_images(doc)
    _strip_tender_page_numbers(doc)  # 清招标原件页码,标书用自己的页码
    _log_unfilled_fields(doc, profile or {})  # 缺字段显式告警(别静默留空)
    # 商务标固定字段收尾:纠正公司名错别字(安徽正气→安徽正奇)+ 核对固定字段一致性。
    from services.commercial_fixed_fields import (
        audit_commercial_fixed_fields,
        enforce_company_name_consistency,
    )

    corrected = enforce_company_name_consistency(doc)
    if corrected:
        logger.warning("商务标公司名错别字已纠正 %d 处(安徽正气→安徽正奇)", corrected)
    for issue in audit_commercial_fixed_fields(doc):
        logger.warning("商务标固定字段核对:%s", issue)
    doc.save(str(output_path))
    return str(output_path)


def convert_appendix_pages_via_cloud(
    tender_pdf_bytes: bytes,
    output_path: str,
) -> str:
    """技术卷附表区(附表一~八) → 福昕可编辑 docx(空模板,数据格留空给投标人填)。

    附表无公司档案真值,**不做字段填**(区别于商务卷)。成功返回路径,失败 raise。
    """
    import os
    import tempfile

    from services.original_docx_format_service import _find_appendix_page_range_in_pdf

    client_id, secret = _foxit_credentials()
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        tmp_pdf.write(tender_pdf_bytes)
        pdf_path = tmp_pdf.name
    try:
        page_range = _find_appendix_page_range_in_pdf(pdf_path)
        if not page_range:
            raise RuntimeError("未能定位技术卷附表区(附表一)")
        slice_bytes = _slice_pdf_to_bytes(pdf_path, page_range)
    finally:
        try:
            os.unlink(pdf_path)
        except OSError:
            pass

    convert_pdf_to_docx_via_foxit(
        slice_bytes, str(output_path), client_id=client_id, secret=secret
    )
    return str(output_path)
