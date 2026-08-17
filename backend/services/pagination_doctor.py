"""排版医生:治商务卷"尾巴页"(一节内容溢出几行到下一页,孤零零一页几个字)。

成因(2026-07-04 实测量化,巢湖商务卷 vs 招标原件):
- 福昕把每个视觉行转成**独立段落**,并给段落塞 3~17pt 的段前/段后距 → 整体行距
  比招标原件松约 15%(基线间距 31pt vs 27pt),一页少装 4~5 行;
- 加上填入真实值(公司名/项目名)比空白下划线长、会换行 → 招标一页装下的函/表,
  我们溢出几行到下一页,形成"一页只有一行字"的尾巴页(实测 p3/p8/p12/p14)。

修法 = 渲染反馈循环(Word 排版师的"孤行收养",自动化):
  ① LibreOffice 渲染成 PDF → 找尾巴页(文字极少、无图片、非新章节起始页);
  ② 把尾巴页内容定位回 docx 段落,向前找到所属节的标题 → 得到"该节段落区间";
  ③ 只对该区间微收:段前/段后距 ×0.8、行距 ×0.95(EXACTLY 也同比),**方向=向招标
    原件的行距回归**,单节累计收紧不超过两轮(≈15%),肉眼几乎无感;
  ④ 重渲验证,尾巴被上一页收养即停;最多 3 轮,总时长有预算上限。

红线:只调间距数值,一个字、一个框、一条线都不动;LibreOffice 不可用/超时/任何异常
→ 静默跳过(文档保持原样,绝不阻断出标)。
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
import tempfile
import time
from pathlib import Path

logger = logging.getLogger(__name__)

# 尾巴页判定:正文字符数低于此值(一两行字);上一页须"像正常内容页"(高于此值)。
_STUB_MAX_CHARS = 220
_PREV_MIN_CHARS = 350
# 节标题(walk-back 的停靠点):一、/（一）/附表X/第X章…
_HEADING_RE = re.compile(r"^([一二三四五六七八九十]+、|（[一二三四五六七八九十]+）|附表|第.{1,4}(章|部分)|[一二三四五六七八九十]+\s*、)")
_MAX_ITERS = 5  # 渲染实测仅~4s/轮(LibreOffice热实例),多验证几轮无妨
_MAX_ROUNDS_PER_SECTION = 3
_TIME_BUDGET_SECONDS = 420


def _soffice() -> str | None:
    for cand in (
        shutil.which("soffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ):
        if cand and Path(cand).exists():
            return cand
    return None


def _render_pdf(soffice: str, docx_path: Path, outdir: Path, timeout: int) -> Path | None:
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
            check=True, capture_output=True, timeout=timeout,
        )
    except Exception:
        logger.warning("排版医生:LibreOffice 渲染失败/超时,跳过", exc_info=True)
        return None
    pdf = outdir / (docx_path.stem + ".pdf")
    return pdf if pdf.exists() else None


def _norm(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def _find_stub_pages(pdf_path: Path) -> list[dict]:
    """找尾巴页:字少、无图、非新节起始、且上一页是正常内容页。返回每个尾巴的定位线索。"""
    import fitz

    stubs: list[dict] = []
    doc = fitz.open(str(pdf_path))
    try:
        prev_chars = 0
        for pg in range(doc.page_count):
            page = doc[pg]
            info = page.get_text("dict")
            lines = [l.strip() for l in page.get_text().splitlines() if l.strip()]
            # 页眉(示范文本字样)、页码脚注、裸页码都不算内容——页码混进 body 会让
            # first_line 变成"第5页/共204页",拿它回 docx 里定位段落必然落空,整个尾巴
            # 直接被跳过(2026-07-30 投标函落款两行溢出没被收养,就是栽在这)。
            body = [
                l for l in lines
                if "招标示范文本" not in l
                and not re.fullmatch(r"第\s*\d+\s*页\s*/?\s*(共\s*\d+\s*页)?", l)
                and not re.fullmatch(r"\d{1,4}", l)
            ]
            chars = len("".join(body))
            # 注意:page.get_images() 是页资源表,LibreOffice 会把全文档图挂到每页 → 恒真误杀。
            # 要用"页面上实际画出的图块"(type=1)判断。
            has_img = any(b.get("type") == 1 for b in info.get("blocks", []))
            if (
                pg > 0
                and not has_img
                and 0 < chars <= _STUB_MAX_CHARS
                and prev_chars >= _PREV_MIN_CHARS
                and body
                and not _HEADING_RE.match(body[0])
            ):
                stubs.append({"page": pg, "first_line": body[0], "last_line": body[-1]})
            prev_chars = chars
        return stubs
    finally:
        doc.close()


def _find_overlap_pages(pdf_path: Path) -> list[dict]:
    """出厂自检:扫"两层文字叠印"页(2026-08-15 马鞍山#216 p3 浮框条款压正文,用户一打开
    满屏叠字)。判定=不同文字块的 span 外框大面积相交(交叠面积>0.35×小框、高>4pt)。
    降噪:双方都是 ≤2 字的碎片(表格个位数字擦边、页码撞线)不算病——招标原样里也有。"""
    import fitz

    result: list[dict] = []
    doc = fitz.open(str(pdf_path))
    try:
        for pg in range(doc.page_count):
            spans: list = []
            for b in doc[pg].get_text("dict").get("blocks", []):
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for s in line.get("spans", []):
                        if s.get("text", "").strip():
                            spans.append((fitz.Rect(s["bbox"]), b.get("number"), s["text"].strip()))
            spans.sort(key=lambda x: x[0].y0)
            hits, sample = 0, ""
            for i in range(len(spans)):
                r1, b1, t1 = spans[i]
                for j in range(i + 1, len(spans)):
                    r2, b2, t2 = spans[j]
                    if r2.y0 >= r1.y1 - 4:
                        break  # 之后的 span 纵向都够不着 r1 了(已按 y0 排序)
                    if b1 == b2:
                        continue
                    inter = r1 & r2
                    if inter.is_empty or inter.height <= 4:
                        continue
                    if inter.get_area() <= 0.35 * min(r1.get_area(), r2.get_area()):
                        continue
                    if len(t1) <= 2 and len(t2) <= 2:
                        continue
                    hits += 1
                    if not sample:
                        sample = (t1 if len(t1) >= len(t2) else t2)[:24]
            if hits:
                result.append({"page": pg + 1, "count": hits, "sample": sample})
        return result
    finally:
        doc.close()


def _locate_section_slice(doc, stub: dict) -> tuple[int, int] | None:
    """把尾巴页首/末行定位回 docx 段落,并向前走到所属节标题 → (起,止) 段落下标区间。"""
    paras = doc.paragraphs
    needle = _norm(stub["first_line"])[:16]
    if len(needle) < 4:
        return None
    hit = None
    for i, p in enumerate(paras):
        if needle in _norm(p.text):
            hit = i
            break
    if hit is None:
        return None
    # 终点:尾巴末行所在段(从 hit 往后найти,窗口 40 段)
    end = hit
    tail_needle = _norm(stub["last_line"])[:16]
    if len(tail_needle) >= 4:
        for j in range(hit, min(hit + 40, len(paras))):
            if tail_needle in _norm(paras[j].text):
                end = j
                break
    # 起点:向前走到节标题(不含标题本身),最多回溯 80 段
    start = max(0, hit - 1)
    for k in range(hit - 1, max(-1, hit - 80), -1):
        t = paras[k].text.strip()
        style_name = getattr(paras[k].style, "name", "") or ""
        if t and (_HEADING_RE.match(t) or style_name.startswith("Heading")):
            start = k + 1
            break
        start = k
    return (start, end)


def _tighten_range(doc, start: int, end: int, round_no: int = 1) -> int:
    """微收 [start,end] 段落的间距。只调数值,不动文字。逐轮加力:
    第1轮:段前/后×0.7;行距 1.5/1.62 这类"福昕撑松档"直接压到 1.15(最大杠杆,
          仍比招标原件行距紧不了多少);其余 float 行距×0.92(下限1.0)。
    第2轮:再收段距×0.7、行距×0.92;**空白段压薄**(整行空白≈31pt → EXACTLY 8pt
          的细缝,视觉仍有分隔感但省下整行)。
    第3轮:同第2轮再来一次(总收紧到位即被渲染验证叫停)。
    """
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    touched = 0
    for p in doc.paragraphs[start : end + 1]:
        pf = p.paragraph_format
        changed = False
        if not p.text.strip():
            # 空白段:第2轮起压成细缝(不删段,保留分隔感)
            if round_no >= 2:
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(8)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                touched += 1
            continue
        if pf.space_before is not None and pf.space_before.pt > 1:
            pf.space_before = Pt(max(0.0, pf.space_before.pt * 0.7))
            changed = True
        if pf.space_after is not None and pf.space_after.pt > 1:
            pf.space_after = Pt(max(0.0, pf.space_after.pt * 0.7))
            changed = True
        ls, rule = pf.line_spacing, pf.line_spacing_rule
        if ls is not None:
            if rule == WD_LINE_SPACING.EXACTLY and hasattr(ls, "pt") and ls.pt > 12.5:
                pf.line_spacing = Pt(max(12.5, ls.pt * 0.92))
                changed = True
            elif isinstance(ls, float) and ls > 1.02:
                if round_no == 1 and ls >= 1.4:
                    pf.line_spacing = 1.15  # 福昕撑松档一步到位
                else:
                    pf.line_spacing = max(1.0, ls * 0.92)
                changed = True
        if changed:
            touched += 1
    return touched


def heal_stub_pages(docx_path: str) -> dict:
    """对最终商务卷做"孤行收养"。就地修改文件;任何失败保持原样。返回报告。"""
    report: dict = {
        "ran": False, "iterations": 0, "stubs_before": None, "stubs_after": None,
        "overlap_pages": None,
    }
    soffice = _soffice()
    if not soffice:
        logger.info("排版医生:未找到 LibreOffice,跳过尾巴页治理")
        return report
    try:
        from docx import Document

        deadline = time.monotonic() + _TIME_BUDGET_SECONDS
        section_rounds: dict[int, int] = {}
        with tempfile.TemporaryDirectory() as td:
            outdir = Path(td)
            src = Path(docx_path)
            last_pdf: Path | None = None
            stale = False  # 上次渲染后 docx 又被改过 → 自检前需重渲
            for it in range(_MAX_ITERS):
                if time.monotonic() > deadline:
                    logger.info("排版医生:时间预算用尽,停在第%d轮", it)
                    break
                pdf = _render_pdf(soffice, src, outdir, timeout=max(30, int(deadline - time.monotonic())))
                if pdf is None:
                    break
                last_pdf, stale = pdf, False
                stubs = _find_stub_pages(pdf)
                if report["stubs_before"] is None:
                    report["stubs_before"] = [s["page"] + 1 for s in stubs]
                report["iterations"] = it
                report["stubs_after"] = [s["page"] + 1 for s in stubs]
                if not stubs:
                    break
                doc = Document(str(src))
                acted = 0
                # 第一手段:尾巴自己(或它前面的空段链)带着"段前分页/硬分页符/另起页分节符"
                # → 直接拆掉那个分页动作,尾巴原地归队。这是"不可以再犯"闸(2026-07-30
                # 用户死命令):不管将来哪段代码、以什么方式把表单尾巴甩下去,出卷前这里
                # 一律兜住,不再依赖修 N 个源头。节标题不在此列(标题分页是版式要求)。
                body_children = list(doc.element.body)
                for stub in sorted(stubs[:8], key=lambda s: -s["page"]):
                    first = str(stub.get("first_line") or "").strip()
                    if not first or _HEADING_RE.match(first):
                        continue
                    from services.blank_page_doctor import (
                        _locate_all,
                        _remove_one_break_before,
                    )

                    for idx in _locate_all(body_children, first, doc):
                        how = _remove_one_break_before(body_children, idx)
                        if how:
                            acted += 1
                            logger.info(
                                "排版医生:第%d页尾巴由多余分页造成,已拆(%s)",
                                stub["page"] + 1, how,
                            )
                            break
                if acted:
                    doc.save(str(src))
                    report["ran"] = True
                    stale = True
                    continue  # 拆完重渲复查,防止和"收间距"手段互相打架
                # 第二手段:没有可拆的分页 → 微收所属节的段距/行距
                for stub in stubs[:8]:
                    sl = _locate_section_slice(doc, stub)
                    if sl is None:
                        continue
                    key = sl[0]
                    if section_rounds.get(key, 0) >= _MAX_ROUNDS_PER_SECTION:
                        continue  # 该节已收到上限,再收肉眼可见,放过
                    section_rounds[key] = section_rounds.get(key, 0) + 1
                    acted += _tighten_range(doc, sl[0], sl[1], round_no=section_rounds[key])
                if acted == 0:
                    break  # 没有可收的了(都到上限/定位不到),停
                doc.save(str(src))
                report["ran"] = True
                stale = True
            # 出厂自检:终稿再扫一遍"文字叠印"(用户 2026-08-15 死命令:出手前自己先看一眼)。
            # 只报不修——已知病因(浮框长条款)在填前 defloat 已治,这里是最后一道眼睛。
            if stale or last_pdf is None:
                last_pdf = _render_pdf(soffice, src, outdir, timeout=120)
            if last_pdf is not None:
                overlaps = _find_overlap_pages(last_pdf)
                report["overlap_pages"] = overlaps
                if overlaps:
                    logger.warning(
                        "出厂自检:发现文字叠印 %s —— 出卷后请人工翻到这些页复核",
                        "; ".join(f"第{o['page']}页×{o['count']}(如:{o['sample']})" for o in overlaps[:6]),
                    )
                else:
                    logger.info("出厂自检:整卷无文字叠印")
        if report["ran"]:
            logger.info(
                "排版医生:尾巴页 %s → %s(%d轮)",
                report["stubs_before"], report["stubs_after"], report["iterations"] + 1,
            )
        return report
    except Exception:
        logger.warning("排版医生失败,文档保持原样(不阻断出标)", exc_info=True)
        return report
