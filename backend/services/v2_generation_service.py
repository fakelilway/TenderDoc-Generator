"""V2 Generation Pipeline — orchestrates the only bid generation architecture.

Wire format extraction → form filling → content writing → audit → export
into a single package for workflow/export.
"""

from __future__ import annotations

import logging
import re
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from typing import Any, Callable

logger = logging.getLogger(__name__)

# Named constants for magic numbers
MAX_KNOWLEDGE_CHUNKS = 10  # Max RAG chunks per content writer call

# Pre-compiled regex patterns for _clean_for_markdown
_RE_HTML_TAG = re.compile(r"<[^>]+>")
_RE_CRLF = re.compile(r"\r\n")
_RE_CR = re.compile(r"\r")
_RE_BLANK_LINES = re.compile(r"\n{3,}")
_RE_PAGE_NUM = re.compile(r"\n\s*\d{1,3}\s*\n")
_RE_UNDERSCORE = re.compile(r"_{3,}")

from schemas.tender import TenderRequirements
from agents.form_filler_agent import (
    FillResult,
    generate_missing_checklist,
)
from agents.content_writer_agent import fill_technical_volume, VolumeFillResult
from services.v2_audit_service import (
    full_audit,
    AuditResult,
    AuditIssue,
)
_RE_CJK_RUN = re.compile(r"[一-鿿A-Za-z0-9]+")


def _cjk_bigrams(text: str) -> set[str]:
    """Character bigrams over CJK/alphanumeric runs.

    Chinese has no word boundaries and the shared tokenizer returns whole runs,
    so exact-token overlap rarely matches. Bigram overlap is a cheap, robust
    similarity signal for matching requirement items to node titles.
    """
    grams: set[str] = set()
    for run in _RE_CJK_RUN.findall(text or ""):
        if len(run) == 1:
            grams.add(run)
            continue
        for i in range(len(run) - 1):
            grams.add(run[i : i + 2])
    return grams


def _rewrite_sections_concurrently(
    rewrite_one: Callable[[str, list[dict[str, str]]], tuple[str, str]],
    targets: list[tuple[str, list[dict[str, str]]]],
    workers: int,
) -> list[tuple[str, str]]:
    """并发跑各节的合规补写,返回 [(节标题, 新正文)]。

    各节补写互不相干,却曾是逐节排队的:2026-07-29 巢湖真卷实测这段串行跑了十几分钟
    (一次只发一个 LLM 请求、CPU 全程空转等回话)。与技术卷正文写作用同一个并发上限
    ``BID_WRITER_CONCURRENCY``,墙上时间约降到 1/并发数。

    单节补写抛异常时只丢这一节(新正文留空,调用方跳过、保留原正文),不连累其它节——
    补写整体本就是 best-effort,出标前的覆盖闸仍会把关。
    """
    if not targets:
        return []

    def _safe(target: tuple[str, list[dict[str, str]]]) -> tuple[str, str]:
        sect, miss = target
        try:
            return rewrite_one(sect, miss)
        except Exception:  # noqa: BLE001 - 单节失败不该拖垮整轮补写
            logger.warning("合规补写单节失败,保留原正文:%s", sect, exc_info=True)
            return sect, ""

    with ThreadPoolExecutor(max_workers=min(workers, len(targets))) as pool:
        return list(pool.map(_safe, targets))


def _distribute_requirement_items(
    titles: list[str],
    items: list[Any],
) -> dict[str, list[dict[str, str]]]:
    """Map each评分项/废标项 to the technical node it best matches by bigram overlap.

    Every item is assigned to exactly one node (its best match, or a catch-all
    node when there is no overlap) so the writer responds to all scored criteria
    without duplicating every item into every section.
    """
    result: dict[str, list[dict[str, str]]] = {t: [] for t in titles}
    if not titles or not items:
        return result

    title_grams = {t: _cjk_bigrams(t) for t in titles}
    # Catch-all for items that match no node title.
    catch_all = next(
        (t for t in titles if ("施工组织" in t or "技术" in t or "组织设计" in t)),
        titles[0],
    )

    for item in items:
        item_title = str(getattr(item, "title", "") or "")
        item_desc = str(getattr(item, "description", "") or "")
        item_grams = _cjk_bigrams(f"{item_title} {item_desc}")
        best_title, best_overlap = None, 0
        for t in titles:
            overlap = len(item_grams & title_grams[t])
            if overlap > best_overlap:
                best_overlap, best_title = overlap, t
        target = best_title or catch_all
        result[target].append({"title": item_title, "description": item_desc})

    return result


def _is_appendix_title(title: str) -> bool:
    """是否为招标规定的附表节(附表一 总体作业计划表 等)——这类是表格,不走 LLM 空写。"""
    return str(title or "").strip().startswith("附表")


_APPENDIX_TITLE_RE = re.compile(r"^附表([一二三四五六七八九十]+)\s*(.+)$")


def _required_appendices(sections: list[dict] | None) -> list[tuple[str, str]]:
    """从技术卷大纲的附表节解析出招标要求的附表清单 [(附表号, 名称)],顺序保持。"""
    out: list[tuple[str, str]] = []
    for s in sections or []:
        m = _APPENDIX_TITLE_RE.match(str((s or {}).get("title", "")).strip())
        if m:
            out.append((m.group(1), m.group(2).strip()))
    return out


def _appendix_markdown(title: str) -> str:
    """把招标附表节渲染成标准 Markdown 表格(由 markdown_to_docx 转成 docx 表格)。

    按附表名关键词匹配公路养护示范文本附表一~五的列式;未知附表给带列占位的通用表。
    满足"招标格式里有表格,生成文件里也要有对应表格"。
    """
    t = str(title or "").strip()
    note = "（按招标文件附表格式填写）"
    if "施工总平面图" in t:
        body = (
            "投标人应递交一份施工总平面图，绘出现场临时设施布置，并附文字说明临时营地、料场、"
            "临时设施、供电、供水、道路、消防等设施的情况和布置。\n\n"
            "【施工总平面图：请在此插入施工总平面布置图（图片 / CAD 图），并附文字说明】"
        )
    elif "作业计划表" in t or "进度计划表" in t:
        body = (
            f"{note}\n\n"
            "| 序号 | 主要工程项目 | 单位 | 工程数量 | 计划开工 | 计划完工 | 总工期(日历天) | 各月进度安排 |\n"
            "|---|---|---|---|---|---|---|---|\n"
            "| 1 | ________ | ____ | ________ | ____年__月 | ____年__月 | ________ | ________ |\n"
            "| 2 | ________ | ____ | ________ | ____年__月 | ____年__月 | ________ | ________ |\n"
            "| 3 | ________ | ____ | ________ | ____年__月 | ____年__月 | ________ | ________ |"
        )
    elif "劳动力计划表" in t:
        body = (
            f"{note}（按工程施工阶段填报各工种投入劳动力人数）\n\n"
            "| 工种 | 施工准备 | 路基/路面施工 | 安全设施施工 | 收尾阶段 | 备注 |\n"
            "|---|---|---|---|---|---|\n"
            "| ________ | ____ | ____ | ____ | ____ | ________ |\n"
            "| ________ | ____ | ____ | ____ | ____ | ________ |"
        )
    elif "临时占地" in t:
        body = (
            f"{note}\n\n"
            "| 桩号 | 用途 | 地类(水田/旱地/果园/荒地等) | 面积(m²) | 需用时间(年月至年月) | 用地位置(左/右, m) |\n"
            "|---|---|---|---|---|---|\n"
            "| ________ | ________ | ________ | ____ | ____年__月至____年__月 | ________ |"
        )
    elif "外供电力" in t or "电力需求" in t:
        body = (
            f"{note}\n\n"
            "| 桩号 | 用电位置(左/右, m) | 用途 | 计划用电量(kW·h) | 需用时间(年月至年月) | 备注 |\n"
            "|---|---|---|---|---|---|\n"
            "| ________ | ________ | ________ | ____ | ____年__月至____年__月 | ________ |"
        )
    else:
        body = f"{note}\n\n【{t}：请按招标文件附表格式填写】"
    return f"## {t}\n\n{body}"


@dataclass
class V2BidPackage:
    """Generated bid package produced by the V2 original-format pipeline."""

    commercial_markdown: str = ""
    technical_markdown: str = ""
    pricing_markdown: str = ""
    combined_markdown: str = ""
    missing_checklist: list[str] = field(default_factory=list)
    audit_result: AuditResult | None = None
    format_docx_path: str | None = None  # Pre-built format DOCX from PDF
    appendix_docx_path: str | None = None  # 技术卷附表(福昕可编辑空表,拼到技术卷末)
    audit_blocked: bool = False  # True when critical audit issues found (content still saved for preview)

    VOLUME_ORDER = ("commercial", "technical", "pricing")
    VOLUME_HEADINGS = {
        "commercial": "商务文件",
        "technical": "技术文件",
        "pricing": "报价文件",
    }

    def volume_map(self) -> dict[str, str]:
        return {
            "commercial": self.commercial_markdown,
            "technical": self.technical_markdown,
            "pricing": self.pricing_markdown,
        }

    @property
    def generation_mode(self) -> str:
        return "v2_format_copy"


def _audit_built_format_docx(docx_path: str) -> list[str]:
    """最低内容体检:复制出来的招标格式章 DOCX 不能是空壳。

    PDF 原格式路径绕过了基于页面的格式审查(原格式模式下 filled_pages 为空),
    一旦 pdf2docx/截图链静默产出一个几乎没有内容的文档,过去会一路绿灯发布。
    这里做**保守**检查:只在产物近乎为空(无有效段落、无表格、无任何图片)时判失败,
    不对版式/对齐做判断(那属 P1 渲染基线的范畴),以免误伤正常产出。返回严重问题列表。
    """
    from docx import Document

    try:
        doc = Document(docx_path)
    except Exception as exc:  # noqa: BLE001 - surface as a content issue
        return [f"格式章 DOCX 无法打开：{exc}"]

    has_text = any(p.text.strip() for p in doc.paragraphs)
    table_count = len(doc.tables)
    inline_images = len(doc.inline_shapes)
    # 整页截图路径用的是浮动图/VML imagedata,不计入 inline_shapes,需扫 body xml。
    try:
        body_xml = doc.element.body.xml
    except Exception:  # noqa: BLE001
        body_xml = ""
    embedded_images = body_xml.count("<a:blip") + body_xml.count("imagedata")

    if not (has_text or table_count or inline_images or embedded_images):
        return ["格式章复制产物为空（无段落/表格/图片）——复制链可能已失败"]
    return []


def _knowledge_chunk_payload(chunk: Any) -> dict[str, Any]:
    """把检索结果(RetrievalResult / dict / str)规整成写作 prompt 认的载荷。

    必须保留 ``metadata``——generator_prompt 靠 ``metadata['document_category']``
    才能把公司历史施组片段标成【公司同类施工方案·以此为骨架:沿用结构/工艺/深度,数据换本项目】;
    旧代码只留 content,导致该深度标注永不触发(料即使接通也被当普通素材)。
    """
    if isinstance(chunk, dict):
        content = str(chunk.get("content", "") or chunk.get("snippet", ""))
        metadata = chunk.get("metadata") or {}
    elif isinstance(chunk, str):
        content, metadata = chunk, {}
    else:
        content = str(getattr(chunk, "content", chunk) or "")
        metadata = getattr(chunk, "metadata", {}) or {}
    return {"content": content, "metadata": metadata}


def _flatten_retrieved_chunks(retrieved: dict[str, list] | None) -> list[Any]:
    """检索结果按"章节标题"归类(见 workflow_service._retrieve_for_outline);
    汇总各节素材、按 chunk_id 去重成一份共享语料,喂给每个写作节点。

    历史致命 bug:旧代码 ``retrieved.get("technical")`` / ``get("施工组织")`` 取的
    key 在该字典里永不存在(键是中文章节标题),检索到的公司施组语料被静默丢弃,
    技术卷退化成 LLM 空写——这正是技术卷偏薄的根因。按 values() 汇总即与键名解耦。
    """
    flattened: list[Any] = []
    seen: set = set()
    for section_chunks in (retrieved or {}).values():
        for chunk in section_chunks or []:
            key = getattr(chunk, "chunk_id", None)
            if key is None:
                key = id(chunk)
            if key in seen:
                continue
            seen.add(key)
            flattened.append(chunk)
    return flattened


def generate_v2_bid_package(
    requirements: TenderRequirements,
    retrieved_chunks_by_section: dict[str, list] | None = None,
    *,
    company_name: str = "",
    tender_text: str = "",
    company_profile: dict[str, str] | None = None,
    original_format_docx_available: bool = False,
    tender_bytes: bytes | None = None,
    confirmed_technical_outline: list[dict] | None = None,
    project_id: int | None = None,
    boq_text: str = "",
) -> V2BidPackage:
    """V2 generation: extract → fill → write → audit.

    If tender_bytes is provided and is a PDF, the format chapter is converted
    directly to DOCX during generation — no separate export step needed.
    """
    from core.config import get_settings

    settings = get_settings()
    company_name = company_name or settings.company_name
    profile = company_profile or _load_company_profile()
    profile["company_name"] = company_name
    retrieved = retrieved_chunks_by_section or {}

    # Build combined profile with project-specific fields from requirements.
    # These must be available BEFORE Phase 0 so the PDF format copy can fill
    # placeholders with Parser-extracted values (质量标准, 安全目标, etc.).
    import re as _re

    _sec = _re.search(
        r"(?:招标项目标段编号|标段编号|招标项目编号|项目编号)[：:\s]*([A-Za-z0-9]{4,40})",
        tender_text or "",
    )
    project_fields = {
        "招标人": str(requirements.tenderer_name or ""),
        "标段编号": _sec.group(1) if _sec else "",
        "项目名称": str(requirements.project_name or ""),
        "工期": str(requirements.planned_duration or ""),
        # 质量:招标优先,留空让 apply_fixed_fields 的 TENDER_FIRST_FALLBACK 兜底填"合格"
        "质量": str(requirements.quality_standard or ""),
        "安全": str(requirements.safety_target or "无安全责任事故发生"),
        "投标有效期": str(requirements.bid_deadline or ""),
        "投标截止时间": str(requirements.bid_deadline or ""),
        "开标时间": str(requirements.bid_deadline or ""),
    }
    combined_profile = {**profile, **project_fields}
    # 法人性别/年龄(从法人身份证 OCR 推导)→ 填法定代表人身份证明表的 性别/年龄 栏。
    combined_profile.update(_legal_rep_pii(str(combined_profile.get("legal_representative", ""))))
    # 商务标固定字段规则(《商务文件固定格式.pdf》定稿)→ 最高优先级覆盖,必须放在
    # 公司档案/招标 parser/法人 OCR 都装配完之后:固定规则 > 公司档案 > parser > AI。
    from services.commercial_fixed_fields import apply_fixed_fields

    apply_fixed_fields(combined_profile)

    # 招标"1.4.4 信誉要求"逐条抽出 → 填"(五)投标人的信誉情况表"(泗沙路实测整表空白,
    # 用户定:按招标1.4.4逐项填响应)。抽不到=空,表保持留白待人工。
    combined_profile["credit_requirement_items"] = _extract_credit_requirement_items(
        tender_text
    )

    # 项目经理/总工简历表字段(选派的人 → 台账结构化信息 + 身份证/毕业证 OCR 兜底)
    pm_for_resume = str(
        combined_profile.get("project_manager_name") or combined_profile.get("project_manager") or ""
    ).strip()
    if pm_for_resume:
        combined_profile["pm_resume"] = build_pm_resume_fields(pm_for_resume, role="项目经理")
    tech_for_resume = str(combined_profile.get("tech_director_name") or "").strip()
    if tech_for_resume:
        combined_profile["tech_resume"] = build_pm_resume_fields(
            tech_for_resume, role="项目技术负责人"
        )

    # ── Phase 0: 招标商务格式章 PDF → 福昕云转可编辑 Word(唯一路径,无降级) ──
    built_format_docx: str | None = None
    built_appendix_docx: str | None = None
    if original_format_docx_available and tender_bytes:
        import tempfile

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = tmp.name
        tmp.close()

        # 按招标原件类型分两条路(靠字节头判定,不依赖上游传扩展名):
        #   Word 招标(PK 开头的 zip) → 原样复制格式章,**全程不碰福昕**、零转换损耗;
        #   PDF 招标            → 福昕云转可编辑 Word。
        # 两条路拿到 docx 之后走的是同一套填充(fill_format_docx)。
        if tender_bytes[:2] == b"PK":
            try:
                from services.original_docx_format_service import (
                    build_original_format_docx,
                )

                built_format_docx = build_original_format_docx(
                    tender_bytes, tmp_path, profile=combined_profile
                )
            except Exception as exc:
                logger.error("Word 招标原格式复制失败", exc_info=True)
                raise ValueError(
                    f"Word 招标文件原格式复制失败,系统不回退近似格式。({exc})"
                ) from exc
        else:
            # 已删 pdf2docx 可编辑 / 整页图+域烧录 / 纯整页截图 三档备胎——福昕失败即硬报错,
            # 绝不降级出不可编辑/近似稿(避免再迷糊"到底走哪条")。
            if (
                str(getattr(settings, "cloud_pdf_convert", "off") or "off").lower()
                != "foxit"
            ):
                raise ValueError(
                    "PDF 招标格式转换仅支持福昕云:请在 .env 配置 CLOUD_PDF_CONVERT=foxit 及 "
                    "FOXIT_CLOUD_CLIENT_ID / FOXIT_CLOUD_SECRET 后重试。"
                    "(若想彻底不用福昕,请向招标代理索取 Word 版招标文件后重新上传)"
                )
            try:
                from services.cloud_pdf_convert import convert_format_pages_via_cloud

                built_format_docx = convert_format_pages_via_cloud(
                    tender_bytes, tmp_path, profile=combined_profile
                )
            except Exception as exc:
                logger.error("福昕云转换失败(无降级路径)", exc_info=True)
                raise ValueError(
                    f"PDF 招标文件原格式复制失败:福昕云转换出错,系统不回退近似格式。({exc})"
                ) from exc

    # Content-level audit of the copied format chapter. The PDF path bypasses
    # the page-based format audit (filled_pages stays empty in original mode), so
    # a silently-empty/broken copy would otherwise ship unchecked. Per the铁律
    # (格式复制失败=硬报错,不输出空壳/近似稿) we fail loudly here.
    if original_format_docx_available and built_format_docx:
        fmt_issues = _audit_built_format_docx(built_format_docx)
        if fmt_issues:
            raise ValueError(
                "V2 生成失败：招标格式章复制产物未通过内容体检"
                "（系统不输出空壳/近似稿，请检查 PDF 格式章或转换链）："
                + "；".join(fmt_issues)
            )

    # 技术卷附表:按招标要求**逐张装配**——命中公司定稿附表(assets/company_appendices)原样拼公司
    # 文件;否则从招标该附表那一页转可编辑表(福昕);都取不到则占位。拼成一个 docx,导出时接到技术卷末。
    # best-effort——失败则技术卷不含附表(投标人另行补),不阻断主流程。
    required_appendices = _required_appendices(confirmed_technical_outline)
    if required_appendices:
        try:
            import tempfile as _tempfile

            from services.appendix_service import build_appendix_docx

            # 只在招标是 PDF 时把原件交给附表装配——它用 fitz 按页找附表,喂 Word 会直接
            # 报错。Word 招标暂只装配公司定稿附表,招标自带附表出占位页(待补:按 DOCX 抠附表)。
            _tender_pdf = None
            if tender_bytes and tender_bytes[:4] == b"%PDF":
                _tp = _tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                _tp.write(tender_bytes)
                _tp.close()
                _tender_pdf = _tp.name
            elif tender_bytes:
                logger.info("Word 招标:附表暂不从招标原件抽取,仅装配公司定稿附表+占位页")
            _ap = _tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            _ap.close()
            built_appendix_docx = build_appendix_docx(
                required_appendices, _ap.name, _tender_pdf
            )
            logger.info("技术卷附表装配:%s", "、".join(f"附表{n}{nm}" for n, nm in required_appendices))
        except Exception:
            logger.warning("技术卷附表装配失败 — 技术卷不含附表", exc_info=True)
            built_appendix_docx = None

    # 福昕原格式模式:格式页直接照抄招标 DOCX,无需文本提取/正则填表。下列容器恒空,
    # 下游装配/审查据此跳过"文本格式页"渲染、只审 prose 正文(v1 文本重建/正则填表路径已删)。
    filled_pages: dict[str, list[tuple[str, str, str]]] = {
        "commercial": [],
        "technical": [],
        "pricing": [],
    }
    fill_results: list[FillResult] = []
    page_pairs: list[tuple[str, str]] = []

    # ── Phase 3: Write prose content ──
    tech_content = ""
    prose_results: VolumeFillResult | None = None
    # A human-confirmed outline (bid_outline_json) drives the technical目录 when
    # present & substantive; otherwise fall back to tender outline / canonical.
    confirmed_sections = _sections_from_confirmed_outline(confirmed_technical_outline)

    # 工程量清单(第五章 BOQ)→ 分部分项占比,驱动技术卷"按占比定详略"+把工程量喂给对应施工方案。
    # best-effort:抽不到/失败返回空 TenderBOQ,写作退回原行为(不阻断出标)。
    try:
        from services import boq_service

        # 优先用本项目上传的工程量清单(另册)全文,占比按真实数量/金额算;没传才从招标正文估。
        boq = boq_service.build_boq(tender_text, boq_text=boq_text)
    except Exception:
        logger.warning("工程量清单抽取异常,技术卷跳过 BOQ 驱动(不阻断)", exc_info=True)
        boq = None

    def _call_content_writer(
        sections: list[dict],
    ) -> tuple[VolumeFillResult | None, str]:
        """Write prose for every technical section in one bounded-concurrency pass.

        Each section still gets its own focused LLM call (with its must-cover
        guidance, distributed 评分项/废标项 and per-section length budget) for
        winning-bid depth (~25 deep sub-sections), but the calls now run
        concurrently inside ``fill_technical_volume`` (capped by
        ``BID_WRITER_CONCURRENCY``), cutting the volume from ~25 min to ~5-6 min."""
        titles = [s["title"] for s in sections]
        # 招标规定的附表(附表一~五等)是表格,不走 LLM 空写——单独渲染成表格模板(见 _recombine)。
        appendix_titles = [t for t in titles if _is_appendix_title(t)]
        prose_titles = [t for t in titles if t not in appendix_titles]
        # FIX(喂料接通):retrieved 按章节标题归类,旧代码取 "technical"/"施工组织"
        # 这两个永不存在的 key → 公司施组语料被静默丢弃、技术卷沦为 LLM 空写。
        # 改为汇总去重后喂每个写作节点,并保留 metadata(让"公司同类施工方案"标注生效)。
        chunks = _flatten_retrieved_chunks(retrieved)
        knowledge_chunks = [
            _knowledge_chunk_payload(c) for c in chunks[:MAX_KNOWLEDGE_CHUNKS]
        ]
        # Distribute scored/废标 criteria across PROSE nodes only(附表不写正文,别把要响应的
        # 评分/废标项分配给附表节而漏写)。
        score_by_title = _distribute_requirement_items(
            prose_titles, requirements.technical_score_items
        )
        invalid_by_title = _distribute_requirement_items(
            prose_titles, requirements.invalid_bid_items
        )
        guidance_by_title = {s["title"]: str(s.get("must_cover", "")) for s in sections}
        min_chars_by_title = {
            s["title"]: int(s.get("target_chars", 0) or 0) for s in sections
        }

        # BOQ 驱动:每节喂"造价占比总览+本节对应清单项+工程量",并按占比调详略(主导加厚、极小压缩)。
        boq_by_title: dict[str, str] = {}
        if boq is not None and not boq.is_empty():
            from services import boq_service

            for t in prose_titles:
                brief = boq_service.section_node_brief(boq, t)
                if brief:
                    boq_by_title[t] = brief
                min_chars_by_title[t] = boq_service.adjust_min_chars(
                    boq,
                    t,
                    min_chars_by_title.get(t, 0) or _CONFIRMED_OUTLINE_TARGET_CHARS,
                )

        result = fill_technical_volume(
            node_titles=prose_titles,
            project_name=requirements.project_name or "投标项目",
            requirements=requirements.model_dump(),
            company_name=company_name,
            knowledge_chunks=knowledge_chunks,
            tender_text=tender_text,
            score_items_by_title=score_by_title,
            invalid_items_by_title=invalid_by_title,
            guidance_by_title=guidance_by_title,
            min_chars_by_title=min_chars_by_title,
            boq_by_title=boq_by_title,
            max_workers=max(1, int(getattr(settings, "bid_writer_concurrency", 5) or 1)),
        )

        # Reassemble in原大纲 order:散文节用 LLM 正文,附表节渲染成招标格式的表格。
        def _recombine() -> str:
            node_by_title = {n.title: n for n in result.nodes}
            parts: list[str] = []
            for s in sections:
                title = s["title"]
                if _is_appendix_title(title):
                    continue  # 附表改为拼成 docx 接到技术卷末(见 build_appendix_docx),正文不再出 markdown 模板
                node = node_by_title.get(title)
                if node is None:
                    continue
                body = _strip_writer_top_level_headings(
                    f"\n## {title}\n\n{node.content}\n"
                )
                parts.append(f"## {title}\n\n{body}")
            return "\n\n".join(parts)

        combined = _recombine()

        # 招标合规定向补写:逐条核对评分点/废标项是否被实质响应,对未响应项重写其归属章节并补足。
        # best-effort——判定/补写异常都不阻断生成,最终仍由出标前的"招标覆盖闸"把关。
        try:
            from agents.content_writer_agent import rewrite_node_for_compliance
            from prompts.generator_prompt import build_node_fill_prompt
            from services.v2_audit_service import evaluate_coverage

            req_dump = requirements.model_dump()
            missing = [
                v
                for v in evaluate_coverage(
                    combined,
                    req_dump.get("invalid_bid_items") or [],
                    req_dump.get("technical_score_items") or [],
                )
                if not v.covered
            ]
            if missing:
                # 键用 (kind, title, description):标题可能重复/为空,只用 title 会把未响应项
                # 路由到错误章节补写。title+description 足以唯一定位(去重已保证唯一)。
                owner_of: dict[tuple[str, str, str], str] = {}
                for kind, by_title in (("废标", invalid_by_title), ("评分", score_by_title)):
                    for sect, its in by_title.items():
                        for it in its:
                            owner_of[(kind, it["title"], it.get("description", ""))] = sect
                missing_by_title: dict[str, list[dict[str, str]]] = {}
                for v in missing:
                    owner = owner_of.get((v.kind, v.title, v.description))
                    # 只补"确实与该技术章节相关"的项(与归属标题有词重叠);与所有技术标题都
                    # 不相关、被 catch_all 强分过来的项(多为商务/资格类,如投标保证金)不硬塞进
                    # 施工方案——交由出标前的全卷覆盖闸核(它对商务+技术+报价整卷判覆盖)。
                    if owner in titles and (
                        _cjk_bigrams(f"{v.title} {v.description}") & _cjk_bigrams(owner)
                    ):
                        missing_by_title.setdefault(owner, []).append(
                            {"kind": v.kind, "title": v.title, "description": v.description}
                        )
                node_by_title = {n.title: n for n in result.nodes}
                repaired = 0

                def _rewrite_one(sect: str, miss: list[dict[str, str]]) -> tuple[str, str]:
                    """重写单节并返回 (节标题, 新正文);失败返回空正文由调用方跳过。"""
                    base_messages = build_node_fill_prompt(
                        node_title=sect,
                        project_name=requirements.project_name or "投标项目",
                        requirements=req_dump,
                        company_name=company_name,
                        knowledge_chunks=knowledge_chunks,
                        tender_text=tender_text,
                        score_items=score_by_title.get(sect),
                        invalid_items=invalid_by_title.get(sect),
                        section_guidance=guidance_by_title.get(sect, ""),
                        target_chars=min_chars_by_title.get(sect, 0),
                        boq_brief=boq_by_title.get(sect, ""),
                    )
                    node = node_by_title[sect]
                    return sect, rewrite_node_for_compliance(
                        base_messages, node.content, sect, miss
                    )

                targets = [
                    (sect, miss)
                    for sect, miss in missing_by_title.items()
                    if sect in node_by_title
                ]
                for sect, new_content in _rewrite_sections_concurrently(
                    _rewrite_one,
                    targets,
                    max(1, int(getattr(settings, "bid_writer_concurrency", 5) or 1)),
                ):
                    node = node_by_title[sect]
                    if new_content and new_content != node.content:
                        node.content = new_content
                        repaired += 1
                if repaired:
                    logger.info("招标合规定向补写:修正 %d 个章节的未响应项", repaired)
                    combined = _recombine()
        except Exception:
            logger.warning(
                "招标合规定向补写异常,跳过(出标前覆盖闸仍会把关)", exc_info=True
            )

        return result, combined

    tech_sections = confirmed_sections or _collect_technical_sections(requirements)
    try:
        prose_results, tech_content = _call_content_writer(tech_sections)
    except Exception as exc:
        logger.error("Content writer failed", exc_info=True)
        raise ValueError(
            "V2 生成失败：施工方案正文生成失败。系统不会输出占位正文，请检查 LLM 配置、" "招标文件文本或知识库资料后重新生成。"
        ) from exc

    # ── Phase 4: Assemble markdown per volume ──

    def _assemble_markdown(vol: str) -> str:
        lines: list[str] = []
        project = requirements.project_name or "投标项目"
        label = V2BidPackage.VOLUME_HEADINGS.get(vol, vol)
        lines.append(f"# {project} {label}\n")
        technical_content_inserted = False
        volume_pages = filled_pages.get(vol, [])

        if volume_pages:
            lines.append("\n<!-- tdg:pagebreak -->\n")
            lines.append(_render_volume_directory(volume_pages))
            lines.append("\n<!-- tdg:pagebreak -->\n")

        for idx, (title, original, filled) in enumerate(volume_pages):
            if vol == "technical" and ("施工" in title or _is_prose_page(title)):
                if tech_content and not technical_content_inserted:
                    lines.append("\n<!-- tdg:pagebreak -->\n")
                    lines.append(
                        f"\n{_add_pagebreaks_before_headings(_clean_for_markdown(tech_content))}\n"
                    )
                    technical_content_inserted = True
                elif not tech_content:
                    content = _clean_for_markdown(filled)
                    content = _RE_UNDERSCORE.sub("________", content)
                    lines.append("\n<!-- tdg:pagebreak -->\n")
                    lines.append(f"\n## {title}\n\n{content}\n")
                continue
            else:
                content = filled

            # Clean content: strip HTML, keep plain markdown
            content = _clean_for_markdown(content)
            content = _RE_UNDERSCORE.sub("________", content)
            content = _render_locked_format_content(
                title,
                original,
                content,
                combined_profile,
            )

            lines.append("\n<!-- tdg:pagebreak -->\n")
            lines.append(f"\n## {title}\n\n{content}\n")

        # Fallback: when volume_pages is empty (original format mode),
        # the for-loop above never runs. Insert tech_content directly
        # so the LLM output is not silently discarded.
        if vol == "technical" and tech_content and not technical_content_inserted:
            lines.append("\n<!-- tdg:pagebreak -->\n")
            lines.append(
                f"\n{_add_pagebreaks_before_headings(_clean_for_markdown(tech_content))}\n"
            )
            technical_content_inserted = True

        return "\n".join(lines)

    commercial_md = _assemble_markdown("commercial")
    technical_md = _assemble_markdown("technical")
    pricing_md = _assemble_markdown("pricing")

    # ── Phase 4b: Enrich commercial/pricing markdown in original format mode ──
    # When original_format_docx_available=True, format pages come from the
    # original DOCX, but the markdown preview still needs compliance content.
    if original_format_docx_available:
        commercial_md = _enrich_commercial_markdown(
            commercial_md,
            requirements,
            combined_profile,
            tender_text=tender_text,
            project_id=project_id,
        )
        # Remove commercial sections from technical markdown if the LLM
        # over-generated them (资格响应, 投标保证金, 项目管理机构 etc.)
        technical_md = _strip_commercial_sections(technical_md)
        notes_md = _build_audit_notes(requirements)
    else:
        notes_md = ""

    # ② 本项目定制插入图(航拍图/本项目图纸):按 target_section 插到技术卷对应节。
    technical_md = _inject_project_images(technical_md, project_id)

    # ── Phase 5: Audit ──
    filled_page_pairs = []
    for vol in ("commercial", "technical", "pricing"):
        for title, original, filled in filled_pages.get(vol, []):
            rendered = _render_locked_format_content(
                title,
                original,
                _clean_for_markdown(filled),
                combined_profile,
            )
            filled_page_pairs.append((title, rendered))

    audit = full_audit(
        pages=page_pairs,
        filled_pages=filled_page_pairs,
        prose_text=tech_content,
        project_name=requirements.project_name or "",
        requirements=requirements.model_dump(),
        filled_fields=_collect_filled_fields(fill_results),
        profile=profile,
        # 招标覆盖校验对"商务+技术+报价"整卷判覆盖:有的废标项在商务卷响应,只看技术卷会误拦。
        coverage_text="\n\n".join(
            t for t in (commercial_md, technical_md, pricing_md) if t
        ),
    )
    audit_blocked = False
    if not audit.passed:
        # Critical issues block the downstream review/export pipeline,
        # but we still return the generated content so the user can preview it.
        has_critical = any(
            issue.severity == "critical"
            for issue in audit.all_issues
        )
        if has_critical:
            logger.warning(
                "Audit found critical issues — blocking review pipeline, "
                "but returning content for preview: %s",
                _format_audit_failure_message(audit),
            )
            audit_blocked = True

    # ── Phase 6: Assemble final package ──
    missing = generate_missing_checklist(fill_results)

    # Use explicit tdg:volume:xxx markers so split_delivery_markdown can recover
    # the three volumes losslessly. Plain "---" separators make it fall back to
    # a heading-keyword heuristic that loses content when volumes share titles.
    from utils.docx_exporter import combine_delivery_volumes

    combined = combine_delivery_volumes(
        doc_title=requirements.project_name or "投标项目",
        volumes={
            "commercial": commercial_md,
            "technical": technical_md,
            "pricing": pricing_md,
        },
        notes=notes_md,
    )

    return V2BidPackage(
        commercial_markdown=commercial_md,
        technical_markdown=technical_md,
        pricing_markdown=pricing_md,
        combined_markdown=combined,
        missing_checklist=missing,
        audit_result=audit,
        format_docx_path=built_format_docx,
        appendix_docx_path=built_appendix_docx,
        audit_blocked=audit_blocked,
    )


def _load_company_profile() -> dict[str, str]:
    """Load company profile, falling back to defaults."""
    try:
        from services.company_profile_service import get_company_profile

        data = get_company_profile()
        profile = data.get("profile", {})
        if isinstance(profile, dict):
            return {str(k): str(v) for k, v in profile.items()}
    except Exception:
        pass
    return {}


def _collect_technical_titles(requirements: TenderRequirements) -> list[str]:
    """Collect construction plan section titles from format tree."""
    titles: list[str] = []
    seen: set[str] = set()

    def add_title(value: str) -> None:
        clean = (value or "").strip()
        if not clean or "投标文件" in clean:
            return
        key = re.sub(r"\s+", "", clean)
        if key in seen:
            return
        seen.add(key)
        titles.append(clean)

    nodes = requirements.format_outline_tree.get("technical", [])
    for node in nodes:
        t = getattr(node, "title", "") or (
            node.get("title", "") if isinstance(node, dict) else ""
        )
        add_title(t)
        children = getattr(node, "children", []) or (
            node.get("children", []) if isinstance(node, dict) else []
        )
        for child in children:
            ct = getattr(child, "title", "") or (
                child.get("title", "") if isinstance(child, dict) else ""
            )
            add_title(ct)
    if not titles:
        titles = ["施工组织设计"]
    return titles


# Tender format trees usually list 技术卷 only as a coarse form requirement
# ("一、施工组织设计 / 二、其他内容"). These are not a content outline — the
# 施工组织设计 is authored by the bidder. When the tree is this thin, expand to
# the canonical deep outline so the technical volume reaches winning-bid depth.
_GENERIC_TECH_TITLES = (
    "施工组织设计", "技术文件", "技术标", "技术方案", "其他内容", "其他材料", "正文",
)


def _collect_technical_sections(requirements: TenderRequirements) -> list[dict]:
    """LEGACY fallback for the technical目录 — only when no confirmed
    ``bid_outline_json`` drives generation (see ``_sections_from_confirmed_outline``).

    Honour the tender's own technical outline at whatever granularity it has
    (列几条就几条). When the tender specifies nothing usable, return a minimal
    neutral shell rather than imposing the detailed canonical outline — 招标各不
    相同，有的技术标目录就很简单，套统一模板反而画蛇添足/不响应。真正的结构由
    人工在大纲编辑器里按本招标定稿。
    """
    titles = _collect_technical_titles(requirements)
    specific = [
        t for t in titles
        if not any(g in t for g in _GENERIC_TECH_TITLES)
    ]
    if specific:
        # Tender provides its own technical outline — honor it as-is.
        return [{"title": t, "must_cover": "", "target_chars": 1500} for t in titles]
    # Nothing specified → minimal neutral shell, NOT a detailed template.
    return [{"title": "施工组织设计", "must_cover": "", "target_chars": 1500}]


# Default per-section length budget for a human-confirmed outline, matching the
# "specific tender outline" path in _collect_technical_sections.
_CONFIRMED_OUTLINE_TARGET_CHARS = 2200


def _sections_from_confirmed_outline(
    confirmed: list[dict] | None,
) -> list[dict] | None:
    """Map a human-confirmed technical outline into content-writer sections.

    ``confirmed`` is the project's ``bid_outline_json`` (a list of
    ``BidSectionOutline`` dicts: title / focus_points / …). When the bidder has
    confirmed a real outline we honour it verbatim — its titles become the
    technical目录 and its ``focus_points`` become each section's must-cover
    guidance — instead of falling back to the canonical hardcoded outline. This
    is the wire that makes "改大纲 → 改生成目录" actually true.

    Returns ``None`` (so the caller falls back to
    :func:`_collect_technical_sections`) when the outline is absent or too thin
    to drive a full technical volume, preserving existing behaviour for projects
    without a real confirmed outline.
    """
    if not confirmed:
        return None
    sections: list[dict] = []
    for item in confirmed:
        if not isinstance(item, dict):
            continue
        title = str(item.get("title", "")).strip()
        if not title:
            continue
        focus = item.get("focus_points") or []
        must_cover = "\n".join(str(p).strip() for p in focus if str(p).strip())
        sections.append(
            {
                "title": title,
                "must_cover": must_cover,
                "target_chars": int(item.get("target_chars", 0) or 0)
                or _CONFIRMED_OUTLINE_TARGET_CHARS,
            }
        )
    # Honour a human-confirmed outline at whatever size it is — a deliberately
    # simple 2-3 section目录 is valid (招标各不相同). Only fall back when empty.
    if not sections:
        return None
    return sections


def _strip_writer_top_level_headings(markdown: str) -> str:
    """Keep format-tree titles as the only top-level headings.

    Content writer calls are scoped to a single node, but models often echo
    ``#`` headings anyway. Remove only ``#`` (H1) headings and ``##`` headings
    that duplicate the node title — preserve meaningful sub-section ``##`` headings
    like "施工方法", "质量保证措施" etc.
    """
    lines = []
    first_heading_stripped = False
    for line in markdown.splitlines():
        # Always strip H1 (# ) — the volume heading is added by _assemble_markdown
        if re.match(r"^\s*#\s+", line) and not re.match(r"^\s*##\s+", line):
            continue
        # Strip only the FIRST ## heading (usually the node title echo),
        # keep all subsequent ## headings as valid sub-sections
        if re.match(r"^\s*##\s+", line) and not first_heading_stripped:
            first_heading_stripped = True
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def _is_prose_page(title: str) -> bool:
    return any(kw in title for kw in ["施工", "方案", "措施", "部署", "计划", "进度", "质量", "安全"])


def _render_locked_format_content(
    title: str,
    original: str,
    content: str,
    profile: dict[str, str],
) -> str:
    """Keep locked commercial/pricing form pages structured.

    Locked commercial/pricing forms must not be approximated. If PDF text
    extraction flattens a required table, the audit layer fails generation
    instead of letting a reconstructed layout masquerade as the tender format.
    """
    if _has_markdown_table(content):
        return content
    if _requires_figure_placeholder(title, original):
        return "【图表占位：请按招标文件要求插入对应组织机构图、进度计划图、施工总平面图或知识库图片资料】"
    return content


def _format_audit_failure_message(audit: AuditResult) -> str:
    critical = [i for i in audit.all_issues if i.severity == "critical"]
    details = "；".join(f"{i.location}: {i.problem}" for i in critical[:8])
    return "V2 生成失败：审查发现严重问题。" f"{details}"


def _render_volume_directory(pages: list[tuple[str, str, str]]) -> str:
    lines = ["## 目 录", ""]
    for title, _original, _filled in pages:
        lines.append(title)
    return "\n".join(lines)


def _add_pagebreaks_before_headings(markdown: str) -> str:
    lines: list[str] = []
    for line in markdown.splitlines():
        if line.startswith("## "):
            if lines and lines[-1].strip() != "<!-- tdg:pagebreak -->":
                lines.extend(["", "<!-- tdg:pagebreak -->", ""])
        lines.append(line)
    return "\n".join(lines)


def _has_markdown_table(content: str) -> bool:
    lines = [line.strip() for line in content.splitlines()]
    for index, line in enumerate(lines[:-1]):
        if (
            line.startswith("|")
            and lines[index + 1].startswith("|")
            and "---" in lines[index + 1]
        ):
            return True
    return False


def _requires_figure_placeholder(title: str, original: str) -> bool:
    text = f"{title}\n{original}"
    return any(
        keyword in text
        for keyword in (
            "组织机构图",
            "框图",
            "施工总平面图",
            "平面布置图",
            "进度计划图",
            "网络图",
            "横道图",
            "附图",
            "图表",
        )
    )


def _collect_filled_fields(results: list[FillResult]) -> list[dict[str, Any]]:
    """Extract filled fields for evidence audit."""
    fields: list[dict[str, Any]] = []
    for r in results:
        for f in r.fields:
            fields.append(
                {
                    "label": f.label,
                    "value": f.value,
                    "matched": f.matched,
                    "profile_key": "",
                }
            )
    return fields


def _clean_for_markdown(text: str) -> str:
    """Strip HTML tags and normalize text for Markdown/DOCX rendering."""
    text = _RE_HTML_TAG.sub("", text)
    text = _RE_CRLF.sub("\n", text)
    text = _RE_CR.sub("\n", text)
    text = _RE_BLANK_LINES.sub("\n\n", text)
    text = _RE_PAGE_NUM.sub("\n", text)
    return text.strip()


def _enrich_commercial_markdown(
    commercial_md: str,
    requirements: TenderRequirements,
    profile: dict[str, str],
    tender_text: str = "",
    project_id: int | None = None,
) -> str:
    """Add compliance response sections to commercial volume in original format mode.

    When the format pages come from the original DOCX, the markdown preview
    would only show a header. This function appends deterministic compliance
    text derived from the parsed requirements so the preview is informative.
    """
    # Only enrich if the markdown is just a header (original format mode)
    lines = commercial_md.splitlines()
    non_empty = [l for l in lines if l.strip() and not l.strip().startswith("#")]
    if non_empty:
        return commercial_md  # Already has content from form filling

    parts = [commercial_md.rstrip()]

    # 商务通读整本招标 → LLM 生成四块商务响应(资格响应/偏离表/声明承诺/投标函一致性)。
    # best-effort:无招标全文或 LLM 失败返回 "",回退到下方模板化资格响应要点(原行为)。
    llm_response_md = ""
    try:
        from services.commercial_response_service import generate_commercial_responses

        llm_response_md = generate_commercial_responses(requirements, tender_text, profile)
    except Exception:
        llm_response_md = ""

    if llm_response_md:
        parts.append(llm_response_md)
    elif requirements.qualification_list:
        # 回退:LLM 不可用(无招标全文/失败)时,仍给出模板化资格响应要点
        parts.append("\n<!-- tdg:pagebreak -->\n")
        parts.append(
            "\n## 附录：商务响应补充说明（系统自动生成，供编制参考；非招标文件格式原文）\n"
        )
        parts.append(
            "\n> 正式商务表格以上方原格式页为准；以下仅为依据招标文件解析自动生成的资格响应要点。\n"
        )
        for item in requirements.qualification_list:
            parts.append(f"\n### {item.title}\n")
            parts.append(f"\n{item.description}\n")
            # Try to match with profile data
            profile_value = _match_profile_field(item.title, profile)
            if profile_value:
                parts.append(f"\n**响应：** {profile_value}\n")

    # 资格符合性 + 投标函三处一致性 硬校验(确定性核对引擎,达标/废标级)。best-effort,
    # 与上方 LLM 软响应并列:软响应给"怎么写",硬校验给"达标没/有没有废标级"。
    if project_id is not None:
        try:
            from services.tender_spec_service import build_conformance_hardcheck_markdown

            hardcheck_md = build_conformance_hardcheck_markdown(project_id)
        except Exception:
            hardcheck_md = ""
        if hardcheck_md:
            parts.append(hardcheck_md)

    # 投标保证金 / 联合体协议书:用户定——保持格式章原样留白,不自动加文字(金额/方式人工填)。
    # (原先这里会追加"## 投标保证金 + 说明文字",现移除。)

    # Project manager section
    parts.append("\n<!-- tdg:pagebreak -->\n")
    parts.append("\n## 项目管理机构\n")
    pm_fields = []
    # CompanyProfile 的字段名是 project_manager_name / project_manager_cert
    # (旧代码取 project_manager / pm_certificate,schema 里不存在 → 恒空)。
    # 保留旧键名作 fallback,避免对可能传入旧键的调用方回归。
    pm_name = profile.get("project_manager_name") or profile.get("project_manager")
    if pm_name:
        pm_fields.append(f"项目经理：{pm_name}")
    pm_cert = profile.get("project_manager_cert") or profile.get("pm_certificate")
    if pm_cert:
        pm_fields.append(f"注册建造师证书：{pm_cert}")
    if profile.get("pm_specialty"):
        pm_fields.append(f"专业：{profile['pm_specialty']}")
    if pm_fields:
        parts.append(f"\n{'，'.join(pm_fields)}。承诺在本项目实施期间在岗，不在其他项目兼职。\n")
    else:
        parts.append(
            "\n拟派项目经理具备相应专业注册建造师证书，在岗承诺详见附件。"
            "安全生产许可证有效且覆盖本项目施工活动。\n"
        )

    # 资格证明材料:从知识库自动插入营业执照/资质/安许/业绩等图片(B)
    evidence_md = _qualification_evidence_markdown()
    if evidence_md:
        parts.append(evidence_md)

    # 法人身份证复印件(正反)→ 法定代表人身份证明 后(招标P120;法人亲签,不附代理人)。
    # 照单附件引擎(2026-07-30)已按表单声明把法人身份证就地插在身份证明表落款后 →
    # 本锚点链让位,否则一卷插两份,且锚点会误命中授权委托书里"附：法定代表人身份证明"
    # 那行、把图楔进委托书(用户实测拍板:授权委托书后不附法人身份证)。
    if not profile.get("_legal_id_inline"):
        legal_rep = str(
            profile.get("legal_representative") or profile.get("法定代表人") or ""
        ).strip()
        legal_id_md = _legal_rep_id_evidence_markdown(legal_rep)
        if legal_id_md:
            parts.append(legal_id_md)

    # 项目经理/总工证件(建造师证/身份证正反/职称/安全/社保)→ 各自资历表后
    pm_name = str(
        profile.get("project_manager_name") or profile.get("project_manager") or ""
    ).strip()
    tech_name = str(profile.get("tech_director_name") or "").strip()
    # 按角色让位:证件已随本人资历表就地插入的角色不再走锚点链;没占到表的角色保留兜底
    cert_inline_roles = set(profile.get("_personnel_certs_inline_roles") or [])
    pm_for_chain = "" if "项目经理" in cert_inline_roles else pm_name
    tech_for_chain = "" if "项目技术负责人" in cert_inline_roles else tech_name
    personnel_cert_md = _personnel_cert_evidence_markdown(pm_for_chain, tech_for_chain)
    if personnel_cert_md:
        parts.append(personnel_cert_md)

    # 本项目选定的类似业绩(多选);未选则空,回退自动行为
    selected_perf_names: list[str] = []
    if project_id is not None:
        try:
            from services import project_service

            sel = project_service.get_selected_performance(project_id).get("selected") or []
            selected_perf_names = [str(s.get("name", "")) for s in sel if s.get("name")]
        except Exception:
            selected_perf_names = []

    # 类似业绩 + 主要人员:从知识库台账/人员证书自动汇总成表(C1,附加参考,不动原表)
    kb_tables_md = _kb_qualification_tables_markdown(selected_perf_names)
    if kb_tables_md:
        parts.append(kb_tables_md)

    # A3:类似业绩证明链。填表阶段已在信息表后**就地插过图**(表→图交替)的角色让位,
    # 防重复;没插成的角色(本卷无它的表/库无证明)保留锚点链兜底,绝不丢证据。
    inline_roles = set(profile.get("_similar_evidence_inline_roles") or [])
    if "bidder" in inline_roles:
        logger.info("投标人业绩证明已随信息表就地插入,跳过其锚点链")
    else:
        perf_evidence_md = _performance_evidence_markdown(
            selected_names=selected_perf_names,
            page_selection=_evidence_page_selection(project_id),
        )
        if perf_evidence_md:
            parts.append(perf_evidence_md)

    # A3c:经理/总工个人业绩证明链——各自"近年完成的类似项目信息表"后附进了表的业绩证明
    if project_id is not None:
        for role in ("pm", "td"):
            if role in inline_roles:
                logger.info("%s业绩证明已就地插入,跳过其锚点链", role)
                continue
            role_md = _role_performance_evidence_markdown(project_id, role)
            if role_md:
                parts.append(role_md)

    return "\n".join(parts)


_NAME_NOISE_TOKENS = (
    "证书", "工程", "公路", "养护", "专管", "劳资", "微信", "图片", "年度",
    "施工员", "材料员", "质量员", "资料员", "机械员", "照片", "人员",
)


def _kb_qualification_tables_markdown(selected_perf_names: list[str] | None = None) -> str:
    """C1:把知识库的业绩台账/主要人员汇总成 markdown 表,作为资格响应附录的参考。

    selected_perf_names 非空时:业绩表只列选中的;为空则汇总近年(原行为)。
    纯附加内容(不填招标原表),无数据则返回空。
    """
    try:
        from services.knowledge_service import (
            list_key_personnel,
            list_performance_records,
        )
    except Exception:
        return ""

    parts: list[str] = []

    try:
        records = list_performance_records(limit=15)
    except Exception:
        records = []
    if records and selected_perf_names:
        from services.performance_archive_service import _norm

        wanted = {_norm(n) for n in selected_perf_names if n}
        records = [r for r in records if _norm(str(r.get("name", ""))) in wanted]
    if records:
        heading = (
            "### 本项目响应类似业绩（已选定）"
            if selected_perf_names
            else "### 近年承建业绩（系统自知识库汇总，请按招标类似业绩要求筛选）"
        )
        lines = [
            "",
            heading,
            "",
            "| 项目名称 | 中标金额 | 年份 | 类型 |",
            "| --- | --- | --- | --- |",
        ]
        for record in records:
            name = (record["name"] or "—").replace("|", "/")[:40]
            lines.append(
                f"| {name} | {record['amount'] or '—'} | "
                f"{record['year'] or '—'} | {record['type'] or '—'} |"
            )
        parts.append("\n".join(lines))

    try:
        personnel = list_key_personnel(limit=40)
    except Exception:
        personnel = []
    personnel = [
        person
        for person in personnel
        if person["name"]
        and 2 <= len(person["name"]) <= 4
        and not any(token in person["name"] for token in _NAME_NOISE_TOKENS)
    ]
    if personnel:
        lines = [
            "",
            "### 主要注册人员（系统自知识库汇总，供编制参考）",
            "",
            "| 姓名 | 持有证书 |",
            "| --- | --- |",
        ]
        for person in personnel[:30]:
            lines.append(f"| {person['name']} | {person['certs'].replace('|', '/')} |")
        parts.append("\n".join(lines))

    if not parts:
        return ""
    return (
        "\n<!-- tdg:pagebreak -->\n"
        "\n## 附录：类似业绩与主要人员（系统按知识库自动汇总，供编制参考）\n"
        + "\n".join(parts)
        + "\n"
    )


def _inject_project_images(technical_md: str, project_id: int | None) -> str:
    """② 本项目定制插入图(航拍图/本项目图纸):按 ``target_section`` 把项目图片插到技术卷
    对应节末尾;未指定或未匹配到节的统一收到末尾"本项目附图"。靠现成 ``{{knowledge_image}}``
    渲染器插入(下游 _render_markdown_body 从 MinIO 取图)。
    """
    if not project_id or not technical_md.strip():
        return technical_md
    try:
        from services.knowledge_service import list_project_insert_images

        images = list_project_insert_images(project_id)
    except Exception:
        return technical_md
    if not images:
        return technical_md

    def _marker(image: dict) -> str:
        caption = str(image.get("caption") or "").replace('"', "")
        return (
            f'\n\n{{{{knowledge_image:document_id={image["document_id"]} '
            f'caption="{caption}" width_cm=14}}}}\n'
        )

    used: set[int] = set()
    parts = re.split(r"(?m)^(##\s.+)$", technical_md)
    rebuilt = [parts[0]]
    idx = 1
    while idx < len(parts):
        heading = parts[idx]
        body = parts[idx + 1] if idx + 1 < len(parts) else ""
        title = heading.lstrip("#").strip()
        block = heading + body
        for image in images:
            target = str(image.get("target_section") or "").strip()
            doc_id = image["document_id"]
            if target and doc_id not in used and (target in title or title in target):
                block += _marker(image)
                used.add(doc_id)
        rebuilt.append(block)
        idx += 2
    result = "".join(rebuilt)

    leftover = [img for img in images if img["document_id"] not in used]
    if leftover:
        result = result.rstrip() + "\n\n## 本项目附图\n"
        for image in leftover:
            result += _marker(image)
    return result


# A2 资格证明材料分组:按真实标书"资格审查"结构成组插**全**公司证件(不止 4 类各 1 张)。
# (组标题, 该组证件类型, 该组最多插几张)。过期证件已在 list_knowledge_image_references 滤掉。
# (组标题, 证件类型, 组上限, 落位锚点)。anchor=招标要求该组插在哪张表后:营业执照/资质/安许/
# 开户/体系→投标人基本情况表后(须知3.5.1);荣誉与信誉→信誉情况表后(3.5.4)。空 anchor=卷尾。
# 单一文件证(营业执照/安许/开户)只插1张(正本优先)——正本+多副本是同一份证,1张够证明。
# 资质/荣誉/专利是多张不同的证,保留较大上限。
# (组标题, 证件类型, 组上限, 落位锚点, 去重模式)。去重模式:
#   one=同一份证只插1张(营业执照/安许/开户,正本+最新有效期优先);
#   specialty=按 专业×分页 各留1张(资质证书:公路工程/市政各1,同专业多次扫描算1;
#     但文件名带 主页/附页/正面/反面/第N页 分页标记的是同一本证的不同页,整套保留——
#     如公路养护资质证书=主页+附页两张,员工反馈"养护证书只放一页"即此病);
#   none=多张不同的证全保留(荣誉/专利/体系)。
# 用户定:基本情况表后只插 营业执照/资质证书(含公路养护,specialty含养护)/安许/开户。
# 体系认证、专利、工法 整本商务卷都不插(已删)。荣誉/信用 落信誉情况表后。
_EVIDENCE_GROUPS: tuple[tuple[str, tuple[str, ...], int, str, str], ...] = (
    ("营业执照", ("营业执照",), 1, "基本情况表", "one"),
    ("企业资质证书", ("资质证书", "施工劳务资质证书"), 16, "基本情况表", "specialty"),
    ("安全生产许可证", ("安全生产许可证",), 1, "基本情况表", "one"),
    # 开户许可证是"正页+注意事项页"两页一本(泗沙路实测只插1页被用户点名),整套插
    ("基本账户开户许可证", ("开户许可证",), 2, "基本情况表", "none"),
    # 荣誉/信用证书**不锚信誉情况表**(2026-07-12 泗沙路用户点名"信誉表后不跟奖项"),
    # 空锚=落卷尾附录区,要用人工取
    ("企业荣誉与信誉证明", ("荣誉证书", "信用证书"), 20, "", "none"),
)


def _extract_credit_requirement_items(tender_text: str | None) -> list[str]:
    """从招标"1.4.4 信誉要求"抽逐条要求文字,供信誉情况表左列。抽不到返回空(表留白)。

    先锚"1.4.4"小节(至 1.4.5/1.5 止),退而求其次锚"信誉要求"字样;逐条按（1）（2）
    编号切分,去空白、掐头尾标点,过滤太短/太长的噪声。"""
    text = re.sub(r"[\t\r]", "", tender_text or "")
    if not text:
        return []

    def _parse(block: str) -> list[str]:
        # "注：情形（1）（2）由投标人自行承诺…"是说明,不是要求条目——在注处收刀,
        # 防止注里引用的（1）（2）编号被误切成条目(泗沙路附录4实测)
        block = re.split(r"注\s*[：:]", block, 1)[0]
        parts = re.split(r"[（(]\s*(\d{1,2})\s*[)）]", block[:3000])
        items: list[str] = []
        for i in range(2, len(parts), 2):
            item = re.sub(r"[\s　]+", "", parts[i]).strip("；;。，,")
            if 4 <= len(item) <= 200:
                items.append(item)
        return items[:12]

    # 逐个候选位置试,取**第一个能解析出条目**的匹配:
    # ①公路养护示范文本把信誉要求放在"附录4 资格审查条件(信誉最低要求)"(泗沙路实测),
    #   前附表only写"信誉要求:见附录4"——附录优先;
    # ②目次里也会出现"1.4.4 信誉要求……页码"(抽不出条目,自动跳过);
    # ③兜底"信誉要求:"必须排除"见附录/见前附表"指路句(否则一路吞到保证金条款)。
    for pattern in (
        r"附录\s*4[^\n]{0,30}信誉[^\n]{0,20}(.*?)(?=附录\s*5|$)",
        r"1\.4\.4\s*信誉要求(.*?)(?=1\.4\.5|1\.5[^0-9]|$)",
        r"信誉要求[:：]?(?!见|详见)(.*?)(?=(?:资格|财务|业绩|人员)要求|1\.4\.5|1\.5[^0-9]|$)",
    ):
        for m in re.finditer(pattern, text, re.S):
            block = m.group(1)
            if re.match(r"\s*[:：]?\s*(见|详见)", block):
                continue  # "信誉要求:见附录4"这类指路句,不是正文
            items = _parse(block)
            if items:
                return items
    return []


def _cert_page_marker(file_name: str) -> str:
    """同一本证的分页标记(主页/附页/正面/反面/背面/第N页);空串=无分页标记。

    带不同分页标记的图属于同一本证的不同页,插入时整套保留、不算重复。
    """
    import re as _re

    m = _re.search(r"(主页|附页|正面|反面|背面|第\d+页)", str(file_name or ""))
    return m.group(1) if m else ""


def _cert_page_order(marker: str) -> int:
    """分页排序:主页/正面在前,附页/反面/背面在后,第N页按N。"""
    import re as _re

    if marker in ("", "主页", "正面"):
        return 0
    m = _re.match(r"第(\d+)页", marker)
    if m:
        return int(m.group(1)) - 1
    return 1  # 附页/反面/背面


def _cert_sort_key(ref: dict) -> tuple:
    """证件排序:正本优先,有效期晚的优先(去重时取这张作代表),同证分页按页序。"""
    import re as _re

    fn = str(ref.get("file_name", ""))
    zheng = 0 if "正本" in fn else 1
    vt = str(ref.get("valid_to", "") or "")
    m = _re.search(r"(20\d{2})", vt)
    year = -int(m.group(1)) if m else 0  # 有效期晚的(年大)排前
    return (zheng, year, _cert_page_order(_cert_page_marker(fn)))


def _qualification_evidence_markdown(limit: int = 80) -> str:
    """生成资格证明材料的图片标记(A2:成组插全公司证件)。

    从知识库按"公司证件 + 确切证件类型"精确选图(不靠模糊评分,免把人员建安证误当公司安许),
    按 _EVIDENCE_GROUPS 分组、每组插全(到组上限),产出 {{knowledge_image:...}} 标记;渲染时
    image_resolver 从 MinIO 取图插入。无匹配返回空,不留死占位。limit=全局张数上限(防失控)。
    """
    try:
        from services.knowledge_service import list_knowledge_image_references

        refs = list_knowledge_image_references("", limit=5000)
    except Exception:
        return ""

    by_type: dict[str, list[dict]] = {}
    for ref in refs:
        if str(ref.get("document_category")) != "公司证件":
            continue
        cert_type = str(ref.get("certificate_type") or "")
        if cert_type:
            by_type.setdefault(cert_type, []).append(ref)  # refs 已 created_at DESC

    blocks: list[str] = []
    seen: set[int] = set()
    for title, cert_types, group_cap, anchor, dedup in _EVIDENCE_GROUPS:
        group_refs: list[dict] = []
        for cert_type in cert_types:
            group_refs.extend(by_type.get(cert_type, []))
        group_refs.sort(key=_cert_sort_key)  # 正本 + 最新有效期优先(去重取代表)
        emitted: list[str] = []
        seen_keys: set[str] = set()  # one/specialty 的去重键
        for ref in group_refs:
            if len(emitted) >= group_cap or len(seen) >= limit:
                break
            doc_id = int(ref.get("document_id", 0) or 0)
            if doc_id <= 0 or doc_id in seen:
                continue
            specialty = str(ref.get("specialty") or "").strip().replace('"', "")
            page = _cert_page_marker(str(ref.get("file_name", "")))
            # 去重:同一份证只取代表张(正本/最新),不同的证才各留;
            # 带分页标记(主页/附页/正反面/第N页)的是同一本证的不同页,整套保留
            if dedup == "one":
                if seen_keys:
                    continue
                seen_keys.add("_")
            elif dedup == "specialty":
                # 键=证件类型×专业×分页:不同类型的证(资质证书 vs 施工劳务资质证书)
                # 或同证不同页都保留;同类型同专业同页的多次扫描才算重复
                cert_type_key = str(ref.get("certificate_type") or "")
                spec_key = f"{cert_type_key}|{specialty or '_'}|{page}"
                if spec_key in seen_keys:
                    continue
                seen_keys.add(spec_key)
            seen.add(doc_id)
            # 图注:证件类型细分(施工劳务,仅specialty组) + 专业(通用不标) + 分页,
            # 如"企业资质证书（公路养护·主页）";none组(荣誉等)保持(1)(2)编号
            type_label = (
                str(ref.get("certificate_type") or "").replace("资质证书", "").strip()
                if dedup == "specialty" else ""
            )
            label_parts = [
                p for p in (type_label, specialty if specialty != "通用" else "", page)
                if p and p not in title
            ]
            spec_label = "·".join(label_parts)
            if spec_label:
                caption = f"{title}（{spec_label}）"
            elif group_cap > 1 and dedup == "none":
                caption = f"{title}（{len(emitted) + 1}）"
            else:
                caption = title
            emitted.append(
                f'\n{{{{knowledge_image:document_id={doc_id} '
                f'anchor="{anchor}" caption="{caption}" width_cm=14}}}}\n'
            )
        if emitted:
            blocks.append(f"\n### {title}\n")
            blocks.extend(emitted)
        if len(seen) >= limit:
            break

    if not blocks:
        return ""
    return (
        "\n<!-- tdg:pagebreak -->\n"
        "\n## 附录：资格证明材料（系统按知识库自动插入，请人工核验/补充）\n"
        + "".join(blocks)
    )


def _legal_rep_pii(legal_rep_name: str) -> dict[str, str]:
    """从法人身份证 OCR 取身份证号,推导 法人性别/法人年龄/法人出生。无则返回 {}。

    填法定代表人身份证明表的 性别/年龄 栏(职务无据可填、留人工)。身份证号第17位奇男偶女、
    7-14位为出生 YYYYMMDD。仅推导、不外泄到正文之外。
    """
    import datetime

    name = (legal_rep_name or "").strip()
    if not name:
        return {}
    try:
        from rag.vector_store import _connect

        with _connect() as conn, conn.cursor() as cur:
            cur.execute(
                "SELECT kc.content FROM knowledge_chunks kc "
                "JOIN documents d ON kc.document_id = d.id "
                "WHERE d.metadata_json->>'owner_name' = %s "
                "AND d.metadata_json->>'certificate_type' LIKE %s LIMIT 30",
                (name, "%身份证%"),
            )
            rows = cur.fetchall()
    except Exception:
        return {}
    today = datetime.date.today()
    for (content,) in rows:
        m = re.search(
            r"(?<![\dXx])(\d{6})(\d{4})(\d{2})(\d{2})(\d{2})(\d)([\dXx])(?![\dXx])",
            content or "",
        )
        if not m:
            continue
        try:
            y, mo, da = int(m.group(2)), int(m.group(3)), int(m.group(4))
            if not (1900 < y <= today.year and 1 <= mo <= 12 and 1 <= da <= 31):
                continue
            gender = "男" if int(m.group(6)) % 2 == 1 else "女"
            age = today.year - y - ((today.month, today.day) < (mo, da))
            return {
                "法人性别": gender,
                "法人年龄": str(age),
                "法人出生": f"{y:04d}-{mo:02d}-{da:02d}",
                "法人身份证号": "".join(m.groups()),
            }
        except (ValueError, IndexError):
            continue
    return {}


def _legal_rep_id_evidence_markdown(legal_rep_name: str) -> str:
    """法人身份证复印件(正反)→ 法定代表人身份证明 后(招标 P120 硬要求)。

    法人亲签路线:只附法定代表人本人身份证(不附委托代理人,本项目未指定代理人)。
    按 owner_name=法人 + 证件类型含'身份证'精确取,正反取前2张,无则返回空。
    """
    name = (legal_rep_name or "").strip()
    if not name:
        return ""
    # 最小覆盖:完整版1张 / 否则正反各1,OCR判面,不重复堆插(完整版+正面+背面)
    try:
        from services import asset_resolver

        cards = asset_resolver.pick_id_card_documents(name)
    except Exception:
        cards = []
    _ID_LABEL = {"both": "（正反面）", "front": "正面", "back": "背面"}
    blocks: list[str] = []
    for r in cards:
        doc_id = int(r.get("document_id", 0) or 0)
        if doc_id <= 0:
            continue
        label = _ID_LABEL.get(str(r.get("side")), "")
        cap = f"法定代表人（{name}）身份证{label}"
        blocks.append(
            f'\n{{{{knowledge_image:document_id={doc_id} '
            f'anchor="法定代表人身份证明" caption="{cap}" width_cm=12}}}}\n'
        )
    if not blocks:
        return ""
    return (
        "\n<!-- tdg:pagebreak -->\n"
        "\n## 附录：法定代表人身份证明材料（系统自动插入，请人工核验）\n"
        + "".join(blocks)
    )


# 项目经理/总工证件:身份证(正反)+ 建造师/职称/安全/毕业证/社保,**每类只取1张代表**
# (同一证扫几十遍是常态,只附一张正本/最新)。cert_types 按优先级排(一级建造师 > 二级),
# 取第一个有的类型的代表张。笼统"人员证件"(大头照)不在内,不会插进去。
# 顺序=用户定的资历表附件序(2026-07-12):身份证(上游单独处理,永远最前)→毕业证→
# 建造师证→建造师证网查截图→安全B证→安全B证网查截图→社保→职称。
# 网查截图库里暂无(员工资料陆续补),命名带"网查"入库即自动跟上;社保/职称看招标要求,
# 库里有才插。⚠️"建造师证网查截图"这类 certificate_type 含"建造师证"子串,匹配时
# 非网查组要跳过带"网查"的类型,防止网查截图被当成证书本体抢先插(见下 _one_person)。
_PERSONNEL_CERT_GROUPS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("毕业证", ("毕业证",)),
    ("注册建造师证", ("一级建造师证", "二级建造师证", "建造师证")),
    ("建造师证网查截图", ("建造师证网查", "建造师网查")),
    ("安全考核证", ("交安证", "建安证", "注册安全工程师证")),
    ("安全B证网查截图", ("安全B证网查", "交安网查", "建安网查", "安全网查")),
    ("社保证明", ("社保",)),
    ("职称证书", ("职称证书",)),
)


def person_cert_documents(role: str, name: str) -> list[tuple[int, str]]:
    """单个人(项目经理/总工)该附的证件图清单 [(document_id, 图注), ...]。

    身份证(正反,OCR判面)永远最前,其后按 _PERSONNEL_CERT_GROUPS 的附件序各取
    1张代表。供 markdown 锚点链与简历表后就地插图共用(同一套挑选规则)。"""
    name = (name or "").strip()
    if not name:
        return []

    out: list[tuple[int, str]] = []
    seen: set[int] = set()

    # 身份证:最小覆盖(完整版1张 / 否则正反各1),OCR判面,不重复堆插
    try:
        from services import asset_resolver

        _ID_LABEL = {"both": "（正反面）", "front": "正面", "back": "背面"}
        for card in asset_resolver.pick_id_card_documents(name):
            doc_id = int(card.get("document_id", 0))
            if doc_id <= 0 or doc_id in seen:
                continue
            seen.add(doc_id)
            label = _ID_LABEL.get(str(card.get("side")), "")
            out.append((doc_id, f"{role}（{name}）身份证{label}"))
    except Exception:
        pass

    # 建造师证/职称/安全/社保:按 owner_name + 证件类型取
    try:
        from services.knowledge_service import list_knowledge_image_references

        refs = list_knowledge_image_references("", limit=5000)
    except Exception:
        refs = []
    def _insertable_image(r: dict) -> bool:
        if str(r.get("image_insertable")) in ("False", "false", "0", "None"):
            return False
        return str(r.get("file_type", "")).lower() in ("jpg", "jpeg", "png")

    by_owner = [
        r
        for r in refs
        if str(r.get("owner_name") or "").strip() == name and _insertable_image(r)
    ]
    for title, cert_types in _PERSONNEL_CERT_GROUPS:
        # 按 cert_types 优先级(一级>二级…)找第一个有的类型,取其正本/最新有效期那张
        chosen = None
        for ct in cert_types:
            cands = [
                r for r in by_owner
                if ct in str(r.get("certificate_type") or "")
                # 非网查组跳过网查截图类型(子串会误命中,如"建造师证网查截图"含"建造师证")
                and ("网查" in ct or "网查" not in str(r.get("certificate_type") or ""))
            ]
            cands.sort(key=_cert_sort_key)
            for r in cands:
                if int(r.get("document_id", 0) or 0) not in seen:
                    chosen = r
                    break
            if chosen:
                break
        if not chosen:
            continue
        doc_id = int(chosen["document_id"])
        seen.add(doc_id)
        out.append((doc_id, f"{role}（{name}）{title}"))
    return out


def _one_person_cert_markdown(role: str, name: str, anchor: str) -> str:
    """单个人(项目经理/总工)的证件插图标记(锚点链兜底用,清单同 person_cert_documents)。"""
    return "".join(
        f'\n{{{{knowledge_image:document_id={doc_id} '
        f'anchor="{anchor}" caption="{caption}" width_cm=12}}}}\n'
        for doc_id, caption in person_cert_documents(role, name)
    )


def build_pm_resume_fields(pm_name: str, role: str = "项目经理") -> dict[str, str]:
    """汇总项目经理/总工简历表要填的字段。

    优先用**人员台账的结构化信息**(最稳):职称、由身份证号算年龄/性别、注册建造师证号+专业;
    台账缺身份证号时才退身份证OCR;学历/毕业学校仍靠毕业证OCR(best-effort)。
    role=拟任职务(项目经理 / 项目技术负责人)。取不到的字段不放进去(留白,不瞎编)。
    """
    import datetime
    import re as _re

    name = (pm_name or "").strip()
    if not name:
        return {}
    fields: dict[str, str] = {"姓名": name, "拟任职务": role}

    # 资历表定稿(2026-07-31 用户提供,员工整理的一人一张定稿)= 最高优先级:
    # "按照我选的人选,从这里面找那个人,按照这个填"。学历/工作年限/毕业学校等
    # 台账没有的字段全靠它;拟任职务仍按选派角色,不照抄文档。
    try:
        from services.curated_resume_service import get_curated_resume

        curated = get_curated_resume(name)
        for k, v in curated.items():
            if k != "拟任职务":
                fields[k] = v
    except Exception:
        pass

    # 台账结构化信息:职称 + 身份证号(→年龄/性别) + 注册建造师证号/专业。比拍照OCR可靠。
    try:
        from services import personnel_roster_service

        member = None
        for m in personnel_roster_service.get_personnel_roster().get("roster", []):
            if str(m.get("name")) == name:
                member = m
                break
        if member:
            # setdefault:资历表定稿在前面已落的字段绝不被台账覆盖(定稿是员工手校的)
            if member.get("title"):
                fields.setdefault("职称", str(member["title"]))
            idn = str(member.get("id_number") or "")
            if len(idn) >= 17 and idn[:17].isdigit():
                fields.setdefault("年龄", str(datetime.date.today().year - int(idn[6:10])))
                fields.setdefault("性别", "男" if int(idn[16]) % 2 == 1 else "女")
            # 注册建造师证号 + 专业:优先取有证号的一级建造师,否则任一有证号的
            certs = member.get("builder_certs") or []
            best = None
            for c in certs:
                if not c.get("cert_no"):
                    continue
                if best is None or "一级" in str(c.get("level") or ""):
                    best = c
            if best:
                fields["建造师证号"] = str(best.get("cert_no") or "")
                if best.get("specialty"):
                    fields.setdefault("专业", str(best["specialty"]))
            if "专业" not in fields and member.get("title_specialty"):
                fields["专业"] = str(member["title_specialty"])
    except Exception:
        pass

    # 身份证 OCR 兜底(台账无身份证号才走):身份证号第7-14位=出生YYYYMMDD,第17位奇男偶女。
    if "年龄" not in fields or "性别" not in fields:
        try:
            from services import asset_resolver
            from services.id_card_ocr import classify_id_card

            for c in asset_resolver.pick_id_card_documents(name):
                cls = classify_id_card(asset_resolver.read_asset_bytes(int(c["document_id"])))
                idn = str(cls.get("id_number") or "")
                if len(idn) >= 17:
                    fields.setdefault("年龄", str(datetime.date.today().year - int(idn[6:10])))
                    fields.setdefault("性别", "男" if int(idn[16]) % 2 == 1 else "女")
                    break
        except Exception:
            pass

    # 毕业证 OCR → 学历/毕业学校/专业(best-effort)。一人可能多张(正/内页),全都读、按页合并。
    try:
        from services.knowledge_service import (
            get_knowledge_document_file_bytes,
            list_knowledge_image_references,
        )
        from utils.file_parser import extract_text_from_image

        refs = list_knowledge_image_references("", limit=5000)
        diplomas = [
            r
            for r in refs
            if str(r.get("owner_name") or "").strip() == name
            and "毕业证" in str(r.get("certificate_type") or "")
            and str(r.get("file_type", "")).lower() in ("jpg", "jpeg", "png")
        ]
        for dipl in diplomas[:3]:
            raw = extract_text_from_image(
                get_knowledge_document_file_bytes(int(dipl["document_id"]))
            ) or ""
            # 去掉全部空白(含换行):真证上"专/科"常被换行切开,只去空格会漏认(王露证实测)
            txt = _re.sub(r"[\s　]+", "", raw)
            for deg in ("博士研究生", "硕士研究生", "博士", "硕士", "本科", "专科", "大专", "中专"):
                if deg in txt:
                    fields.setdefault("学历", deg)
                    break
            else:
                # "专（高职）科"这类中间插注的写法(李刚证实测)也算专科
                if _re.search(r"专[（(][^）)]{0,6}[）)]科", txt):
                    fields.setdefault("学历", "专科")
            # 学校必须是真校名(xx大学/xx学院):遍历所有候选,跳过"普通高等学校"等模板套话;
            # 前缀最少2字,"安徽大学"这类短校名才不漏。OCR去空白后正文会粘连("准予毕业
            # 校长某某某安徽大学"),校名里混进这类词=脏值,宁可留白待人工也不填(不瞎编)
            _SCHOOL_NOISE = ("普通", "高等", "全日制", "本校", "毕业", "校长", "准予", "证书", "批准", "教学")
            for m in _re.finditer(r"([一-龥]{2,12}(?:大学|学院))", txt):
                school = m.group(1)
                if not any(w in school for w in _SCHOOL_NOISE):
                    fields.setdefault("毕业学校", school)
                    break
            # 专业 OCR 噪声大,只在"专业："后紧跟明确专业名(以工程/管理/技术等收尾)才填
            m = _re.search(r"专业[:：]\s*([一-龥]{2,10}(?:工程|管理|技术|科学|设计|建筑))", txt)
            if m:
                fields.setdefault("专业", m.group(1))
            if "学历" in fields and "毕业学校" in fields:
                break  # 两样都齐了,后面的页不用再OCR
    except Exception:
        pass

    return fields


def _personnel_cert_evidence_markdown(pm_name: str, tech_name: str) -> str:
    """选派的项目经理 + 总工的证件扫描件,落到各自资历表后(资格审查硬内容)。"""
    sections: list[str] = []
    pm_md = _one_person_cert_markdown("项目经理", pm_name, "项目经理资历表")
    if pm_md:
        sections.append(pm_md)
    tech_md = _one_person_cert_markdown("项目技术负责人", tech_name, "项目总工资历表")
    if tech_md:
        sections.append(tech_md)
    if not sections:
        return ""
    return (
        "\n<!-- tdg:pagebreak -->\n"
        "\n## 附录：项目经理 / 项目技术负责人证件材料（系统自动插入，请人工核验）\n"
        + "".join(sections)
    )


# A3b 业绩证明链:每个类似业绩附 中标通知书 + 合同 + 交工验收 扫描(import_performance_evidence
# 已入库,document_category=业绩证明、metadata 带 performance_project/evidence_type/evidence_seq)。
_PERF_EVIDENCE_ORDER = ("中标通知书", "合同", "交工验收")
_PERF_PER_TYPE_CAP = {"中标通知书": 2, "合同": 2, "交工验收": 4}


def evidence_images_for_project(
    project_name: str,
    rows: list[tuple] | None = None,
    page_selection: dict[str, list[int]] | None = None,
    full: bool = False,
) -> list[tuple[int, str]]:
    """某个业绩项目该插的证明图清单 [(document_id, 图注), ...],中标→合同→交工按序。

    与 _build_performance_evidence_md 同一套选片规则(人工选页优先且不设上限,
    否则默认每类取前几张);供 similar_project_fill_service 在每张信息表后就地插图。
    rows 不传则现查库;查不到/没有返回空列表。

    ``full=True`` = 全量模式(2026-07-30 用户拍板"文件夹里所有图片都要附"):
    该业绩库里有的图**一张不落**全给——不设每类上限、不看人工选页、"其他"类
    (无关键词文件/PDF转图页)也排在三大类之后一并给。供资格审查三表用。
    """
    from services.performance_archive_service import _norm

    if rows is None:
        rows = _query_performance_evidence_rows()
    wanted = _norm(str(project_name or ""))
    if not wanted:
        return []
    by_type: dict[str, list[tuple[int, int]]] = {}
    for doc_id, proj, etype, _year, seq in rows:
        if _norm(str(proj or "")) != wanted or not etype:
            continue
        try:
            seq_i = int(seq) if seq is not None else 0
        except (TypeError, ValueError):
            seq_i = 0
        try:
            by_type.setdefault(str(etype), []).append((seq_i, int(doc_id)))
        except (TypeError, ValueError):
            continue

    chosen_ids = None if full else (page_selection or {}).get(wanted)
    title = str(project_name).replace('"', "")[:48]
    # 全量模式把库里出现的其它类型(其他/PDF页等)接在三大类后面,顺序稳定
    type_order = list(_PERF_EVIDENCE_ORDER)
    if full:
        type_order += sorted(t for t in by_type if t not in _PERF_EVIDENCE_ORDER)
    out: list[tuple[int, str]] = []
    for etype in type_order:
        all_imgs = sorted(by_type.get(etype, []))
        if full:
            imgs = all_imgs
        elif chosen_ids is not None:
            imgs = [(s, d) for s, d in all_imgs if d in chosen_ids]
        else:
            imgs = all_imgs[: _PERF_PER_TYPE_CAP[etype]]
        for j, (_seq, doc_id) in enumerate(imgs, 1):
            label = "业绩证明" if etype == "其他" else etype
            cap = f"{title}-{label}" + (f"（{j}）" if len(imgs) > 1 else "")
            out.append((doc_id, cap))
    return out


def _build_performance_evidence_md(
    rows: list[tuple],
    limit_projects: int = 6,
    anchor: str = "类似项目情况表",
    label: str = "类似业绩",
    heading: str | None = None,
    page_selection: dict[str, list[int]] | None = None,
) -> str:
    """纯函数:业绩证明行 [(doc_id, 项目, 类型, 年, 序号)] → 按项目成组的插图 markdown。

    选片规则:有完整链(中标+交工)优先,再按年份新→旧取前 limit_projects 个项目;每个项目
    每类按序号取前几张(中标2/合同2/交工4)。anchor 决定这组图落到格式章哪张表后
    (见 generation_service._ANCHOR_SPECS);label/heading 供角色业绩链定制标题。
    page_selection={归一化项目名:[doc_id,...]}:该项目人工选过页 → 只插勾中的那几页、
    **绕开每类张数上限**(员工意见7:盖章页排在第5张也能进);空列表=该项目不附图。
    """
    projects: dict[str, dict] = {}
    for doc_id, proj, etype, year, seq in rows:
        if not proj or not etype:
            continue
        try:
            doc_id = int(doc_id)
        except (TypeError, ValueError):
            continue
        bucket = projects.setdefault(proj, {"year": 0, "by_type": {}})
        try:
            bucket["year"] = max(bucket["year"], int(year or 0))
        except (TypeError, ValueError):
            pass
        try:
            seq_i = int(seq) if seq is not None else 0
        except (TypeError, ValueError):
            seq_i = 0
        bucket["by_type"].setdefault(etype, []).append((seq_i, doc_id))

    def has_chain(bucket: dict) -> bool:
        return {"中标通知书", "交工验收"} <= set(bucket["by_type"])

    ordered = sorted(
        projects.items(),
        key=lambda kv: (has_chain(kv[1]), kv[1]["year"]),
        reverse=True,
    )
    from services.performance_archive_service import _norm

    page_selection = page_selection or {}
    blocks: list[str] = []
    for i, (proj, bucket) in enumerate(ordered[:limit_projects], 1):
        title = str(proj).replace('"', "")[:48]
        seg = [f"\n### {label} {i}：{title}\n"]
        any_img = False
        chosen_ids = page_selection.get(_norm(str(proj)))
        for etype in _PERF_EVIDENCE_ORDER:
            all_imgs = sorted(bucket["by_type"].get(etype, []))
            if chosen_ids is not None:
                # 人工选页:只留勾中的,不设上限(盖章页排多后都能进)
                imgs = [(s, d) for s, d in all_imgs if d in chosen_ids]
            else:
                imgs = all_imgs[: _PERF_PER_TYPE_CAP[etype]]
            for j, (_seq, doc_id) in enumerate(imgs, 1):
                cap = f"{title}-{etype}" + (f"（{j}）" if len(imgs) > 1 else "")
                seg.append(
                    f'\n{{{{knowledge_image:document_id={doc_id} '
                    f'anchor="{anchor}" caption="{cap}" width_cm=14}}}}\n'
                )
                any_img = True
        if any_img:
            blocks.append("".join(seg))

    if not blocks:
        return ""
    if heading is None:
        heading = (
            "\n## 附录：类似业绩证明材料"
            "（中标通知书·合同·交工验收，系统自动插入，请人工核验筛选）\n"
        )
    return "\n<!-- tdg:pagebreak -->\n" + heading + "".join(blocks)


def _query_performance_evidence_rows() -> list[tuple]:
    """取全部可插的业绩证明行 (doc_id, 项目, 类型, 年, 序号);查询异常返回空。"""
    try:
        from rag.vector_store import get_db_connection

        with get_db_connection() as conn, conn.cursor() as cursor:
            cursor.execute(
                "SELECT id, metadata_json->>'performance_project', "
                "metadata_json->>'evidence_type', metadata_json->>'project_year', "
                "metadata_json->>'evidence_seq' FROM documents "
                "WHERE project_id IS NULL "
                "AND metadata_json->>'document_category' = '业绩证明' "
                "AND coalesce(metadata_json->>'image_insertable', '') <> 'false'"
            )
            return list(cursor.fetchall())
    except Exception:
        return []


def _evidence_page_selection(project_id: int | None) -> dict[str, list[int]]:
    """读本项目的业绩证明选页(员工意见7),键归一化项目名。取不到返回空(走默认规则)。"""
    if project_id is None:
        return {}
    try:
        from services import project_service
        from services.performance_archive_service import _norm

        project = project_service._fetch_project(project_id)
        stored = project.get("selected_evidence_pages") or {}
        return {
            _norm(str(k)): [int(d) for d in (v or [])]
            for k, v in stored.items()
        }
    except Exception:
        return {}


def _performance_evidence_markdown(
    limit_projects: int = 6,
    selected_names: list[str] | None = None,
    page_selection: dict[str, list[int]] | None = None,
) -> str:
    """A3b:从知识库取业绩证明扫描,按项目成组插中标通知书/合同/交工验收。无数据返回空。

    selected_names 非空时:只插**选中的**业绩(按项目名匹配),招标要几个就插几个,不再一股脑全堆。
    为空时:回退原行为(取前 limit_projects 个,兼容未做业绩选择的项目)。
    page_selection:人工选过页的业绩只插勾中的那几页(见 _build_performance_evidence_md)。
    """
    rows = _query_performance_evidence_rows()
    if not rows:
        return ""

    if selected_names:
        from services.performance_archive_service import _norm

        wanted = {_norm(n) for n in selected_names if n}
        rows = [r for r in rows if _norm(str(r[1] or "")) in wanted]
        limit_projects = max(limit_projects, len(wanted))  # 选了几个就放几个进去
    return _build_performance_evidence_md(
        rows, limit_projects, page_selection=page_selection
    )


# 角色业绩证明链(员工意见第10条):经理/总工个人的"近年完成的类似项目信息表"后
# 也要附该表所列业绩的证明扫描。role → (人员键, 中文, 落位锚点)。
_ROLE_EVIDENCE = {
    "pm": ("project_manager", "项目经理", "项目经理类似项目表"),
    "td": ("tech_director", "项目总工", "项目总工类似项目表"),
}


def _role_performance_evidence_markdown(project_id: int, role: str) -> str:
    """经理/总工业绩证明链:插"进了此人业绩表的那几条业绩"的证明扫描,锚定其个人节后。

    进表口径与 similar_project_fill_service._records_for_role 一致(全部人工手选):
    勾中哪几条附哪几条;没勾(None/[])/没选派人=表留白,图也不附。
    取不到数据一律返回空串,绝不拦生成。
    """
    person_key, role_label, anchor = _ROLE_EVIDENCE[role]
    try:
        from services import project_service

        person = (
            project_service.get_selected_personnel(project_id).get("selected") or {}
        ).get(person_key) or {}
        person_name = str(person.get("name") or "").strip()
        if not person_name:
            return ""
        chosen = project_service.get_selected_role_performance(project_id, role)[
            "selected"
        ]
        names = [str(it.get("name") or "") for it in (chosen or []) if it.get("name")]
        names = [n for n in names if n]
        if not names:
            return ""

        from services.performance_archive_service import _norm

        wanted = {_norm(n) for n in names}
        rows = [
            r for r in _query_performance_evidence_rows()
            if _norm(str(r[1] or "")) in wanted
        ]
        return _build_performance_evidence_md(
            rows,
            limit_projects=len(names),
            anchor=anchor,
            label=f"{role_label}类似业绩",
            heading=(
                f"\n## 附录：{role_label}（{person_name}）类似业绩证明材料"
                "（中标通知书·合同·交工验收，插于其类似项目信息表后）\n"
            ),
            page_selection=_evidence_page_selection(project_id),
        )
    except Exception:
        logger.warning("%s业绩证明链生成失败,跳过(不拦生成)", role_label, exc_info=True)
        return ""


def _match_profile_field(title: str, profile: dict[str, str]) -> str:
    """Try to find a matching profile value for a requirement title."""
    title_lower = title.lower()
    # CompanyProfile 真实字段名放首位(qualification_grade / safety_license_no /
    # credit_code),旧的不存在键名保留作 fallback,避免改动破坏其它传入路径。
    mapping = {
        "资质": ["qualification_grade", "qualification_level", "qualification", "资质等级"],
        "营业执照": ["credit_code", "business_license", "license", "营业执照号"],
        "安全生产": ["safety_license_no", "safety_license", "safety_production_license", "安全生产许可证"],
        "财务": ["financial", "财务状况"],
        "业绩": ["performance", "业绩"],
        "信誉": ["credit", "信誉"],
    }
    for keyword, keys in mapping.items():
        if keyword in title_lower:
            for key in keys:
                val = profile.get(key, "")
                if val:
                    return val
    return ""


def _build_audit_notes(requirements: TenderRequirements) -> str:
    """Build the notes volume with audit correction items from requirements.

    This generates the 审查修正说明 section with compliance responses
    derived from the parsed invalid_bid_items, qualification requirements,
    and other compliance fields.
    """
    parts = ["## 审查修正说明\n"]

    # Individual compliance items from requirements fields
    compliance_items: list[tuple[str, str]] = []

    # Project manager certificate
    if any("项目经理" in (item.title + item.description) for item in requirements.qualification_list):
        compliance_items.append((
            "project_manager_certificate",
            "在项目管理机构或人员配置章节明确项目经理注册建造师证书、专业、有效期和在岗承诺。",
        ))

    # Safety production license
    if any("安全" in (item.title + item.description) for item in requirements.qualification_list):
        compliance_items.append((
            "safety_production_license",
            "在资格响应或安全文明施工章节写明安全生产许可证有效、覆盖本项目施工活动。",
        ))

    # Bid bond
    compliance_items.append((
        "bid_bond",
        "补充投标保证金或保函提交方式、金额、有效期和到账/开具时间响应。",
    ))

    # Qualification level
    if requirements.qualification_list:
        compliance_items.append((
            "qualification_level",
            "在资格响应章节明确企业施工资质等级，并与招标文件要求保持一致。",
        ))

    # Schedule response
    if requirements.planned_duration:
        compliance_items.append((
            "schedule_response",
            "在进度计划章节补充总工期、关键节点、资源投入和延期风险控制措施。",
        ))

    # Quality response
    if requirements.quality_standard:
        compliance_items.append((
            "quality_response",
            "在质量保证章节补充质量目标、三检制、隐蔽工程验收和资料归档措施。",
        ))

    # Invalid bid items
    for idx, item in enumerate(requirements.invalid_bid_items, start=1):
        compliance_items.append((
            f"invalid_bid_item_{idx}",
            item.description or item.title,
        ))

    # Render
    for key, desc in compliance_items:
        parts.append(f"\n-针对 `{key}`：{desc}\n")

    return "\n".join(parts)


_COMMERCIAL_SECTION_KEYWORDS = (
    "资格响应",
    "投标保证金",
    "项目管理机构",
    "资格审查",
    "投标函",
    "诚信投标",
    "法定代表人",
    "授权委托",
)


def _canonical_tech_titles() -> set[str]:
    from prompts.construction_plan_outline import CONSTRUCTION_PLAN_OUTLINE

    return {re.sub(r"\s+", "", str(s["title"])) for s in CONSTRUCTION_PLAN_OUTLINE}


def _strip_commercial_sections(technical_md: str) -> str:
    """Remove commercial compliance sections that the LLM may have
    over-generated into the technical volume.

    In original format mode, commercial content (资格响应, 投标保证金,
    项目管理机构 etc.) belongs in the commercial volume, not the
    technical one. The LLM sometimes "helpfully" generates these in
    the construction plan output — strip them here.

    A canonical technical section title is never stripped, so e.g.
    "项目管理机构与岗位职责" (a legit 施工组织设计 section) is not removed by
    the "项目管理机构" commercial keyword.
    """
    tech_titles = _canonical_tech_titles()
    lines = technical_md.splitlines()
    kept: list[str] = []
    skipping = False
    for line in lines:
        stripped = line.strip()
        # Check if this line starts a ## heading that matches a commercial section
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            is_canonical_tech = re.sub(r"\s+", "", heading) in tech_titles
            if not is_canonical_tech and any(
                kw in heading for kw in _COMMERCIAL_SECTION_KEYWORDS
            ):
                skipping = True
                continue
            else:
                skipping = False
        elif stripped.startswith("# ") and not stripped.startswith("## "):
            # H1 heading — stop skipping
            skipping = False
        if not skipping:
            kept.append(line)
    return "\n".join(kept)
