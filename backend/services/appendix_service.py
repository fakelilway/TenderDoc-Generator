"""技术卷附表装配。

按招标"投标文件格式"要求的每一张附表,逐张决定来源(顺序与招标一致):
  ① 命中**公司定稿附表**(backend/assets/company_appendices/)→ 原样拼公司文件(已填好,不改格式);
  ② 否则从**招标 PDF 里该附表那一页**转可编辑表(福昕)→ 人工填,与招标一模一样;
  ③ 实在取不到 → 占位页,提示人工从招标复制该附表。
所有片段用 docxcompose 拼成一个附表 docx,由 generation_service 接到技术卷末(另起页)。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

_APPENDIX_DIR = Path(__file__).resolve().parents[1] / "assets" / "company_appendices"

# 公司定稿附表:招标附表名含关键词 → 用该文件(原样搬入,已填好)。
_COMPANY_APPENDICES: list[tuple[str, str]] = [
    ("施工总平面图", "施工总平面图.docx"),
    ("劳动力计划表", "劳动力计划表.docx"),
    ("临时占地计划表", "临时占地计划表.docx"),
    ("外供电力需求计划表", "外供电力需求计划表.docx"),
]

_CN_NUM = "一二三四五六七八九十"


def match_company_appendix(name: str) -> Path | None:
    """招标附表名命中公司定稿附表则返回其文件路径,否则 None。"""
    n = str(name or "")
    for kw, fn in _COMPANY_APPENDICES:
        if kw in n:
            p = _APPENDIX_DIR / fn
            if p.exists():
                return p
    return None


def _number_company_appendix(src: Path, num: str, name: str, out: Path) -> Path:
    """复制公司定稿附表,并给其首个标题前补"附表X "(让目录与正文页都显示招标的附表编号)。

    公司文件标题是 Heading 2、文字如"劳动力计划表"(无编号);补成"附表三 劳动力计划表",
    既进目录(带编号),正文页标题也带编号,与招标一致。原文件不动,只改副本。
    """
    from docx import Document

    doc = Document(str(src))
    heading = next(
        (p for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip()),
        None,
    )
    target = heading or next((p for p in doc.paragraphs if p.text.strip()), None)
    if target and not target.text.strip().startswith("附表"):
        prefix = f"附表{num} "
        # 补到首个有文字的 run(标题可能拆成多 run),保住原文字的格式;都没文字才退整段
        run = next((r for r in target.runs if r.text), None)
        if run is not None:
            run.text = prefix + run.text
        elif target.runs:
            target.runs[0].text = prefix + target.runs[0].text
        else:
            target.text = prefix + target.text
    doc.save(str(out))
    return out


def _find_single_appendix_page_range(
    pdf_path: str, num: str, name: str
) -> tuple[int, int] | None:
    """定位单张附表(附表X 名称)在招标 PDF 的零基右开页范围:从它所在页起,到下一张附表止。"""
    import fitz

    title_re = re.compile(re.escape(f"附表{num}") + r".{0,8}" + re.escape(name[:4]))
    any_appendix_re = re.compile(r"附表[" + _CN_NUM + r"]")
    doc = fitz.open(pdf_path)
    try:
        compacts = [re.sub(r"\s+", "", doc[i].get_text()) for i in range(doc.page_count)]
        start = next((i for i, c in enumerate(compacts) if title_re.search(c)), None)
        if start is None:
            return None
        end = doc.page_count
        for i in range(start + 1, doc.page_count):
            if any_appendix_re.search(compacts[i][:60]):
                end = i
                break
        return (start, max(start + 1, end))
    finally:
        doc.close()


def _tender_appendix_docx(tender_pdf_path: str, num: str, name: str) -> Path | None:
    """从招标 PDF 把单张附表那一页转成可编辑 docx(福昕)。失败返回 None。"""
    from core.config import get_settings
    from services.cloud_pdf_convert import (
        _slice_pdf_to_bytes,
        convert_pdf_to_docx_via_foxit,
    )

    if getattr(get_settings(), "cloud_pdf_convert", "off") != "foxit":
        return None
    try:
        pr = _find_single_appendix_page_range(tender_pdf_path, num, name)
        if not pr:
            return None
        slice_bytes = _slice_pdf_to_bytes(tender_pdf_path, pr)
        out = Path(tender_pdf_path).with_name(f"_tender_fb_{num}.docx")
        convert_pdf_to_docx_via_foxit(slice_bytes, str(out))
        return out if out.exists() else None
    except Exception:
        logger.warning("招标附表%s单张福昕转换失败", num, exc_info=True)
        return None


def _placeholder_docx(num: str, name: str, out_path: Path) -> Path:
    """占位页:招标要求但既无公司定稿表、又取不到招标原表时,提示人工从招标复制。"""
    from docx import Document

    doc = Document()
    doc.add_heading(f"附表{num} {name}", level=2)
    doc.add_paragraph(
        f"【请按招标文件中“附表{num} {name}”的格式由投标人填写，须与招标文件保持一致。】"
    )
    doc.save(str(out_path))
    return out_path


def build_appendix_docx(
    required: list[tuple[str, str]],
    out_path: str,
    tender_pdf_path: str | None = None,
) -> str | None:
    """按招标要求的附表清单(顺序)装配一个附表 docx。

    ``required``: [(附表号, 名称)],如 [("一","总体作业计划表"),("三","劳动力计划表")]。
    每张:公司定稿表 → 招标原表(福昕) → 占位。docxcompose 拼成一个 docx(片段间另起页)。
    返回 out_path;无任何附表返回 None。best-effort,单张失败用占位兜底、不抛错。
    """
    from docx import Document
    from docxcompose.composer import Composer

    if not required:
        return None

    tmp = Path(out_path).parent
    pieces: list[Path] = []
    for num, name in required:
        company = match_company_appendix(name)
        if company:
            # 公司定稿表标题无编号(如"劳动力计划表"),补成"附表三 劳动力计划表"再拼——
            # 否则目录(Word TOC 域抓标题)与正文页都只剩表名、丢了招标的"附表X"编号。
            numbered = _number_company_appendix(company, num, name, tmp / f"_co_{num}.docx")
            pieces.append(numbered)
            logger.info("附表%s %s → 公司定稿文件 %s(已补编号)", num, name, company.name)
            continue
        tender = (
            _tender_appendix_docx(tender_pdf_path, num, name) if tender_pdf_path else None
        )
        if tender:
            pieces.append(tender)
            logger.info("附表%s %s → 招标原表(福昕转)", num, name)
            continue
        pieces.append(_placeholder_docx(num, name, tmp / f"_ph_{num}.docx"))
        logger.info("附表%s %s → 占位(人工从招标复制)", num, name)

    if not pieces:
        return None

    master = Document(str(pieces[0]))
    composer = Composer(master)
    for p in pieces[1:]:
        master.add_page_break()
        composer.append(Document(str(p)))
    composer.save(str(out_path))
    return out_path
