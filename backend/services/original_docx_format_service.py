from __future__ import annotations

import logging
import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


FORMAT_CHAPTER_RE = re.compile(r"第[一二三四五六七八九十百\d]+章\s*(?:投标文件格式|响应文件格式)")
NEXT_CHAPTER_RE = re.compile(r"第[一二三四五六七八九十百\d]+章")
# A real next-chapter heading sits at the very start of a page (optionally after
# the page number), e.g. "106第八章评标办法". Anchoring here avoids treating a
# mid-sentence cross-reference ("…招标文件第二章…的要求") as a chapter boundary.
_NEXT_CHAPTER_HEAD_RE = re.compile(r"^\d{0,5}第[一二三四五六七八九十百\d]+章")
# Start of the 技术文件/报价文件 卷 inside the "投标文件格式" 章, e.g.
# "（标段名称）施工招标投标文件（技术文件）". The 商务卷 copy ends here.
_OTHER_VOLUME_START_RE = re.compile(r"投标文件\s*[（(]\s*(?:技术|报价|经济)")
# 技术卷附表区起始:页首(可带页码)就是"附表一 …"的那页才算起点,避免命中正文里
# 引用"附表一"的页或附表目录页(注:不可套 _skip_toc_pages——它会把稀疏附表误判成 TOC、丢表)。
_APPENDIX_START_RE = re.compile(r"^\d{0,5}附表一")
FORMAT_BODY_MARKERS = (
    "投标文件（商务文件）",
    "投标文件（技术文件）",
    "投标文件（报价文件）",
    "响应文件格式",
    "一、投标函",
)

# Maximum pages to scan after a chapter heading for TOC content
MAX_TOC_PAGES = 5


def build_original_format_docx(
    tender_docx_bytes: bytes,
    output_path: str | Path,
    *,
    profile: dict[str, Any] | None = None,
) -> str:
    """招标本身是 Word 时的格式章路径:原样复制 + 走同一套填充。**全程不碰福昕**。

    This is intentionally not a markdown reconstruction path. Uses copy-then-prune
    (copy the source file, remove body children outside the format chapter) so the
    chapter keeps the tender's own Word XML AND its related parts — page
    headers/footers, embedded images, numbering — which a deepcopy into a fresh
    Document would drop. Merged cells, borders, underlines, alignment, spacing all
    survive.

    复制完调用与福昕路径**完全相同**的 :func:`fill_format_docx`(填字段/填表/就地插证件
    与业绩扫描件/格式体检)。2026-07-29 之前这里只做 _replace_known_fields+清红章,
    导致 Word 招标即使被识别也产不出可用商务卷。
    """
    source = Document(BytesIO(tender_docx_bytes))
    elements = list(source.element.body)
    start = _find_format_start(elements)
    if start is None:
        raise ValueError("未能在 DOCX 招标文件中定位“投标文件格式”章节，不能原样复制。")
    end = _find_format_end(elements, start)
    keep = set(range(start, end))

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(tender_docx_bytes)  # preserves headers/footers/images/parts

    target = Document(str(path))
    body = target.element.body
    for index, child in enumerate(list(body)):
        if child.tag == qn("w:sectPr"):
            continue  # keep the body section properties (and its header/footer refs)
        if index not in keep:
            body.remove(child)
    target.save(str(path))

    return fill_format_docx(str(path), profile or {}, from_foxit=False)


_CHAPTER_TITLE_RE = re.compile(r"^第[一二三四五六七八九十百]+章投标文件格式$")


def _strip_leading_chapter_title(document: Any) -> int:
    """删掉卷首的"第X章 投标文件格式"——那是招标书自己的章标题,不是投标文件内容。

    2026-07-31 用户炸锅:成品第一页顶着"第九章 投标文件格式"("这不是吓我的头吗")。
    只看**卷首前6段**(封面之前),整段恰好是章标题才删;正文/目录里引用它的句子不受影响。
    分两段写的("第九章"+"投标文件格式")也逮;带分节符的段只清文字保留段。
    """
    paras = list(document.paragraphs)[:6]
    removed = 0
    i = 0
    while i < len(paras):
        p = paras[i]
        t = re.sub(r"[\s　]+", "", p.text)
        if not t:
            i += 1
            continue
        victims = []
        if _CHAPTER_TITLE_RE.match(t):
            victims = [p]
        elif re.fullmatch(r"第[一二三四五六七八九十百]+章", t):
            for q in paras[i + 1: i + 3]:
                if re.sub(r"[\s　]+", "", q.text) == "投标文件格式":
                    victims = [p, q]
                    break
        if not victims:
            break  # 碰到第一段实质内容(封面),收工——绝不深入正文
        for v in victims:
            el = v._p
            pPr = el.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                for run in v.runs:  # 分节符段:只清字,版式骨架留着
                    run.text = ""
            else:
                el.getparent().remove(el)
            removed += 1
        break
    if removed:
        logger.info("已删卷首招标章标题 %d 段(第X章 投标文件格式)", removed)
    return removed


def fill_format_docx(
    output_path: str, profile: dict[str, Any], *, from_foxit: bool = True
) -> str:
    """把公司真值填进格式章 docx(原地改写),返回该路径。

    这是商务卷的**唯一**填充流程,福昕路径(PDF招标)与原样复制路径(Word招标)共用它——
    两条路只在"怎么拿到格式章 docx"上不同,拿到之后填的东西一模一样。

    ``from_foxit=False``(招标本身是 Word)时**跳过所有福昕伪影矫正**:切标签归位、
    劈句合并、假两栏、幽灵图那些 healer 是专治转换伪影的,原生 Word 没这些毛病,
    跑了会帮倒忙——实测 split_paragraphs 会把"第八章 投标文件格式"和下一行
    "投标文件（商务文件）"错误并成一段。格式体检改跑只含通用项的原生版。
    """
    from docx import Document

    from services.docx_format_doctor import (
        run_format_doctor,
        run_format_doctor_native_docx,
        run_format_doctor_prefill,
    )

    doc = Document(str(output_path))
    _strip_leading_chapter_title(doc)  # 卷首"第九章 投标文件格式"是招标自己的章标题,投标文件以封面开头
    if from_foxit:
        _drop_spurious_stream_tables(doc)
        _normalize_split_labels(doc)  # 理顺福昕切开的两字标签(性 别→性别),须在填值前:标签干净才填得上
        # 填前体检:孤字归位(福昕把两列区右列标签劈成"性…"+"别："两半 → 拼回"性别："),
        # 同样须在填值前:标签完整,性别/职务这类空才填得上。
        run_format_doctor_prefill(doc)
    # 招标原件的招标人/代理红章必须在**就地插图之前**清:红占比判定的前提是"文档里
    # 还没有我们插的证件/业绩扫描件"(2026-07-12 七连修把插图挪进了填表阶段,此调用
    # 相应从卷尾前移到这,守住绝不误删证据图的红线)。
    _strip_seal_images(doc)
    # "近年完成的类似项目"六种节(投标人/项目经理/项目总工 × 资格审查/详细评审):
    # 按选派经理名下的业绩信息表记录原样填 汇总情况表+一业绩一张的详细信息表(克隆/裁剪)。
    # 必须在通用填表之前:真值先落格,总工节留白的表进 handled 名单、通用逻辑绕行。
    from services.similar_project_fill_service import fill_similar_project_sections

    similar_result = fill_similar_project_sections(doc, profile)
    # 公司组件照搬:员工做好的成品(组织结构框图/项目管理机构图)按首格锚整表替换空框,
    # 一模一样(字体实化)。须在通用填表前:成品表绝不允许再被填值。
    from services.company_component_service import fill_company_components

    fill_company_components(doc)
    _replace_known_fields(doc, profile)
    _fill_textbox_placeholders(doc, profile)  # 浮动文本框里的占位符(致：（招标人名称）等),正文替换够不着
    _fill_basic_info_subfields(doc, profile)  # 基本情况表专项:法人/技术负责人职称电话+员工总数(须在通用表格填充前)
    _fill_credit_status_table(doc, profile)  # 信誉情况表:按招标1.4.4逐条填响应(2026-07-12)
    _fill_subcontract_table(doc, profile)  # 拟分包情况表:无分包计划→填"无"并收空行(2026-07-29)
    _fill_known_table_cells(doc, profile)
    _fill_inline_labeled_blanks(doc, profile)  # 投标函内联空:工程质量/安全目标/工期/经营期限/法人联系电话
    _fill_authorization_letter(doc, profile)  # 授权委托书"本人___（姓名）系"→法人名
    _fill_establish_segmented(doc, profile)  # 法人证明"成立时间：__年__月__日"分段填
    _fill_bid_date_today(doc)  # 投标/签署日期落款 → 标书制作当天
    _fill_entrust_lines(doc, profile)  # 新款投标函"我方将委托__同志"→选派经理/总工(2026-08-05)
    _fill_personnel_table(doc, profile)
    _fill_performance_table(  # 投标人业绩情况表 → 填选中的类似业绩项目名(按节处理过的表绕行)
        doc, profile, skip_tables=similar_result.get("handled_tables")
    )
    _fill_resume_tables(  # 项目经理 + 总工简历表(各填选派那个人的台账信息)
        doc,
        profile.get("pm_resume") or {},
        profile.get("tech_resume") or {},
        profile=profile,  # 单表双人克隆+证件就地插图要回写让位标志
    )
    _attach_declared_id_scans(doc, profile)  # 表单自己写的"附:身份证扫描件"照单附上(2026-07-30)
    # 格式体检:所有值都填完后,对整份文档做"只修格式、不改文字"的合理化
    # (第一版治填空槽下划线断裂:值 run 未继承槽的下划线 → 线画一半)。
    # 传 profile 当白名单:只认我们自己填的值,防止给"年/月/日"等招标原文标签误加线。
    if from_foxit:
        run_format_doctor(doc, profile)
    else:
        run_format_doctor_native_docx(doc, profile)
    # (红章清理已前移到填表/插图之前,见上)
    _strip_tender_page_numbers(doc)  # 清招标原件页码,标书用自己的页码
    _log_unfilled_fields(doc, profile)  # 缺字段显式告警(别静默留空)
    # 商务标固定字段收尾:纠正公司名错别字(安徽正气→安徽正奇)+ 核对固定字段一致性。
    from services.commercial_fixed_fields import (
        audit_commercial_fixed_fields,
        enforce_company_name_consistency,
    )

    corrected = enforce_company_name_consistency(doc)
    if corrected:
        logger.warning("商务标公司名错别字已纠正 %d 处(安徽正气→安徽正奇)", corrected)
    for issue in audit_commercial_fixed_fields(doc):
        logger.warning("商务标固定字段核对:%s", issue)
    doc.save(str(output_path))
    return str(output_path)


def _clear_document_body(document: Document) -> None:
    body = document.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _is_toc_line(text: str) -> bool:
    """A table-of-contents entry: dot/leader run, or leaders followed by a page no."""
    return bool(
        re.search(r"[.．·…]{4,}", text)
        or re.search(r"[.．·…]{2,}\s*\d{1,4}\s*$", text)
    )


def _find_format_start(elements: list[Any]) -> int | None:
    """Locate the start of the format chapter body.

    The chapter heading appears multiple times — in the TOC, in 须知 references,
    and as the actual chapter title. Use the LAST non-TOC heading (the real
    chapter usually follows the TOC and the body references), mirroring the PDF
    path's "use the last match". Falls back to the first body-form marker.
    """
    format_headings: list[int] = []
    first_body_marker: int | None = None
    for index, element in enumerate(elements):
        text = _element_text(element)
        if not text:
            continue
        compact = re.sub(r"\s+", "", text)
        if FORMAT_CHAPTER_RE.search(compact) and not _is_toc_line(text):
            format_headings.append(index)
        if first_body_marker is None and any(
            marker in compact for marker in FORMAT_BODY_MARKERS
        ):
            first_body_marker = index
    if format_headings:
        return format_headings[-1]
    return first_body_marker


def _find_format_end(elements: list[Any], start: int) -> int:
    for index in range(start + 1, len(elements)):
        text = _element_text(elements[index])
        if not text:
            continue
        compact = re.sub(r"\s+", "", text)
        if NEXT_CHAPTER_RE.match(compact) and not FORMAT_CHAPTER_RE.search(compact):
            return index
    return len(elements)


def _element_text(element: Any) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def _replace_known_fields(document: Document, profile: dict[str, Any]) -> None:
    replacements = _known_replacements(profile)
    if not any(replacements.values()):
        return  # No non-empty replacement targets — skip iteration entirely
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
        _shrink_slot_before_grown_value(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _shrink_slot_before_grown_value(paragraph: Any, replacements: dict[str, str]) -> int:
    """token 替换让行变长时,把它前面的空白槽等量瘦身,总行宽回到模板原样。

    "系＿＿＿＿（投标人名称）的法定代表人":（投标人名称）6字被公司全名10字顶掉后
    整行多出4字宽,Word 里折行、"法定代表人。"孤零零掉到下一行(用户截图实测)。
    值写槽家族的第三处——前两处是落款日期/成立时间。规则:值比 token 长出 N 字且
    紧前面是 ≥N+2 字的纯空白 run,就从槽里砍掉 N 个空白,行宽不变;槽不够宽不砍。
    """
    shrunk = 0
    for token, value in replacements.items():
        if not value or len(value) <= len(token):
            continue
        grow = len(value) - len(token)
        runs = paragraph.runs
        for i, run in enumerate(runs):
            if value not in run.text or i == 0:
                continue
            prev = runs[i - 1]
            blank = prev.text
            if blank.strip() == "" and len(blank) >= grow + 2:
                prev.text = blank[grow:]
                shrunk += 1
            break
    return shrunk


def _fill_textbox_placeholders(document: Document, profile: dict[str, Any]) -> int:
    """浮动文本框里的占位符替换(实测:福昕把部分抬头如"致：（招标人名称）"转成浮动
    文本框,document.paragraphs 看不见 → 正文替换够不着,占位符原样留在成品里)。

    只对 w:txbxContent 内的段落跑同一套 _known_replacements 文字替换——**不动框、
    不动线条、不增删任何元素**,只把框里的占位文字换成真值。返回替换的段数。
    """
    from docx.text.paragraph import Paragraph

    replacements = {k: v for k, v in _known_replacements(profile).items() if v}
    if not replacements:
        return 0
    filled = 0
    for txbx in document.element.body.iter(qn("w:txbxContent")):
        for p_el in txbx.iter(qn("w:p")):
            para = Paragraph(p_el, document)
            before = para.text
            if not before.strip():
                continue
            _replace_in_paragraph(para, replacements)
            if para.text != before:
                filled += 1
    if filled:
        logger.info("浮动文本框占位符替换:%d 段", filled)
    return filled


def _known_replacements(profile: dict[str, Any]) -> dict[str, str]:
    project_name = str(profile.get("项目名称") or profile.get("project_name") or "")
    tenderer = str(profile.get("招标人") or profile.get("tenderer_name") or "")
    company = str(profile.get("company_name") or profile.get("投标人") or "")
    duration = str(profile.get("工期") or profile.get("planned_duration") or "")
    quality = str(profile.get("质量") or profile.get("quality_standard") or "")
    safety = str(profile.get("安全") or profile.get("safety_target") or "")
    deadline = str(profile.get("投标有效期") or profile.get("投标截止时间") or profile.get("bid_deadline") or "")
    legal_rep = str(profile.get("legal_representative") or "")
    mapping = {
        "（招标人）": tenderer,
        "（招标人名称）": tenderer,
        "受益人（招标人）名称": tenderer,
        "(招标人名称)": tenderer,
        "（招标项目名称）": project_name,
        "（项目名称）": project_name,
        "（标段名称）": project_name,
        "（标段编号）": str(profile.get("标段编号") or profile.get("section_no") or ""),
        "（投标人名称）": company,
        "（工期）": duration,
        "（计划工期）": duration,
        "（质量标准）": quality,
        "（质量要求）": quality,
        "（质量目标）": quality,
        "（安全目标）": safety,
        "（安全生产目标）": safety,
        "（投标有效期）": deadline,
        "（投标截止时间）": deadline,
        "（开标时间）": deadline,
    }
    if legal_rep:
        # 授权委托书正文"本人（姓名）系…的法定代表人"——这里的（姓名）即法人,可填。
        # 仅锚定"本人"前缀,绝不全局替换（姓名）(人员表每行（姓名）不能都填成法人)。
        mapping["本人（姓名）"] = f"本人{legal_rep}"
        # 新款身份证明"姓名：（法定代表人姓名）"(2026-08-05 外地格式试跑)
        mapping["（法定代表人姓名）"] = legal_rep
        mapping["本人 （姓名）"] = f"本人 {legal_rep}"
    return mapping


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace placeholder text while preserving per-run formatting.

    Instead of collapsing all runs into the first run (which destroys bold,
    italic, font size, etc.), this function performs replacements within
    each run individually. If a placeholder spans multiple runs, we fall
    back to paragraph-level replacement only for that specific placeholder.
    """
    if not paragraph.runs:
        return
    original = paragraph.text
    updated = original
    for source, target in replacements.items():
        if target:
            updated = updated.replace(source, target)
    if updated == original:
        return

    # Try per-run replacement first to preserve formatting
    changed = False
    for run in paragraph.runs:
        run_text = run.text
        new_text = run_text
        for source, target in replacements.items():
            if target:
                new_text = new_text.replace(source, target)
        if new_text != run_text:
            run.text = new_text
            changed = True

    if changed and paragraph.text == updated:
        return  # Per-run replacement succeeded — formatting preserved

    # Fallback: a placeholder spans multiple runs; collapse into first run
    first_run = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first_run.text = updated


# ── PDF original copy ──────────────────────────────────────────────────
# fitz (PyMuPDF) is imported lazily inside functions to avoid the ~2s module
# load overhead when only the DOCX path is used.


# Labels whose blank is segmented or date-like → never auto-fill(分段年/月/日 空)。
# 标签→值的映射统一走 _table_label_value(最长匹配 + 宽泛主体键防误填),不再用单字短键
# ("名"/"址")误把 法人/地址 乱配 —— PDF 烧录路径与可编辑表格路径共用同一套稳健映射。
_FILL_SKIP_KEYWORDS = ("成立", "日期", "经营期限", "别", "龄", "务", "年", "月", "日")


# ── 可编辑路径(pdf2docx)的表格自动填 ──────────────────────────────────
# 整页截图烧录只认下划线;可编辑路径产出真实 Word 表格(如投标人基本情况表),
# 这里按"标签单元格 → 同行右侧第一个空值格"填入公司档案。用**精确整标签**匹配
# (不用 _FILL_FIELD_ALIASES 的单字短键),避免技术负责人行的"姓名"被误填成法定
# 代表人。只写空格、不覆盖、不改表结构 → 保真。
_TABLE_FILL_LABELS: tuple[tuple[str, str], ...] = (
    ("投标人名称", "company_name"),
    ("投标人", "company_name"),
    ("注册地址", "registered_address"),
    ("单位性质", "company_type"),
    ("统一社会信用代码", "credit_code"),
    ("信用代码", "credit_code"),
    ("注册资本", "registered_capital"),
    ("企业资质等级", "qualification_grade"),
    ("资质等级", "qualification_grade"),
    ("投标人响应资质", "qualification_grade"),
    ("响应资质", "qualification_grade"),
    # 投标人基本情况表:关联企业情况(股东及股权/下属企业/同法人单位,固定字段直填)
    ("投标人关联企业情况", "affiliated_companies"),
    ("关联企业情况", "affiliated_companies"),
    ("关联企业", "affiliated_companies"),
    ("法定代表人", "legal_representative"),
    ("项目经理", "project_manager_name"),
    ("注册建造师", "project_manager_name"),
    ("安全生产许可证", "safety_license_no"),
    ("开户银行", "bank_name"),
    ("银行账号", "bank_account"),
    ("联系人", "contact_person"),
    ("经营范围", "business_scope"),
    ("成立时间", "establish_date"),
    ("成立日期", "establish_date"),
    # 扩标签别名:招标各家措辞不一(企业名称≠投标人名称…),覆盖常见同义写法,避免
    # 标签对不上而留空。靠最长匹配保证"注册地址">"注册地"、不会被短别名抢走。
    ("企业名称", "company_name"),
    ("公司名称", "company_name"),
    ("单位名称", "company_name"),
    ("投标人全称", "company_name"),
    ("法定代表", "legal_representative"),
    ("法人代表", "legal_representative"),
    ("住所", "registered_address"),
    ("注册地", "registered_address"),
    ("通讯地址", "registered_address"),
    ("联系地址", "registered_address"),
    ("联系电话", "contact_phone"),
    ("电话", "contact_phone"),
    ("手机", "contact_phone"),
    ("开户行", "bank_name"),
    ("开户银行名称", "bank_name"),
    ("账号", "bank_account"),
    # 项目级可推导字段:从招标解析得到,经 v2_generation_service.project_fields 以**中文键**
    # 并入 combined_profile(profile["项目名称"]/["工期"])。这里把表格标签映射到这些中文键,
    # 让基本情况表/投标函里的"项目名称""工期"格直接按招标值填上(不需公司档案)。
    ("项目名称", "项目名称"),
    ("工程名称", "项目名称"),
    ("标段名称", "项目名称"),
    ("项目工期", "工期"),
    ("计划工期", "工期"),
    ("总工期", "工期"),
    ("工期", "工期"),
    # 基本情况表补充字段(配合 company_profile 扩展字段;法人/技术负责人的职称电话走专项)
    ("邮政编码", "postal_code"),
    ("传真", "fax"),
    ("电子邮件", "email"),
    ("技术负责人", "tech_director_name"),
    ("高级职称人员", "senior_title_count"),
    ("中级职称人员", "mid_title_count"),
    ("初级职称人员", "junior_title_count"),
    ("技工", "technician_count"),
)
# 这些是"另一个人/需按行归属判断/格内填"的标签,通用"标签→右邻"逻辑绝不当作待填项。
# 技术职称/员工总人数 由 _fill_basic_info_subfields 专项按行上下文/格内处理(不走通用)。
_TABLE_FILL_SKIP = (
    "项目总工", "技术职称", "员工总人数",
)
# 子标签:行标签与值格之间的小标题(如 法定代表人 | 姓名 | [值]),右扫时跳过去找值格。
_TABLE_SUBLABELS = frozenset(
    {"姓名", "职称", "级别", "证号", "证书名称", "专业", "养老保险", "电话"}
)
# 留白/占位字符集。整格仅由这些组成 → 视为"空、该填"(下划线/省略号/点线签字线)。
# 关键取舍(经对抗验证):
#   · 空白/下划线/省略号 单个即算占位(它们从不作真实值用);
#   · 点线/破折号家族(. － — · 等)易与"无/不适用/小数/序号"混淆,要求连续≥2 才算占位,
#     孤立单个不判空(中文表里单格 "—"/"-" 常是人工填的"无");
#   · 句号 。(U+3002)永不判空(真标点,误判会覆盖真值);
#   · 含任何汉字/数字/字母 → 非空(绝不覆盖真值)。
_BLANK_SOLO_RE = re.compile(r"^[\s　_＿…‥]+$")
_BLANK_LEADER_CHARS = ".．·・‧․-－—–"
_BLANK_LEADER_RE = re.compile(
    rf"^[\s　_＿…‥{re.escape(_BLANK_LEADER_CHARS)}]+$"
)


def _is_blank_or_placeholder(text: str) -> bool:
    """整格仅由留白/占位字符组成 → 视为"空、该填"。详见 _BLANK_* 注释。"""
    s = (text or "").strip()
    if not s:
        return True
    if "。" in s:  # 句号是真标点,绝不当占位(防覆盖"…以上。"这类真内容)
        return False
    if _BLANK_SOLO_RE.match(s):
        return True
    if _BLANK_LEADER_RE.match(s):
        # 点线/破折号至少出现 2 个才算占位线,挡住孤立 "—"/"-"/"·"=「无/不适用」
        leader_count = sum(s.count(ch) for ch in _BLANK_LEADER_CHARS)
        return leader_count >= 2
    return False

# 宽泛"主体"键:标签里**含**它 ≠ 就该填它的名字。"投标人响应资质 / 项目经理身份证
# 号码 / …荣誉 / …业绩"含主体词但问的是别的字段——绝不能拿公司名/项目经理名去填(实测
# 122 商务卷正栽在这:项目经理身份证号被填成"江舟"、牵头人信用代码被填成公司名)。仅当
# 标签是"要名字"(键本身,或键+名称/姓名等名字尾巴)时才填,否则留空待人工。
# 注:不再按"保证金/保函"关键词留白——保证金承诺函里"投标人：(盖单位章)"、日期、邮编/电话/传真
# 都该正常填(用户实测要填);保证金金额/方式本无对应档案字段,本就填不上,无需特意留白。
_LEAVE_BLANK_KEYWORDS = ()

# "联合体协议书"整章留白(单独投标无联合体);"保函示范文本"整章留白——那是给**银行**
# 开保函用的模板(申请人/受益人/开立人及其地址电话都是银行方或业主方信息),我们无差别
# 把裸"地址/电话"填成投标人自己的=乱填他方数据(埇桥p15实测:受益人地址被填成我司地址)。
# 注意只圈"示范文本/格式如下"这类模板标题;各类承诺函有自己的编号标题、会正常退出并照填。
_BLANK_SECTION_TITLES = ("联合体协议", "联合体共同", "保函示范文本", "保函格式", "保函（格式）")
_BLANK_SECTION_END_RE = __import__("re").compile(
    r"^[（(][一二三四五六七八九十百]+[)）]"  # （七）...
    r"|^[一二三四五六七八九十百]+[、.，]"  # 七、...
    r"|^第[一二三四五六七八九十]+[章节部]"  # 第七章
    r"|^\d+[、.]"  # 7. / 7、
)


def _blank_zone_step(para_text: str, in_blank: bool) -> bool:
    """状态机:按段顺序判断当前是否在"留白章节"(联合体协议书/投标保证金)内。

    遇到留白章节标题→进入;遇到下一个编号章节标题→退出。供填充循环实时跳过
    (不用 lxml 元素 id——它跨遍历不稳定)。
    """
    t = (para_text or "").strip()
    short = 0 < len(t) < 30
    if short and any(k in t for k in _BLANK_SECTION_TITLES):
        return True
    if (
        in_blank
        and short
        and _BLANK_SECTION_END_RE.search(t)
        and not any(k in t for k in _BLANK_SECTION_TITLES)
    ):
        return False
    return in_blank


_BROAD_ENTITY_KEYS = frozenset({"投标人", "项目经理", "注册建造师"})
_NAME_TAILS = frozenset({"名称", "姓名", "名", "全称", "单位名称"})
# 判定"是否只剩名字尾巴"前,先剥掉这些主体周围的修饰词/装饰。
_ENTITY_MODIFIERS = (
    "独立", "或联合体", "联合体", "牵头人", "或", "单位", "本",
    "盖章", "公章", "签字", "（", "）", "(", ")",
)


def _label_to_profile_key(label: str) -> str:
    """标签 → 该填哪个档案字段(profile key);非可填标签返回 ""。

    取所有命中键里**最长(最具体)**的(让"统一社会信用代码"胜过宽泛"投标人",修
    first-in-list-order 误配);宽泛主体键(投标人/项目经理)只在标签"要名字"时才认,
    "投标人响应资质/项目经理身份证号"等子字段返回 ""(不乱填)。
    """
    norm = (
        label.replace(" ", "").replace("　", "").replace("：", "").replace(":", "").strip()
    )
    if not norm or any(skip in norm for skip in _TABLE_FILL_SKIP):
        return ""
    if any(kw in norm for kw in _LEAVE_BLANK_KEYWORDS):
        return ""  # 联合体/保证金 留白
    matches = [(key, pkey) for key, pkey in _TABLE_FILL_LABELS if key in norm]
    if not matches:
        return ""
    key, profile_key = max(matches, key=lambda kp: len(kp[0]))
    if key in _BROAD_ENTITY_KEYS:
        remainder = norm.replace(key, "", 1)
        for modifier in _ENTITY_MODIFIERS:
            remainder = remainder.replace(modifier, "")
        if remainder and remainder not in _NAME_TAILS:
            return ""
    return profile_key


def _table_label_value(label: str, profile: dict[str, Any]) -> str:
    profile_key = _label_to_profile_key(label)
    if not profile_key:
        return ""
    return str(profile.get(profile_key, "") or "").strip()


def unfilled_known_fields(
    document: Any, profile: dict[str, Any]
) -> list[tuple[str, str]]:
    """扫表格里"认得的标签、但公司档案无值 → 没法填"的字段,供显式告警(别静默留空)。

    返回 ``[(标签文本, profile_key), …]`` 去重。让出标前能提示"商务卷这些格因档案缺
    XX 字段没填上",而不是用户出完才发现一片空格。
    """
    seen: dict[str, str] = {}
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                profile_key = _label_to_profile_key(cell.text)
                if not profile_key or profile_key in seen:
                    continue
                if not str(profile.get(profile_key, "") or "").strip():
                    seen[profile_key] = cell.text.strip().replace("\n", " ")[:24]
    return [(label, profile_key) for profile_key, label in seen.items()]


def _log_unfilled_fields(document: Any, profile: dict[str, Any]) -> None:
    """显式告警:商务卷里认得的标签但档案无值 → 不静默留空,出标前留痕提示补档案。"""
    missing = unfilled_known_fields(document, profile)
    if missing:
        logger.warning(
            "商务卷有 %d 个可填字段因公司档案缺值未填(请在「公司档案」补全后重出):%s",
            len(missing),
            "、".join(label for label, _ in missing),
        )


def _set_cell_value(cell: Any, value: str) -> None:
    """Write a value into a table cell with a form-fitting font (宋体五号).

    新建空格默认会继承偏大的字号,在窄列(职务/证号 等)里被裁成"项目经""皖";
    显式设宋体 10.5pt 让 4 字职务/长证号能容下(过长则自动换行)。
    """
    from docx.shared import Pt

    paragraph = cell.paragraphs[0]
    runs = paragraph.runs
    lines = str(value).split("\n")
    if runs:
        run = runs[0]
        run.text = lines[0]
        for extra in runs[1:]:
            extra.text = ""
    else:
        run = paragraph.add_run(lines[0])
    # 多行值(如关联企业情况的(1)(2)(3)):同 run 内用换行符接续,保持同一套字体设置
    for extra_line in lines[1:]:
        run.add_break()
        run.add_text(extra_line)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "SimSun")
    # 允许长串(如 17 位证号)在窄列里折行,而不是溢出被裁成"皖"
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:wordWrap")) is None:
        word_wrap = OxmlElement("w:wordWrap")
        word_wrap.set(qn("w:val"), "0")
        p_pr.append(word_wrap)


def _row_label_at(
    cells: Any, i: int, profile: dict[str, Any]
) -> tuple[str, int]:
    """从 cells[i] 起识别标签,支持被逐字拆进相邻短格的"碎标签"。

    中文公文表常把"投标人："逐字对齐成 `投|标|人：` 三格,单格匹配不到。
    这里把从 i 起连续的短格(≤3 字、非空)拼起来再匹配已知标签;一旦命中即返回
    ``(value, label_end_index)``(label_end 为标签占用的最后一格下标)。未命中返回
    ``("", i)``。只匹配 ``_TABLE_FILL_LABELS`` 里已有的标签,故不会凭空造出新标签。
    """
    text = cells[i].text.strip()
    value = _table_label_value(text, profile)
    if value:
        return value, i
    # 碎标签重组:仅拼接 ≤3 字的相邻短格(标签碎片如 投/标/人:),遇到值格/长文本即停。
    joined = text
    for k in range(i + 1, min(i + 4, len(cells))):
        if cells[k]._tc is cells[i]._tc:
            break  # 合并单元格,不算相邻碎片
        nxt = cells[k].text.strip()
        if not nxt or len(nxt) > 3:
            break
        joined += nxt
        v = _table_label_value(joined, profile)
        if v:
            return v, k
    return "", i


# ── 投标人基本情况表 专项填充 ─────────────────────────────────────────────
# 这张表 9 列、含"法定代表人|姓名|值|技术职称|值|电话|值"多子标签行,通用"标签→右邻"搞不定:
# ① 法人/技术负责人行的 技术职称/电话 必须按"本行归属"取值(否则通用会把"电话"当公司
# 联系电话 contact_phone 误填);② "员工总人数："是合并标签格、值要写格内冒号后。
# 故本函数须在 _fill_known_table_cells **之前**调用,先精确占位,通用逻辑再填其余空格。
_BASIC_INFO_PERSON_ROWS: dict[str, dict[str, str]] = {
    "法定代表人": {"技术职称": "legal_rep_title", "电话": "legal_rep_phone"},
    "技术负责人": {"技术职称": "tech_director_title", "电话": "tech_director_phone"},
}


def _append_to_cell(cell: Any, value: str) -> None:
    """在单元格已有文本(如"员工总人数：")后追加值,不覆盖标签。"""
    paras = cell.paragraphs
    if paras and paras[0].runs:
        paras[0].runs[-1].text = paras[0].runs[-1].text + value
    elif paras:
        paras[0].add_run(value)


def _fill_subcontract_table(document: Any, profile: dict[str, Any] | None = None) -> int:
    """"六、拟分包项目情况表"没分包计划时按招标要求填"无",并收掉多余空行。

    招标这张表自带 19 行空数据行 + 备注列写着"若无分包计划，则投标人应在本表填写'无'"。
    正奇不分包,交出去却是整张空表、白花花占两页(2026-07-29 用户实测:"这个表你也是
    空的啊 空白巨大 然后占了两页")。规则:数据行**全空**才动手(有内容=真有分包计划,
    绝不碰),首行各列填"无"、合计行数值填"无",其余全空数据行删掉。返回处理的表数。
    """
    from docx.table import Table
    from docx.oxml.ns import qn as _qn

    handled = 0
    for table in document.tables:
        if not table.rows:
            continue
        header = [re.sub(r"[\s　]+", "", c.text) for c in table.rows[0].cells]
        if not any("拟分包的工程项目" in h for h in header):
            continue
        # 备注列(整列纵向合并写着那条注)不算数据列,不填也不判空
        note_cols = {i for i, h in enumerate(header) if h.startswith("备注")}
        data_rows = list(table.rows[1:])
        if not data_rows:
            continue
        # 合计行:首格含"合计"——它是表尾结构行,不算空数据行、不删
        def _is_total(row: Any) -> bool:
            return "合计" in re.sub(r"[\s　]+", "", row.cells[0].text)

        body_rows = [r for r in data_rows if not _is_total(r)]
        total_rows = [r for r in data_rows if _is_total(r)]

        def _row_empty(row: Any) -> bool:
            return all(
                not row.cells[i].text.strip()
                for i in range(len(row.cells))
                if i not in note_cols
            )

        if not all(_row_empty(r) for r in body_rows):
            continue  # 已填了真实分包内容 → 一个字都不动
        if not body_rows:
            continue
        first = body_rows[0]
        for i in range(len(first.cells)):
            if i not in note_cols:
                _set_cell_value(first.cells[i], "无")
        for row in total_rows:  # 合计行的数值格也填"无"(前两格是合并的标签)
            for i in range(len(row.cells)):
                if i not in note_cols and not row.cells[i].text.strip():
                    _set_cell_value(row.cells[i], "无")
        handled += 1
        logger.info("拟分包项目情况表:无分包计划,已按招标要求填“无”")
    return handled


def _fill_credit_status_table(document: Any, profile: dict[str, Any]) -> int:
    """"(五)投标人的信誉情况表"按招标1.4.4逐条填(泗沙路实测整表空白,用户定)。

    左列=招标信誉要求原文,右列=响应"无此类情形"。只对**数据行全空**的表动手
    (有预印内容=招标自带示例,不碰);行不够就克隆末行扩,多余空行留白无妨。
    返回填的行数。"""
    items = [str(x).strip() for x in (profile.get("credit_requirement_items") or []) if str(x).strip()]
    if not items:
        return 0
    from copy import deepcopy as _dc

    filled = 0
    for table in document.tables:
        if len(table.rows) < 2 or len(table.rows[0].cells) < 2:
            continue
        header = re.sub(r"[\s　]+", "", " ".join(c.text for c in table.rows[0].cells))
        if "项目" not in header or "投标人情况说明" not in header:
            continue
        data_rows = list(table.rows[1:])
        if any(any(c.text.strip() for c in r.cells) for r in data_rows):
            continue  # 有预印内容,别覆盖
        while len(table.rows) - 1 < len(items):
            table._tbl.append(_dc(table.rows[-1]._tr))
        for i, item in enumerate(items):
            row = table.rows[1 + i]
            c0, cl = row.cells[0], row.cells[-1]
            if c0._tc is cl._tc:
                # 整行合并成一格(个别模板):要求和响应写同一格,防止自己盖自己
                _set_cell_value(c0, f"{item}：无此类情形")
            else:
                _set_cell_value(c0, item)
                _set_cell_value(cl, "无此类情形")
            filled += 1
        break  # 只填第一张信誉表
    return filled


def _fill_basic_info_subfields(document: Any, profile: dict[str, Any]) -> int:
    """投标人基本情况表专项:法人/技术负责人行的 技术职称·电话(按行归属),员工总人数(格内)。"""
    filled = 0
    for table in document.tables:
        full = " ".join(c.text for row in table.rows for c in row.cells)
        if "投标人名称" not in full or "员工总人数" not in full:
            continue
        rel_done = False  # 关联企业整格重写每表只做一次(标签格竖向合并会在多行重现)
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            submap = _BASIC_INFO_PERSON_ROWS.get(cells[0].text.strip())
            if submap:
                n = len(cells)
                for i in range(1, n):
                    key = submap.get(cells[i].text.strip())
                    if not key:
                        continue
                    val = str(profile.get(key, "") or "").strip()
                    if not val:
                        continue
                    for j in range(i + 1, n):
                        if cells[j]._tc is cells[i]._tc:
                            continue
                        if _is_blank_or_placeholder(cells[j].text.strip()):
                            _set_cell_value(cells[j], val)
                            filled += 1
                        break
            emp = str(profile.get("employee_total", "") or "").strip()
            if emp:
                for cell in cells:
                    t = cell.text.strip()
                    if t.startswith("员工总人数") and ("：" in t or ":" in t):
                        after = t.split("：")[-1].split(":")[-1].strip()
                        if not after:  # 冒号后还没值才填,避免重复追加
                            _append_to_cell(cell, emp)
                            filled += 1
                        break
            # 关联企业情况格(2026-07-12 用户定稿):格里是招标预印的提示文字,且常被福昕
            # 转得叠行乱码——整格**重写**成定稿全文(含股东股权三条,分行)。预印提示不足惜:
            # 定稿文本自带完整提示+答案。只认标签行,别的格不碰。
            rel = str(profile.get("affiliated_companies", "") or "").strip()
            if rel and not rel_done and "关联企业" in cells[0].text:
                vcell = cells[-1]
                if vcell._tc is not cells[0]._tc and vcell.text.strip() != rel:
                    for extra_p in list(vcell.paragraphs)[1:]:
                        extra_p._p.getparent().remove(extra_p._p)
                    _set_cell_value(vcell, rel)
                    filled += 1
                    rel_done = True
            # 项目经理(或注册建造师)格:基本情况表此处填"人数",覆盖误填的人选名(语义=数量,非人名)
            pmc = str(profile.get("project_manager_count", "") or "").strip()
            if pmc:
                for ci, cell in enumerate(cells):
                    ct = cell.text.strip()
                    if ct.startswith("项目经理") and "职称" not in ct:
                        for j in range(ci + 1, len(cells)):
                            if cells[j]._tc is cell._tc:
                                continue
                            if cells[j].text.strip() in ("其中", "其 中"):
                                continue
                            _set_cell_value(cells[j], pmc)
                            filled += 1
                            break
                        break
    return filled


def _fill_establish_segmented(document: Any, profile: dict[str, Any]) -> int:
    """填"成立时间/日期：__年__月__日"这种分段空(法定代表人身份证明段)。

    establish_date 形如 2011-07-05 → 拆成 年=2011 / 月=7 / 日=5,分别填到"年/月/日"
    前最近的空槽(制表符/空格)。只动"成立时间/日期"标签之后、首组年月日,避免误填地址等处。
    """
    raw = str(profile.get("establish_date", "") or "").strip()
    m = re.match(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", raw)
    if not m:
        return 0
    parts = {"年": m.group(1), "月": str(int(m.group(2))), "日": str(int(m.group(3)))}
    filled = 0
    for para in document.paragraphs:
        if "成立时间" not in para.text and "成立日期" not in para.text:
            continue
        runs = para.runs
        started = False
        # 与 _fill_bid_date_today 同一个坑同一个方子(2026-07-30 用户实测身份证明页
        # 成立时间"点线和数字叠一起"):槽与"年"之间常夹一个不带下划线的纯空格 run,
        # 只记"最近的空 run"会把值写进夹缝空格 → 带线空槽原样留着、值叠在点线上。
        # 优先写**带下划线**的空槽,没有带线槽才退回普通空 run。
        last_slot = None
        last_underlined_slot = None
        for i, r in enumerate(runs):
            t = r.text
            if not started:
                if any(k in t for k in ("成立时间", "立时间", "成立日期", "立日期")):
                    started = True
                continue
            stripped = t.strip()
            if stripped == "":  # 空格/制表符槽
                last_slot = i
                if r.font.underline:
                    last_underlined_slot = i
            elif stripped[0] in parts:  # 命中 年/月/日
                slot = last_underlined_slot if last_underlined_slot is not None else last_slot
                if slot is not None and runs[slot].text.strip() == "":
                    runs[slot].text = parts[stripped[0]]
                    filled += 1
                last_slot = last_underlined_slot = None
            else:
                last_slot = last_underlined_slot = None  # 遇到其它实质文本,断开,避免跨太远误填
    return filled


# 投标日期=标书制作当天:这些是"别的日期",绝不当投标日期填(成立/有效期/截止/开标/出生/签发)。
_BID_DATE_SKIP = ("成立", "有效", "截止", "开标", "出生", "签发", "注册", "到期", "起止")


_ENTRUST_RE = re.compile(
    r"我方将委托([\s　_＿]*)同志作为本工程的(项目经理|现场总工程师|项目总工程师|总工程师|项目总工|技术负责人)"
)


def _fill_entrust_lines(document: Any, profile: dict[str, Any]) -> int:
    """新款投标函"我方将委托____同志作为本工程的项目经理/现场总工程师"(2026-08-05 外地格式)。

    空槽在"同志"之前:按句尾角色填选派的经理/总工姓名。只动空槽,已有名字不覆盖。
    """
    pm = str((profile.get("pm_resume") or {}).get("姓名") or "").strip()
    td = str((profile.get("tech_resume") or {}).get("姓名") or "").strip()
    if not pm and not td:
        return 0
    filled = 0
    for para in document.paragraphs:
        text = para.text
        m = _ENTRUST_RE.search(text)
        if not m or m.group(1).strip("_＿ 　"):
            continue
        name = pm if m.group(2) == "项目经理" else td
        if not name:
            continue
        # 字符→run 定位,把槽段替换成姓名(保留槽的前后文)
        owner: list[tuple[int, int]] = []
        runs = para.runs
        for ri, r in enumerate(runs):
            for li in range(len(r.text)):
                owner.append((ri, li))
        a, b = m.start(1), m.end(1)
        if a >= b:  # 无槽(委托和同志贴着) → 在"同志"前插入
            ri, li = owner[m.end(1)] if m.end(1) < len(owner) else (len(runs) - 1, len(runs[-1].text))
            runs[ri].text = runs[ri].text[:li] + name + runs[ri].text[li:]
        else:
            ra, la = owner[a]
            rb, lb = owner[b - 1]
            if ra == rb:
                runs[ra].text = runs[ra].text[:la] + name + runs[ra].text[lb + 1:]
            else:
                runs[ra].text = runs[ra].text[:la] + name
                for mid in range(ra + 1, rb):
                    runs[mid].text = ""
                runs[rb].text = runs[rb].text[lb + 1:]
        filled += 1
    return filled


def _fill_bid_date_today(document: Any, today: Any = None) -> int:
    """把投标/签署日期的"__年__月__日"空槽填成标书制作当天(用户定)。

    只动:① 含"日期"标签的段(排除成立/有效期等);② 整段就是"年 月 日"+空槽的落款日期行。
    成立日期另由 _fill_establish_segmented 用 establish_date 填,这里跳过,不冲突。
    """
    import datetime as _dt

    d = today or _dt.date.today()
    parts = {"年": str(d.year), "月": str(d.month), "日": str(d.day)}
    filled = 0
    in_blank = False  # 联合体协议书/保证金章节内的日期不填(实时跟踪,lxml的id不稳)
    prev_nonempty = ""  # 上一个非空段(识别福昕把"日期"劈成 '日'段+'期：…'段 的落款变体)
    for para in document.paragraphs:
        text = para.text
        # 福昕落款常见劈法:上一段只有一个"日",本段是"期： 年 月 日"——合起来才是"日期"标签
        split_date = prev_nonempty == "日" and re.sub(r"[\s　]", "", text).startswith("期")
        if text.strip():
            prev_nonempty = text.strip()
        in_blank = _blank_zone_step(text, in_blank)
        if in_blank:
            continue
        if any(s in text for s in _BID_DATE_SKIP):
            continue
        has_ymd = "年" in text and "月" in text and "日" in text
        if not has_ymd:
            continue
        residue = re.sub(r"[年月日\s　_．.、]", "", text)
        # 落款日期行:整段只剩年月日+空槽(residue空),或明确含"日期"标签
        # ("日 期："中间被塞空格也算——按去空白后的文本认);或上述"日/期："劈开变体。
        compact = re.sub(r"[\s　]", "", text)
        if not ("日期" in compact or residue == "" or (split_date and residue in ("期：", "期:"))):
            continue
        runs = para.runs

        def _is_label_ri(idx: int) -> bool:
            """runs[idx] 的首字'日'是"日 期"标签的开头(而非日期单位)——本 run 去空格后
            紧跟'期',或下一个可见 run 以'期'开头。误判会把"日"的数字填到标签前(实测 p#182)。"""
            s = re.sub(r"[\s　]", "", runs[idx].text)
            if len(s) >= 2 and s[0] == "日" and s[1] == "期":
                return True
            if s == "日":
                for j in range(idx + 1, len(runs)):
                    nxt = re.sub(r"[\s　]", "", runs[j].text)
                    if nxt:
                        return nxt[0] == "期"
            return False

        # ① 拆 run 结构(福昕转换件常见):"__"是空 run、"年"是另一 run → 填前面的空槽。
        # 槽与单位之间常夹一个**不带下划线**的纯空格 run(招标原件实测:
        # [U]'      ' + [-]' ' + '年'),只记"最近的空 run"会把值写进那个夹缝空格里 →
        # 带线空槽原样留着、整行凭空变长,右对齐落款行的字就挤成一坨(2026-07-29 用户截图)。
        # 故优先写**带下划线**的空槽,没有带线槽时才退回普通空 run。
        last_slot = None  # 最近的空 run(兜底)
        last_underlined_slot = None  # 最近的带下划线空 run(真正的填空槽)
        for i, r in enumerate(runs):
            stripped = r.text.strip()
            if stripped == "":
                last_slot = i
                if r.font.underline:
                    last_underlined_slot = i
            elif stripped[0] in parts:
                if stripped[0] == "日" and _is_label_ri(i):
                    # "日 期"标签的"日",不是单位,别往前面槽里塞数字
                    last_slot = last_underlined_slot = None
                    continue
                slot = last_underlined_slot if last_underlined_slot is not None else last_slot
                if slot is not None and runs[slot].text.strip() == "":
                    runs[slot].text = parts[stripped[0]]
                    filled += 1
                last_slot = last_underlined_slot = None
            else:
                last_slot = last_underlined_slot = None
        # ② 单 run 结构:"日期：__年_月_日"整串在一个 run → 正则就地把空槽换成今天
        for r in runs:
            t = r.text
            if not any(k in t for k in ("年", "月", "日")):
                continue
            new_t = re.sub(r"(?<![\d标])([_＿\s　]+)年", parts["年"] + "年", t)
            new_t = re.sub(r"(?<![\d])([_＿\s　]+)月", parts["月"] + "月", new_t)
            new_t = re.sub(r"(?<![\d])([_＿\s　]+)日", parts["日"] + "日", new_t)
            if new_t != t:
                r.text = new_t
                filled += 1
    return filled


# "近年完成的类似项目信息表"竖表特征标签:首列命中 ≥3 个即认定。这种表说的是
# **过去干过的工程**,项目名称/电话/项目经理都不是本项目的值——通用填表器碰它必错
# (实测:项目名称被填成本次招标、发包人电话被填成联系人手机)。识别后整表绕行,
# 由 similar_project_fill_service 按选派经理名下的业绩记录原样填。
_SIMILAR_DETAIL_KEYS = (
    "项目名称", "发包人名称", "发包人地址", "开工日期", "交工日期",
    "承担的工作", "项目描述", "监理单位",
)


def _is_similar_project_detail_table(table: Any) -> bool:
    """是否"类似项目信息表"式竖表(一业绩一张,标签|值两列)。"""
    try:
        labels = {
            re.sub(r"[\s　]+", "", row.cells[0].text)
            for row in table.rows
            if row.cells
        }
    except Exception:
        return False
    hits = sum(1 for key in _SIMILAR_DETAIL_KEYS if any(key in lb for lb in labels))
    return hits >= 3


def _fill_known_table_cells(document: Any, profile: dict[str, Any]) -> int:
    """基本情况表式网格自动填:标签格 → 同行右侧第一个空值格。

    只写空格、绝不覆盖、不改表结构(保真);未知标签和第二个人的子标签保持空白。
    支持标签被逐字拆进相邻格的"碎标签"(见 _row_label_at)。返回填入的格数。
    "类似项目信息表"竖表整表绕行(它填的是历史工程,不是本项目字段)。
    """
    if not any(profile.get(key) for _, key in _TABLE_FILL_LABELS):
        return 0
    filled = 0
    for table in document.tables:
        if _is_similar_project_detail_table(table):
            continue
        for row in table.rows:
            cells = row.cells
            n = len(cells)
            i = 0
            while i < n:
                value, label_end = _row_label_at(cells, i, profile)
                if not value:
                    i += 1
                    continue
                for j in range(label_end + 1, n):
                    target = cells[j]
                    if target._tc is cells[i]._tc:
                        continue  # 同一个合并单元格(含碎标签所在格)
                    text = target.text.strip()
                    if not _is_blank_or_placeholder(text) and _table_label_value(
                        text, profile
                    ):
                        break  # 右邻又是个已知标签,本标签不填
                    if text in _TABLE_SUBLABELS:
                        continue  # 子标签(姓名/职称等),跳过去找它后面的值格
                    if _is_blank_or_placeholder(text):
                        _set_cell_value(target, value)
                        filled += 1
                    break
                i = label_end + 1
    return filled


# 段落内联填空:"标签：<tab/下划线占位>" 这种写在正文里的空(投标函常用
# "工程质量：__，安全目标：__，工期：__日历天")。表格填空与 token 替换都不管它——
# 这是诊断早标注的"段落内联"缺口。只填白名单标签、只动 tab/占位 run、不覆盖已有值。
_INLINE_LABELS: tuple[tuple[str, str], ...] = (
    ("工程质量", "质量"),
    ("质量评定", "质量"),  # 新款投标函"交工/竣工验收的质量评定：__"(2026-08-05 外地格式试跑)
    ("质量目标", "质量"),
    ("质量要求", "质量"),
    ("质量标准", "质量"),
    ("安全目标", "安全"),
    ("安全生产目标", "安全"),
    ("计划工期", "工期"),
    ("总工期", "工期"),
    ("工期", "工期"),
    ("项目名称", "项目名称"),
    ("工程名称", "项目名称"),
    # 正文表单(非签署块)里的可填名称/地址:法定代表人身份证明"投标人名称："等。
    # ⚠️已移除裸"投标人"/"法定代表人":它们只出现在签署块（盖单位章）/（签字）行,
    # 必须留人工盖章签字(用户明确要求,撤 c274972)。法人名走"姓名："语境特判 +
    # 基本情况表的表格路径,不在这里印到签字位。
    ("投标人名称", "company_name"),
    ("企业名称", "company_name"),
    ("公司名称", "company_name"),
    ("单位名称", "company_name"),
    ("投标人", "company_name"),  # 正文表单 + 签署块"投标人：__（盖单位章）"都填公司名(名字打前、标记保留)
    ("法定代表人或其委托代理人", "legal_representative"),
    ("法定代表人", "legal_representative"),  # 含签署块"法定代表人：__（签字或盖章）",名字打前、标记保留
    ("单位性质", "company_type"),
    ("企业性质", "company_type"),
    ("注册地址", "registered_address"),
    ("通讯地址", "registered_address"),
    ("联系地址", "registered_address"),
    ("住所", "registered_address"),
    ("地址", "registered_address"),  # 裸"地 址："(最长匹配保证 注册地址>地址,不抢具体标签)
    ("投标人响应资质", "qualification_grade"),
    ("响应资质", "qualification_grade"),
    ("经营期限", "business_term"),
    ("营业期限", "business_term"),
    # 基本账户/开户许可证段的正文空(原先只在表格里能填,正文段漏填):邮编/传真/电话/邮箱。
    # 法人语境的电话/联系方式由 _inline_value_for 的 id_proof 分支优先处理,不会被这里抢。
    ("邮政编码", "postal_code"),
    ("邮编", "postal_code"),
    ("传真", "fax"),
    ("电子邮件", "email"),
    ("电子邮箱", "email"),
    ("联系电话", "contact_phone"),
    ("电话", "contact_phone"),
)
_INLINE_DELIMS = "，,。.；;、\n\r"
# 模板里"标签：__单位"已带单位时,去掉值里重复的尾随单位(工期：90日历天日历天 → 90日历天)。
_INLINE_UNITS = ("日历天", "个月", "万元", "天", "元", "%")

# 紧跟槽位的签字/盖章标记 → 该槽永远留人工(撤 c274972 代签代盖)。这条单一规则同时
# 满足:①签署块不印名 ②正文表单(无此标记)照填。
_SIGN_MARKERS = (
    "（签字）", "(签字)", "（签 字）", "（签名）",
    "（盖单位章）", "(盖单位章)", "（盖章）", "(盖章)", "（盖公章）",
    "（签章）", "(签章)", "（公章）", "(公章)",
)
# 正文表单里成组出现的同级小标签:判定"标签：[空]下一个标签："→当前槽为空、可填。
_FORM_SIBLING_LABELS = (
    "投标人名称", "法定代表人", "姓名", "性别", "年龄", "职务", "职称",
    "联系电话", "联系方式", "电话", "传真", "邮政编码", "邮编", "地址", "住所",
    "联系人", "电子邮件", "电子邮箱", "成立时间", "经营期限", "单位性质",
    "身份证号码", "身份证号",
)
# 槽位留白字符(空格/制表符/下划线/省略号/点线)。冒号后由这些组成的一段=待填槽。
_SLOT_CHARS = frozenset(" 　\t_＿…‥.．·・‧․-－—–")


def _looks_like_next_label(s: str) -> bool:
    """s 开头是否为"同级小标签 + 冒号"(如 性别：/性　别：)——判定前一个槽为空、可填。

    标签字之间常被排版拉开大空格("性　　别"),先去掉前若干字符里的空白再比对,
    否则"姓名"的填充会越过空格串进"性别",出现"许明英性别："。
    """
    # 去掉开头一段的空白(含全角)后再比对,容忍 "性  别：" 这类拉开的标签
    compact = s.replace(" ", "").replace("　", "").replace("\t", "")
    for lbl in _FORM_SIBLING_LABELS:
        if compact.startswith(lbl):
            rest = compact[len(lbl):]
            if rest[:1] in ("：", ":"):
                return True
    return False


def _iter_fillable_with_idproof(document: Any, track_idproof: bool):
    """按文档体顺序产出 (段落, 是否处于"法定代表人身份证明"章节)。含表格单元格段落。

    该章节里"姓名：性别：…职务："身份声明块的 姓名/性别/年龄=法人、可填;章节外(尤其人员表)
    绝不填,避免歧义误填(实测 122 卷"姓名"被误填法人)。以含"法定代表人身份证明"的短标题行为
    界,遇下一章节标题(协议书/保函/声明/资格审查)或 25 段/表后收口。
    **不跨遍历比对元素身份**——lxml 每次遍历新建代理、id() 不稳,故 section flag 随本遍历当场
    算出(否则段落版偶中、表格版必失效)。track_idproof=False(无法人名)时一律产出 False。
    """
    from docx.oxml.ns import qn as _qn
    from docx.table import Table as _Table
    from docx.text.paragraph import Paragraph as _Para

    inside = False
    count = 0
    for child in document.element.body.iterchildren():
        if child.tag == _qn("w:p"):
            p = _Para(child, document)
            if track_idproof:
                t = p.text.strip()
                # "授权委托书或法定代表人身份证明"是合并章标题,会把授权委托书段也圈进身份证明
                # 语境→委托代理人的身份证号/联系方式被误填成法人值。排除它,只认纯身份证明子标题。
                if "法定代表人身份证明" in t and "授权委托书" not in t and len(t) < 40:
                    inside, count = True, 0
                    yield p, True
                    continue
                if inside:
                    count += 1
                    if count > 25 or (
                        re.match(r"^[一二三四五六七八九（(]", t)
                        and any(w in t for w in ("协议书", "投标保", "声明函", "资格审查"))
                    ):
                        inside = False
            yield p, inside
        elif child.tag == _qn("w:tbl"):
            if track_idproof and inside:
                count += 1
            for row in _Table(child, document).rows:
                for cell in row.cells:
                    for cp in cell.paragraphs:
                        yield cp, inside


def _inline_value_for(
    seg: str,
    profile: dict[str, Any],
    para_text: str = "",
    id_proof_context: bool = False,
) -> str:
    """seg=冒号前刚累计的文字;其结尾是已知内联标签且档案有值 → 返回值,否则 ""。

    特例:正文"姓名："仅在法定代表人身份证明语境(本段含"的法定代表人",或处于身份证明
    章节的身份声明块)才填法人名,避免把人员表/其他人的"姓名"误填成法人。
    """
    norm = seg.replace(" ", "").replace("　", "")
    # 保证金/保函 留白:只看**标签本身**(如"投标保证金金额：")。绝不看整段——投标函正文
    # 里顺带提到"投标保证金"会把整段的质量/工期等空也误留白(实测回归)。章节级留白另由
    # _blank_zone_step 处理。
    if any(kw in norm for kw in _LEAVE_BLANK_KEYWORDS):
        return ""
    if norm.endswith("姓名") and (id_proof_context or "的法定代表人" in para_text):
        return str(profile.get("legal_representative", "") or "").strip()
    # 法定代表人身份证明表的 性别/年龄(从法人身份证 OCR 推导,见 v2._legal_rep_pii)。
    # 职务无据可填 → 不在此返回,留人工。仅在身份证明语境(章节归属 或 同段含"的法定
    # 代表人")填,避免误填人员表。
    _idp = id_proof_context or "的法定代表人" in para_text
    if _idp and norm.endswith("性别"):
        return str(profile.get("法人性别", "") or "").strip()
    if _idp and norm.endswith("年龄"):
        return str(profile.get("法人年龄", "") or "").strip()
    # 职务/联系方式/身份证号:OCR 取不到,由固定字段规则供值(见 commercial_fixed_fields)。
    # 仅在身份证明语境填,避免人员表"职务"列、委托代理人"联系方式/身份证号"被误填成法人值。
    if _idp and norm.endswith("职务"):
        return str(profile.get("法人职务", "") or "").strip()
    if _idp and (
        norm.endswith("联系方式")
        or norm.endswith("联系电话")
        or norm.endswith("手机号码")
        or norm.endswith("手机号")
    ):
        return str(profile.get("法人联系方式", "") or "").strip()
    if _idp and (norm.endswith("身份证号") or norm.endswith("身份证号码")):
        return str(profile.get("法人身份证号", "") or "").strip()
    best_label = ""
    for label, _key in _INLINE_LABELS:
        if norm.endswith(label) and len(label) > len(best_label):
            best_label = label
    if not best_label:
        return ""
    key = next(k for lbl, k in _INLINE_LABELS if lbl == best_label)
    return str(profile.get(key, "") or "").strip()


def _iter_fillable_paragraphs(document: Any):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _fill_inline_labeled_blanks(document: Any, profile: dict[str, Any]) -> int:
    """填 "标签：<槽位>" 这类段落内联空(投标函工程质量/安全/工期、法代身份证明
    投标人名称/姓名 等)。返回填入数。

    **不依赖 run 边界**:把整段拼成全文 + 记每个字符归属哪个 run,在全文上识别
    "标签：<槽位>"。槽位 = 冒号后的一段留白(tab/下划线/省略号/点线),或空槽(冒号后
    直接接分隔符/括号/下一个同级小标签)。命中后把该段在所属 run 里替成值。
    **签字/盖章守卫**:槽位后若紧跟（签字）/（盖单位章）等标记,一律跳过留人工(撤
    c274972,不代签代盖)。标签须紧贴冒号、跨分隔符作废;只填空槽、绝不覆盖真值;单位去重。
    """
    has_legal_rep = bool(str(profile.get("legal_representative") or "").strip())
    has_data = any(profile.get(key) for _lbl, key in _INLINE_LABELS) or has_legal_rep
    if not has_data:
        return 0
    filled = 0
    in_blank = False  # 联合体协议书/保证金 章节整段留白(实时跟踪)
    for paragraph, id_proof_ctx in _iter_fillable_with_idproof(document, has_legal_rep):
        runs = paragraph.runs
        if not runs:
            continue
        para_text = paragraph.text
        in_blank = _blank_zone_step(para_text, in_blank)
        if in_blank:
            continue
        # 全文 + 每字符 (run下标, run内下标)
        chars: list[str] = []
        owner: list[tuple[int, int]] = []
        for ri, run in enumerate(runs):
            for li, ch in enumerate(run.text):
                chars.append(ch)
                owner.append((ri, li))
        s = "".join(chars)
        n = len(s)
        # (run下标, run内起, run内止, 值):把 run.text[起:止] 替成值(起==止 即插入)
        edits: list[tuple[int, int, int, str]] = []
        seg = ""
        i = 0
        while i < n:
            ch = s[i]
            if ch in "：:":
                label_seg = seg  # 当前标签(冒号前),用于签字处守卫
                value = _inline_value_for(seg, profile, para_text, id_proof_ctx)
                seg = ""
                if value:
                    j = i + 1
                    while j < n and s[j] in " 　":  # 跳过冒号后的空格/全角空格
                        j += 1
                    k = j
                    while k < n and s[k] in _SLOT_CHARS:  # 吃掉留白槽(tab/下划线/点线)
                        k += 1
                    # 用户定:法定代表人签字/盖章处留白(人工签),不打名字;投标人(盖单位章)照填公司名。
                    _lbl = label_seg.replace(" ", "").replace("　", "")
                    _is_legal_rep = _lbl.endswith(("法定代表人", "法人代表", "委托代理人"))
                    _after = s[k : k + 12]
                    _is_sign_spot = any(m in _after for m in _SIGN_MARKERS) or (
                        "签字" in _after or "签名" in _after
                    )
                    if _is_legal_rep and _is_sign_spot:
                        i += 1
                        continue  # 法人签字处留白
                    # 用户拍板(签字盖章块):投标人/法定代表人在「（盖单位章）」「（签字或盖章）」前
                    # 打上公司名/姓名,标记原样保留——填的是冒号后留白槽 s[j:k],标记在 s[k:] 不动,
                    # 故不再因签字/盖章标记跳过(区别于曾被否的 c274972"代签代盖":那是把标记替换掉)。
                    # 日期(年/月/日)无标签映射,仍留空给人工。
                    had_blank = k > j
                    # 冒号后那段留白已被上面的 j 循环吃掉,所以"槽是空的"这件事要靠
                    # j 是否前进过来判断(had_blank 只在还剩下划线/点线时为真)。
                    had_whitespace = j > i + 1
                    empty_ok = (
                        k >= n
                        or s[k] in _INLINE_DELIMS
                        or s[k] in "（("
                        or _looks_like_next_label(s[k:])
                        # "工期：＿＿日历天。"——留白后紧跟单位词(日历天/万元/个月)说明
                        # 这就是个空槽,不是已填的值。2026-07-29 用户实测:同一句里质量、
                        # 安全因为槽后是"；"分隔符填上了,唯独工期因为后面跟"日历天"被
                        # 误判成"已有真实值"而跳过,交出去就是空着的。
                        or (
                            had_whitespace
                            and any(s[k:].startswith(u) for u in _INLINE_UNITS)
                        )
                    )
                    if not (had_blank or empty_ok):
                        i += 1
                        continue  # 槽位后是真实值,不覆盖
                    after = s[k : k + 14].lstrip()
                    for unit in _INLINE_UNITS:  # 模板已带单位则去重
                        if after.startswith(unit) and value.endswith(unit):
                            value = value[: -len(unit)].strip()
                            break
                    # 冒号后连排≥2个空格视为书写槽的一部分,填值时一并吃掉让值紧贴
                    # 冒号(员工反馈第2条:长空格把值推开、打乱原有版式);单个空格当
                    # 正常间隔保留。
                    lead = j - (i + 1)
                    span_start = i + 1 if lead >= 2 else j
                    # 同一行后面还跟着下一个小标签("姓名：__性别：__")时,值比槽短不能把
                    # 槽整个吃掉——否则"许明英"和"性 别"挤成一团(2026-07-30 身份证明页实测)。
                    # 槽剩多少宽度就补多少空格,行内各标签的位置纹丝不动;行尾/分隔符前的槽
                    # 不补(值后直接接标点没有挤字问题,补了反而拖长行)。
                    pad = ""
                    if _looks_like_next_label(s[k:]):
                        slot_w = k - span_start
                        if slot_w > len(value):
                            pad = " " * (slot_w - len(value))
                    value_padded = value + pad
                    if span_start == k:  # 空槽 → 原位插入值
                        if k < n:
                            ri, li = owner[k]
                        else:
                            ri, li = len(runs) - 1, len(runs[-1].text)
                        edits.append((ri, li, li, value))
                    else:  # 留白槽 → 替换 s[span_start:k]
                        rj, lj = owner[span_start]
                        rk, lk = owner[k - 1]
                        if rj == rk:
                            edits.append((rj, lj, lk + 1, value_padded))
                        else:  # 跨 run:值进首段,清掉其余留白段
                            edits.append((rj, lj, len(runs[rj].text), value_padded))
                            for mid in range(rj + 1, rk):
                                edits.append((mid, 0, len(runs[mid].text), ""))
                            edits.append((rk, 0, lk + 1, ""))
                    filled += 1
                    i = k
                    continue
            elif ch in _INLINE_DELIMS:
                seg = ""
            else:
                seg += ch
            i += 1
        # 每个 run 从后往前应用,避免前面的替换移位后面的下标
        for ri, a, b, val in sorted(edits, key=lambda e: (e[0], -e[1])):
            run = runs[ri]
            text = run.text
            run.text = text[:a] + val + text[b:]
    return filled


# 项目经理(项目技术负责人)简历表的标签 → resume 字段键。只在简历表内填,不外溢。
_PM_RESUME_LABELS: tuple[tuple[str, str], ...] = (
    ("拟在本标段", "拟任职务"),  # 拟在本标段工程担任职务
    ("毕业学校", "毕业学校"),
    ("姓名", "姓名"),
    ("年龄", "年龄"),
    ("学历", "学历"),
    ("技术职称", "职称"),  # 模板常写"技术职称",startswith("职称")够不着(2026-07-31 定稿接入)
    ("职称", "职称"),
    ("性别", "性别"),
    ("类似施工经验", "类似施工经验年限"),  # 须在裸"专业"之前——有些模板此格写"类似专业经验"
    ("工作年限", "工作年限"),
    ("获奖情况", "获奖情况"),
    ("专业", "专业"),
    ("注册建造师", "建造师证号"),     # 注册建造师执业资格证书号
    ("建造师执业资格", "建造师证号"),
    ("建造师注册证", "建造师证号"),
)


def _fill_one_resume_table(table: Any, resume: dict[str, Any]) -> int:
    """把一个人的 resume 填进一张简历表:按标签找单元格、填其右邻空格(只填空格)。返回填入格数。"""
    if not resume or not resume.get("姓名"):
        return 0
    filled = 0
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            ctext = cell.text.strip()
            key = next(
                (k for lbl, k in _PM_RESUME_LABELS if ctext.startswith(lbl) or ctext == lbl),
                None,
            )
            if not key:
                continue
            val = str(resume.get(key, "") or "").strip()
            if not val:
                continue
            for j in range(i + 1, len(cells)):
                if cells[j]._tc is cell._tc:
                    continue
                if _is_blank_or_placeholder(cells[j].text.strip()):
                    _set_cell_value(cells[j], val)
                    filled += 1
                break
    return filled


def _resume_table_role(table_text: str, heading_text: str) -> str:
    """猜一张简历表属于项目经理还是总工(技术负责人):看表内文字 + 紧邻上方标题。"""
    ctx = (heading_text or "") + " " + (table_text or "")
    if any(k in ctx for k in ("技术负责人", "总工")):
        return "tech"
    if "项目经理" in ctx:
        return "pm"
    return ""


def _fill_resume_tables(
    document: Any,
    pm_resume: dict[str, Any] | None,
    tech_resume: dict[str, Any] | None = None,
    profile: dict[str, Any] | None = None,
) -> int:
    """填项目经理 + 总工两张简历表(各含"拟在本标段")。

    按"表内/上方标题里的角色字样"把每张表分给对应的人。角色判不出的表只做两种有把握的
    推断:项目经理还没占表时首张无名表归他(单表模板常见);总工还没占表时末张无名表归他
    (总工表惯例排后)。**其余判不出的表一律留空给人工**——绝不默认填项目经理,防止把
    项目经理信息灌进其他人员的简历表(员工反馈第11条的真实事故)。
    只填空格,绝不动已填内容。返回填入格数。

    单表双人(2026-07-12 泗沙路实测):招标只给一张"项目经理和项目总工资历表"而两个角色
    都已选派 → **克隆一张**(另起一页),表1=项目经理、表2=总工;随后各自的证件扫描件
    就地插在各自表后(设 profile['_personnel_certs_inline'] 让 markdown 证件链让位)。
    """
    from copy import deepcopy as _dc

    from docx.oxml.ns import qn as _qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    pm_resume = pm_resume or {}
    tech_resume = tech_resume or {}

    # 按文档体顺序收集简历表 + 其紧邻上方标题(用于判角色);顺带留全文供"合用表"判定。
    # **"汇总表"绝不算简历表**(巢湖实测:"(八)其他管理和技术人员汇总表"表头含"拟在本
    # 标段工程任职",被当成简历表后江舟简历+证件全灌错位);"其他…人员资历表"是
    # 其他人员的地盘,经理/总工永远不许占(张冠李戴红线)。
    targets: list[tuple[Any, str]] = []
    target_texts: list[str] = []
    target_is_others: list[bool] = []
    last_heading = ""
    for child in document.element.body.iterchildren():
        if child.tag == _qn("w:p"):
            t = Paragraph(child, document).text.strip()
            if t and len(t) <= 30:
                last_heading = t
        elif child.tag == _qn("w:tbl"):
            tb = Table(child, document)
            full = " ".join(c.text for row in tb.rows for c in row.cells)
            if "拟在本标段" not in full:  # 唯一锚:只认简历表
                continue
            if "汇总" in last_heading or "汇总表" in full[:80]:
                continue  # 汇总表不是简历表
            targets.append((tb, _resume_table_role(full, last_heading)))
            target_texts.append(f"{last_heading} {full}")
            target_is_others.append("其他" in last_heading)
    if not targets:
        return 0

    has_pm = bool(pm_resume.get("姓名"))
    has_tech = bool(tech_resume.get("姓名"))

    # 经理/总工的候选表=非"其他人员"的简历表。**唯一候选是"合用表"**(标题/表文同时
    # 点名项目经理和总工,如"(五)项目经理（项目总工）简历表"/"(六)拟委任的项目经理和
    # 项目总工资历表")且两个角色都选派了 → 克隆一张:原表=项目经理,克隆=总工。
    # 专属单表(只提一个角色)绝不克隆,防第11条张冠李戴回归。
    pmtech_idx = [i for i, other in enumerate(target_is_others) if not other]
    src_text = target_texts[pmtech_idx[0]] if len(pmtech_idx) == 1 else ""
    combined_title = "项目经理" in src_text and ("总工" in src_text or "技术负责人" in src_text)
    if len(pmtech_idx) == 1 and has_pm and has_tech and combined_title:
        from docx.oxml import OxmlElement

        src_i = pmtech_idx[0]
        src_tbl = targets[src_i][0]
        note = _resume_adjacent_note_p(src_tbl._tbl)
        tail = note if note is not None else src_tbl._tbl
        new_tbl_el = _dc(src_tbl._tbl)
        pb = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pPr.append(OxmlElement("w:pageBreakBefore"))
        pb.append(pPr)
        tail.addnext(pb)
        pb.addnext(new_tbl_el)
        if note is not None:
            note_copy = _dc(note)
            for sect in note_copy.findall(f".//{qn('w:sectPr')}"):  # 剥分节符,防凭空多节
                sect.getparent().remove(sect)
            new_tbl_el.addnext(note_copy)
        keep = [
            (targets[i], target_is_others[i])
            for i in range(len(targets))
            if i != src_i
        ]
        targets = [(src_tbl, "pm"), (Table(new_tbl_el, document), "tech")] + [t for t, _ in keep]
        target_is_others = [False, False] + [o for _, o in keep]

    assign: list[tuple[Any, dict[str, Any] | None]] = []
    pm_assigned = False
    tech_assigned = False
    for tb, role in targets:
        if role == "tech":
            assign.append((tb, tech_resume if has_tech else None))
            tech_assigned = tech_assigned or has_tech
        elif role == "pm":
            assign.append((tb, pm_resume if has_pm else None))
            pm_assigned = pm_assigned or has_pm
        else:  # 判不出的表先留空,下面只做两种有把握的推断
            assign.append((tb, None))
    unassigned = [
        i for i, (_, role) in enumerate(targets)
        if role == "" and not target_is_others[i]  # 其他人员的表绝不进推断池
    ]
    # 推断①:项目经理没占到表 → 首张无名表归他(单表/表头无角色字样的模板)
    if has_pm and not pm_assigned and unassigned:
        idx = unassigned.pop(0)
        assign[idx] = (targets[idx][0], pm_resume)
    # 推断②:选了总工但没占到表 → 末张无名表归他(总工表惯例排项目经理后)
    if has_tech and not tech_assigned and unassigned:
        idx = unassigned.pop()
        assign[idx] = (targets[idx][0], tech_resume)
    # 其余无名表保持留空:宁可人工补,不张冠李戴(员工反馈第11条)

    filled = 0
    inline_roles: set[str] = set()
    certs_done: set[str] = set()  # 同一人只插一套(资格审查+详细评审双表时防重复)
    for tb, resume in assign:
        if resume and resume.get("姓名"):
            # 成品模版优先(2026-08-01 用户拍板"整表照搬"):有这个人的成品表就整表替换,
            # 拟任职务按选派角色改、经历填勾选业绩;没有才退回逐格填空白表
            tpl_tb = _replace_with_template_resume(document, tb, resume, profile)
            if tpl_tb is not None:
                tb = tpl_tb
                filled += 1
            else:
                filled += _fill_one_resume_table(tb, resume)
            # 证件扫描件**立刻**跟在本人资历表后(泗沙路实测两人证件挤一处冲突)
            name = str(resume.get("姓名"))
            if name not in certs_done:
                got = _insert_person_certs_after_table(document, tb, resume)
                if got:
                    certs_done.add(name)
                    inline_roles.add(str(resume.get("拟任职务") or "项目经理"))
    if inline_roles and profile is not None:
        # **按角色**让位:谁的证件就地插成了,才掐谁的 markdown 证件链;没占到表的
        # 角色保留锚点链兜底(否则总工没表时他整套证件会静默消失,对抗审查修正)
        profile["_personnel_certs_inline_roles"] = sorted(inline_roles)
    return filled


def _fill_template_experience_rows(
    table: Any, name: str, records: list[dict[str, Any]]
) -> int:
    """成品资历表"经 历"区的空行 ← 该人被勾选的业绩(2026-08-01 用户:"经历就填勾选的业绩")。

    列序按模版:时间 | 参加过的类似工程项目名称 | 担任职务 | 发包人及联系电话。
    担任职务按 47表 记录里他在那个工程的角色(经理/总工)填;只填空行、行不够就装多少填多少。
    """
    header_i = None
    for i, row in enumerate(table.rows):
        joined = re.sub(r"[\s　]+", "", " ".join(c.text for c in row.cells))
        if "时间" in joined and "担任职务" in joined:
            header_i = i
            break
    if header_i is None or not records:
        return 0
    filled = 0
    ri = header_i + 1
    for rec in records:
        # 找下一个可用空行(碰到 获奖/在岗/备注 就停,那是表尾结构行)
        while ri < len(table.rows):
            row = table.rows[ri]
            first = re.sub(r"[\s　]+", "", row.cells[0].text)
            if any(k in first for k in ("获奖", "在岗", "备注")):
                return filled
            if all(not c.text.strip() for c in row.cells):
                break
            ri += 1
        if ri >= len(table.rows):
            break
        row = table.rows[ri]
        cells = row.cells
        start = str(rec.get("start_date") or "").strip()
        end = str(rec.get("end_date") or "").strip()
        when = f"{start}-{end}" if (start and end) else str(rec.get("project_year") or "")
        role_in_proj = (
            "项目经理" if str(rec.get("project_manager") or "").strip() == name else "项目总工"
        )
        owner = str(rec.get("owner_name") or "").strip()
        phone = str(rec.get("owner_phone") or "").strip()
        contact = f"{owner}/{phone}" if (owner and phone) else (owner or phone)
        values = [when, str(rec.get("project_name") or ""), role_in_proj, contact]
        # 合并格去重后按列序写
        seen = []
        for c in cells:
            if not seen or seen[-1]._tc is not c._tc:
                seen.append(c)
        for cell, val in zip(seen, values):
            if val:
                _set_cell_value(cell, val)
        filled += 1
        ri += 1
    return filled


def _replace_with_template_resume(
    document: Any,
    host_tb: Any,
    resume: dict[str, Any],
    profile: dict[str, Any] | None,
) -> Any | None:
    """招标空白资历表 ← 选派人选的**成品模版表整表照搬**(2026-08-01 用户拍板)。

    有成品:整表替换(字体实化防豆腐块),"拟在本标段工程任职"改成实际选派角色,
    "经历"空行填该人被勾选的业绩,返回新 Table;没成品返回 None(调用方退回字段填空)。
    """
    name = str((resume or {}).get("姓名") or "").strip()
    role = str((resume or {}).get("拟任职务") or "").strip() or "项目经理"
    if not name:
        return None
    try:
        from docx.table import Table as _Table

        from services.company_component_service import _default_fonts_of, _solidify_fonts
        from services.curated_resume_service import get_template_table_el

        tpl_el = get_template_table_el(name)
        if tpl_el is None:
            logger.warning("选派人选「%s」没有成品资历表模版,退回字段填空(缺少简历)", name)
            return None
        # 字体实化:模版文档的默认字体不随行,不实化会在宿主里退成默认字(豆腐块风险)
        try:
            from io import BytesIO

            from core.config import settings as _st
            from utils.minio_client import minio_client as _mc
            from docx import Document as _Doc

            tpl_doc = _Doc(BytesIO(_mc.download_bytes(_st.minio_bucket, "curated/resume_templates.docx")))
            ascii_f, ea_f = _default_fonts_of(tpl_doc)
            _solidify_fonts(tpl_el, ascii_f, ea_f)
        except Exception:  # noqa: BLE001 - 实化失败不拦替换
            pass
        host_tb._tbl.addnext(tpl_el)
        host_tb._tbl.getparent().remove(host_tb._tbl)
        new_tb = _Table(tpl_el, document._body)
        # 拟任职务按实际选派角色改(模版里有人预填的是"项目总工")
        for row in new_tb.rows:
            cells = row.cells
            for i, c in enumerate(cells):
                if "拟在本标段" in re.sub(r"[\s　]+", "", c.text) and i + 1 < len(cells):
                    tail_role = "项目总工" if ("总工" in role or "技术负责" in role) else "项目经理"
                    _set_cell_value(cells[i + 1], tail_role)
                    break
        # 经历 ← 该角色勾选的业绩
        rec_key = (
            "similar_projects_td"
            if ("总工" in role or "技术负责" in role)
            else "similar_projects_pm"
        )
        records = list((profile or {}).get(rec_key) or [])
        got = _fill_template_experience_rows(new_tb, name, records)
        logger.info("资历表整表照搬:%s(%s),经历填 %d 条勾选业绩", name, role, got)
        return new_tb
    except Exception:  # noqa: BLE001
        logger.warning("资历表整表照搬失败(%s),退回字段填空", name, exc_info=True)
        return None


def _resume_adjacent_note_p(tbl_el: Any) -> Any | None:
    """资历表紧随其后的"注：…"段(中间只允许空段);没有返回 None。"""
    nxt = tbl_el.getnext()
    while nxt is not None and nxt.tag == qn("w:p"):
        text = "".join(nxt.itertext()).strip()
        if text.startswith(("注：", "注:", "注1")):
            return nxt
        if text:
            return None
        nxt = nxt.getnext()
    return None


def _insert_person_certs_after_table(document: Any, table: Any, resume: dict[str, Any]) -> int:
    """把这个人的证件扫描件就地插在他的资历表(含注)之后,返回插图张数。逐图容错。"""
    name = str((resume or {}).get("姓名") or "").strip()
    role = str((resume or {}).get("拟任职务") or "").strip() or "项目经理"
    if not name:
        return 0
    inserted = 0
    try:
        from services.generation_service import _insert_image_after
        from services.v2_generation_service import person_cert_documents

        items = person_cert_documents(role, name)
        if not items:
            return 0
        note = _resume_adjacent_note_p(table._tbl)
        from services.similar_project_fill_service import _extend_anchor_past_note_block

        anchor = _extend_anchor_past_note_block(note if note is not None else table._tbl)
        for doc_id, caption in items:
            try:
                anchor = _insert_image_after(anchor, document, doc_id, caption, 12.0)
                inserted += 1
            except Exception:
                logger.warning("资历表后插证件图失败(doc %s),跳过该图", doc_id, exc_info=True)
    except Exception:
        logger.warning("资历表后就地插证件整体失败,退回锚点链兜底", exc_info=True)
    return inserted


_ATTACH_SCAN_RE = re.compile(r"(附|本页后附)[：:]*.{0,22}身份证.{0,10}扫描件")
# 表单尾巴行(签字/日期/证号):附件要插在整个落款之后,不能楔进表单中间
_FORM_TAIL_RE = re.compile(
    r"^(投标人|法定代表人|项目经理|项目总工|委托代理人|身份证号|日期)[：:（(]"
)


def _attach_declared_id_scans(document: Any, profile: dict[str, Any]) -> int:
    """按表单里自己写的"附：××身份证正反面扫描件"就地附上扫描件,返回插图张数。

    2026-07-30 用户拍脸质问"招标文件里面说了本页后面附哪些东西,你为什么不附":
    法定代表人身份证明、授权委托书、项目经理/总工承诺书页脚都印着"附：…身份证…扫描件",
    此前这些声明只是被原样照抄、图从没插过。规则:
    - 谁的证:行里点名 法定代表人→法人;项目经理→选派经理;项目总工/技术负责人→选派总工;
      只提"代理人"的跳过(法人直投无代理人,留人工)。
    - 插哪:该表单落款(投标人/法定代表人/日期行)之后,不楔进表单中间。
    - 插什么:pick_id_card_documents 的最小覆盖(正反合一1张,否则正+反各1张),OCR判面。
    - 取不到图:不插不删,声明行原样留着(人工补),日志告警。
    """
    pm_name = str((profile.get("pm_resume") or {}).get("姓名") or "").strip()
    td_name = str((profile.get("tech_resume") or {}).get("姓名") or "").strip()
    legal = str(profile.get("legal_representative") or "").strip()

    def _owner_of(line: str) -> tuple[str, str] | None:
        t = re.sub(r"[\s　]+", "", line)
        if "代理人" in t:
            # 授权委托书的声明行(附：法定代表人身份证明+代理人身份证扫描件):法人直投
            # 无代理人,整行留人工——授权委托书后**不附**法人身份证(2026-07-30 用户拍板)
            return None
        if "法定代表人" in t:
            return ("法定代表人", legal) if legal else None
        if "项目经理" in t:
            return ("项目经理", pm_name) if pm_name else None
        if "项目总工" in t or "技术负责人" in t or "总工" in t:
            return ("项目总工", td_name) if td_name else None
        return None  # 定位不到人 → 留人工

    inserted = 0
    try:
        from services.asset_resolver import pick_id_card_documents
        from services.generation_service import _insert_image_after
    except Exception:
        logger.warning("照单附件:依赖不可用,跳过", exc_info=True)
        return 0

    paras = list(document.paragraphs)
    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text or not _ATTACH_SCAN_RE.search(re.sub(r"[\s　]+", "", text)):
            continue
        owner = _owner_of(text)
        if owner is None:
            logger.info("照单附件:声明「%s」无法定位到人(或属代理人),留人工", text[:40])
            continue
        role, name = owner
        try:
            assets = pick_id_card_documents(name)
        except Exception:
            logger.warning("照单附件:取 %s 身份证失败", name, exc_info=True)
            continue
        assets = [a for a in (assets or []) if a.get("document_id")]
        if not assets:
            logger.warning("照单附件:资料库没有 %s(%s) 的身份证图,声明行留人工", name, role)
            continue
        # 锚点:从声明行向后越过表单尾巴(签字/日期/空段),附件跟在整个表单之后。
        # 落款标签常被排版拉开("投  标  人："),必须去空白后再匹配,否则锚点提前
        # 停在声明行、图插进落款前面(2026-07-30 第6页实测)。
        anchor_p = para
        for j in range(i + 1, min(i + 10, len(paras))):
            nxt = re.sub(r"[\s　]+", "", paras[j].text)
            if not nxt or _FORM_TAIL_RE.match(nxt):
                if nxt:
                    anchor_p = paras[j]
                continue
            break
        anchor = anchor_p._p
        side_cn = {"front": "正面", "back": "背面", "both": "正反面"}
        for asset in assets:
            caption = f"{role}身份证（{side_cn.get(str(asset.get('side') or ''), '扫描件')}）"
            try:
                anchor = _insert_image_after(
                    anchor, document, int(asset["document_id"]), caption, 10.0
                )
                inserted += 1
                if role == "法定代表人":
                    # 让位标志:法人身份证已按声明就地插好,卷尾锚点链不得再插一份
                    # (老锚点还会误命中授权委托书里的"附：法定代表人身份证明"行)
                    profile["_legal_id_inline"] = True
            except Exception:
                logger.warning("照单附件:插 %s 失败,跳过该图", caption, exc_info=True)
    if inserted:
        logger.info("照单附件:按表单声明插入身份证扫描件 %d 张", inserted)
    return inserted


def _fill_pm_resume_table(document: Any, resume: dict[str, Any]) -> int:
    """向后兼容入口:只填项目经理简历(单人)。多人/总工见 _fill_resume_tables。"""
    return _fill_resume_tables(document, resume, None)


def _fill_authorization_letter(document: Any, profile: dict[str, Any]) -> int:
    """授权委托书:"本人 ___（姓名）系 ___（投标人名称）的法定代表人" 里第一个（姓名）填法人名。

    占位形态是"空白在前、括号提示在后"(___（姓名）),通用内联引擎只认"标签：___"故不覆盖。
    严格语境锚定:只在含"法定代表人"+"委托/代理人"的段、且匹配"本人…（姓名）系"才填,绝不全局
    填'（姓名）'(否则人员表/项目经理简历的姓名会被误填成法人)。委托代理人那个（姓名）留白(法人亲签)。
    (投标人名称已由 _known_replacements 的 '（投标人名称）'→company 处理,这里不碰。)
    """
    import re as _re
    from collections import defaultdict

    legal_rep = str(profile.get("legal_representative") or "").strip()
    if not legal_rep:
        return 0
    filled = 0
    for paragraph in document.paragraphs:
        runs = paragraph.runs
        if not runs:
            continue
        owner: list[tuple[int, int]] = []
        for ri, run in enumerate(runs):
            for li in range(len(run.text)):
                owner.append((ri, li))
        s = "".join(r.text for r in runs)
        if "法定代表人" not in s or ("委托" not in s and "代理人" not in s):
            continue
        m = _re.search(r"本人[ \t　_＿]*（\s*姓\s*名\s*）系", s)
        if not m:
            continue
        bg = m.start() + 2  # "本人" 之后
        paren = s.find("（", m.start())
        be = paren  # 空白槽 = s[bg:be]
        if s[bg:be].strip():  # 槽里已有实质内容(已填) → 不覆盖
            continue
        edits: list[tuple[int, int, int, str]] = []
        if be > bg:  # 有空白槽:首空白字符替成法人名,其余空白清掉
            ri0, li0 = owner[bg]
            edits.append((ri0, li0, li0 + 1, legal_rep))
            for k in range(bg + 1, be):
                rk, lk = owner[k]
                edits.append((rk, lk, lk + 1, ""))
        else:  # 无槽:插到"本人"之后
            ri0, li0 = owner[m.start() + 1]
            edits.append((ri0, li0 + 1, li0 + 1, legal_rep))
        by_run: dict[int, list[tuple[int, int, str]]] = defaultdict(list)
        for ri, a, b, v in edits:
            by_run[ri].append((a, b, v))
        for ri, ops in by_run.items():
            t = runs[ri].text
            for a, b, v in sorted(ops, reverse=True):
                t = t[:a] + v + t[b:]
            runs[ri].text = t
        filled += 1
    return filled


def _fill_personnel_table(document: Any, profile: dict[str, Any]) -> bool:
    """填"项目管理机构人员组成表"的项目经理 + 总工两行(职务/姓名/职称/证号),从选派取。

    列表头驱动的多列表(职务|姓名|职称|证书名称|级别|证号|专业|养老保险|备注)。
    只填空格、首数据行已有人则整表留人工、不改表结构。返回是否填了。
    """
    pm_name = str(profile.get("project_manager_name") or "").strip()
    tech_name = str(profile.get("tech_director_name") or "").strip()
    if not pm_name and not tech_name:
        return False
    # 要填的人:(职务, 姓名, 证号, 职称)。总工职务=项目技术负责人(用户定)。
    people: list[tuple[str, str, str, str]] = []
    if pm_name:
        people.append((
            "项目经理", pm_name,
            str(profile.get("project_manager_cert") or "").strip(),
            str(profile.get("project_manager_title") or "").strip(),
        ))
    if tech_name:
        people.append((
            "项目技术负责人", tech_name,
            str(profile.get("tech_director_cert") or "").strip(),
            str(profile.get("tech_director_title") or "").strip(),
        ))

    for table in document.tables:
        try:
            rows = table.rows
            n_cols = len(table.columns)
            if len(rows) < 2 or n_cols < 3:
                continue

            def header(col: int, _rows=rows, _table=table) -> str:
                return " ".join(
                    _table.cell(r, col).text.strip() for r in range(min(2, len(_rows)))
                )

            headers = [header(c) for c in range(n_cols)]
            col_role = next((c for c, h in enumerate(headers) if "职务" in h), None)
            col_name = next((c for c, h in enumerate(headers) if "姓名" in h), None)
            col_cert = next(
                (c for c, h in enumerate(headers) if "证号" in h or "证书号" in h), None
            )
            col_title = next((c for c, h in enumerate(headers) if "职称" in h), None)
            if col_role is None or col_name is None:
                continue
            # 必须是"人员组成"表(含证书/资格证明列),别误填别的两列表
            if not any(("证书" in h or "资格证明" in h or "证号" in h) for h in headers):
                continue
            # 执业资格证明跨行、子列在第2行 → 表头占2行,数据从第3行起
            two_row_header = any(
                ("证书名称" in headers[c] or "级别" in headers[c] or "证号" in headers[c])
                for c in range(n_cols)
            )
            data_r = 2 if two_row_header else 1
            if data_r >= len(rows):
                continue

            if not _is_blank_or_placeholder(table.cell(data_r, col_name).text.strip()):
                continue  # 首数据行已有人,整表留人工
            filled_any = False
            ri = data_r
            for role, name, cert, title in people:
                # 找下一个姓名空的行
                while ri < len(rows) and not _is_blank_or_placeholder(
                    table.cell(ri, col_name).text.strip()
                ):
                    ri += 1
                if ri >= len(rows):
                    break
                _set_cell_value(table.cell(ri, col_name), name)
                if _is_blank_or_placeholder(table.cell(ri, col_role).text.strip()):
                    _set_cell_value(table.cell(ri, col_role), role)
                if col_cert is not None and cert and _is_blank_or_placeholder(
                    table.cell(ri, col_cert).text.strip()
                ):
                    _set_cell_value(table.cell(ri, col_cert), cert)
                if col_title is not None and title and _is_blank_or_placeholder(
                    table.cell(ri, col_title).text.strip()
                ):
                    _set_cell_value(table.cell(ri, col_title), title)
                filled_any = True
                ri += 1
            if filled_any:
                return True
        except Exception:
            continue
    return False


def _fill_performance_table(
    document: Any, profile: dict[str, Any], skip_tables: set | None = None
) -> int:
    """填"投标人业绩情况表"(业绩序号|项目名称（合同名称）|备注)的项目名称列,用选中的类似业绩。

    只填空格、首数据行已有内容则整表留人工(保真)。返回填入的行数。
    选中业绩来自 profile['selected_performance'](_apply_selected_project_manager 注入,已去重)。
    skip_tables: similar_project_fill_service 已按节归属人处理过的表(w:tbl 元素),
    这里绕行——尤其总工节留白的汇总表,不能拿项目经理的业绩名去凑。
    """
    perf = profile.get("selected_performance") or []
    names: list[str] = []
    seen: set[str] = set()
    for p in perf:
        nm = str(p.get("name", "")).strip()
        if nm and nm not in seen:  # 函数内也去重(双保险,防勾重)
            seen.add(nm)
            names.append(nm)
    if not names:
        return 0

    filled = 0
    for table in document.tables:
        try:
            if skip_tables and table._tbl in skip_tables:
                continue  # 已由 similar_project_fill_service 按节归属处理(含"留白"决定)
            rows = table.rows
            n_cols = len(table.columns)
            if len(rows) < 2 or n_cols < 2:
                continue
            headers = [table.cell(0, c).text.strip() for c in range(n_cols)]
            # 必须是"业绩情况表":含"业绩序号"列 + "项目名称/合同名称"列
            if not any("业绩序号" in h for h in headers):
                continue
            col_name = next(
                (c for c, h in enumerate(headers)
                 if "项目名称" in h or "合同名称" in h or "工程名称" in h),
                None,
            )
            col_seq = next((c for c, h in enumerate(headers) if "业绩序号" in h or "序号" in h), None)
            if col_name is None:
                continue
            # 首数据行项目名称已有内容 → 留人工
            if not _is_blank_or_placeholder(table.cell(1, col_name).text.strip()):
                continue
            i = 0
            for ri in range(1, len(rows)):
                if i >= len(names):
                    break
                cell = table.cell(ri, col_name)
                t = cell.text.strip()
                if t and t not in ("……", "…", "...", "．．．"):
                    continue  # 该行已有项目名,跳过
                _set_cell_value(cell, names[i])
                # 序号列若是 …… 或空,补成序号数字
                if col_seq is not None:
                    st = table.cell(ri, col_seq).text.strip()
                    if not st or st in ("……", "…", "..."):
                        _set_cell_value(table.cell(ri, col_seq), str(i + 1))
                i += 1
                filled += 1
        except Exception:
            continue
    return filled


def _image_is_seal(blob: bytes) -> bool:
    """True if the image is a red ink seal (公章/印章).

    印章=红印泥,整图红像素占比高(实测招标人/代理公章 22%~30%);证件扫描、附表
    线框图、页面截图近 0%(黑字白底,即便角上有小红章也 <2%)。阈值 5% 干净区分。
    """
    from io import BytesIO

    try:
        from PIL import Image

        im = Image.open(BytesIO(blob)).convert("RGBA")
    except Exception:
        return False
    w, h = im.size
    if w * h == 0:
        return False
    if w * h > 14400:  # 大图抽样提速,不影响占比估计
        im = im.resize((min(w, 120), min(h, 120)))
    pixels = list(im.getdata())
    red = sum(
        1
        for r, g, b, a in pixels
        if a > 40 and r > 110 and r - g > 40 and r - b > 40
    )
    return bool(pixels) and red / len(pixels) > 0.05


def _strip_tender_page_numbers(document: Any) -> int:
    """清掉福昕从招标原件搬进来的页码:① 清空页脚 ② 删正文里"纯数字"页码段(如 171)。

    投标文件应用自己的页码,不该带招标文件的。返回清掉的元素数。
    """
    import re as _re

    cleared = 0
    # ① 页脚清空(招标原件页码常在页脚)
    try:
        for section in document.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            for para in footer.paragraphs:
                for run in list(para.runs):
                    run._r.getparent().remove(run._r)
                    cleared += 1
    except Exception:
        pass
    # ② 正文里整段只有数字(可带"第 X 页/共 X 页")的页码行 → 删
    page_re = _re.compile(r"^第?\s*\d{1,4}\s*页?(\s*/?\s*共?\s*\d{1,4}\s*页?)?$")
    for para in list(document.paragraphs):
        t = para.text.strip()
        if t and (t.isdigit() and len(t) <= 4 or page_re.match(t)):
            el = para._p
            el.getparent().remove(el)
            cleared += 1
    return cleared


def _strip_seal_images(document: Any) -> int:
    """删除从招标原件复制进来的红色印章图(招标人/代理公章)。

    投标文件不该带招标人/代理的章,投标人自己的章必须人工手盖 → 任何自动出现的印章
    都要清掉。在**格式章构建阶段**调用(此时还没追加资格证明附录,故绝不会误删证件
    扫描件)。判据见 :func:`_image_is_seal`(整图红占比)。返回删除的图章引用数。
    """
    seal_rids = set()
    for rid, part in document.part.related_parts.items():
        if "image" not in (getattr(part, "content_type", "") or ""):
            continue
        blob = getattr(part, "blob", b"")
        if blob and _image_is_seal(blob):
            seal_rids.add(rid)
    if not seal_rids:
        return 0

    removed = 0
    blip_tag, embed_attr = qn("a:blip"), qn("r:embed")
    try:
        vml_tag, vml_id = qn("v:imagedata"), qn("r:id")
    except Exception:  # pragma: no cover - 'v' 命名空间缺失时退化为只处理 drawing
        vml_tag = vml_id = None
    for container in (qn("w:drawing"), qn("w:pict")):
        for element in list(document.element.iter(container)):
            rid = next(
                (b.get(embed_attr) for b in element.iter(blip_tag)), None
            )
            if rid is None and vml_tag is not None:
                rid = next(
                    (img.get(vml_id) for img in element.iter(vml_tag)), None
                )
            if rid in seal_rids:
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
                    removed += 1
    return removed


def _table_has_real_borders(table: Any) -> bool:
    """True if any cell declares a real (non-nil) border.

    pdf2docx gives真实表格(基本情况表等)逐格 w:tcBorders=single;而把填空行误判出的
    "假表"边框为空。据此区分真表 vs 流式误判表。
    """
    sides = ("w:top", "w:bottom", "w:start", "w:end", "w:left", "w:right")
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            if tc_pr is None:
                continue
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                continue
            for side in sides:
                element = borders.find(qn(side))
                if element is not None and (
                    element.get(qn("w:val")) or ""
                ).lower() not in ("", "nil", "none"):
                    return True
    return False


def _drop_spurious_stream_tables(document: Any) -> int:
    """Flatten pdf2docx 的"假表"回普通段落。

    pdf2docx 偶尔把"标签+一段宽下划线空白"的填空行误判成 2 格无线表。这类
    (逻辑格数≤2 且 无真边框)还原成段落;真网格表(有边框,如基本情况表/图说明)保留。
    返回还原的表数。
    """
    from docx.oxml import OxmlElement

    dropped = 0
    for table in list(document.tables):
        if len(table.rows) * len(table.columns) > 2:
            continue
        if _table_has_real_borders(table):
            continue
        text = " ".join(
            cell.text.strip()
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        )
        tbl = table._tbl
        paragraph = OxmlElement("w:p")
        if text:
            run = OxmlElement("w:r")
            node = OxmlElement("w:t")
            node.set(qn("xml:space"), "preserve")
            node.text = text
            run.append(node)
            paragraph.append(run)
        tbl.addprevious(paragraph)
        tbl.getparent().remove(tbl)
        dropped += 1
    return dropped


# 福昕把招标里两字标签的字间空格照搬进来("性 别""电 话"),既丑、又害填空——填值靠"认标签",
# 标签带空格就对不上、旁边的空填不进去。把这些**已知标签**去掉字间空格,理顺回"性别""电话"。
# 只对**去空格后恰好等于已知标签**的格动手(精确等值,非包含),绝不碰投标人填的值(值几乎不会
# 恰好等于某个标签词)。放在填值之前跑 → 标签干净 → 填空更准更多。
_SPLIT_LABEL_WHITELIST = frozenset(
    {
        "性别", "年龄", "职务", "职称", "技术职称", "姓名", "电话", "传真", "联系人",
        "联系方式", "邮政编码", "电子邮件", "法定代表人", "技术负责人", "注册地址",
        "注册资本", "成立日期", "成立时间", "开户银行", "开户许可证", "账号", "账户",
        "员工总人数", "统一社会信用代码", "企业资质等级", "经营范围", "其中", "备注",
        "项目经理", "项目负责人", "高级职称人员", "中级职称人员", "初级职称人员", "技工",
        "学历", "专业", "毕业学校", "投标人名称", "投标人", "拟在本标段", "序号", "名称",
        "数量", "金额", "日期", "单位", "级别", "证号", "证书名称", "养老保险",
    }
)


def _normalize_split_labels(document: Any) -> int:
    """理顺福昕切开的两字标签("性 别"→"性别")。只动去空格后恰好=已知标签的段,值不碰。返回理顺数。"""

    def _fix_paragraph(paragraph: Any) -> int:
        text = paragraph.text
        de = re.sub(r"[\s　]+", "", text)
        if not de or de == text or de not in _SPLIT_LABEL_WHITELIST:
            return 0
        runs = paragraph.runs
        if not runs:
            return 0
        runs[0].text = de  # 保住首个 run 的格式,其余清空
        for run in runs[1:]:
            run.text = ""
        return 1

    fixed = 0
    for paragraph in document.paragraphs:
        fixed += _fix_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fixed += _fix_paragraph(paragraph)
    return fixed


def _find_format_page_range_in_pdf(pdf_path: str) -> tuple[int, int] | None:
    """Find zero-based, end-exclusive format chapter page range."""
    import fitz

    doc = fitz.open(pdf_path)
    try:
        # Extract all page texts once — avoids re-extraction per page
        page_texts: list[str] = [
            doc[page_num].get_text() for page_num in range(doc.page_count)
        ]
        page_compacts: list[str] = [re.sub(r"\s+", "", text) for text in page_texts]

        # Find ALL occurrences of the format chapter heading — use the LAST one
        # (PDFs often have a TOC reference early; the actual chapter body is later)
        matches: list[int] = []
        for page_num, compact in enumerate(page_compacts):
            if FORMAT_CHAPTER_RE.search(compact):
                matches.append(page_num)

        if not matches:
            return None

        # Use the last match as the actual chapter start
        format_start = matches[-1]
        # Skip TOC pages following the chapter heading
        format_start = _skip_toc_pages(doc, format_start)

        # Find the end: the next real chapter heading, OR the start of the
        # 技术文件/报价文件 卷 (the format chapter packs all three volumes; the
        # 商务卷 copy must stop where 技术/报价 begins — those are生成/外部, not copied).
        format_end = doc.page_count
        for page_num in range(format_start + 1, doc.page_count):
            compact = page_compacts[page_num]
            if _looks_like_next_chapter_page(compact) or _looks_like_other_volume_start(
                compact
            ):
                format_end = page_num
                break

        return (format_start, max(format_start + 1, format_end))
    finally:
        doc.close()


def _find_appendix_page_range_in_pdf(pdf_path: str) -> tuple[int, int] | None:
    """定位技术卷附表区(附表一…)的零基、右开页范围。

    附表模板落在"投标文件格式"章的技术文件段;从"附表一"起,到报价文件卷起始或下一章止。
    商务范围(_find_format_page_range_in_pdf)已在技术卷起始处收尾、不含附表,故附表单独定位。
    """
    import fitz

    doc = fitz.open(pdf_path)
    try:
        compacts = [
            re.sub(r"\s+", "", doc[i].get_text()) for i in range(doc.page_count)
        ]
        start = next(
            (i for i, c in enumerate(compacts) if _APPENDIX_START_RE.search(c)), None
        )
        if start is None:
            return None
        end = doc.page_count
        for i in range(start + 1, doc.page_count):
            if _looks_like_other_volume_start(compacts[i]) or _looks_like_next_chapter_page(
                compacts[i]
            ):
                end = i
                break
        return (start, max(start + 1, end))
    finally:
        doc.close()


def _looks_like_next_chapter_page(compact_text: str) -> bool:
    """Detect a NEW chapter heading at the START of a PDF page.

    A real chapter heading sits at the very start of the page (optionally after a
    leading page number), e.g. "106第八章评标办法". A "第X章" embedded mid-sentence
    is a cross-reference (e.g. "投标人应根据招标文件第二章…的要求") and must NOT end
    the format chapter — searching the whole page head for it falsely cut the
    chapter and dropped later 商务 form pages (实测招标#122：p105 的
    "…招标文件第二章…" 把资格审查后半 + 八、其他资料切掉)。
    """
    head = compact_text[:40]
    if not _NEXT_CHAPTER_HEAD_RE.match(head):
        return False
    return not FORMAT_CHAPTER_RE.search(head)


def _looks_like_other_volume_start(compact_text: str) -> bool:
    """Detect the start of the 技术文件/报价文件 卷 within the format chapter.

    The "投标文件格式" 章 packs all three volumes; these volume headings (not
    "第X章") mark where the 商务卷 ends — 技术卷由 LLM 生成、报价卷外部,不照抄。
    """
    return bool(_OTHER_VOLUME_START_RE.search(compact_text[:60]))


def _skip_toc_pages(pdf: fitz.Document, from_page: int) -> int:
    """Skip TOC pages after finding the format chapter heading."""
    # Check next few pages — if they're TOC (contain "........" dots), skip them
    for offset in range(MAX_TOC_PAGES):
        page_num = from_page + offset
        if page_num >= pdf.page_count:
            return from_page
        text = pdf[page_num].get_text()
        # 封面页("投标人：____（盖单位章）年月日")是投标文件正文的第一页,必须保留。
        # 它字少、无正文标记,历史逻辑把它当过场页跳掉,导致整个商务卷丢了封面(用户实测)。
        # 目录页绝不含"盖单位章",以此把封面认出来、从封面起。
        if "盖单位章" in text or "盖单位公章" in text or "盖公章" in text:
            return page_num
        # If page contains actual form content markers, it's not TOC
        if any(marker in text for marker in FORMAT_BODY_MARKERS):
            return page_num
        # If page has lots of dot leaders, it's TOC — skip
        if text.count("........") + text.count("……") + text.count("....") > 3:
            continue
        # If page has substantial unique text, it's probably content
        if len(set(text)) > 100:
            return page_num
    return from_page
