"""格式体检(format doctor):福昕转换+填值完成后,对整份 docx 做"格式合理化"修复。

设计红线(用户定,2026-07-02):**只修格式、绝不改文字内容**,且修完仍与招标原样一致。
每个 healer 只允许调整格式属性(如下划线),不得增删改任何字符;每段修完做"全段文字
逐字相同"校验;修复失败静默跳过,绝不阻断出标。

第一版 healer:治"下划线画了一半"。招标原文的填空槽是"带下划线的空白(空格/制表符)",
系统把值填进槽后,值 run 未继承下划线 → 值没线、槽两头残留短线,视觉上线断成两截。
两条规则(先白名单后兜底,都要求值紧挨着带线 run——那是填空槽的确定特征):
- 白名单:值文本(忽略空白)等于 profile 里我们自己填的值(公司名/项目名/法人名/工期…)
  → 补线;值和"："连在一个 run(如抬头"招标人名称：")→ 拆 run 只给值上线。
- 夹心兜底:「带线空白 + 无线文字 + 带线空白」的中间文字,≥6字、不含冒号、
  非"（…）"提示语 → 认作填进槽的值,补线。(短标签如"年/月/日"被线夹着是招标原样,
  长度阈值把它们挡在外面——不加这道闸就会给原文标签误加线。)

后续发现新的格式毛病,按同样契约往 _HEALERS 里加新 healer,别再散落各处打补丁。
"""

from __future__ import annotations

import logging
import re
from copy import deepcopy
from typing import Any, Callable, Iterable

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# profile 里"会被填进文档"的值的键(中英都认);其余键不当白名单,免得误伤。
_VALUE_KEYS = (
    "项目名称", "project_name",
    "招标人", "tenderer_name",
    "company_name", "投标人", "公司名称",
    "legal_representative",
    "工期", "planned_duration",
    "质量", "quality_standard",
    "安全", "safety_target",
    "投标有效期", "bid_validity",
    "注册地址", "address", "registered_address",
    "tech_director_name", "project_manager_name",
    "单位性质", "company_type",
)

_WS_RE = re.compile(r"\s+")
_HINT_RE = re.compile(r"^（.*）[。，；]?$")  # （盖单位章）（签字或盖章）等原文提示语
_LABEL_TAILS = ("：", ":")


def _norm(text: str) -> str:
    """忽略空白比对(福昕爱在数字/英文两侧塞空格)。"""
    return _WS_RE.sub("", text or "")


def fill_values_from_profile(profile: dict[str, Any] | None) -> list[str]:
    """白名单:我们自己填进文档的值(去空白,≥2字),长值在前(优先配最长)。"""
    values: set[str] = set()
    for key in _VALUE_KEYS:
        v = _norm(str((profile or {}).get(key) or ""))
        if len(v) >= 2:
            values.add(v)
    return sorted(values, key=len, reverse=True)


def _iter_all_paragraphs(document: Any) -> Iterable[Any]:
    """遍历正文 + 所有表格单元格里的段落。"""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _is_underlined(run: Any) -> bool:
    return bool(run.font.underline)


def _is_blank(run: Any) -> bool:
    return run.text != "" and run.text.strip() == ""


def _match_value_prefix(text: str, value_norm: str) -> int | None:
    """value(已去空白)是否为 text 忽略空白后的前缀;是则返回其在 text 中的结束下标。"""
    vi = 0
    for i, ch in enumerate(text):
        if ch.isspace():
            continue
        if vi < len(value_norm) and ch == value_norm[vi]:
            vi += 1
            if vi == len(value_norm):
                return i + 1
        else:
            return None
    return None


def _split_run_at(run: Any, idx: int) -> bool:
    """把 run 在字符 idx 处拆成两个 run(文字逐字保留)。仅处理单 w:t 的常规 run。"""
    r = run._r
    ts = r.findall(qn("w:t"))
    if len(ts) != 1:
        return False
    text = ts[0].text or ""
    if not 0 < idx < len(text):
        return False
    tail = deepcopy(r)
    ts[0].text = text[:idx]
    ts[0].set(qn("xml:space"), "preserve")
    tail_t = tail.findall(qn("w:t"))[0]
    tail_t.text = text[idx:]
    tail_t.set(qn("xml:space"), "preserve")
    r.addnext(tail)
    return True


def _heal_whitelist_values(paragraph: Any, values: list[str]) -> int:
    """白名单修复:无线的已填值紧挨带线 run → 补线;"值："同 run 则拆开只给值上线。"""
    healed = 0
    runs = paragraph.runs
    n = len(runs)
    for i, run in enumerate(runs):
        if _is_underlined(run) or not run.text.strip():
            continue
        prev_u = i > 0 and _is_underlined(runs[i - 1])
        next_u = i + 1 < n and _is_underlined(runs[i + 1])
        if not (prev_u or next_u):
            continue  # 不挨着线,不是填空槽
        norm = _norm(run.text)
        for v in values:
            if norm == v:
                run.font.underline = True
                healed += 1
                break
            # "值+尾巴"连在一个 run(抬头"招标人名称："、委托书"公司名的法定代表人，现委托")
            # → 拆 run 只给值上线,尾巴(招标原文)不动。值≥4字才拆,防短值误配。
            end = _match_value_prefix(run.text, v)
            if end is not None and len(v) >= 4:
                if _split_run_at(run, end):
                    run.font.underline = True  # 拆完 run 仍指向前半(值)
                    healed += 1
                break
    return healed


# 章节标题两级:大节"一、投标函…"/子表"（一）投标人基本情况表"。短行、无句读才算标题。
_SECTION_TITLE_L1_RE = re.compile(r"^[一二三四五六七八九十]+、")
_SECTION_TITLE_L2_RE = re.compile(r"^（[一二三四五六七八九十]+）")


def _title_level(text: str) -> int:
    """0=非标题;1=大节(一、);2=子表(（一）)。标题=短行(≤30字)且无句号逗号。"""
    t = text.strip()
    if not t or len(t) > 30 or any(c in t for c in "。，；"):
        return 0
    if _SECTION_TITLE_L1_RE.match(t):
        return 1
    if _SECTION_TITLE_L2_RE.match(t):
        return 2
    return 0


def heal_section_title_page_breaks(document: Any, profile: dict[str, Any] | None = None) -> int:
    """给正文章节标题(一、二、…/（一）（二）…)补"段前分页",每个新章节另起一页。

    泗沙路实测:法人身份证图后"三、联合体协议书"直接接排在同一页——原版靠福昕分节符
    分页,插图/填值撑版后标题就飘到半页腰上。补 pageBreakBefore 一劳永逸(本就在页首的
    标题加了也无副作用)。**目录区绝不能加**:目录的特征是标题行**连排**(中间没有正文/
    表格),连排 ≥3 行的一整串按目录跳过——两级目录(一、下挂（一）（二）)也逮得住;
    正文里"一、大节"紧跟"（一）子节"只连排2行,照加。返回补分页的标题数。
    """
    from docx.text.paragraph import Paragraph

    # 按文档体顺序记 (元素, 级别, 是否有内容):表格算"有内容的非标题",天然隔断连排
    items: list[tuple[Any, int, bool]] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, document).text.strip()
            items.append((child, _title_level(text), bool(text)))
        elif child.tag == qn("w:tbl"):
            items.append((child, 0, True))

    # 连排分组:相邻标题(允许中间夹空段)归同一串;正文/表格隔断。串长≥3=目录,整串跳过。
    in_toc: set[int] = set()
    run: list[int] = []
    for i, (_el, lvl, has_content) in enumerate(items):
        if lvl:
            run.append(i)
        elif has_content:
            if len(run) >= 3:
                in_toc.update(run)
            run = []
        # 空段:透明,不断串
    if len(run) >= 3:
        in_toc.update(run)

    # 紧跟在大节标题后面的第一个子节标题(中间只有空段)不再分页:大节标题会被独自
    # 撂在一张几乎全空的页上,子节内容跑到下一页——2026-07-29 用户实测"五、项目管理机构"
    # 单独占一页、组织机构图另起一页,并列为"空白页太多"的一个来源。大节和它的第一个
    # 子节本就该同页起排。
    glued: set[int] = set()
    prev_title_idx: int | None = None
    for i, (_el, lvl, has_content) in enumerate(items):
        if lvl == 1:
            prev_title_idx = i
        elif lvl > 1:
            if prev_title_idx is not None:
                glued.add(i)
            prev_title_idx = None
        elif has_content:
            prev_title_idx = None
        # 空段:透明,不打断"大节紧跟子节"的判定

    healed = 0
    for i, (el, lvl, _has) in enumerate(items):
        if not lvl or i in in_toc or i in glued or el.tag != qn("w:p"):
            continue
        pPr_el = el.find(qn("w:pPr"))
        if pPr_el is None:
            pPr_el = OxmlElement("w:pPr")
            el.insert(0, pPr_el)
        if pPr_el.find(qn("w:pageBreakBefore")) is not None:
            continue
        pb = OxmlElement("w:pageBreakBefore")
        # CT_PPr 子元素顺序:pageBreakBefore 必须排在 pStyle/keepNext/keepLines 之后
        anchor_el = None
        for tag in ("w:pStyle", "w:keepNext", "w:keepLines"):
            found = pPr_el.find(qn(tag))
            if found is not None:
                anchor_el = found
        if anchor_el is not None:
            anchor_el.addnext(pb)
        else:
            pPr_el.insert(0, pb)
        healed += 1
    return healed


# 招标模板页眉行:整行恰好是"××招标示范文本（XXXX年版）"(前缀≤20字)
_TEMPLATE_HEADER_RE = re.compile(r"^.{0,20}招标示范文本[（(]\d{4}年版[）)]$")


def heal_template_header_lines(document: Any, profile: dict[str, Any] | None = None) -> int:
    """删掉从招标 PDF 抄进来的**模板页眉行**(如"公路养护施工招标示范文本（2023年版）")。

    交通部示范文本每页顶上印着这行页眉,福昕转 Word 时把它当正文文字一页一条搬进来
    (巢湖商务卷实测40条)——投标文件不该带招标书自己的页眉。这是"绝不改文字"红线的
    **唯一例外**(2026-07-12 用户拍板"修"):删的是招标模板自己的页眉装饰,非本卷内容。
    规则从严,三道闸:① 去空白后整段**恰好**是"××招标示范文本（XXXX年版）"(正文里
    引用它的句子带上下文,不会整段匹配);② 同一文本全卷出现≥3次才动手(页眉特征=反复
    出现,防误删偶发单处引用);③ 段里带分节符的只清文字保留段(分节符动了会乱版)。

    招标本身是 Word 时,这行是**真页眉**(w:hdr 部件,巢湖实测 48 个节都有),不在正文里,
    上面三道闸够不着 → 另外扫一遍各节的页眉部件同样清掉(2026-07-29)。
    返回清掉的行数(正文行 + 页眉行)。
    """
    healed = _clear_template_headers_in_parts(document)
    body_paras = list(document.paragraphs)
    norm_of: dict[int, str] = {}
    counts: dict[str, int] = {}
    for para in body_paras:
        norm = re.sub(r"[\s　]+", "", para.text)
        if norm and _TEMPLATE_HEADER_RE.match(norm):
            norm_of[id(para)] = norm
            counts[norm] = counts.get(norm, 0) + 1

    for para in body_paras:
        norm = norm_of.get(id(para))
        if not norm or counts[norm] < 3:
            continue
        p_el = para._p
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            for run in para.runs:  # 分节符段:只清字,段和分节符留着
                run.text = ""
        else:
            p_el.getparent().remove(p_el)
        healed += 1
    return healed


def _clear_template_headers_in_parts(document: Any) -> int:
    """清各节页眉部件里的招标模板页眉行(只清文字,页眉部件本身留着不动版式)。

    只认整行**恰好**是"××招标示范文本（XXXX年版）"的,别的页眉(页码、公司自己的
    页眉)一律不碰。返回清掉的行数。
    """
    healed = 0
    for section in getattr(document, "sections", []):
        for attr in ("header", "first_page_header", "even_page_header"):
            part = getattr(section, attr, None)
            if part is None:
                continue
            try:
                paragraphs = list(part.paragraphs)
            except Exception:  # noqa: BLE001 - 缺失/损坏的页眉部件直接跳过
                continue
            for para in paragraphs:
                norm = re.sub(r"[\s　]+", "", para.text)
                if not norm or not _TEMPLATE_HEADER_RE.match(norm):
                    continue
                for run in para.runs:
                    run.text = ""
                healed += 1
    return healed


def heal_filled_value_char_squeeze(document: Any, profile: dict[str, Any] | None = None) -> int:
    """剥掉"我们填的值"所在 run 上的福昕挤压字距(2026-08-15 马鞍山#216 封面实测病)。
    病:福昕为把占位空白压进下划线槽,在 run 的 rPr 里写 w:spacing(实测 -119 =
    每字倒挤近 6pt)。槽空着时看不出来;填进"安徽正奇建设有限公司"后 10 个字互相
    叠约四成,WPS 打开就像公司名"印了两遍"(用户截图的封面双影)。
    修:run 文本命中填值白名单、且 |spacing|≥20(每字≥1pt)→ 删 spacing 属性。
    只动我们自己填的值;招标原有的拉宽标题(装饰性正字距)不在白名单,永不碰。"""
    values = fill_values_from_profile(profile)
    if not values:
        return 0
    healed = 0
    for r in document.element.body.iter(qn("w:r")):
        text = _norm("".join(t.text or "" for t in r.findall(qn("w:t"))))
        if len(text) < 2:
            continue
        if not any(v == text or (v in text and len(text) - len(v) <= 4) for v in values):
            continue
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            continue
        sp = rpr.find(qn("w:spacing"))
        if sp is None:
            continue
        try:
            val = abs(int(sp.get(qn("w:val")) or "0"))
        except ValueError:
            continue
        if val < 20:
            continue
        rpr.remove(sp)
        healed += 1
    if healed:
        logger.info("格式体检:%d 个填值 run 的福昕挤压字距已剥除(防叠字双影)", healed)
    return healed


def heal_underline_slots(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治"下划线画了一半"。先白名单,后夹心兜底;每段守"文字逐字不变"。"""
    values = fill_values_from_profile(profile)
    healed = 0
    for paragraph in _iter_all_paragraphs(document):
        before = paragraph.text
        healed += _heal_whitelist_values(paragraph, values)

        runs = paragraph.runs  # 白名单可能拆过 run,重取
        n = len(runs)
        i = 0
        while i < n:
            # 左侧:一个或多个 带线空白
            if not (_is_underlined(runs[i]) and _is_blank(runs[i])):
                i += 1
                continue
            left_end = i
            while left_end + 1 < n and _is_underlined(runs[left_end + 1]) and _is_blank(runs[left_end + 1]):
                left_end += 1
            # 中间:一个或多个 无线、有文字 的值 run
            mid_start = left_end + 1
            mid_end = mid_start - 1
            while (
                mid_end + 1 < n
                and runs[mid_end + 1].text.strip()
                and not _is_underlined(runs[mid_end + 1])
            ):
                mid_end += 1
            if mid_end < mid_start:
                i = left_end + 1
                continue
            # 右侧:必须紧跟 带线空白(槽的另一半边),才认定中间是槽里的值
            right = mid_end + 1
            if right < n and _is_underlined(runs[right]) and _is_blank(runs[right]):
                mid_text = _norm("".join(runs[k].text for k in range(mid_start, mid_end + 1)))
                # 闸:太短(年/月/日等原文标签)、带冒号(标签)、带句读(值+原文连排,如
                # "公司名的法定代表人，现委托"——整串上线会给招标原文误加线)、（…）提示语 → 不动
                if (
                    len(mid_text) >= 6
                    and not any(t in mid_text for t in _LABEL_TAILS)
                    and not any(t in mid_text for t in ("，", "。", "；", ","))
                    and not _HINT_RE.match(mid_text)
                ):
                    for k in range(mid_start, mid_end + 1):
                        runs[k].font.underline = True
                        healed += 1
                i = right
            else:
                i = mid_end + 1

        if paragraph.text != before:  # 铁律:一个字都不许变(只改 rPr/拆 run,理论到不了这)
            logger.error("格式体检:段落文字被改动,违约!before=%r after=%r", before[:50], paragraph.text[:50])
    return healed


# ── 填空前 healer:孤字归位(拼回被福昕劈成两半的两字标签) ──────────────────
# 病(真实文件实测,巢湖v7 身份证明两列区):招标原文两列布局
#   姓 名：许明英 | 性 别：女
# 福昕转完变成:左列一段"姓 名： 许明英[性]年 龄： 50[职]"(右列标签头字粘在值后),
# "别："/"务："各自孤零零成段。→ 视觉错乱,且标签残缺导致"性别/职务"的空永远填不上。
# 修:孤字(独立成 run 的单字)搬回下文"别："段首拼成"性别：",原位换成换行(恢复两行布局)。
# 全文档字符一个不多一个不少,只是把福昕劈乱的字归位;须在**填值之前**跑(标签完整才填得上)。

# 可拼回的两字标签(第一字=孤字,第二字+冒号=下文段首)。
_REJOIN_LABELS = (
    "性别", "职务", "职称", "电话", "传真", "姓名", "年龄", "学历", "专业", "备注",
)


def heal_orphan_split_labels(document: Any, profile: dict[str, Any] | None = None) -> int:
    """孤字归位:X 粘在填空槽后 + 下文段首是"Y：" 且 XY 是已知两字标签 → X 搬回 Y 前。

    "粘在槽后"的精确特征 = 孤字 run 的前一个可见 run **带下划线**——填空前是带线空白槽、
    填空后是带线的值,两种状态都算(本 healer 跑在填值前,此时槽还是空白;曾错误地额外要求
    前 run 有非空文字=已填值,导致填前永不触发,萧县实测暴露)。段首单字(如"姓 名："的
    "姓",k==0)前面没有带线 run,绝不误搬。
    """
    from docx.oxml import OxmlElement

    paras = list(_iter_all_paragraphs(document))
    healed = 0
    for idx, paragraph in enumerate(paras):
        runs = [r for r in paragraph.runs if r.text]
        for k, run in enumerate(runs):
            ch = run.text
            if len(ch) != 1 or not ("一" <= ch <= "鿿") or _is_underlined(run):
                continue
            if k == 0 or not _is_underlined(runs[k - 1]):
                continue  # 前面不是带线 run(空槽/已填值均可) → 不是"粘在槽后"的孤字
            # 在后面最多4个非空段里找"另一半":段首正是 Y：且 X+Y 是已知标签
            seen = 0
            for j in range(idx + 1, len(paras)):
                target = paras[j]
                ttext = target.text
                if not ttext.strip():
                    continue
                seen += 1
                if seen > 4:
                    break
                if len(ttext) >= 2 and ttext[1] in ("：", ":") and (ch + ttext[0]) in _REJOIN_LABELS:
                    t_run = target.runs[0] if target.runs else None
                    if t_run is None or not t_run.text.startswith(ttext[0]):
                        break
                    t_ts = t_run._r.findall(qn("w:t"))
                    if len(t_ts) != 1:
                        break
                    # 搬字:X 接到目标段首;原 run 清字、原位补换行(恢复两行布局;段尾孤字则不补)
                    t_ts[0].text = ch + (t_ts[0].text or "")
                    t_ts[0].set(qn("xml:space"), "preserve")
                    for t_el in run._r.findall(qn("w:t")):
                        run._r.remove(t_el)
                    if k < len(runs) - 1:
                        run._r.append(OxmlElement("w:br"))
                    healed += 1
                    break
    return healed


_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def heal_defloat_long_textboxes(document: Any, profile: dict[str, Any] | None = None) -> int:
    """把福昕做成"漂浮文本框"的**长条款**拆回正文流(2026-08-15 马鞍山#216实测重病)。

    病:这份招标福昕转出52个浮动框,其中整段承诺条款(如"（6）我单位承诺:若我单位拟派…")
    被钉死在页面固定位置;正文一旦因填值/排版微收而流动,就滑到钉死的框底下 → 两层文字
    叠印(用户截图,一眼致命)。
    修:含 ≥40 个非空白字符的**纯文本**浮动框 → 把框内段落原样搬到宿主段落之后(正文流,
    跟着排版走,物理上不可能再叠印),删掉浮件本体。短小浮件(致:/盖章框/日期线等定位装饰)
    不动。⚠️ mc:AlternateContent 里 Choice(drawing)+Fallback(pict) 是**同一个框的两份拷贝**,
    必须整体处理一次,否则条款会被搬两遍。字符守恒:只搬位置,不增不减。
    """
    body = document.element.body
    moved = 0

    def _box_text(holder) -> str:
        return re.sub(r"\s", "", "".join(t.text or "" for t in holder.iter(qn("w:t"))))

    def _host_paragraph(node):
        cur = node
        while cur is not None and cur.tag != qn("w:p"):
            cur = cur.getparent()
        return cur

    def _removable(node):
        """要删的浮件本体:优先删所在的 w:r(run),删不到就删自己。"""
        cur = node
        while cur is not None and cur.tag != qn("w:r"):
            cur = cur.getparent()
        return cur if cur is not None else node

    # 1) mc:AlternateContent(drawing+pict 双份拷贝)整体处理
    holders: list = list(body.iter(f"{{{_MC_NS}}}AlternateContent"))
    seen_ac = set(id(h) for h in holders)
    # 2) 不在 AlternateContent 里的裸 drawing/pict
    for tag in ("w:drawing", "w:pict"):
        for h in body.iter(qn(tag)):
            cur = h.getparent()
            inside_ac = False
            while cur is not None:
                if id(cur) in seen_ac:
                    inside_ac = True
                    break
                cur = cur.getparent()
            if not inside_ac:
                holders.append(h)

    for holder in holders:
        boxes = holder.findall(".//" + qn("w:txbxContent"))
        if not boxes:
            continue
        if holder.tag == f"{{{_MC_NS}}}AlternateContent":
            choice = holder.find(f"{{{_MC_NS}}}Choice")
            src_boxes = (choice.findall(".//" + qn("w:txbxContent")) if choice is not None else None) or boxes[:1]
        else:
            src_boxes = boxes
        if len(src_boxes) != 1:
            # 多个文本框组成的浮件=图形组合(组织机构框图等),格子文字加起来再长也不是条款,
            # 拆了会把图毁成一串散段(马鞍山#216实测:董事会/总经理…89字差点丢)——一律不碰。
            continue
        text = "".join(_box_text(b) for b in src_boxes)
        if len(text) < 40:
            continue
        if any(next(iter(b.iter(qn(t2))), None) is not None for b in boxes for t2 in ("w:drawing", "w:pict")):
            continue  # 框里还套图,不碰
        host = _host_paragraph(holder)
        if host is None or host.getparent() is None:
            continue
        insert_after = host
        for b in src_boxes:
            for p_el in list(b.findall(qn("w:p"))):
                newp = deepcopy(p_el)
                # 拆回正文流后清掉框内段落可能带的绝对定位属性
                pPr = newp.find(qn("w:pPr"))
                if pPr is not None:
                    fp = pPr.find(qn("w:framePr"))
                    if fp is not None:
                        pPr.remove(fp)
                insert_after.addnext(newp)
                insert_after = newp
        target = _removable(holder)
        if target.getparent() is not None:
            target.getparent().remove(target)
            moved += 1
    if moved:
        logger.info("格式体检(填前):%d 个长条款浮动框已拆回正文流(防叠印)", moved)
    return moved


_PREFILL_HEALERS: tuple[tuple[str, Callable[[Any, dict[str, Any] | None], int]], ...] = (
    ("defloat_long_textboxes", heal_defloat_long_textboxes),
    ("orphan_split_labels", heal_orphan_split_labels),
)

def heal_filler_blank_runs(document: Any, profile: dict[str, Any] | None = None) -> int:
    """压缩福昕塞的大段连续空段(为凑原 PDF 版面塞的填充回车)。

    实测(用户截图):表单签名后跟着 20+ 个空段,而下一节本就由"分节符(下一页)"另起新页
    → 这些空段自己占满了一整页纯空白。规则:连续 ≥4 个**真空段**(无文字/无图/无换行符/
    无制表位/不带分节符——日期空槽那类"只有下划线制表位"的段绝不算空)压缩到剩 2 个。
    删的是零字符的空段,不碰任何内容。返回删除的空段数。
    """
    body = document.element.body

    def _is_truly_blank(p_el) -> bool:
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            return False  # 分节符段,动不得
        for tag in ("w:drawing", "w:pict", "w:br", "w:tab"):
            if next(iter(p_el.iter(qn(tag))), None) is not None:
                return False
        for t in p_el.iter(qn("w:t")):
            if (t.text or "").strip():
                return False
        return True

    removed = 0
    run: list = []
    for child in list(body) + [None]:
        if child is not None and child.tag == qn("w:p") and _is_truly_blank(child):
            run.append(child)
            continue
        if len(run) >= 4:
            for extra in run[2:]:  # 留2个保分隔感
                body.remove(extra)
                removed += 1
        run = []
    if removed:
        logger.info("格式体检:压缩福昕填充空段 %d 个", removed)
    return removed


def _iter_all_paragraphs_deep(container: Any):
    """正文 + 所有表格(含嵌套)单元格里的段落。福昕的表单常有嵌套表,须递归。"""
    for p in container.paragraphs:
        yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs_deep(cell)


def heal_line_spacing(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治福昕的"叠行/行距忽紧忽松"。只改行距属性,一个字不动。

    福昕转 PDF→Word 时给部分段落设了"固定行高(lineRule=exact)+行高小于单倍"
    → 固定行高比字还矮,多行文字被压得叠在一起(用户实测:联合体协议书/投标函密集段
    第二行贴着第一行)。规则:
    · 固定行高(exact) → 改成自动行高(auto);行高值不足单倍(240)的提到 240,保证不叠;
    · 过松(auto 且 line>300,即 >1.25 倍)的正文行 → 收回单倍,消除忽紧忽松。
    行距是纯格式属性,改它不违反"绝不改文字"红线。返回调整的段数。
    """
    fixed = 0
    for p in _iter_all_paragraphs_deep(document):
        ppr = p._p.find(qn("w:pPr"))
        if ppr is None:
            continue
        sp = ppr.find(qn("w:spacing"))
        if sp is None:
            continue
        rule = sp.get(qn("w:lineRule"))
        line = sp.get(qn("w:line"))
        line_i = int(line) if (line and line.isdigit()) else None
        if rule == "exact":
            sp.set(qn("w:lineRule"), "auto")
            if line_i is not None and line_i < 240:
                sp.set(qn("w:line"), "240")
            fixed += 1
        elif rule == "auto" and line_i is not None and line_i > 300:
            sp.set(qn("w:line"), "240")
            fixed += 1
    if fixed:
        logger.info("格式体检:捋匀福昕叠行/过松行距 %d 段", fixed)
    return fixed


def _usable_width_twips(document: Any) -> int:
    """页面文字区宽度(twips)=页宽-左右边距。多节取最小(最保守);拿不到用 A4 常见值兜底。"""
    widths = []
    try:
        for sec in document.sections:
            w = int(sec.page_width.twips - sec.left_margin.twips - sec.right_margin.twips)
            if w > 2000:
                widths.append(w)
    except Exception:
        pass
    return min(widths) if widths else 9026  # A4(11906) - 边距(1440*2)


def _est_text_twips(text: str) -> int:
    """粗估文本渲染宽度(twips):CJK≈290/字,西文数字≈145/字。**故意估宽**——
    估窄了会再折行(实测),估宽了只是缩进/下划线槽略短,右对齐观感无碍。"""
    return sum(290 if ord(ch) > 0x2E80 else 145 for ch in text if ch != "\t")


def heal_signature_wrap(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治落款块折行:「投标人：公司名（盖单位章）」「（签章）」被挤到下一行。

    根因(用户实测,埇桥商务卷):福昕为复刻原 PDF 落款靠右的位置,给这些段落设了
    巨大左缩进(实测 5486 twips≈9.7cm)+右对齐。招标原样是空下划线,短,排得下;
    值(公司名/28字地址)填进去后超宽 → "（盖单位章）""（签章）"折行、地址断三行。

    修法:这些段落是**右对齐**的——左缩进砍小后文字自动贴右、观感不变,但一行排得下。
    制表位画得太远(槽尾+"（签章）"超页宽)的,把制表位收进来。只动缩进/制表位,
    文字一个不动(红线)。返回调整的段数。
    """
    avail = _usable_width_twips(document)
    fixed = 0
    for p in document.paragraphs:
        text = p.text
        if not text.strip():
            continue
        ppr = p._p.find(qn("w:pPr"))
        if ppr is None:
            continue
        jc = ppr.find(qn("w:jc"))
        if jc is None or jc.get(qn("w:val")) not in ("right", "end"):
            continue
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            continue
        try:
            left = int(ind.get(qn("w:left")) or 0)
            right = int(ind.get(qn("w:right")) or 0)
        except ValueError:
            continue
        if left < 1500:
            continue  # 没有福昕式大缩进,不是本病
        # 行需要的宽度:有制表位时 = 最远制表位 + 制表位之后的文字;否则 = 全文
        tabs_el = ppr.find(qn("w:tabs"))
        tab_positions = []
        if tabs_el is not None:
            for t in tabs_el.findall(qn("w:tab")):
                try:
                    tab_positions.append(int(t.get(qn("w:pos")) or 0))
                except ValueError:
                    pass
        if "\t" in text and tab_positions:
            tail = text.rsplit("\t", 1)[1]
            need = max(tab_positions) + _est_text_twips(tail)
        else:
            need = _est_text_twips(text)
        if left + need + right <= avail - 60:
            continue  # 排得下,原样保真
        # 折行了:左缩进直接砍到 0。段落是右对齐,砍缩进只是给足排版空间,
        # 文字照样贴右、观感不变;字宽是估算值,留一部分缩进赌"刚好够"会再折(实测)。
        ind.set(qn("w:left"), "0")
        # 左缩进归零还不够(制表位自己画出页外) → 把最远制表位收进来,余量给足
        if need + right > avail - 60 and tab_positions and "\t" in text:
            tail = text.rsplit("\t", 1)[1]
            head = text.rsplit("\t", 1)[0]
            new_pos = max(
                _est_text_twips(head) + 300,
                avail - right - _est_text_twips(tail) - 400,
            )
            farthest = max(
                tabs_el.findall(qn("w:tab")),
                key=lambda t: int(t.get(qn("w:pos")) or 0),
            )
            if new_pos < int(farthest.get(qn("w:pos")) or 0):
                farthest.set(qn("w:pos"), str(new_pos))
        fixed += 1
    if fixed:
        logger.info("格式体检:治落款折行(砍福昕大缩进/收制表位) %d 段", fixed)
    return fixed


# 句子在这些字符上收尾才算"说完了";其余结尾(逗号/裸汉字/"在")都是被福昕劈开的半句
_SENTENCE_END = "。？！；;：…" + "”\"』】)）"
# 新条目/新表单行的开头模式:绝不能把它并进上一段
_NEW_ITEM_RE = re.compile(
    r"^\s*(?:[0-9０-９]+\s*[．.、)）]|（[0-9０-９一二三四五六七八九十]+）"
    r"|[一二三四五六七八九十]+\s*[、.．]|附[表件]|第[一二三四五六七八九十百0-9]+[章节条部]"
    r"|[注致][：:]|投\s*标\s*人|法\s*定\s*代\s*表\s*人|地\s*址|网\s*址|电\s*话"
    r"|传\s*真|邮\s*政\s*编\s*码|开\s*立\s*人)"
)
# "短标签+冒号"开头(单位性质：/成立时间：/性别：…)=表单行,同样不许被吞并
_LABEL_START_RE = re.compile(r"^\s*[一-龥]{1,6}\s*[：:]")


def _is_form_line(text: str) -> bool:
    """短且带冒号的"标签：值"表单行。它天生不带句号,绝不能被句子合并当成
    "没说完的话"往后吞——实测回归:"投标人：公司名"吞掉"单位性质：…"、
    "地址：…"连吞"成立时间/经营期限",法定代表人身份证明整页挤成一坨。"""
    t = (text or "").strip()
    return bool(t) and len(t) <= 40 and ("：" in t or ":" in t)


def _para_alignment(p_el: Any) -> str | None:
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        return None
    jc = ppr.find(qn("w:jc"))
    return jc.get(qn("w:val")) if jc is not None else None


def _para_left_indent(p_el: Any) -> int:
    ppr = p_el.find(qn("w:pPr"))
    ind = ppr.find(qn("w:ind")) if ppr is not None else None
    try:
        return int(ind.get(qn("w:left")) or 0) if ind is not None else 0
    except ValueError:
        return 0


def _p_text(p_el: Any) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def _unify_run_format(run_el: Any, host_rpr: Any) -> None:
    """把并入 run 的字体/字号/加粗统一成宿主段的(同一句话不能两种脸);
    下划线/颜色等保留——填空值的下划线不能被抹掉。host_rpr=None 表示宿主继承默认,
    则删掉并入 run 的这三项让它同样继承默认。"""
    rpr = run_el.find(qn("w:rPr"))
    for tag in ("w:rFonts", "w:sz", "w:szCs", "w:b", "w:bCs"):
        host_val = host_rpr.find(qn(tag)) if host_rpr is not None else None
        cur = rpr.find(qn(tag)) if rpr is not None else None
        if host_val is None:
            if cur is not None:
                rpr.remove(cur)
        else:
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run_el.insert(0, rpr)
                cur = None
            if cur is not None:
                rpr.remove(cur)
            rpr.append(deepcopy(host_val))


def _in_textbox(el: Any) -> bool:
    """元素是否在浮动文本框(txbxContent)内——那里的排版不归我们管。"""
    anc = el.getparent()
    while anc is not None:
        if anc.tag.endswith("}txbxContent"):
            return True
        anc = anc.getparent()
    return False


def heal_split_paragraphs(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治福昕"一句话劈成几段"(用户实测投标函:一句被切三截、中间夹空行、粗细不一)。

    福昕把原 PDF 的视觉行转成独立段落:段尾停在逗号/裸字上、下一段接着说。规则:
    左对齐正文段结尾不是句末标点 → 把下一段并回来(runs 搬家,文字一个不丢),
    中间的零字符空段删掉,并入文字的字体/字号/加粗统一成宿主段的(下划线保留)。
    绝不并:编号/表单标签开头的下一段、居中/右对齐段、大缩进标题段、带分节符段、表格。
    正文和表格单元格内都治(埇桥p8/p9实测:附录表值格里一句话碎成4段夹空行)。
    """
    merged = 0
    containers = [document.element.body]
    for tc in document.element.body.iter(qn("w:tc")):
        if not _in_textbox(tc):
            containers.append(tc)
    for container in containers:
        merged += _merge_split_in(container)
    if merged:
        logger.info("格式体检:合并福昕劈开的半句 %d 处", merged)
    return merged


def _merge_split_in(body: Any) -> int:
    merged = 0
    children = list(body.iterchildren())
    i = 0
    while i < len(children):
        el = children[i]
        if el.tag != qn("w:p") or el.getparent() is None:
            i += 1
            continue
        text = _p_text(el).rstrip()
        if (
            not text.strip()
            or text[-1] in _SENTENCE_END
            or text.lstrip().startswith("致")  # "致：xxx"抬头独立成行,不许吞下一段
            or _is_form_line(text)  # 表单行不当宿主往后吞(短+带冒号)
            or _para_alignment(el) in ("center", "right", "end")
            or _para_left_indent(el) > 1500  # 视觉居中的标题段
            or el.find(f".//{qn('w:sectPr')}") is not None
        ):
            i += 1
            continue
        # 宿主段末个带文字 run 的格式 = 本句的"标准脸"
        host_rpr = None
        for r in reversed(el.findall(qn("w:r"))):
            if _p_text(r).strip():
                host_rpr = r.find(qn("w:rPr"))
                break
        # 往后找可并的段,最多接 8 截
        joins = 0
        j = i + 1
        while joins < 8 and j < len(children):
            nxt = children[j]
            if nxt.tag != qn("w:p"):
                break  # 表格/其他元素,这句到头了
            if nxt.find(f".//{qn('w:sectPr')}") is not None:
                break
            nxt_text = _p_text(nxt)
            if not nxt_text.strip():
                # 夹在半句中间的空段:零字符的删掉,含空白字符的把空白并进句里
                if not nxt_text:
                    has_content = any(
                        next(iter(nxt.iter(qn(t))), None) is not None
                        for t in ("w:drawing", "w:pict", "w:br", "w:tab")
                    )
                    if has_content:
                        break
                    body.remove(nxt)
                    j += 1
                    continue
                for r in list(nxt.findall(qn("w:r"))):
                    el.append(r)
                body.remove(nxt)
                j += 1
                continue
            if (
                _NEW_ITEM_RE.match(nxt_text)
                or _LABEL_START_RE.match(nxt_text)  # "单位性质：/性别：…"表单行不许被吞
                or _para_alignment(nxt) in ("center", "right", "end")
                or _para_left_indent(nxt) > 1500
            ):
                break
            # 并:runs 搬进宿主,统一字体,删空壳
            for r in list(nxt.findall(qn("w:r"))):
                _unify_run_format(r, host_rpr)
                el.append(r)
            body.remove(nxt)
            merged += 1
            joins += 1
            j += 1
            new_text = _p_text(el).rstrip()
            if new_text and new_text[-1] in _SENTENCE_END:
                break  # 句子说完了
        i = j if joins else i + 1
    return merged


def heal_midsentence_breaks(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治福昕的"句中硬换行":一段里句子说到一半插 w:br 断行(还常连打两个变出空行)。

    (5)条实测:整条和后续几条都在同一段里,句中"…其他管[br][br]理和技术人员…"——
    段落合并治不了它。规则:左对齐正文段里,**前文没说完**(累计文字结尾不是句末标点)
    的换行符删掉,句子自然连排;句末标点后的换行保留(它是条目分隔);"致：…"抬头行的
    换行保留;分页符(type=page)和浮动文本框内部绝不碰。正文和表格单元格都治
    (埇桥p8/p9实测:附录表值格里句中断行碎成渣)。只删换行符,文字一个不动。
    """
    removed = 0
    for p_el in document.element.body.iter(qn("w:p")):
        if _in_textbox(p_el):
            continue
        if (
            _para_alignment(p_el) in ("center", "right", "end")
            or _para_left_indent(p_el) > 1500
            or p_el.find(f".//{qn('w:sectPr')}") is not None
        ):
            continue
        # 只走段落直属的 run(不钻进浮动文本框 w:drawing/txbxContent)
        prefix = ""
        for r in list(p_el.iterchildren(qn("w:r"))):
            for child in list(r.iterchildren()):
                if child.tag == qn("w:t"):
                    prefix += child.text or ""
                elif child.tag == qn("w:br"):
                    br_type = child.get(qn("w:type"))
                    if br_type and br_type != "textWrapping":
                        continue  # 分页/分栏符,神圣不可侵犯
                    tail = prefix.rstrip()
                    # 当前行(上一个换行之后的文字)是"致：/致:"抬头或表单行 → 换行是
                    # 格式,保留(实测回归:"致：xx中心"和"我公司…"连成一句;
                    # "投标人：公司名"和"单位性质：…"挤成一行)
                    cur_line = prefix.rsplit("\n", 1)[-1].strip()
                    if cur_line.startswith("致") or _is_form_line(cur_line):
                        prefix += "\n"
                        continue
                    if not tail or tail[-1] not in _SENTENCE_END:
                        r.remove(child)
                        removed += 1
                    else:
                        prefix += "\n"
    if removed:
        logger.info("格式体检:删福昕句中硬换行 %d 个", removed)
    return removed


def heal_phantom_images(document: Any, profile: dict[str, Any] | None = None) -> int:
    """删福昕的"幽灵小图":1x1 像素透明 PNG 被拉成 19x7pt 嵌在句子中间,把行切断。

    埇桥实测:同一张 86 字节透明图(rId10)全文埋了 10 处,投标函(5)条"…其他管[图]
    理和技术人员…"的断行元凶就是它——文字/换行/控制符全查过,最后 fitz 页面对象
    坐标锁定它。判定双保险:显示尺寸微小(高<10pt 且宽<26pt)**且**图片字节<1KB,
    才认定是残渣;真图(证照扫描几十 KB 起)绝不误伤。删图不删字。
    """
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    r_embed = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    removed = 0
    try:
        rels = document.part.rels
    except Exception:
        rels = {}
    for d in list(document.element.body.iter(qn("w:drawing"))):
        node = d.find(f"{{{wp_ns}}}inline")
        if node is None:
            node = d.find(f"{{{wp_ns}}}anchor")
        if node is None:
            continue
        ext = node.find(f"{{{wp_ns}}}extent")
        if ext is None:
            continue
        try:
            cx_pt = int(ext.get("cx")) / 914400 * 72
            cy_pt = int(ext.get("cy")) / 914400 * 72
        except (TypeError, ValueError):
            continue
        if cy_pt >= 10 or cx_pt >= 26:
            continue  # 不是微型图
        if "".join(t.text or "" for t in d.iter(qn("w:t"))).strip():
            continue  # 带文字的(文本框)不碰
        blip = d.find(f".//{{{a_ns}}}blip")
        if blip is None:
            continue  # 无图纯形状(边框线等)不碰
        rid = blip.get(r_embed)
        try:
            blob_len = len(rels[rid].target_part.blob) if rid in rels else 10**9
        except Exception:
            blob_len = 10**9
        if blob_len >= 1024:
            continue  # 真图不碰
        holder = d.getparent()  # w:r
        if holder is not None and holder.getparent() is not None:
            holder.remove(d)
            if not _p_text(holder).strip() and holder.find(qn("w:br")) is None:
                holder.getparent().remove(holder)  # 空壳 run 一并清掉
            removed += 1
    if removed:
        logger.info("格式体检:删福昕幽灵小图 %d 个", removed)
    return removed


_SIG_SECTION_MARKERS = ("盖单位章", "法定代表人", "签字或盖章", "签章", "日期", "日 期", "日  期")

# 封面日期行:"2026年7月8日" 或 未填的 "__年__月__日"
_COVER_DATE_RE = re.compile(r"^[\s\d_＿]*年[\s\d_＿]*月[\s\d_＿]*日\s*$")


def _flatten_cols(cols: Any) -> None:
    """把分栏元素拉回单栏:num=1,删掉每栏 <w:col> 定义。"""
    cols.set(qn("w:num"), "1")
    for c in list(cols.findall(qn("w:col"))):
        cols.remove(c)
    cols.set(qn("w:equalWidth"), "1")


def _center_paragraph(p_el: Any) -> None:
    """段落居中,并清掉左/首行/右缩进(福昕给封面标题留的巨缩进)。"""
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p_el.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), "center")
    ind = ppr.find(qn("w:ind"))
    if ind is not None:
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:firstLine"), "0")
        ind.set(qn("w:right"), "0")


def _set_space_before(p_el: Any, twips: int) -> None:
    """设段前距(封面竖向铺开用)。spacing 须在 pPr 里 rPr 之后、ind/jc 之前的合适位置;
    已有 spacing 就改属性,没有则新建插到 pPr 头部(rPr 之后)。"""
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p_el.insert(0, ppr)
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        rpr = ppr.find(qn("w:rPr"))
        if rpr is not None:
            rpr.addnext(sp)
        else:
            ppr.insert(0, sp)
    sp.set(qn("w:before"), str(twips))


def heal_cover_columns(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治福昕把商务/技术卷**封面**切成假两栏:标题被塞进窄左栏(整页只用一半宽、挤成4行)、
    "标段招标""（盖单位章）"甩到右栏飘着——用户实测"乱七八糟"。

    根因同落款(福昕拿两栏并排复刻原 PDF 的横向位置):原版封面本是
        （招标项目名称）____标段招标      ← 一行,项目名+标段招标并排
        投标人：____（盖单位章）          ← 一行
    福昕转出来把每行的左右两半各塞进两栏的左/右栏。项目名短时(原版占位符)一行排得下;
    换成真项目名(长)就在窄左栏里挤成三四行,右栏的"标段招标"孤零零飘着。

    修法(只在**明确识别出标准封面**时动手,非封面一律不碰):
      ① 封面区 = 文档开头到第一个"下一页"分节(通常是目录起点)之前,且必须同时含
         "标段招标"+"投标文件"+"盖单位章"三特征,否则判定不是标准封面直接返回;
      ② 封面区内每个两栏小节:两段非空的=被拆开的一对,把右半 run 合并进左半、删中间空段、
         拉直单栏、整行居中(还原"项目名____标段招标"一行);一段非空的直接拉直;
      ③ 合并会让内容上移挤在页顶 → 给"投标文件"行和日期行补段前距,让封面竖向铺开(对齐原版)。
    只挪 run/改分栏对齐/加间距,一个字不改(红线)。返回处理的封面小节数。
    """
    body = document.element.body
    kids = list(body.iterchildren())

    def _sectpr_of(ch: Any) -> Any:
        if ch.tag == qn("w:sectPr"):
            return ch
        if ch.tag == qn("w:p"):
            ppr = ch.find(qn("w:pPr"))
            if ppr is not None:
                return ppr.find(qn("w:sectPr"))
        return None

    # ① 定位封面结束:第一个非 continuous 分节(默认/nextPage=翻页,通常目录起点)
    cover_end = None
    for i, ch in enumerate(kids):
        sectpr = _sectpr_of(ch)
        if sectpr is None:
            continue
        typ = sectpr.find(qn("w:type"))
        typv = typ.get(qn("w:val")) if typ is not None else "nextPage"
        if typv != "continuous":
            cover_end = i
            break
    if cover_end is None:
        return 0
    cover_paras = [
        kids[j] for j in range(0, cover_end + 1) if kids[j].tag == qn("w:p")
    ]
    cover_text = "".join(_p_text(p) for p in cover_paras)
    if not ("标段招标" in cover_text and "盖单位章" in cover_text and "投标文件" in cover_text):
        return 0  # 不是标准封面,红线:不碰

    fixed = 0
    prev = -1
    for i in range(0, cover_end + 1):
        sectpr = _sectpr_of(kids[i])
        if sectpr is None:
            continue
        cols = sectpr.find(qn("w:cols"))
        try:
            num = int(cols.get(qn("w:num")) or "1") if cols is not None else 1
        except ValueError:
            num = 1
        if cols is None or num < 2:
            prev = i
            continue
        seg_paras = [
            kids[j] for j in range(prev + 1, i + 1) if kids[j].tag == qn("w:p")
        ]
        nonempty = [p for p in seg_paras if _p_text(p).strip()]
        if len(nonempty) == 2:
            # 被拆开的一对:右半合并进左半
            host, donor = nonempty[0], nonempty[1]
            for r in donor.findall(qn("w:r")):
                host.append(r)  # lxml append 把 run 从 donor 移到 host 行尾
            _center_paragraph(host)
            hi, di = seg_paras.index(host), seg_paras.index(donor)
            for k in range(hi + 1, di):  # 删两半之间的空 br 段
                mid = seg_paras[k]
                if not _p_text(mid).strip() and mid.getparent() is not None:
                    mid.getparent().remove(mid)
            _flatten_cols(cols)
            fixed += 1
        elif len(nonempty) <= 1:
            _flatten_cols(cols)
            fixed += 1
        prev = i

    # ③ 竖向铺开:投标文件行、日期行补段前距(合并后内容会挤在页顶)
    if fixed:
        for p in cover_paras:
            if p.getparent() is None:
                continue
            t = _p_text(p).strip()
            if t == "投标文件":
                _set_space_before(p, 2400)
            elif _COVER_DATE_RE.match(t):
                _set_space_before(p, 2400)

    if fixed:
        logger.info("格式体检:理顺福昕给封面误造的假两栏 %d 节", fixed)
    return fixed


def heal_signature_columns(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治福昕把落款/日期行单独塞进"假两栏"小节,导致日期在整页有大片空白时还折行。

    根因(巢湖实测,查了半天才逮到):福昕复刻原 PDF 时,会给"日期：__年__月__日"
    这一行单独起一个 `<w:sectPr>` 且设成 `<w:cols w:num="2">`——第一栏只有 5336
    twips(≈267pt)宽。日期文字左缩进 151pt 起、排到 319pt 正好撞上这窄栏的右边界,
    "日"就被挤到下一行。跟字宽/缩进/下划线全无关,纯是这个假两栏把整页 488pt 压成 267pt。

    修法:凡是**辖段很少(≤8 段非空)且内容含落款/日期专属标记**的多栏小节,一律拉回
    单栏(删掉 `<w:col>` 子元素、num 置 1)。范围极窄:大段真两栏正文(辖段多)一律不碰,
    只逮福昕给落款/日期误造的这种小节。只改分栏,一个字不动(红线)。返回改回单栏的节数。
    """
    body = document.element.body
    fixed = 0
    prev = -1
    kids = list(body.iterchildren())
    for i, ch in enumerate(kids):
        sectpr = None
        if ch.tag == qn("w:sectPr"):
            sectpr = ch
        elif ch.tag == qn("w:p"):
            ppr = ch.find(qn("w:pPr"))
            if ppr is not None:
                sectpr = ppr.find(qn("w:sectPr"))
        if sectpr is None:
            continue
        cols = sectpr.find(qn("w:cols"))
        try:
            num = int(cols.get(qn("w:num")) or "1") if cols is not None else 1
        except ValueError:
            num = 1
        if cols is None or num < 2:
            prev = i
            continue
        # 该节所辖:prev+1 .. i,统计非空段与文字
        seg_paras = [
            kids[j] for j in range(prev + 1, i + 1) if kids[j].tag == qn("w:p")
        ]
        nonempty = [p for p in seg_paras if _p_text(p).strip()]
        seg_text = "".join(_p_text(p) for p in nonempty)
        has_sig = any(m in seg_text for m in _SIG_SECTION_MARKERS)
        if has_sig and len(nonempty) <= 8:
            cols.set(qn("w:num"), "1")
            for c in list(cols.findall(qn("w:col"))):
                cols.remove(c)
            cols.set(qn("w:equalWidth"), "1")
            fixed += 1
        prev = i
    if fixed:
        logger.info("格式体检:拉直福昕给落款/日期误造的假两栏 %d 节", fixed)
    return fixed


def heal_signature_block_layout(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治福昕把落款3行(投标人/法定代表人/日期)搞乱,只认落款专属标记,范围极窄。

    招标原样(巢湖实测):
        投 标 人：____（盖单位章）
        法定代表人：____（签字或盖章）
        日  期：__年__月__日
    福昕转出来:①投标人行和法定代表人行被并进**同一段**(填了公司名后还折行);
    ②"日 期："被切成"日"一段、"期：__年__月__日"另一段(中间还夹空段)。
    两处修复只挪 run/断段,一个字不改:
      ① 一段里同时含 盖单位章 + 法定代表人 + (签字或盖章|签章) → 在"法定代表人"run处拆成两段;
      ② 整段就是"日" + 后面(跳空段)首段以"期"打头 → 把"期…"并回"日"段同一行。
    """
    body = document.element.body
    avail = _usable_width_twips(document)
    fixed = 0

    # ① 拆:投标人行 | 法定代表人行
    split_groups: list[list[Any]] = []
    for p_el in list(body.iterchildren(qn("w:p"))):
        text = _p_text(p_el)
        if not (
            "盖单位章" in text
            and "法定代表人" in text
            and ("签字或盖章" in text or "签章" in text)
        ):
            continue
        runs = p_el.findall(qn("w:r"))
        split_idx = next(
            (
                i
                for i, r in enumerate(runs)
                if "".join(t.text or "" for t in r.findall(qn("w:t"))).startswith("法定")
            ),
            None,
        )
        if not split_idx:  # None 或 0 都不拆
            continue
        new_p = OxmlElement("w:p")
        ppr = p_el.find(qn("w:pPr"))
        if ppr is not None:
            new_p.append(deepcopy(ppr))
        for r in runs[split_idx:]:
            new_p.append(r)  # lxml append 把 r 从原段移走
        # 投标人行尾残留的纯空白 run 去掉
        for r in reversed(p_el.findall(qn("w:r"))):
            rt = "".join(t.text or "" for t in r.findall(qn("w:t")))
            if rt.strip() == "" and r.find(qn("w:tab")) is None and r.find(qn("w:br")) is None:
                p_el.remove(r)
            else:
                break
        p_el.addnext(new_p)
        split_groups.append([p_el, new_p])
        fixed += 1

    # ② 合:"日" + "期：…" → 同一行
    kids = list(body.iterchildren())
    for i, el in enumerate(kids):
        if el.tag != qn("w:p") or _p_text(el).strip() != "日":
            continue
        j = i + 1
        while j < len(kids) and kids[j].tag == qn("w:p") and not _p_text(kids[j]).strip():
            j += 1
        if (
            j >= len(kids)
            or kids[j].tag != qn("w:p")
            or not _p_text(kids[j]).strip().startswith("期")
        ):
            continue
        host, donor = el, kids[j]
        for r in donor.findall(qn("w:r")):
            host.append(r)  # 期… 并回"日"段行尾
        for k in range(i + 1, j + 1):
            if kids[k].getparent() is not None:
                kids[k].getparent().remove(kids[k])
        fixed += 1

    # ③ 落款块统一左缩进:同一块的 投标人/法代/日期 对齐到同一 x,且缩到能排下最长行
    #    只动"真会折行"的块(某行在当前缩进下排不下),整块对齐;不折的块(如封面)一律不碰。
    _DATE_ONLY_RE = re.compile(r"^\s*\d{4}\s*年\s*\d{1,2}\s*月\s*\d{0,2}\s*日?\s*$")

    def _ntab(p_el: Any) -> int:  # 制表符是独立 w:tab 元素,不在 w:t 文字里,单独数
        return sum(len(r.findall(qn("w:tab"))) for r in p_el.findall(qn("w:r")))

    def _line_need(p_el: Any) -> int:  # 该行文字宽(制表符按紧凑槽宽1000估)
        return _est_text_twips(_p_text(p_el)) + _ntab(p_el) * 1000

    def _left_of(p_el: Any) -> int:
        ppr = p_el.find(qn("w:pPr"))
        ind = ppr.find(qn("w:ind")) if ppr is not None else None
        try:
            return int(ind.get(qn("w:left")) or 0) if ind is not None else 0
        except ValueError:
            return 0

    def _is_sig_line(p_el: Any) -> bool:
        t = _p_text(p_el)
        ts = t.lstrip()
        return (
            ("盖单位章" in ts and ts[:4].startswith("投"))
            or (("签字或盖章" in ts or "签章" in ts) and ts.startswith("法定"))
            or ts.startswith(("日期", "日 期", "日  期"))
            or bool(_DATE_ONLY_RE.match(t))  # 光日期行"2026年 7月 8日"(填好的,无标签)
        )

    sig_paras = list(body.iterchildren(qn("w:p")))
    i = 0
    while i < len(sig_paras):
        if not _is_sig_line(sig_paras[i]):
            i += 1
            continue
        block = [sig_paras[i]]
        k = i + 1
        while k < len(sig_paras):
            if _is_sig_line(sig_paras[k]):
                block.append(sig_paras[k])
                k += 1
            elif not _p_text(sig_paras[k]).strip():
                k += 1  # 跳块内空段
            else:
                break
        # 只在**至少一行会折**时才动整块。留 400 余量:字宽是估算,宁可多修不可漏(实测
        # 埇桥"日期：2026年7月8日"带标签的行就卡在临界点上折了)。
        wraps = any(_left_of(p) + _line_need(p) > avail - 400 for p in block)
        if wraps:
            fixed += 1
            max_need = max(_line_need(p) for p in block)
            # 目标缩进:既排得下最长行(再留500余量),又不超过块里原最小缩进(不把任何行往右推)
            target = max(0, min(min(_left_of(p) for p in block), avail - max_need - 500))
            for p in block:
                ppr = p.find(qn("w:pPr"))
                if ppr is None:
                    ppr = OxmlElement("w:pPr")
                    p.insert(0, ppr)
                ind = ppr.find(qn("w:ind"))
                if ind is None:
                    ind = OxmlElement("w:ind")
                    ppr.append(ind)
                ind.set(qn("w:left"), str(target))
                # 日期行的制表符默认跳页面绝对制表位→撑宽;给紧凑自定义制表位
                ntab = _ntab(p)
                if ntab and _p_text(p).lstrip().startswith(("日期", "日 期", "日  期")):
                    old = ppr.find(qn("w:tabs"))
                    if old is not None:
                        ppr.remove(old)
                    tabs = OxmlElement("w:tabs")
                    for m in range(ntab):
                        tb = OxmlElement("w:tab")
                        tb.set(qn("w:val"), "left")
                        tb.set(qn("w:pos"), str(target + 1200 + m * 1150))
                        tabs.append(tb)
                    anchor_el = ppr.find(qn("w:spacing")) or ppr.find(qn("w:ind"))
                    if anchor_el is not None:
                        anchor_el.addprevious(tabs)
                    else:
                        ppr.append(tabs)
        i = k

    if fixed:
        logger.info("格式体检:理顺福昕搞乱的落款3行 %d 处", fixed)
    return fixed


# ↓↓↓ 成品级 healer:只在**拼卷后的成品商务卷**上跑(核对表/业绩证据图是拼卷才加的,
# 格式副本阶段没有,所以不进 _HEALERS,由 run_format_doctor_assembled 单独调)。

# 合规自查核对表列宽(7列,合计≈9765twips):核对项/出处/招标要求/我方取值 给足文字宽,
# 判定/处置(✅一致、填/留空)收窄,备注最宽。等宽会把长文本挤成4字/行的细条(用户实测)。
_CHECKLIST_HEADERS = ("核对项", "判定", "处置")
_CHECKLIST_COL_WIDTHS = (1150, 1500, 1450, 1750, 900, 780, 2235)


def heal_checklist_table_widths(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治合规自查核对表(markdown 转表默认等宽)的窄列:按内容重分配列宽,长文本列不再竖成细条。

    只认表头含 核对项+判定+处置 且正好 7 列的那张系统自查表,范围极窄;别的表一律不碰。
    只改列宽(tblLayout=fixed + gridCol + tcW),一个字不动(红线)。返回修的表数。
    """
    fixed = 0
    for table in document.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if not all(h in header for h in _CHECKLIST_HEADERS):
            continue
        if len(table.columns) != len(_CHECKLIST_COL_WIDTHS):
            continue
        widths = _CHECKLIST_COL_WIDTHS
        tbl = table._tbl
        tblpr = tbl.find(qn("w:tblPr"))
        if tblpr is None:
            tblpr = OxmlElement("w:tblPr")
            tbl.insert(0, tblpr)
        lay = tblpr.find(qn("w:tblLayout"))
        if lay is None:
            lay = OxmlElement("w:tblLayout")
            tblpr.append(lay)
        lay.set(qn("w:type"), "fixed")
        tw = tblpr.find(qn("w:tblW"))
        if tw is None:
            tw = OxmlElement("w:tblW")
            tblpr.append(tw)
        tw.set(qn("w:w"), str(sum(widths)))
        tw.set(qn("w:type"), "dxa")
        grid = tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
                gc.set(qn("w:w"), str(w))
        for row in table.rows:
            for cell, w in zip(row.cells, widths):
                tcpr = cell._tc.find(qn("w:tcPr"))
                if tcpr is None:
                    tcpr = OxmlElement("w:tcPr")
                    cell._tc.insert(0, tcpr)
                tcw = tcpr.find(qn("w:tcW"))
                if tcw is None:
                    tcw = OxmlElement("w:tcW")
                    tcpr.append(tcw)
                tcw.set(qn("w:w"), str(w))
                tcw.set(qn("w:type"), "dxa")
        fixed += 1
    if fixed:
        logger.info("格式体检:重排合规核对表列宽 %d 张", fixed)
    return fixed


_EVIDENCE_CAPTION_MARKERS = (
    "中标通知书", "合同", "证书", "营业执照", "资质", "安全生产许可", "业绩证明",
)


def _has_big_image(el: Any) -> bool:
    """段落含"大图"(显示高>150pt=业绩扫描件,非页眉细线/句中小图)。"""
    if el.tag != qn("w:p"):
        return False
    for ext in el.iter():
        if ext.tag.endswith("}extent"):
            try:
                if int(ext.get("cy") or 0) / 12700 > 150:
                    return True
            except (TypeError, ValueError):
                pass
    return False


def heal_evidence_caption_binding(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治业绩证据"扫描图+图注"被分页拆散/图注错配:给图段加 keepNext,让图与它下方图注同页。

    结构(实测):一张业绩扫描图(大图段) 紧跟一行图注(如"…合同（1）")。分页时图注会漂到
    别页、或某页出现"上一张图+下一张图注"的错配。给**图段**加 keepNext,把图和它的图注绑住。
    只认紧跟证据图注(含中标通知书/合同/证书等标记)的大图段,范围极窄。只加段属性,不改文字。
    """
    body = document.element.body
    kids = list(body.iterchildren())
    fixed = 0
    for i, el in enumerate(kids):
        if not _has_big_image(el):
            continue
        j = i + 1
        while (
            j < len(kids)
            and kids[j].tag == qn("w:p")
            and not _p_text(kids[j]).strip()
            and not _has_big_image(kids[j])
        ):
            j += 1
        if j >= len(kids) or kids[j].tag != qn("w:p"):
            continue
        cap = _p_text(kids[j]).strip()
        if not cap or len(cap) > 120:
            continue
        if not any(m in cap for m in _EVIDENCE_CAPTION_MARKERS):
            continue
        for k in range(i, j):  # 图段+中间空段都 keepNext,链住到图注
            if kids[k].tag != qn("w:p"):
                continue
            ppr = kids[k].find(qn("w:pPr"))
            if ppr is None:
                ppr = OxmlElement("w:pPr")
                kids[k].insert(0, ppr)
            if ppr.find(qn("w:keepNext")) is None:
                ppr.insert(0, OxmlElement("w:keepNext"))
        fixed += 1
    if fixed:
        logger.info("格式体检:绑定业绩证据图与图注(keepNext) %d 处", fixed)
    return fixed


def heal_table_row_integrity(document: Any, profile: dict[str, Any] | None = None) -> int:
    """给所有表格行加"禁止跨页拆分"(w:cantSplit):一行的内容不再被劈到两页
    (员工反馈第3条"一页内容被分段成多页"最扎眼的形态)。装不下的整行挪到下页;
    比整页还高的行 Word 会照常拆,不会溢出丢内容。返回加固的行数。"""
    fixed = 0
    for table in document.tables:
        for row in table.rows:
            tr = row._tr
            trPr = tr.find(qn("w:trPr"))
            if trPr is None:
                trPr = OxmlElement("w:trPr")
                tr.insert(0, trPr)
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(OxmlElement("w:cantSplit"))
                fixed += 1
    if fixed:
        logger.info("格式体检:表格行禁止跨页拆分(cantSplit) %d 行", fixed)
    return fixed


def run_format_doctor_assembled(document: Any) -> dict[str, int]:
    """成品级体检:只跑拼卷后成品才需要的 healer(核对表列宽、业绩图注绑定、
    表格行防拆页)。逐个容错,单个崩不阻断出标。"""
    report: dict[str, int] = {}
    for name, healer in (
        ("template_header_lines", heal_template_header_lines),
        ("checklist_table_widths", heal_checklist_table_widths),
        ("evidence_caption_binding", heal_evidence_caption_binding),
        ("table_row_integrity", heal_table_row_integrity),
    ):
        try:
            report[name] = healer(document, None)
        except Exception:
            logger.warning("成品级格式体检 healer %s 失败,跳过", name, exc_info=True)
            report[name] = 0
    fixed = {k: v for k, v in report.items() if v}
    if fixed:
        logger.info("成品级格式体检修复: %s", fixed)
    return report


def heal_idproof_column_pairs(document: Any, profile: dict[str, Any] | None = None) -> int:
    """治福昕把身份证明的两栏拆成前后段:姓名/年龄一段(内含换行),性别/职务各自孤段。

    2026-08-05 用户实测:填完后"性别：女"悬在姓名和年龄之间的右侧、"职务：总经理"又
    低半行,Word里还改不了(段落缩进定位)。规则极窄:段落文本形如"姓名：X⏎年龄：Y",
    其后4段内出现孤段"性别：Z"/"职务：W"(短文本) → 把它们并回对应行
    (姓名行尾接"　性　别：Z",年龄行尾接"　职　务：W"),孤段删除。文字一字不丢。
    """
    from docx.oxml import OxmlElement

    healed = 0
    paras = list(_iter_all_paragraphs(document))
    for idx, para in enumerate(paras):
        t = re.sub(r"[\s　]+", "", para.text)
        if not ("姓名：" in t and "年龄：" in t) or "性别" in t:
            continue
        runs = [r for r in para.runs]
        br_pos = None  # 姓名行结尾的换行 run
        for k, r in enumerate(runs):
            if r._r.find(qn("w:br")) is not None:
                br_pos = k
                break
        if br_pos is None:
            continue
        # 找孤段 性别/职务(≤4段内,允许中间空段)
        victims: dict[str, tuple[Any, str]] = {}
        for nxt in paras[idx + 1: idx + 5]:
            nt = re.sub(r"[\s　]+", "", nxt.text)
            if not nt:
                continue
            m = re.match(r"^(性别|职务)[：:](.{1,12})$", nt)
            if m and m.group(1) not in victims:
                victims[m.group(1)] = (nxt, m.group(2))
            elif not m:
                break
        if not victims:
            continue

        def _mk_run(text: str, template_r: Any) -> Any:
            new_r = OxmlElement("w:r")
            rpr = template_r._r.find(qn("w:rPr"))
            if rpr is not None:
                from copy import deepcopy as _dc2

                new_r.append(_dc2(rpr))
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve")
            t_el.text = text
            new_r.append(t_el)
            return new_r

        if "性别" in victims:
            _p, val = victims["性别"]
            ref = runs[br_pos - 1] if br_pos > 0 else runs[br_pos]
            # 插在换行 run 之前 → 落在姓名行行尾
            runs[br_pos]._r.addprevious(_mk_run(f"　　性　别：{val}", ref))
            _p._p.getparent().remove(_p._p)
            healed += 1
        if "职务" in victims:
            _p, val = victims["职务"]
            ref = runs[-1]
            para._p.append(_mk_run(f"　　职　务：{val}", ref))
            _p._p.getparent().remove(_p._p)
            healed += 1
    if healed:
        logger.info("格式体检:身份证明两栏并回 %d 处(性别/职务归位)", healed)
    return healed


# (名称, healer)。healer 契约:输入 (document, profile),返回修复数;只改格式,绝不改文字。
_HEALERS: tuple[tuple[str, Callable[[Any, dict[str, Any] | None], int]], ...] = (
    ("template_header_lines", heal_template_header_lines),
    ("section_title_page_breaks", heal_section_title_page_breaks),
    ("underline_slots", heal_underline_slots),
    ("filled_value_char_squeeze", heal_filled_value_char_squeeze),
    ("filler_blank_runs", heal_filler_blank_runs),
    ("phantom_images", heal_phantom_images),
    ("split_paragraphs", heal_split_paragraphs),
    ("midsentence_breaks", heal_midsentence_breaks),
    ("cover_columns", heal_cover_columns),
    ("signature_columns", heal_signature_columns),
    ("signature_block_layout", heal_signature_block_layout),
    ("line_spacing", heal_line_spacing),
    ("signature_wrap", heal_signature_wrap),
    ("idproof_column_pairs", heal_idproof_column_pairs),
)


# 招标本身就是 Word(原样复制路径)时能跑的 healer——**只留与福昕无关的**。
# 上面 _HEALERS 里绝大多数是修福昕转换伪影的(劈句/句中硬换行/假两栏/幽灵图/填充空段…),
# 原生 Word 文档没有这些毛病,跑了反而帮倒忙:实测 split_paragraphs 会把
# "第八章 投标文件格式" 和下一行 "投标文件（商务文件）" 错误地并成一段。
# 这几个是真通用的:清模板页眉行、章节另起一页(用户点名要的)、填空槽下划线补齐、
# 压缩大段连续空段。最后一个名字里带"福昕"是历史原因,实为**招标模板本身**的毛病:
# 巢湖 Word 招标原件的格式章里就塞着 226 个空段(最长一串 37 个)用来凑版面,填值插图后
# 版面一变,这些空段就变成整页整页的白纸(2026-07-29 用户实测"空白页太多")。
_NATIVE_DOCX_HEALERS: tuple[tuple[str, Callable[[Any, dict[str, Any] | None], int]], ...] = (
    ("template_header_lines", heal_template_header_lines),
    ("section_title_page_breaks", heal_section_title_page_breaks),
    ("underline_slots", heal_underline_slots),
    ("filler_blank_runs", heal_filler_blank_runs),
)


def run_format_doctor_native_docx(
    document: Any, profile: dict[str, Any] | None = None
) -> dict[str, int]:
    """原生 Word 招标的格式体检:只跑与福昕无关的 healer。逐个容错,绝不阻断。"""
    report: dict[str, int] = {}
    for name, healer in _NATIVE_DOCX_HEALERS:
        try:
            report[name] = healer(document, profile)
        except Exception:
            logger.warning("格式体检(原生docx) healer %s 失败,跳过", name, exc_info=True)
            report[name] = 0
    fixed = {k: v for k, v in report.items() if v}
    if fixed:
        logger.info("格式体检(原生docx)修复: %s", fixed)
    return report


def run_format_doctor_prefill(document: Any) -> dict[str, int]:
    """填值前的体检(标签修复类):孤字归位等。逐个容错,绝不阻断。"""
    report: dict[str, int] = {}
    for name, healer in _PREFILL_HEALERS:
        try:
            report[name] = healer(document, None)
        except Exception:
            logger.warning("格式体检(填前) healer %s 失败,跳过", name, exc_info=True)
            report[name] = 0
    fixed = {k: v for k, v in report.items() if v}
    if fixed:
        logger.info("格式体检(填前)修复: %s", fixed)
    return report


def run_format_doctor(document: Any, profile: dict[str, Any] | None = None) -> dict[str, int]:
    """跑全部 healer,返回 {名称: 修复数}。逐个容错:单个 healer 崩不阻断其余/出标。"""
    report: dict[str, int] = {}
    for name, healer in _HEALERS:
        try:
            report[name] = healer(document, profile)
        except Exception:
            logger.warning("格式体检 healer %s 失败,跳过", name, exc_info=True)
            report[name] = 0
    fixed = {k: v for k, v in report.items() if v}
    if fixed:
        logger.info("格式体检修复: %s", fixed)
    return report
